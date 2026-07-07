"""
SOC Platform v4
───────────────
• Auto-fetches incidents every 30 seconds once token is verified
• Dashboard tab: severity bars, status bars, latest incidents, assignee chart
• Cleaner sidebar with connection status + refresh countdown
• Incidents tab with severity filter + auto-refresh indicator
• Chat stub (wire LangChain inside chat_respond)
• ChromaDB semantic search & sync
"""

import re
import streamlit as st
import streamlit.components.v1 as components
import requests
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
import time
import os
import base64
from datetime import datetime, timedelta
from typing import Optional
from collections import Counter
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor

try:
    from dotenv import load_dotenv, set_key, find_dotenv
    DOTENV_OK = True
except ImportError:
    DOTENV_OK = False

try:
    import chromadb
    CHROMA_OK = True
except ImportError:
    CHROMA_OK = False

# ── SOC Triage Agent (LangChain) ──────────────────────────────────────────────
from soc_triage_agent import CiscoLLMConfig, soc_triage_chat_respond, _TRIAGE_TRIGGER

def _normalise_llm_url(url: str) -> str:
    """
    Ensure the base URL is safe for ChatOpenAI.
    - Strips trailing slashes
    - Removes any embedded model path (e.g. /models/fdtn-ai/...)
    - Guarantees the URL ends with /v1
    """
    from urllib.parse import urlparse
    url = url.strip().rstrip("/")
    parsed = urlparse(url)
    # If someone pasted a model-specific URL, reduce it to just host + /v1
    if "/models/" in parsed.path:
        url = f"{parsed.scheme}://{parsed.netloc}/v1"
    elif not parsed.path.endswith("/v1"):
        url = url + "/v1"
    return url

def get_cisco_cfg() -> CiscoLLMConfig:
    """
    Build CiscoLLMConfig from session state at call time so that
    credentials entered in the sidebar are always picked up.
    """
    raw_url = st.session_state.get("cisco_url", "").strip()
    url     = _normalise_llm_url(raw_url) if raw_url else "https://api-inference.huggingface.co/v1"
    return CiscoLLMConfig(
        base_url    = url,
        api_key     = st.session_state.get("cisco_key",   "").strip() or "changeme",
        model       = st.session_state.get("cisco_model", "").strip()
                      or "fdtn-ai/Foundation-Sec-8B-Reasoning",
        temperature = 0.0,
        max_tokens  = 1024,
        timeout     = 300,
    )

# ══════════════════════════════════════════════════════════════════════════════
# .ENV  — load persisted credentials on every startup
# ══════════════════════════════════════════════════════════════════════════════
ENV_FILE = Path(__file__).parent / ".env"

def _clean(val: str) -> str:
    """Strip whitespace and stray quotes that dotenv sometimes leaves."""
    return val.strip().strip("'\"").strip()

def env_load() -> dict:
    """Read credentials from .env."""
    if DOTENV_OK and ENV_FILE.exists():
        load_dotenv(ENV_FILE, override=True)
    return {
        "host":          _clean(os.environ.get("NW_HOST",       "")),
        "username":      _clean(os.environ.get("NW_USERNAME",   "")),
        "password":      _clean(os.environ.get("NW_PASSWORD",   "")),
        "nw_cert_path":  _clean(os.environ.get("NW_CERT_PATH",  "")),
        # Cisco LLM
        "cisco_url":     _clean(os.environ.get("CISCO_LLM_URL",   "")),
        "cisco_key":     _clean(os.environ.get("CISCO_LLM_KEY",   "")),
        "cisco_model":   _clean(os.environ.get("CISCO_LLM_MODEL", "")),
    }

def env_save(host: str, username: str, password: str) -> None:
    """Persist NetWitness credentials to .env."""
    if not DOTENV_OK:
        return
    try:
        ENV_FILE.touch(exist_ok=True)
        set_key(str(ENV_FILE), "NW_HOST",     host.strip())
        set_key(str(ENV_FILE), "NW_USERNAME", username.strip())
        # base64-encode password to avoid special char issues
        set_key(str(ENV_FILE), "NW_PASSWORD",
                base64.b64encode(password.strip().encode()).decode("ascii"))
    except Exception:
        pass  # .env locked on Windows — credentials still active this session

def nw_cert_env_save(cert_path: str) -> None:
    """Persist the TLS cert path to .env.
    Silently skips if .env is locked (common on Windows) — cert path
    is still stored in session state and works for the current session.
    """
    if not DOTENV_OK:
        return
    try:
        ENV_FILE.touch(exist_ok=True)
        set_key(str(ENV_FILE), "NW_CERT_PATH", cert_path.strip())
    except Exception:
        pass  # .env locked on Windows — session state still holds the path

def nw_cert_env_clear() -> None:
    if not DOTENV_OK:
        return
    try:
        if ENV_FILE.exists():
            set_key(str(ENV_FILE), "NW_CERT_PATH", "")
    except Exception:
        pass

def cisco_env_save(url: str, key: str, model: str) -> None:
    """Persist Cisco LLM credentials to .env."""
    if not DOTENV_OK:
        return
    try:
        ENV_FILE.touch(exist_ok=True)
        set_key(str(ENV_FILE), "CISCO_LLM_URL",   url.strip())
        set_key(str(ENV_FILE), "CISCO_LLM_MODEL", model.strip())
        # base64-encode key to avoid special char issues
        set_key(str(ENV_FILE), "CISCO_LLM_KEY",
                base64.b64encode(key.strip().encode()).decode("ascii"))
    except Exception:
        pass

def env_clear() -> None:
    if not DOTENV_OK:
        return
    try:
        if ENV_FILE.exists():
            set_key(str(ENV_FILE), "NW_HOST",     "")
            set_key(str(ENV_FILE), "NW_USERNAME", "")
            set_key(str(ENV_FILE), "NW_PASSWORD", "")
    except Exception:
        pass

def cisco_env_clear() -> None:
    if not DOTENV_OK:
        return
    try:
        if ENV_FILE.exists():
            set_key(str(ENV_FILE), "CISCO_LLM_URL",   "")
            set_key(str(ENV_FILE), "CISCO_LLM_KEY",   "")
            set_key(str(ENV_FILE), "CISCO_LLM_MODEL", "")
    except Exception:
        pass

def nw_login(host: str, username: str, password: str) -> tuple[bool, str, str]:
    """
    Login with username/password → returns (ok, message, access_token).
    POST /rest/api/auth/userpass with form-encoded credentials.
    Auto-retries with verify=False if cert verification fails.
    """
    if not host.strip():
        return False, "Host URL is empty — enter https://192.168.x.x", ""
    host = host.strip().strip("'\"").strip()  # remove any stray quotes
    if not username.strip():
        return False, "Username is empty.", ""
    if not password.strip():
        return False, "Password is empty.", ""

    def _attempt(verify):
        return requests.post(
            f"{host.rstrip('/')}/rest/api/auth/userpass",
            data={"username": username, "password": password},
            headers={
                "Content-Type": "application/x-www-form-urlencoded; charset=ISO-8859-1",
                "Accept":       "application/json;charset=UTF-8",
            },
            timeout=15,
            verify=verify,
        )

    try:
        verify = nw_tls_verify()
        try:
            r = _attempt(verify)
        except requests.exceptions.SSLError:
            # Cert didn't work — clear it and retry without verification
            st.session_state.nw_cert_path = ""
            r = _attempt(False)

        if r.status_code == 200:
            data  = r.json()
            token = data.get("accessToken") or data.get("access_token") or ""
            if token:
                return True, "NetWitness connected", token
            return False, f"Login OK but no token in response: {str(data)[:100]}", ""
        else:
            return False, f"HTTP {r.status_code}: {r.text[:200]}", ""
    except requests.exceptions.ConnectionError as e:
        return False, f"Cannot reach {host} — check VPN/network: {str(e)[:120]}", ""
    except requests.exceptions.Timeout:
        return False, "Login timed out — check GP VPN is connected.", ""
    except Exception as e:
        return False, f"Login error: {e}", ""

# ══════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Security Dashboard",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════════════════════════
# CSS
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Share+Tech+Mono&display=swap');

:root {
  --bg:      #04080F;
  --bg1:     #080F1A;
  --bg2:     #0C1624;
  --border:  #162840;
  --accent:  #00D4FF;
  --green:   #0AF0A0;
  --warn:    #FFB700;
  --danger:  #FF3B3B;
  --orange:  #FF7700;
  --muted:   #5A80A0;
  --text:    #C8DCF0;
  --mono:    'Share Tech Mono', monospace;
  --sans:    'Inter', sans-serif;
}

html, body, [class*="css"] {
  font-family: var(--sans);
  background: var(--bg);
  color: var(--text);
  font-size: 14px;
  line-height: 1.6;
}
.main { background: var(--bg); padding-top: 0.5rem; }

/* ── Sidebar ── */
section[data-testid="stSidebar"] > div:first-child {
  background: linear-gradient(180deg, #040810 0%, #06101A 100%);
  border-right: 1px solid #0E1E30;
  padding-top: 1rem;
}

/* ── Headings ── */
h1,h2,h3 {
  font-family: var(--sans);
  font-weight: 700;
  color: var(--accent);
  margin-bottom: 0.4rem;
}
h1 { font-size: 1.4rem !important; letter-spacing: 0.5px; }
h2 { font-size: 1.1rem !important; }
h3 { font-size: 0.95rem !important; }

/* ── Metric cards ── */
[data-testid="metric-container"] {
  background: linear-gradient(135deg, #070D18 0%, #0A1628 100%);
  border: 1px solid var(--border);
  border-radius: 12px;
  padding: 18px 20px;
  box-shadow: 0 2px 16px rgba(0,0,0,0.4);
  transition: all 0.2s;
}
[data-testid="metric-container"]:hover {
  border-color: #1E4060;
  transform: translateY(-2px);
  box-shadow: 0 4px 24px rgba(0,212,255,0.08);
}
[data-testid="metric-container"] label {
  color: var(--muted);
  font-size: 0.72rem;
  font-weight: 600;
  letter-spacing: 0.5px;
  text-transform: uppercase;
  font-family: var(--sans);
}
[data-testid="metric-container"] [data-testid="stMetricValue"] {
  color: var(--accent);
  font-family: var(--mono);
  font-size: 1.8rem;
  font-weight: 400;
}

/* ── Buttons ── */
.stButton > button {
  background: linear-gradient(90deg, #052030, #083050);
  color: var(--accent);
  border: 1px solid #0E4A6A;
  border-radius: 8px;
  font-family: var(--sans);
  font-size: 0.78rem;
  font-weight: 500;
  padding: 8px 16px;
  transition: all 0.15s;
}
.stButton > button:hover {
  background: linear-gradient(90deg, #083050, #0E4A6A);
  box-shadow: 0 0 18px rgba(0,212,255,0.2);
  color: #fff;
  border-color: var(--accent);
  transform: translateY(-1px);
}
.stButton > button:active { transform: translateY(0); }

/* ── Inputs ── */
.stTextInput > div > div > input,
.stTextArea textarea {
  background: #060E1A !important;
  border: 1px solid #1A3050 !important;
  color: var(--text) !important;
  border-radius: 8px;
  font-family: var(--sans);
  font-size: 0.88rem;
  padding: 10px 14px !important;
}
.stTextInput > div > div > input:focus,
.stTextArea textarea:focus {
  border-color: var(--accent) !important;
  box-shadow: 0 0 0 3px rgba(0,212,255,0.1) !important;
}
.stTextInput label, .stTextArea label, .stSelectbox label {
  color: var(--text) !important;
  font-size: 0.82rem !important;
  font-weight: 500 !important;
  margin-bottom: 4px !important;
}
.stSelectbox > div > div {
  background: #060E1A !important;
  border: 1px solid #1A3050 !important;
  color: var(--text) !important;
  border-radius: 8px;
}

/* ── Tabs ── */
.stTabs [data-baseweb="tab-list"] {
  background: transparent;
  border-bottom: 1px solid var(--border);
  gap: 4px;
}
.stTabs [data-baseweb="tab"] {
  font-family: var(--sans);
  font-size: 0.82rem;
  font-weight: 500;
  color: var(--muted);
  border: none;
  padding: 10px 20px;
  background: transparent;
  border-radius: 8px 8px 0 0;
  transition: all 0.15s;
}
.stTabs [data-baseweb="tab"]:hover {
  color: var(--text);
  background: rgba(0,212,255,0.04);
}
.stTabs [aria-selected="true"] {
  color: var(--accent) !important;
  border-bottom: 2px solid var(--accent) !important;
  background: rgba(0,212,255,0.06) !important;
  font-weight: 600 !important;
}

/* ── Cards ── */
.card {
  background: linear-gradient(135deg, #070D18 0%, #08111E 100%);
  border: 1px solid var(--border);
  border-radius: 12px;
  padding: 16px 20px;
  margin: 6px 0;
  transition: all 0.15s;
}
.card:hover {
  border-color: #1E4060;
  transform: translateX(3px);
  box-shadow: 0 2px 16px rgba(0,0,0,0.3);
}
.card-critical { border-left: 4px solid var(--danger) !important; }
.card-high     { border-left: 4px solid var(--orange) !important; }
.card-medium   { border-left: 4px solid var(--warn)   !important; }
.card-low      { border-left: 4px solid var(--green)  !important; }

/* ── Badges ── */
.badge {
  padding: 3px 10px;
  border-radius: 20px;
  font-size: 0.7rem;
  font-weight: 600;
  display: inline-block;
  font-family: var(--sans);
}
.badge-critical { background:#200808; color:var(--danger);  border:1px solid #4A1010; }
.badge-high     { background:#201200; color:var(--orange);  border:1px solid #4A2A00; }
.badge-medium   { background:#181400; color:var(--warn);    border:1px solid #3A3000; }
.badge-low      { background:#001810; color:var(--green);   border:1px solid #004025; }
.badge-info     { background:#001828; color:var(--accent);  border:1px solid #003850; }

/* ── Chat bubbles ── */
.bubble-user {
  background: #080F1A;
  border-left: 3px solid #005C8A;
  padding: 12px 16px;
  border-radius: 0 10px 10px 0;
  margin: 8px 0;
  font-size: 0.9rem;
  line-height: 1.6;
}
.bubble-agent {
  background: #040A12;
  border-left: 3px solid var(--accent);
  padding: 12px 16px;
  border-radius: 0 10px 10px 0;
  margin: 8px 0;
  font-size: 0.9rem;
  line-height: 1.6;
}
.bubble-label {
  font-family: var(--sans);
  font-size: 0.65rem;
  font-weight: 600;
  letter-spacing: 0.5px;
  margin-bottom: 5px;
  color: var(--accent);
  text-transform: uppercase;
}

/* ── Status dots ── */
.dot { display:inline-block; width:8px; height:8px; border-radius:50%; margin-right:7px; vertical-align:middle; }
.dot-green  { background:var(--green);  box-shadow:0 0 8px var(--green);  animation:pulse 2s infinite; }
.dot-red    { background:var(--danger); box-shadow:0 0 6px var(--danger); }
.dot-yellow { background:var(--warn);   box-shadow:0 0 6px var(--warn);   animation:pulse 1s infinite; }
@keyframes pulse { 0%,100%{opacity:1} 50%{opacity:.35} }

/* ── Section labels ── */
.sec-label {
  font-family: var(--sans);
  font-size: 0.7rem;
  font-weight: 600;
  color: #4A7090;
  text-transform: uppercase;
  letter-spacing: 1px;
  margin: 20px 0 10px;
  padding-bottom: 6px;
  border-bottom: 1px solid #0E1E2E;
}

/* ── Stat mini cards ── */
.stat-mini {
  background: #070D18;
  border: 1px solid var(--border);
  border-radius: 10px;
  padding: 14px 16px;
  text-align: center;
  font-family: var(--sans);
}
.stat-mini .val { font-size: 1.5rem; font-weight: 700; color: var(--accent); }
.stat-mini .lbl { font-size: 0.68rem; font-weight: 500; color: var(--muted); margin-top: 3px; text-transform: uppercase; }

/* ── Tooltips / info boxes ── */
.info-box {
  background: #060E1A;
  border: 1px solid #1A3050;
  border-radius: 10px;
  padding: 14px 18px;
  font-size: 0.82rem;
  color: var(--text);
  line-height: 1.6;
  margin: 8px 0;
}
.info-box .title {
  font-weight: 600;
  color: var(--accent);
  margin-bottom: 6px;
  font-size: 0.85rem;
}

/* ── Expanders ── */
.streamlit-expanderHeader {
  font-family: var(--sans) !important;
  font-size: 0.85rem !important;
  font-weight: 500 !important;
  color: var(--text) !important;
}

/* ── Scrollbar ── */
hr { border-color: #0E1E2E; margin: 1.2rem 0; }
::-webkit-scrollbar { width: 5px; }
::-webkit-scrollbar-track { background: var(--bg); }
::-webkit-scrollbar-thumb { background: #1E3050; border-radius: 3px; }
::-webkit-scrollbar-thumb:hover { background: #2A4060; }

/* ── Sidebar logo / app name ── */
.app-logo {
  text-align: center;
  padding: 10px 0 16px;
  border-bottom: 1px solid #0E1E2E;
  margin-bottom: 4px;
}
.app-logo .name {
  font-size: 1rem;
  font-weight: 700;
  color: var(--accent);
  letter-spacing: 1px;
}
.app-logo .sub {
  font-size: 0.65rem;
  color: var(--muted);
  margin-top: 2px;
}
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE  — seed from .env on first load
# ══════════════════════════════════════════════════════════════════════════════
_env = env_load()
# Decode base64 password
if _env["password"]:
    try:
        _env["password"] = base64.b64decode(_env["password"].encode()).decode("utf-8")
    except Exception:
        pass
# Decode base64 Cisco API key
if _env.get("cisco_key"):
    try:
        _env["cisco_key"] = base64.b64decode(_env["cisco_key"].encode()).decode("utf-8")
    except Exception:
        pass

DEFAULTS = {
    "nw_host":          _env["host"],
    "nw_username":      _env["username"],
    "nw_password":      _env["password"],
    "nw_token":         "",       # refreshed on every login
    "nw_verified":      False,
    "nw_msg":           "",
    "nw_working_ep":    "",       # endpoint that worked
    "nw_working_auth":  {},       # auth headers that worked
    "nw_incidents_path":"/rest/api/incidents",   # ← configurable endpoint path
    "nw_auth_style":    "NetWitness-Token",       # ← configurable auth header style
    "endpoint_scan_results": [],   # ← persisted scanner results
    "nw_cert_path":     _env.get("nw_cert_path", ""),  # ← path to CA/server cert for TLS verification
    "incidents":        [],
    "last_fetch":       None,
    "last_full_fetch":  None,   # ← last time a full (non-incremental) fetch ran
    "last_fetch_mode":  None,   # ← "full" | "incremental", shown in diagnostics
    "chat_history":     [],
    "chat_incident":    None,
    "pending_auto_triage": False,   # set by "🩺 Triage" button — auto-runs the pipeline
    "jump_to_ask_tab":     False,   # set by "🩺 Triage" button — switches to Ask a Question tab
    "chroma_client":    None,
    "chroma_col":       None,
    "search_results":   [],
    "_startup_done":    False,
    # ── File upload ──────────────────────────────────────────
    "uploaded_incident": None,
    "uploaded_filename": "",
    # ── Cisco Foundation LLM ─────────────────────────────────
    "cisco_url":         _env.get("cisco_url",   ""),
    "cisco_key":         _env.get("cisco_key",   ""),
    "cisco_model":       _env.get("cisco_model", ""),
    "cisco_connected":   bool(_env.get("cisco_key", "").strip()),
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

REFRESH_INTERVAL     = 30                      # seconds
INCREMENTAL_OVERLAP  = timedelta(minutes=5)    # clock-skew / indexing-lag buffer
FULL_RESYNC_INTERVAL = timedelta(minutes=10)   # periodic ground-truth resync


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def normalise_sev(inc: dict) -> str:
    raw = str(inc.get("riskScore") or inc.get("severity") or "").upper()
    try:
        s = int(raw)
        return "CRITICAL" if s >= 90 else "HIGH" if s >= 70 else "MEDIUM" if s >= 40 else "LOW"
    except ValueError:
        return raw if raw in ("CRITICAL","HIGH","MEDIUM","LOW") else "LOW"

def gp_is_reachable() -> tuple[bool, str]:
    """
    Check if the NW server is reachable via GlobalProtect VPN.
    Simply tries to open a TCP socket to the host:port.
    Returns (reachable, message).
    """
    import socket
    from urllib.parse import urlparse
    host = st.session_state.nw_host.rstrip("/")
    if not host:
        return False, "No host configured."
    try:
        parsed = urlparse(host)
        hostname = parsed.hostname or host
        port     = parsed.port or (443 if parsed.scheme == "https" else 80)
        sock = socket.create_connection((hostname, port), timeout=5)
        sock.close()
        return True, f"VPN reachable — {hostname}:{port}"
    except OSError:
        return False, (
            "Cannot reach server — GlobalProtect may not be connected.\n"
            "Connect GP VPN and try again."
        )

def nw_headers(include_content_type: bool = False) -> dict:
    """
    Build request headers based on the currently selected auth style.
    include_content_type should only be True for POST/PATCH requests that
    send a JSON body. GET requests must NOT include Content-Type or NW
    returns 400 'Unsupported format was supplied'.
    """
    token = st.session_state.nw_token.strip()
    style = st.session_state.get("nw_auth_style", "NetWitness-Token")

    base = {
        "Accept": "application/json;charset=UTF-8",
    }
    if include_content_type:
        base["Content-Type"] = "application/json;charset=UTF-8"

    if style == "NetWitness-Token":
        base["NetWitness-Token"] = token
    elif style == "Bearer":
        base["Authorization"] = f"Bearer {token}"
    elif style == "Cookie":
        base["Cookie"] = f"access_token={token}"
    elif style == "Both":
        base["Authorization"] = f"Bearer {token}"
        base["Cookie"]        = f"access_token={token}"
    else:
        base["NetWitness-Token"] = token

    return base


def nw_tls_verify():
    """
    Returns the value to pass as requests(verify=...).
    - If a valid cert path is configured AND the file exists → try it.
    - If the cert causes any SSL error → silently fall back to False.
    - Default is False (suppressed via urllib3.disable_warnings above).
    """
    cert_path = st.session_state.get("nw_cert_path", "").strip()
    if cert_path and Path(cert_path).is_file():
        return cert_path
    return False


def nw_incidents_url(host: str) -> str:
    """Build the incidents endpoint URL from the configurable path."""
    path = st.session_state.get("nw_incidents_path", "/rest/api/incidents").strip()
    if not path.startswith("/"):
        path = "/" + path
    return f"{host.rstrip('/')}{path}"

def nw_verify_token() -> tuple[bool, str]:
    host  = st.session_state.nw_host.rstrip("/")
    token = st.session_state.nw_token.strip()
    if not host or not token:
        return False, "Enter both Host URL and token."

    def _attempt(verify):
        return requests.get(
            nw_incidents_url(host),
            headers=nw_headers(),
            params={"pageSize": 1, "pageNumber": 0},
            timeout=15,
            verify=verify,
        )
    try:
        try:
            r = _attempt(nw_tls_verify())
        except requests.exceptions.SSLError:
            st.session_state.nw_cert_path = ""
            r = _attempt(False)

        if r.status_code == 200:
            return True, "NetWitness connected"
        elif r.status_code == 403:
            return False, (
                "Access Denied — add 'integration-server.api.access' "
                "permission to the Administrators role in NW."
            )
        elif r.status_code == 401:
            return False, "Token rejected — login again."
        else:
            return False, f"Unexpected response — HTTP {r.status_code}"
    except requests.exceptions.Timeout:
        return False, "Timed out — check GP VPN is connected."
    except Exception as e:
        return False, f"Cannot reach server: {e}"

def nw_fetch_incidents(
    limit: int | None = None, since: str | None = None
) -> tuple[bool, list, str]:
    """
    Returns (ok, items, diagnostic_message).
    ok=True even if items is empty (empty just means no incidents in NW right now).
    ok=False means a real error occurred (auth, network, etc).

    limit=None (default) means load every incident the API has — pagination
    keeps going until the API reports hasNext=False. MAX_PAGES is a safety
    valve, not a real-world cap, in case an API bug ever returns hasNext=True
    forever.

    since=None (default) means a full fetch — every incident NetWitness has.
    Pass an ISO8601 cutoff (see incremental_since()) to ask NetWitness for
    only incidents created/updated after that point — used by the periodic
    auto-refresh so it isn't re-fetching + re-enriching the entire incident
    history (and every incident's full alert history) every 30s.
    """
    if not st.session_state.nw_verified or not st.session_state.nw_token:
        return False, [], "Not authenticated."
    host = st.session_state.nw_host.rstrip("/")
    all_items = []
    page = 0
    diag = ""
    MAX_PAGES = 1000   # 1000 * pageSize(100) = 100,000 incidents ceiling
    # Include a wide date range by default — some NW versions return 400 without it
    since = since or "2020-01-01T00:00:00.000Z"
    try:
        while True:
            r = requests.get(
                nw_incidents_url(host),
                headers=nw_headers(),
                params={"pageSize": 100, "pageNumber": page, "since": since},
                timeout=30,
                verify=False,
            )
            if r.status_code == 200:
                data      = r.json()
                items     = data.get("items", [])
                total_api = data.get("totalItems", "?")
                all_items.extend(items)
                has_next  = data.get("hasNext", False)
                reached_limit = limit is not None and len(all_items) >= limit
                if not has_next or reached_limit or page >= MAX_PAGES:
                    diag = (
                        f"API reports {total_api} total incident(s) — "
                        f"fetched {len(all_items)} across {page+1} page(s)."
                    )
                    break
                page += 1
            elif r.status_code == 401:
                st.session_state.nw_verified = False
                st.session_state.nw_msg = "Token expired — login again."
                return False, [], "401 Unauthorized — token expired or invalid."
            elif r.status_code == 403:
                return False, [], (
                    "403 Forbidden — account lacks integration-server.api.access permission."
                )
            else:
                return False, [], f"HTTP {r.status_code}: {r.text[:200]}"
        # Fetch associated alerts/logs for each incident — in parallel.
        # This used to be a serial for-loop doing one blocking request per
        # incident (up to 15s timeout each), so on a list of N incidents the
        # whole fetch took N * request-time. Since this runs on every
        # auto-refresh and on every rerun after one is due (sending a chat
        # message, uploading a file — any Streamlit interaction reruns the
        # whole script), that serial loop is what made the entire UI look
        # frozen. Fetching concurrently bounds the wall-clock time to
        # roughly one request instead of N.
        clean_path = st.session_state.get("nw_incidents_path", "/rest/api/incidents").strip()
        if clean_path.endswith("/list"):
            clean_path = clean_path[:-5]

        def _fetch_alerts(inc: dict) -> None:
            inc_id = str(inc.get("id") or inc.get("incidentId") or "").strip()
            if not inc_id:
                inc["alerts"] = []
                return
            alerts_url = f"{host.rstrip('/')}{clean_path}/{inc_id}/alerts"
            # Paginate fully — this used to fetch only pageNumber=0, so any
            # incident with more than 100 alerts silently lost the rest.
            collected: list = []
            a_page = 0
            try:
                while a_page < MAX_PAGES:
                    r_alerts = requests.get(
                        alerts_url,
                        headers=nw_headers(),
                        params={"pageSize": 100, "pageNumber": a_page},
                        timeout=10,
                        verify=False,
                    )
                    if r_alerts.status_code != 200:
                        break
                    a_data = r_alerts.json()
                    collected.extend(a_data.get("items", []))
                    if not a_data.get("hasNext", False):
                        break
                    a_page += 1
            except Exception:
                pass
            # Tag every alert with its parent incident ID so alerts are
            # traceable back to the incident they came from.
            for a in collected:
                a["incident_id"] = inc_id
            inc["alerts"] = collected

        if all_items:
            with ThreadPoolExecutor(max_workers=min(16, len(all_items))) as pool:
                list(pool.map(_fetch_alerts, all_items))

        return True, all_items, diag
    except requests.exceptions.Timeout:
        return False, [], "Request timed out — check VPN/network."
    except Exception as e:
        return False, [], f"Exception: {e}"

# ══════════════════════════════════════════════════════════════════════════════
# SQLITE  — permanent incident log
# ══════════════════════════════════════════════════════════════════════════════
import sqlite3
import json as _json

DB_FILE = Path(__file__).parent / "soc_incidents.db"

def db_connect() -> sqlite3.Connection:
    con = sqlite3.connect(str(DB_FILE), check_same_thread=False)
    con.row_factory = sqlite3.Row
    return con

def db_init() -> None:
    """Create tables if they don't exist yet."""
    with db_connect() as con:
        con.execute("""
            CREATE TABLE IF NOT EXISTS incidents (
                id          TEXT PRIMARY KEY,
                title       TEXT,
                severity    TEXT,
                status      TEXT,
                assignee    TEXT,
                alert_count INTEGER,
                created     TEXT,
                updated     TEXT,
                raw_json    TEXT,
                first_seen  TEXT,
                last_seen   TEXT
            )
        """)
        con.execute("""
            CREATE TABLE IF NOT EXISTS fetch_log (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                fetched_at  TEXT,
                count       INTEGER
            )
        """)
        con.commit()

def db_upsert_incidents(incidents: list) -> int:
    """Upsert a list of incident dicts — new rows inserted, existing rows updated."""
    now = datetime.now().isoformat(timespec="seconds")
    rows = []
    for inc in incidents:
        inc_id = str(inc.get("id") or inc.get("incidentId") or "").strip()
        if not inc_id:
            continue
        rows.append((
            inc_id,
            inc.get("title") or inc.get("name") or "",
            normalise_sev(inc),
            str(inc.get("status") or ""),
            str(inc.get("assignee") or ""),
            int(inc.get("alertCount") or inc.get("numAlerts") or 0),
            str(inc.get("created") or inc.get("createdDate") or "")[:19],
            str(inc.get("updated") or inc.get("lastUpdated") or "")[:19],
            _json.dumps(inc),
            now,   # first_seen  — INSERT only
            now,   # last_seen   — always updated
        ))
    if not rows:
        return 0
    with db_connect() as con:
        con.executemany("""
            INSERT INTO incidents
                (id, title, severity, status, assignee, alert_count,
                 created, updated, raw_json, first_seen, last_seen)
            VALUES (?,?,?,?,?,?,?,?,?,?,?)
            ON CONFLICT(id) DO UPDATE SET
                title       = excluded.title,
                severity    = excluded.severity,
                status      = excluded.status,
                assignee    = excluded.assignee,
                alert_count = excluded.alert_count,
                updated     = excluded.updated,
                raw_json    = excluded.raw_json,
                last_seen   = excluded.last_seen
        """, rows)
        con.execute(
            "INSERT INTO fetch_log (fetched_at, count) VALUES (?,?)",
            (now, len(rows)),
        )
        con.commit()
    return len(rows)

def db_load_incidents(
    severity: str = "ALL",
    status:   str = "ALL",
    search:   str = "",
    limit:    int = 500,
) -> list[dict]:
    """Query incidents from SQLite with optional filters."""
    clauses, params = [], []
    if severity != "ALL":
        clauses.append("severity = ?"); params.append(severity)
    if status != "ALL":
        clauses.append("status = ?");   params.append(status)
    if search.strip():
        clauses.append("(title LIKE ? OR assignee LIKE ? OR id LIKE ?)")
        params += [f"%{search}%", f"%{search}%", f"%{search}%"]
    where = ("WHERE " + " AND ".join(clauses)) if clauses else ""
    with db_connect() as con:
        rows = con.execute(
            f"SELECT * FROM incidents {where} ORDER BY last_seen DESC LIMIT ?",
            params + [limit],
        ).fetchall()
    return [dict(r) for r in rows]

def db_stats() -> dict:
    """Return aggregate stats from the permanent log."""
    with db_connect() as con:
        total   = con.execute("SELECT COUNT(*) FROM incidents").fetchone()[0]
        by_sev  = dict(con.execute(
            "SELECT severity, COUNT(*) FROM incidents GROUP BY severity"
        ).fetchall())
        by_stat = dict(con.execute(
            "SELECT status, COUNT(*) FROM incidents GROUP BY status"
        ).fetchall())
        fetches = con.execute("SELECT COUNT(*) FROM fetch_log").fetchone()[0]
        last_f  = con.execute(
            "SELECT fetched_at FROM fetch_log ORDER BY id DESC LIMIT 1"
        ).fetchone()
    return {
        "total":   total,
        "by_sev":  by_sev,
        "by_stat": by_stat,
        "fetches": fetches,
        "last_fetch": last_f[0] if last_f else "—",
    }

def db_export_csv() -> str:
    """Return all incidents as a CSV string for download."""
    import csv, io
    rows = db_load_incidents(limit=100_000)
    if not rows:
        return ""
    buf = io.StringIO()
    w   = csv.DictWriter(buf, fieldnames=[
        "id","title","severity","status","assignee",
        "alert_count","created","updated","first_seen","last_seen",
    ])
    w.writeheader()
    for r in rows:
        w.writerow({k: r.get(k,"") for k in w.fieldnames})
    return buf.getvalue()

# ══════════════════════════════════════════════════════════════════════════════
# STARTUP — runs once per session, silently connects & fetches
# ══════════════════════════════════════════════════════════════════════════════

# Initialise DB on every startup (no-op if tables already exist)
db_init()
# ══════════════════════════════════════════════════════════════════════════════
# PIPELINE DATABASE — per-stage SQLite + ChromaDB, fully inline (no 2nd app)
# ══════════════════════════════════════════════════════════════════════════════
PIPELINE_STAGES = [
    "alerts_to_triage",
    "post_triage_investigate",
    "post_triage_no_investigate",
    "initial_ticket",
    "pending_ticket_report",
    "finalized_report",
]
PIPELINE_LABELS = {
    "alerts_to_triage":           "Alerts to Triage",
    "post_triage_investigate":    "Post-Triage · Needs Investigation",
    "post_triage_no_investigate": "Post-Triage · No Investigation Needed",
    "initial_ticket":             "Initial Ticket Generation",
    "pending_ticket_report":      "Pending Ticket / Report Generation",
    "finalized_report":           "Finalized Report",
}
PIPELINE_ICONS = {
    "alerts_to_triage":           "🚨",
    "post_triage_investigate":    "🔍",
    "post_triage_no_investigate": "✅",
    "initial_ticket":             "🎫",
    "pending_ticket_report":      "⏳",
    "finalized_report":           "📄",
}
PIPELINE_COLORS = {
    "alerts_to_triage":           "#FF3B3B",
    "post_triage_investigate":    "#FF7700",
    "post_triage_no_investigate": "#0AF0A0",
    "initial_ticket":             "#00D4FF",
    "pending_ticket_report":      "#FFB700",
    "finalized_report":           "#A78BFA",
}
PIPELINE_DB_FILE = Path(__file__).parent / "soc_pipeline.db"

def _pl_con():
    con = sqlite3.connect(str(PIPELINE_DB_FILE), check_same_thread=False)
    con.row_factory = sqlite3.Row
    return con

def pipeline_db_init():
    with _pl_con() as c:
        for s in PIPELINE_STAGES:
            c.execute(f"""CREATE TABLE IF NOT EXISTS {s} (
                id TEXT PRIMARY KEY, incident_id TEXT, title TEXT,
                severity TEXT, stage TEXT, created_at TEXT,
                summary TEXT, raw_json TEXT)""")
        c.commit()

def pipeline_count(stage):
    try:
        with _pl_con() as c:
            return c.execute(f"SELECT COUNT(*) FROM {stage}").fetchone()[0]
    except Exception:
        return 0

def pipeline_load(stage, limit=300):
    try:
        with _pl_con() as c:
            rows = c.execute(
                f"SELECT * FROM {stage} ORDER BY created_at DESC LIMIT ?", (limit,)
            ).fetchall()
        return [dict(r) for r in rows]
    except Exception:
        return []

def pipeline_insert(stage, record):
    import uuid as _uuid
    rec_id = str(record.get("id") or record.get("unc") or _uuid.uuid4())[:64]
    now = datetime.now().isoformat(timespec="seconds")
    try:
        with _pl_con() as c:
            c.execute(
                f"INSERT OR REPLACE INTO {stage} "
                "(id,incident_id,title,severity,stage,created_at,summary,raw_json) "
                "VALUES (?,?,?,?,?,?,?,?)",
                (rec_id,
                 str(record.get("incident_id") or record.get("incidentId") or ""),
                 str(record.get("title") or record.get("name") or ""),
                 str(record.get("severity") or record.get("classification") or ""),
                 stage, now,
                 str(record.get("summary") or record.get("description") or "")[:500],
                 _json.dumps(record)))
            c.commit()
    except Exception:
        pass
    return rec_id

def _pl_chroma_col(stage):
    if not CHROMA_OK or not st.session_state.get("chroma_client"):
        return None
    try:
        return st.session_state.chroma_client.get_or_create_collection(
            name=f"pipeline_{stage}", metadata={"hnsw:space": "cosine"})
    except Exception:
        return None

def pipeline_chroma_insert(stage, record):
    col = _pl_chroma_col(stage)
    if col is None:
        return
    import uuid as _uuid
    rec_id = str(record.get("id") or record.get("unc") or _uuid.uuid4())[:64]
    title = str(record.get("title") or "")
    summary = str(record.get("summary") or record.get("description") or "")
    doc = f"{title}\n{summary}".strip() or "no content"
    try:
        col.upsert(documents=[doc], ids=[rec_id],
                   metadatas=[{"stage": stage,
                                "severity": str(record.get("severity") or ""),
                                "created": datetime.now().isoformat(timespec="seconds")}])
    except Exception:
        pass

def pipeline_chroma_count(stage):
    col = _pl_chroma_col(stage)
    if col is None:
        return 0
    try:
        return col.count()
    except Exception:
        return 0

def pipeline_chroma_search(stage, query, n=5):
    col = _pl_chroma_col(stage)
    if col is None or not query.strip():
        return []
    try:
        total = col.count()
        if total == 0:
            return []
        res = col.query(query_texts=[query], n_results=min(n, total))
        return [{"id": res["ids"][0][i], "doc": d,
                 "score": round((1 - res["distances"][0][i]) * 100, 1),
                 "meta": res["metadatas"][0][i]}
                for i, d in enumerate(res["documents"][0])]
    except Exception as e:
        return [{"id": "err", "doc": str(e), "score": 0, "meta": {}}]

def pipeline_chroma_all(stage):
    col = _pl_chroma_col(stage)
    if col is None:
        return []
    try:
        data = col.get(include=["documents", "metadatas"])
        return [{"id": data["ids"][i],
                 "doc": (data["documents"] or [""])[i],
                 "meta": (data["metadatas"] or [{}])[i]}
                for i in range(len(data["ids"]))]
    except Exception:
        return []

def pipeline_insert_full(stage, record):
    """Insert into SQLite and ChromaDB (if connected)."""
    rec_id = pipeline_insert(stage, record)
    pipeline_chroma_insert(stage, record)
    return rec_id

def pipeline_delete(stage, rec_id):
    try:
        with _pl_con() as c:
            c.execute(f"DELETE FROM {stage} WHERE id=?", (rec_id,))
            c.commit()
    except Exception:
        pass
    col = _pl_chroma_col(stage)
    if col:
        try:
            col.delete(ids=[rec_id])
        except Exception:
            pass

pipeline_db_init()

# ──────────────────────────────────────────────────────────────────────────────
# EXPORT HELPERS  — generate downloadable bytes from a pipeline record
# ──────────────────────────────────────────────────────────────────────────────

def _make_csv_bytes(row: dict) -> bytes:
    """
    Build a CSV file from a pipeline SQLite row.
    Flattens the raw_json payload if present so the sheet has all fields.
    Returns UTF-8 bytes suitable for st.download_button.
    """
    import csv, io
    # Start with the fixed SQLite columns
    base = {
        "id":          row.get("id", ""),
        "incident_id": row.get("incident_id", ""),
        "title":       row.get("title", ""),
        "severity":    row.get("severity", ""),
        "stage":       row.get("stage", ""),
        "created_at":  row.get("created_at", ""),
        "summary":     row.get("summary", ""),
    }
    # Merge in any extra keys from raw_json (flat scalars only)
    try:
        extra = _json.loads(row.get("raw_json") or "{}")
        if isinstance(extra, dict):
            for k, v in extra.items():
                if k not in base and not isinstance(v, (dict, list)):
                    base[k] = str(v)
    except Exception:
        pass
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=list(base.keys()))
    writer.writeheader()
    writer.writerow(base)
    return buf.getvalue().encode("utf-8")


def _make_docx_bytes(row: dict) -> bytes:
    """
    Build a professional SOC Initial Ticket .docx from a pipeline record.
    Uses python-docx (pip install python-docx). Falls back to a plain-text
    .txt wrapped as bytes if the library is unavailable.
    """
    import io
    title      = row.get("title", "Untitled Ticket")
    inc_id     = row.get("incident_id", "—")
    severity   = row.get("severity", "—")
    stage      = row.get("stage", "initial_ticket")
    created_at = row.get("created_at", "—")
    summary    = row.get("summary", "No summary available.")

    # Try to parse raw_json for richer data
    extra = {}
    try:
        extra = _json.loads(row.get("raw_json") or "{}")
        if not isinstance(extra, dict):
            extra = {}
    except Exception:
        pass

    try:
        from docx import Document as _DocxDocument
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = _DocxDocument()

        # ── Page margins ─────────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        # ── Title ────────────────────────────────────────────
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title_para.runs:
            run.font.color.rgb = RGBColor(0x00, 0xD4, 0xFF)
            run.font.size = Pt(16)

        doc.add_paragraph()

        # ── Metadata table ───────────────────────────────────
        meta_fields = [
            ("Ticket ID",     row.get("id", "—")),
            ("Incident ID",   inc_id),
            ("Severity",      severity),
            ("Stage",         stage),
            ("Generated At",  created_at),
            ("Assignee",      str(extra.get("assignee") or "Unassigned")),
            ("Status",        str(extra.get("status") or "NEW")),
        ]
        tbl = doc.add_table(rows=len(meta_fields), cols=2)
        tbl.style = "Table Grid"
        for i, (label, value) in enumerate(meta_fields):
            tbl.cell(i, 0).text = label
            tbl.cell(i, 1).text = str(value)
            # Bold the label
            for run in tbl.cell(i, 0).paragraphs[0].runs:
                run.bold = True

        doc.add_paragraph()

        # ── Summary section ───────────────────────────────────
        doc.add_heading("Triage Summary", level=2)
        doc.add_paragraph(summary)

        # ── IOC / Alert section if present ───────────────────
        alert_count = extra.get("alertCount") or extra.get("numAlerts") or extra.get("alert_count")
        if alert_count:
            doc.add_paragraph()
            doc.add_heading("Alert Count", level=2)
            doc.add_paragraph(str(alert_count))

        description = extra.get("description") or extra.get("raw_log") or ""
        if description and description != summary:
            doc.add_paragraph()
            doc.add_heading("Description / Raw Log", level=2)
            # Cap at 3000 chars to keep document manageable
            doc.add_paragraph(str(description)[:3000])

        # ── Footer note ───────────────────────────────────────
        doc.add_paragraph()
        footer_para = doc.add_paragraph(
            f"Auto-generated by SOC Platform · {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in footer_para.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x3A, 0x60, 0x7A)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        # python-docx not installed — return UTF-8 text as fallback
        lines = [
            f"SOC INITIAL TICKET",
            f"==================",
            f"Title:      {title}",
            f"Ticket ID:  {row.get('id', '—')}",
            f"Incident:   {inc_id}",
            f"Severity:   {severity}",
            f"Stage:      {stage}",
            f"Created:    {created_at}",
            f"",
            f"SUMMARY",
            f"-------",
            summary,
            f"",
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        ]
        return "\n".join(lines).encode("utf-8")


# Step 1 — load persisted incidents from SQLite immediately
# so the UI is never empty even before the live fetch completes
if not st.session_state.incidents:
    _db_rows = db_load_incidents(limit=500)
    if _db_rows:
        st.session_state.incidents = [
            _json.loads(r["raw_json"]) for r in _db_rows if r.get("raw_json")
        ]

# Step 2 — silently auto-verify & auto-fetch if .env has credentials
# This runs once per session (tracked by _startup_done flag)
if not st.session_state._startup_done and _env["host"] and _env["username"] and _env["password"]:
    st.session_state.nw_host     = _env["host"]
    st.session_state.nw_username = _env["username"]
    st.session_state.nw_password = _env["password"]
    ok, msg, token = nw_login(_env["host"], _env["username"], _env["password"])
    if ok:
        st.session_state.nw_token    = token
        st.session_state.nw_verified = True
        st.session_state.nw_msg      = msg
        ok2, items, _diag = nw_fetch_incidents()
        if ok2:
            st.session_state.incidents       = items
            st.session_state.last_fetch      = datetime.now()
            st.session_state.last_full_fetch = datetime.now()
            st.session_state.last_fetch_mode = "full"
            db_upsert_incidents(items)
    else:
        st.session_state.nw_msg = msg
    st.session_state._startup_done = True

def incremental_since(last_fetch: datetime) -> str:
    """
    ISO8601 cutoff for an incremental refresh — everything since the last
    successful fetch, minus a 5-minute overlap to tolerate clock skew and
    NetWitness indexing lag. Incidents re-fetched inside that overlap just
    get upserted again by id — harmless, not duplicated.
    """
    cutoff = last_fetch - INCREMENTAL_OVERLAP
    return cutoff.strftime("%Y-%m-%dT%H:%M:%S.000Z")


def _merge_incidents(cached: list, fresh: list) -> list:
    """
    Upsert `fresh` incidents into `cached` by id, keeping everything else
    untouched. Used for incremental refreshes, where `fresh` is only the
    new/updated subset NetWitness returned for the `since` window — not
    the full incident set, so a plain replace would drop everything older.
    """
    by_id = {
        str(inc.get("id") or inc.get("incidentId") or f"_unkeyed_{i}"): inc
        for i, inc in enumerate(cached)
    }
    for inc in fresh:
        key = str(inc.get("id") or inc.get("incidentId") or "").strip()
        if key:
            by_id[key] = inc
    return list(by_id.values())


def maybe_auto_fetch():
    if not st.session_state.nw_verified:
        return
    now  = datetime.now()
    last = st.session_state.last_fetch
    if last is None or (now - last).total_seconds() >= REFRESH_INTERVAL:
        # We don't actually know whether NetWitness's `since` filter matches
        # on incident creation time or last-updated time — the API isn't
        # documented clearly enough to bet on it. If it's creation-time only,
        # a purely incremental refresh would silently miss status/alert
        # changes on older incidents. So: incremental fetches most cycles
        # (cheap — only new/updated incidents come back), but force a full
        # ground-truth resync every FULL_RESYNC_INTERVAL regardless, which
        # caps how stale the cache can ever get at a known worst case.
        last_full  = st.session_state.last_full_fetch
        needs_full = last_full is None or (now - last_full) >= FULL_RESYNC_INTERVAL
        since = None if needs_full else incremental_since(last)

        ok, items, _diag = nw_fetch_incidents(since=since)
        if ok:
            st.session_state.incidents = (
                items if since is None
                else _merge_incidents(st.session_state.incidents, items)
            )
            st.session_state.last_fetch      = now
            st.session_state.last_fetch_mode = "full" if since is None else "incremental"
            if since is None:
                st.session_state.last_full_fetch = now
            db_upsert_incidents(items)   # ← persist every fetch to SQLite

def chroma_connect(path: str = "./chroma_db") -> tuple[bool, str]:
    if not CHROMA_OK:
        return False, "chromadb not installed — run: pip install chromadb"
    try:
        client = chromadb.PersistentClient(path=path)
        col    = client.get_or_create_collection(
            name="soc_incidents", metadata={"hnsw:space": "cosine"}
        )
        st.session_state.chroma_client = client
        st.session_state.chroma_col    = col
        return True, f"Ready — {col.count()} vectors"
    except Exception as e:
        return False, str(e)

def chroma_sync(incidents: list) -> tuple[int, str]:
    col = st.session_state.chroma_col
    if col is None:
        return 0, "Connect ChromaDB first."
    docs, ids, metas = [], [], []
    for inc in incidents:
        inc_id = str(inc.get("id") or inc.get("incidentId") or "").strip()
        if not inc_id:
            continue
        title   = inc.get("title") or inc.get("name") or ""
        summary = inc.get("summary") or inc.get("description") or ""
        docs.append(f"{title}\n{summary}".strip() or "no content")
        ids.append(inc_id)
        metas.append({
            "severity": str(inc.get("riskScore") or inc.get("severity") or ""),
            "status":   str(inc.get("status") or ""),
            "created":  str(inc.get("created") or inc.get("createdDate") or "")[:19],
        })
    if not docs:
        return 0, "No valid incidents to store."
    col.upsert(documents=docs, ids=ids, metadatas=metas)
    return len(docs), f"Synced {len(docs)} incidents."

def chroma_search(query: str, n: int = 5) -> list:
    col = st.session_state.chroma_col
    if col is None or not query.strip() or col.count() == 0:
        return []
    try:
        res = col.query(query_texts=[query], n_results=min(n, col.count()))
        return [{"id": res["ids"][0][i], "text": doc,
                 "score": round((1 - res["distances"][0][i]) * 100, 1),
                 "meta": res["metadatas"][0][i]}
                for i, doc in enumerate(res["documents"][0])]
    except Exception as e:
        return [{"id":"err","text":str(e),"score":0,"meta":{}}]

def chat_respond(user_msg: str, incident: Optional[dict] = None) -> str:
    """
    ── SOC TRIAGE AGENT (LangChain) ─────────────────────────
    Powered by soc_triage_agent.py.
    Trigger words: triage, analyse, analyze, ioc, classify, ticket, investigate
    When triggered with an incident selected, runs the full pipeline:
      Phase 1 → IOC Checklists (Availability / Confidentiality / Integrity)
      Phase 2 → Risk Rating Methodology
      Phase 3 → SOC Classification Template
    Outputs: metakeys_payload + UNC ticket (both queued for downstream agents).
    All other messages fall back to plain SOC analyst Q&A via LangChain.
    ─────────────────────────────────────────────────────────
    """
    return soc_triage_chat_respond(user_msg, incident, llm_config=get_cisco_cfg())


# ══════════════════════════════════════════════════════════════════════════════
# AUTO-FETCH on every render
# ══════════════════════════════════════════════════════════════════════════════
maybe_auto_fetch()


# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:

    st.markdown(
        '<div class="app-logo">'
        '<div class="name">🛡️ Security Dashboard</div>'
        '<div class="sub">Powered by NetWitness</div>'
        '</div>',
        unsafe_allow_html=True,
    )
    st.markdown("---")

    # ── Connection status card ─────────────────────────────────
    if st.session_state.nw_verified:
        last      = st.session_state.last_fetch
        last_str  = last.strftime("%H:%M:%S") if last else "—"
        elapsed   = int((datetime.now() - last).total_seconds()) if last else 0
        remaining = max(REFRESH_INTERVAL - elapsed, 0)
        pct       = min(elapsed / REFRESH_INTERVAL, 1.0)
        bar_color = "var(--accent)" if pct < 0.8 else "var(--warn)"

        # Check GP status via the existing connection state
        _gp_status = (
            '<span class="dot dot-green"></span>'
            '<span style="color:var(--green);font-size:0.58rem">GP VPN ACTIVE</span>'
            if st.session_state.nw_verified else
            '<span class="dot dot-yellow"></span>'
            '<span style="color:var(--warn);font-size:0.58rem">GP VPN STATUS UNKNOWN</span>'
        )

        st.markdown(
            f'<div style="background:#040C14;border:1px solid #0E2030;'
            f'border-radius:7px;padding:11px 13px">'
            f'<div style="font-family:var(--mono);font-size:0.68rem">'
            f'<span class="dot dot-green"></span>Connected ✓</div>'
            f'<div style="margin-top:4px">{_gp_status}</div>'
            f'<div style="font-family:var(--mono);font-size:0.58rem;'
            f'color:var(--muted);margin-top:3px">{st.session_state.nw_msg}</div>'
            f'<div style="font-family:var(--mono);font-size:0.56rem;'
            f'color:#1A4A62;margin-top:6px">'
            f'Synced {last_str} &nbsp;·&nbsp; refresh in {remaining}s</div>'
            f'<div style="font-family:var(--mono);font-size:0.55rem;'
            f'color:#2A5A78;margin-top:4px">'
            f'📡 {st.session_state.nw_incidents_path} · {st.session_state.nw_auth_style}</div>'
            f'<div style="background:#060E1A;border-radius:2px;height:3px;'
            f'width:100%;margin-top:7px;overflow:hidden">'
            f'<div style="width:{pct*100:.0f}%;height:100%;border-radius:2px;'
            f'background:{bar_color};transition:width 1s linear"></div>'
            f'</div></div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            '<div style="background:#0A0608;border:1px solid #2A1010;'
            'border-radius:7px;padding:11px 13px;font-family:var(--mono);font-size:0.68rem">'
            '<span class="dot dot-red"></span>Not Connected<br>'
            '<span style="font-size:0.57rem;color:var(--muted)">'
            'Please enter your login details below</span></div>',
            unsafe_allow_html=True,
        )

    # ── NetWitness credentials ─────────────────────────────────
    st.markdown('<div class="sec-label">🔌  Connection</div>', unsafe_allow_html=True)

    # If already auto-connected from .env, show a clean status + update option
    if st.session_state.nw_verified and _env["username"]:
        st.markdown(
            '<div style="background:#041208;border:1px solid #0A3020;border-radius:5px;'
            'padding:9px 12px;font-family:var(--mono);font-size:0.62rem;margin-bottom:8px">'
            '<span class="dot dot-green"></span>'
            '<strong style="color:var(--green)">AUTO-CONNECTED FROM .ENV</strong><br>'
            '<span style="color:var(--muted);font-size:0.58rem">'
            f'Logged in as <strong>{st.session_state.nw_username}</strong> · '
            'token refreshed on startup.</span>'
            '</div>',
            unsafe_allow_html=True,
        )
        with st.expander("🔑 Update credentials"):
            host_in = st.text_input("Host URL", value=st.session_state.nw_host, key="sb_host")
            user_in = st.text_input("Username", value=st.session_state.nw_username, key="sb_user")
            pass_in = st.text_input("Password", value="", type="password",
                                     placeholder="Enter password…", key="sb_pass")
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.6rem;'
                'color:var(--muted);margin:6px 0 2px">— or paste token directly —</div>',
                unsafe_allow_html=True,
            )
            token_paste_in = st.text_area(
                "Paste Token",
                value="",
                placeholder="Paste a fresh accessToken (eyJ…) to skip re-login",
                height=80,
                key="sb_token_paste",
                label_visibility="collapsed",
            )
            cv, cs, cd = st.columns(3)
            if cv.button("🔌 Login", use_container_width=True, key="sb_verify"):
                st.session_state.nw_host     = host_in
                st.session_state.nw_username = user_in
                st.session_state.nw_password = pass_in

                raw_paste = token_paste_in.strip()
                if raw_paste:
                    # Use pasted token directly
                    st.session_state.nw_token = raw_paste
                    with st.spinner("Verifying token…"):
                        ok, msg = nw_verify_token()
                    st.session_state.nw_verified = ok
                    st.session_state.nw_msg      = msg
                    if not ok:
                        st.session_state.nw_token = ""
                else:
                    # Fall back to username/password
                    with st.spinner("Logging in…"):
                        ok, msg, token = nw_login(host_in, user_in, pass_in)
                    st.session_state.nw_verified = ok
                    st.session_state.nw_msg      = msg
                    if ok:
                        st.session_state.nw_token = token

                if st.session_state.nw_verified:
                    ok2, items, _diag = nw_fetch_incidents()
                    if ok2:
                        st.session_state.incidents  = items
                        st.session_state.last_fetch = datetime.now()
                        db_upsert_incidents(items)
                st.rerun()
            if cs.button("💾 Save", use_container_width=True, key="sb_save"):
                if pass_in:
                    env_save(host_in, user_in, pass_in)
                    st.success("Saved to .env")
                else:
                    st.warning("Enter password to save.")
            if cd.button("✕ Clear", use_container_width=True, key="sb_clear"):
                env_clear()
                st.session_state.update(
                    nw_token="", nw_username="", nw_password="",
                    nw_host="", nw_verified=False, nw_msg="",
                    incidents=[], last_fetch=None, _startup_done=False,
                )
                st.rerun()
    else:
        # Show last login error if there is one
        _last_err = st.session_state.get("nw_msg", "")
        if _last_err and not st.session_state.nw_verified:
            st.markdown(
                f'<div style="background:#1A0505;border:1px solid #5A1010;border-radius:5px;'
                f'padding:8px 11px;font-family:var(--mono);font-size:0.6rem;'
                f'color:#FF6B6B;margin-bottom:8px">'
                f'❌ Last error: {_last_err}</div>',
                unsafe_allow_html=True,
            )
        if DOTENV_OK and not _env["username"]:
            st.markdown(
                '<div style="background:#0A0800;border:1px solid #3A3000;border-radius:5px;'
                'padding:8px 11px;font-family:var(--mono);font-size:0.6rem;'
                'color:#FFB700;margin-bottom:8px">'
                '⚠ No credentials in .env<br>'
                '<span style="color:var(--muted)">Enter below & click 💾 Save</span></div>',
                unsafe_allow_html=True,
            )
        host_in = st.text_input("Host URL", value=st.session_state.nw_host,
                                 placeholder="https://192.168.x.x")
        if host_in != st.session_state.nw_host:
            st.session_state.nw_host     = host_in
            st.session_state.nw_verified = False

        # ── Login method toggle ────────────────────────────────
        login_method = st.radio(
            "Login method",
            ["Username / Password", "Paste Token"],
            horizontal=True,
            label_visibility="collapsed",
        )

        if login_method == "Username / Password":
            user_in = st.text_input("Username", value=st.session_state.nw_username,
                                     placeholder="admin")
            pass_in = st.text_input("Password", value="", type="password",
                                     placeholder="Enter password…")
            token_in = ""
        else:
            user_in  = st.session_state.nw_username
            pass_in  = ""
            token_in = st.text_area(
                "NetWitness Token",
                value="",
                placeholder="Paste your accessToken (eyJ…) here",
                height=100,
                label_visibility="collapsed",
            )
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.58rem;'
                'color:var(--muted);margin-top:-8px;margin-bottom:6px">'
                '⚡ Tokens expire — re-paste when you get a 401</div>',
                unsafe_allow_html=True,
            )

        cv, cs, cd = st.columns(3)
        if cv.button("🔌 Login", use_container_width=True):
            st.session_state.nw_username = user_in
            st.session_state.nw_password = pass_in

            # ── Token-paste path ──────────────────────────────
            if login_method == "Paste Token":
                raw_token = token_in.strip()
                if not raw_token:
                    st.error("❌ Paste a token first.")
                elif not host_in.strip():
                    st.error("❌ Enter the Host URL.")
                else:
                    st.session_state.nw_token = raw_token
                    with st.spinner("Verifying token…"):
                        ok, msg = nw_verify_token()
                    st.session_state.nw_verified = ok
                    st.session_state.nw_msg      = msg
                    if ok:
                        st.success(f"✅ {msg}")
                        ok2, items, _diag = nw_fetch_incidents()
                        if ok2:
                            st.session_state.incidents  = items
                            st.session_state.last_fetch = datetime.now()
                            db_upsert_incidents(items)
                            st.success(f"Fetched {len(items)} incidents")
                        else:
                            st.warning("Token accepted but no incidents fetched yet.")
                    else:
                        st.session_state.nw_token = ""
                        st.error(f"❌ {msg}")
                st.rerun()

            # ── Username / Password path ───────────────────────
            else:
                with st.spinner("Logging in…"):
                    try:
                        ok, msg, token = nw_login(host_in, user_in, pass_in)
                    except Exception as _ve:
                        ok, msg, token = False, f"Exception: {_ve}", ""
                st.session_state.nw_verified = ok
                st.session_state.nw_msg      = msg
                if ok:
                    st.session_state.nw_token = token
                    st.success(f"✅ {msg}")
                    ok2, items, _diag = nw_fetch_incidents()
                    if ok2:
                        st.session_state.incidents  = items
                        st.session_state.last_fetch = datetime.now()
                        db_upsert_incidents(items)
                        st.success(f"Fetched {len(items)} incidents")
                    else:
                        st.warning("Connected but no incidents fetched yet.")
                else:
                    st.error(f"❌ {msg}")
                st.rerun()

        if cs.button("💾 Save", use_container_width=True,
                     help="Save to .env — auto-connects on next startup"):
            if host_in and user_in and pass_in:
                env_save(host_in, user_in, pass_in)
                st.success("Saved — will auto-connect on next startup", icon="💾")
            else:
                st.warning("Enter host, username and password first.")

        if cd.button("✕ Clear", use_container_width=True):
            env_clear()
            st.session_state.update(
                nw_token="", nw_username="", nw_password="",
                nw_host="", nw_verified=False, nw_msg="",
                incidents=[], last_fetch=None, _startup_done=False,
            )
            st.rerun()

    # ── TLS Certificate (Option B — verified HTTPS) ─────────────
    st.markdown('<div class="sec-label">🔒  Security Certificate</div>', unsafe_allow_html=True)

    # Auto-clear bad cert if last message was an SSL error
    _last_msg = st.session_state.get("nw_msg", "")
    if "SSL error" in _last_msg and st.session_state.get("nw_cert_path", ""):
        st.warning(
            f"⚠️ The uploaded cert caused an SSL error — it has been removed. "
            f"Revert to browser export or try again.\n\n`{_last_msg}`"
        )
        st.session_state.nw_cert_path = ""

    _cert_active = st.session_state.get("nw_cert_path", "").strip()
    _cert_valid  = bool(_cert_active) and Path(_cert_active).is_file()

    if _cert_valid:
        st.markdown(
            f'<div style="font-family:var(--mono);font-size:0.62rem;color:var(--green);'
            f'margin-bottom:6px"><span class="dot dot-green"></span>'
            f'Verifying against: {Path(_cert_active).name}</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--warn);'
            'margin-bottom:6px"><span class="dot dot-yellow"></span>'
            'No cert configured — TLS verification skipped (insecure)</div>',
            unsafe_allow_html=True,
        )

    cert_upload = st.file_uploader(
        "Upload server/CA certificate (.pem / .crt)",
        type=["pem", "crt", "cer"],
        key="cert_uploader",
        label_visibility="collapsed",
    )
    if cert_upload is not None:
        certs_dir = Path(__file__).parent / "certs"
        certs_dir.mkdir(exist_ok=True)
        cert_dest = certs_dir / cert_upload.name
        cert_dest.write_bytes(cert_upload.getvalue())
        st.session_state.nw_cert_path = str(cert_dest)
        nw_cert_env_save(str(cert_dest))
        st.success(f"✅ Saved {cert_upload.name} — re-login to apply verified TLS")
        st.rerun()

    cert_path_in = st.text_input(
        "…or enter an existing cert path",
        value=st.session_state.get("nw_cert_path", ""),
        placeholder="/path/to/netwitness-ca.pem",
        key="cert_path_text",
    )
    ccert1, ccert2 = st.columns(2)
    if ccert1.button("💾 Use this path", use_container_width=True):
        if cert_path_in.strip() and Path(cert_path_in.strip()).is_file():
            st.session_state.nw_cert_path = cert_path_in.strip()
            nw_cert_env_save(cert_path_in.strip())
            st.success("✅ Cert path set — re-login to apply")
        else:
            st.error("File not found at that path.")
        st.rerun()
    if ccert2.button("✕ Remove cert", use_container_width=True):
        st.session_state.nw_cert_path = ""
        nw_cert_env_clear()
        st.info("Reverted to verify=False (insecure)")
        st.rerun()

    # ── Foundation LLM (HuggingFace) ──────────────────────────
    st.markdown('<div class="sec-label">🤖  AI Settings</div>', unsafe_allow_html=True)

    # Connection status indicator
    if st.session_state.cisco_connected:
        st.markdown(
            '<div style="background:#041208;border:1px solid #0A3020;border-radius:5px;'
            'padding:8px 11px;font-family:var(--mono);font-size:0.62rem;margin-bottom:8px">'
            '<span class="dot dot-green"></span>'
            '<strong style="color:var(--green)">LLM CONFIGURED</strong><br>'
            f'<span style="color:var(--muted);font-size:0.58rem">'
            f'Model: {st.session_state.cisco_model or "—"}</span>'
            '</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            '<div style="background:#0A0608;border:1px solid #2A1010;border-radius:5px;'
            'padding:8px 11px;font-family:var(--mono);font-size:0.62rem;margin-bottom:8px">'
            '<span class="dot dot-red"></span>'
            '<span style="color:var(--danger)">NOT CONFIGURED</span><br>'
            '<span style="font-size:0.57rem;color:var(--muted)">Enter your HF token below</span>'
            '</div>',
            unsafe_allow_html=True,
        )

    cisco_url_in = st.text_input(
        "Endpoint URL",
        value=st.session_state.cisco_url,
        placeholder="https://api-inference.huggingface.co/v1",
        key="sb_cisco_url",
    )
    cisco_key_in = st.text_input(
        "HuggingFace Token",
        value="",
        type="password",
        placeholder="hf_xxxxxxxxxxxxxxxxxxxx",
        key="sb_cisco_key",
    )
    cisco_model_in = st.text_input(
        "Model name",
        value=st.session_state.cisco_model,
        placeholder="fdtn-ai/Foundation-Sec-8B-Reasoning",
        key="sb_cisco_model",
    )

    cl1, cl2, cl3 = st.columns(3)

    if cl1.button("✅ Apply", use_container_width=True, key="cisco_apply"):
        if not cisco_url_in.strip():
            st.error("❌ Enter the endpoint URL.")
        elif not cisco_key_in.strip():
            st.error("❌ Enter your HF token.")
        else:
            st.session_state.cisco_url       = cisco_url_in.strip()
            st.session_state.cisco_key       = cisco_key_in.strip()
            st.session_state.cisco_model     = (
                cisco_model_in.strip() or "fdtn-ai/Foundation-Sec-8B-Reasoning"
            )
            st.session_state.cisco_connected = True
            st.success("✅ LLM configured!")
            st.rerun()

    if cl2.button("💾 Save", use_container_width=True, key="cisco_save"):
        if cisco_url_in.strip() and cisco_key_in.strip():
            st.session_state.cisco_url       = cisco_url_in.strip()
            st.session_state.cisco_key       = cisco_key_in.strip()
            st.session_state.cisco_model     = (
                cisco_model_in.strip() or "fdtn-ai/Foundation-Sec-8B-Reasoning"
            )
            st.session_state.cisco_connected = True
            cisco_env_save(
                cisco_url_in.strip(),
                cisco_key_in.strip(),
                st.session_state.cisco_model,
            )
            st.success("💾 Saved to .env")
            st.rerun()
        else:
            st.warning("Enter URL and HF token first.")

    if cl3.button("✕ Clear", use_container_width=True, key="cisco_clear"):
        st.session_state.cisco_url       = ""
        st.session_state.cisco_key       = ""
        st.session_state.cisco_model     = ""
        st.session_state.cisco_connected = False
        cisco_env_clear()
        st.rerun()

    # ── ChromaDB ───────────────────────────────────────────────
    st.markdown('<div class="sec-label">🧠  Knowledge Base</div>', unsafe_allow_html=True)

    chroma_path = st.text_input("Persist path", value="./chroma_db")
    cc1, cc2 = st.columns(2)

    if cc1.button("🗄️ Connect", use_container_width=True):
        ok, msg = chroma_connect(chroma_path)
        if ok: st.success(msg, icon="✅")
        else:  st.error(msg)
        st.rerun()

    if st.session_state.chroma_col is not None:
        count = st.session_state.chroma_col.count()
        st.markdown(
            f'<div style="font-family:var(--mono);font-size:0.64rem;'
            f'color:var(--green);margin-top:4px">'
            f'<span class="dot dot-green"></span>{count} vectors stored</div>',
            unsafe_allow_html=True,
        )
        if cc2.button("⬆️ Sync", use_container_width=True):
            if st.session_state.incidents:
                n, msg = chroma_sync(st.session_state.incidents)
                st.success(msg) if n else st.error(msg)
            else:
                st.warning("No incidents loaded yet.")

    st.markdown("---")
    st.markdown(
        f'<div style="font-family:var(--mono);font-size:0.54rem;'
        f'color:#1A3A52;text-align:center;line-height:1.9">'
        f'v4 · AUTO-REFRESH {REFRESH_INTERVAL}s<br>'
        f'{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</div>',
        unsafe_allow_html=True,
    )


# ══════════════════════════════════════════════════════════════════════════════
# AUTO-RERUN scheduler
# ══════════════════════════════════════════════════════════════════════════════
if st.session_state.nw_verified and st.session_state.last_fetch:
    elapsed   = (datetime.now() - st.session_state.last_fetch).total_seconds()
    remaining = max(REFRESH_INTERVAL - elapsed, 0)
    if remaining <= 1:
        time.sleep(1)
        st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# HEADER METRICS
# ══════════════════════════════════════════════════════════════════════════════
incidents = st.session_state.incidents
total     = len(incidents)
active    = len([i for i in incidents
                 if str(i.get("status","")).upper()
                 not in ("CLOSED","RESOLVED","REMEDIATED")])
by_sev    = Counter(normalise_sev(i) for i in incidents)
vectors   = st.session_state.chroma_col.count() if st.session_state.chroma_col else 0
last_sync = st.session_state.last_fetch.strftime("%H:%M:%S") if st.session_state.last_fetch else "—"
db_total  = db_stats()["total"]

st.markdown(
    '<div style="font-family:var(--mono);font-size:1.2rem;color:var(--accent);'
    'letter-spacing:3px;margin-bottom:6px">🛡️ SOC COMMAND CENTER</div>',
    unsafe_allow_html=True,
)

m1,m2,m3,m4,m5,m6,m7 = st.columns(7)
m1.metric("📋 Live",          total)
m2.metric("🔴 Active",        active)
m3.metric("🚨 Critical",      by_sev.get("CRITICAL", 0))
m4.metric("🟠 High",          by_sev.get("HIGH", 0))
m5.metric("🗄️ Vectors",      vectors)
m6.metric("🔄 Last Sync",     last_sync)
m7.metric("💾 DB Total",      db_total)

st.markdown("---")


# ══════════════════════════════════════════════════════════════════════════════
# TABS
# ══════════════════════════════════════════════════════════════════════════════
_connected = st.session_state.nw_verified
_inc_count = len(incidents)
st.markdown(
    f'<div style="display:flex;align-items:center;justify-content:space-between;'
    f'padding:10px 0 16px">'
    f'<div>'
    f'<div style="font-size:1.3rem;font-weight:700;color:var(--accent)">🛡️ Security Dashboard</div>'
    f'<div style="font-size:0.8rem;color:var(--muted);margin-top:2px">'
    f'{"✅ Connected — " + str(_inc_count) + " alerts loaded" if _connected else "⚠️ Not connected — please log in using the left panel"}'
    f'</div></div>'
    f'<div style="font-size:0.72rem;color:var(--muted);text-align:right">'
    f'{datetime.now().strftime("%A, %d %B %Y")}</div>'
    f'</div>',
    unsafe_allow_html=True,
)

tab_dash, tab_inc, tab_chat, tab_chroma, tab_log, tab_pipeline = st.tabs([
    "📊  Overview",
    "🚨  Security Alerts",
    "💬  Ask a Question",
    "🧠  Knowledge Base",
    "📁  History",
    "🔁  Data Pipeline",
])

# Streamlit's st.tabs has no server-side "set active tab" API — the click
# from the "🩺 Triage" button is simulated client-side by finding the tab
# button whose label contains "Ask a Question" and clicking it. One-shot:
# the flag is cleared immediately so this doesn't refire on later reruns.
if st.session_state.jump_to_ask_tab:
    st.session_state.jump_to_ask_tab = False
    components.html(
        """
        <script>
        (function() {
            function clickAskTab() {
                const doc = window.parent.document;
                const tabs = doc.querySelectorAll('button[role="tab"], [data-baseweb="tab"]');
                for (const t of tabs) {
                    if (t.innerText && t.innerText.includes("Ask a Question")) {
                        t.click();
                        return true;
                    }
                }
                return false;
            }
            let attempts = 0;
            const timer = setInterval(() => {
                attempts++;
                if (clickAskTab() || attempts > 20) {
                    clearInterval(timer);
                }
            }, 100);
        })();
        </script>
        """,
        height=0,
    )

SEV_COLORS = {
    "CRITICAL": "#FF3B3B",
    "HIGH":     "#FF7700",
    "MEDIUM":   "#FFB700",
    "LOW":      "#0AF0A0",
}
STATUS_COLORS = {
    "NEW":         "#00D4FF",
    "ASSIGNED":    "#FFB700",
    "IN_PROGRESS": "#FF7700",
    "CLOSED":      "#3A607A",
    "RESOLVED":    "#0AF0A0",
    "REMEDIATED":  "#0AF0A0",
}


# ─────────────────────────────────────────────────────────────
# TAB 1 — DASHBOARD
# ─────────────────────────────────────────────────────────────
with tab_dash:
    # ── Live diagnostic banner ─────────────────────────────────
    if st.session_state.nw_verified:
        _last = st.session_state.last_fetch
        _mode = st.session_state.last_fetch_mode
        _mode_str = {"full": "🌐 full", "incremental": "⚡ incremental"}.get(_mode, "—")
        _diag_str = (
            f"🕐 Last fetch: {_last.strftime('%H:%M:%S') if _last else 'never'} "
            f"({_mode_str}) · "
            f"Incidents in session: {len(incidents)} · "
            f"Host: {st.session_state.nw_host}"
        )
        st.markdown(
            f'<div style="background:#04090F;border:1px solid #0E2030;border-radius:5px;'
            f'padding:7px 12px;font-family:var(--mono);font-size:0.6rem;'
            f'color:var(--muted);margin-bottom:10px">'
            f'📡 {_diag_str}</div>',
            unsafe_allow_html=True,
        )
        rc1, rc2 = st.columns([1, 4])
        if rc1.button("🔄 Refresh Data", use_container_width=True, help="Forces a full resync, not incremental"):
            ok_r, items_r, diag_r = nw_fetch_incidents()
            if ok_r:
                st.session_state.incidents       = items_r
                st.session_state.last_fetch      = datetime.now()
                st.session_state.last_full_fetch = datetime.now()
                st.session_state.last_fetch_mode = "full"
                db_upsert_incidents(items_r)
                st.success(f"✅ {diag_r}")
            else:
                st.error(f"❌ Fetch failed: {diag_r}")
            st.rerun()

    if not incidents:
        if st.session_state.nw_verified:
            st.warning(
                "✅ Connected successfully, but there are **no security alerts** to show right now. "
                "This is normal if everything is quiet. If you expected to see data, please "
                "contact your IT administrator to check your account permissions."
            )
        else:
            # ── Connection Test Panel ──────────────────────────
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.65rem;'
                'color:var(--muted);letter-spacing:2px;margin-bottom:12px">'
                '🔧  Getting Started — Connection Setup</div>',
                unsafe_allow_html=True,
            )
            st.markdown(
                '<div style="font-family:var(--sans);font-size:0.78rem;'
                'color:var(--muted);margin-bottom:14px">'
                "Let's get you connected. Follow the steps below.</div>",
                unsafe_allow_html=True,
            )

            host  = st.session_state.nw_host.strip()
            token = st.session_state.nw_token.strip()

            # ── Step 1: Check host is set ──────────────────────
            s1_ok = bool(host)
            st.markdown(
                f'<div style="display:flex;align-items:center;gap:10px;'
                f'font-family:var(--mono);font-size:0.65rem;'
                f'padding:8px 12px;margin-bottom:6px;border-radius:4px;'
                f'background:{"#041A0A" if s1_ok else "#1A0505"};'
                f'border-left:3px solid {"#00E676" if s1_ok else "#FF5252"}">'
                f'{"✅" if s1_ok else "❌"} '
                f'<strong>Step 1 — Server Address</strong> &nbsp;'
                f'<span style="color:var(--muted)">'
                f'{"Set to: " + host if s1_ok else "Not set — enter https://192.168.x.x in the sidebar"}'
                f'</span></div>',
                unsafe_allow_html=True,
            )

            # ── Step 2: Check token is set ─────────────────────
            s2_ok = bool(token)
            token_preview = (token[:12] + "…" + token[-6:]) if len(token) > 20 else token
            st.markdown(
                f'<div style="display:flex;align-items:center;gap:10px;'
                f'font-family:var(--mono);font-size:0.65rem;'
                f'padding:8px 12px;margin-bottom:6px;border-radius:4px;'
                f'background:{"#041A0A" if s2_ok else "#1A0505"};'
                f'border-left:3px solid {"#00E676" if s2_ok else "#FF5252"}">'
                f'{"✅" if s2_ok else "❌"} '
                f'<strong>Step 2 — Login Session</strong> &nbsp;'
                f'<span style="color:var(--muted)">'
                f'{"Token present: " + token_preview if s2_ok else "No token — login with credentials in the sidebar"}'
                f'</span></div>',
                unsafe_allow_html=True,
            )

            # ── Step 3: Live connectivity test ─────────────────
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;'
                'color:var(--muted);margin:10px 0 6px">Step 3 — Live Tests</div>',
                unsafe_allow_html=True,
            )
            t3a, t3b, t3c = st.columns(3)

            # Test A: Can we reach the host at all?
            if t3a.button("🌐 Check Network", use_container_width=True,
                          disabled=not s1_ok, key="ping_host"):
                with st.spinner("Reaching host…"):
                    try:
                        r = requests.get(host, timeout=8, verify=False)
                        st.session_state["_test_ping"] = (True, f"HTTP {r.status_code} — host reachable")
                    except requests.exceptions.ConnectionError:
                        st.session_state["_test_ping"] = (False, "Connection refused — check VPN/host")
                    except requests.exceptions.Timeout:
                        st.session_state["_test_ping"] = (False, "Timed out — check VPN")
                    except Exception as e:
                        st.session_state["_test_ping"] = (False, str(e)[:100])

            # Test B: Does the auth endpoint respond?
            if t3b.button("🔑 Check Login Page", use_container_width=True,
                          disabled=not s1_ok, key="test_auth"):
                with st.spinner("Testing auth endpoint…"):
                    try:
                        r = requests.post(
                            f"{host.rstrip('/')}/rest/api/auth/userpass",
                            data={"username": "test", "password": "test"},
                            headers={"Content-Type": "application/x-www-form-urlencoded"},
                            timeout=8, verify=False,
                        )
                        if r.status_code == 200:
                            st.session_state["_test_auth"] = (True, "Auth endpoint OK — credentials accepted")
                        elif r.status_code in (401, 403):
                            st.session_state["_test_auth"] = (True, f"Auth endpoint reachable (HTTP {r.status_code} — wrong test creds, expected)")
                        else:
                            st.session_state["_test_auth"] = (False, f"Unexpected HTTP {r.status_code}: {r.text[:80]}")
                    except Exception as e:
                        st.session_state["_test_auth"] = (False, str(e)[:100])

            # Test C: Does the incidents endpoint respond with the current token?
            if t3c.button("📋 Check Data Access", use_container_width=True,
                          disabled=not (s1_ok and s2_ok), key="test_incidents"):
                with st.spinner("Testing incidents endpoint…"):
                    try:
                        r = requests.get(
                            nw_incidents_url(host),
                            headers=nw_headers(),
                            params={"pageSize": 1, "pageNumber": 0},
                            timeout=10, verify=False,
                        )
                        body = r.text[:200]
                        if r.status_code == 200:
                            total = r.json().get("totalItems", "?")
                            st.session_state["_test_inc"] = (True, f"✅ HTTP 200 — {total} incident(s) in NW")
                        elif r.status_code == 401:
                            st.session_state["_test_inc"] = (False, "401 Unauthorised — token expired, login again")
                        elif r.status_code == 403:
                            st.session_state["_test_inc"] = (False, "403 Forbidden — account needs 'integration-server.api.access' permission in NW Admin → Roles")
                        elif r.status_code == 400:
                            st.session_state["_test_inc"] = (False, f"400 Bad Request — {body}")
                        else:
                            st.session_state["_test_inc"] = (False, f"HTTP {r.status_code}: {body}")
                    except Exception as e:
                        st.session_state["_test_inc"] = (False, str(e)[:120])

            # Show test results
            for key, label in [("_test_ping","🌐 Ping"), ("_test_auth","🔑 Auth"), ("_test_inc","📋 Incidents")]:
                result = st.session_state.get(key)
                if result:
                    ok, msg = result
                    st.markdown(
                        f'<div style="font-family:var(--mono);font-size:0.62rem;'
                        f'padding:6px 12px;margin:3px 0;border-radius:4px;'
                        f'background:{"#041A0A" if ok else "#1A0505"};'
                        f'border-left:3px solid {"#00E676" if ok else "#FF5252"}">'
                        f'{label}: {msg}</div>',
                        unsafe_allow_html=True,
                    )

            # ── Step 4: One-click fix ──────────────────────────
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;'
                'color:var(--muted);margin:12px 0 6px">Step 4 — Quick Fix</div>',
                unsafe_allow_html=True,
            )
            fa, fb = st.columns(2)
            if fa.button("🔌 Connect Automatically",
                         use_container_width=True, key="quick_relogin",
                         disabled=not (bool(_env.get("host")) and bool(_env.get("username")) and bool(_env.get("password")))):
                with st.spinner("Logging in…"):
                    ok, msg, tok = nw_login(_env["host"], _env["username"], _env["password"])
                if ok:
                    st.session_state.nw_token    = tok
                    st.session_state.nw_verified = True
                    st.session_state.nw_msg      = msg
                    ok2, items2, diag2 = nw_fetch_incidents()
                    if ok2:
                        st.session_state.incidents  = items2
                        st.session_state.last_fetch = datetime.now()
                        db_upsert_incidents(items2)
                    st.rerun()
                else:
                    st.error(f"❌ {msg}")

            if fb.button("🗑️ Reset",
                         use_container_width=True, key="clear_tests"):
                for k in ["_test_ping", "_test_auth", "_test_inc"]:
                    st.session_state.pop(k, None)
                st.rerun()
    else:
        col_sev, col_status, col_recent = st.columns([1.1, 1.1, 2.2])

        # ── Severity bars ──────────────────────────────────────
        with col_sev:
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.65rem;'
                'color:var(--muted);letter-spacing:2px;margin-bottom:12px">'
                '🎯  Alert Priority Breakdown</div>', unsafe_allow_html=True,
            )
            sev_total = sum(by_sev.values()) or 1
            for s in ["CRITICAL","HIGH","MEDIUM","LOW"]:  # shown with friendly labels below
                cnt   = by_sev.get(s, 0)
                pct   = cnt / sev_total
                color = SEV_COLORS[s]
                st.markdown(
                    f'<div style="margin:8px 0">'
                    f'<div style="display:flex;justify-content:space-between;'
                    f'font-family:var(--mono);font-size:0.63rem;margin-bottom:4px">'
                    f'<span style="color:{color}">{{"CRITICAL":"🔴 Critical","HIGH":"🟠 High","MEDIUM":"🟡 Medium","LOW":"🟢 Low"}}.get(s, s)</span>'
                    f'<span style="color:var(--muted)">{cnt}</span></div>'
                    f'<div style="background:#0A1420;border-radius:3px;height:7px">'
                    f'<div style="width:{pct*100:.1f}%;height:100%;border-radius:3px;'
                    f'background:{color};box-shadow:0 0 8px {color}50"></div>'
                    f'</div></div>',
                    unsafe_allow_html=True,
                )

        # ── Status bars ────────────────────────────────────────
        with col_status:
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.65rem;'
                'color:var(--muted);letter-spacing:2px;margin-bottom:12px">'
                '📌  Current Status</div>', unsafe_allow_html=True,
            )
            status_counts = Counter(
                str(i.get("status") or "UNKNOWN").upper() for i in incidents
            )
            status_total = sum(status_counts.values()) or 1
            for status, cnt in sorted(status_counts.items(), key=lambda x: -x[1])[:6]:
                pct   = cnt / status_total
                color = STATUS_COLORS.get(status, "#3A607A")
                st.markdown(
                    f'<div style="margin:8px 0">'
                    f'<div style="display:flex;justify-content:space-between;'
                    f'font-family:var(--mono);font-size:0.63rem;margin-bottom:4px">'
                    f'<span style="color:{color}">{status}</span>'
                    f'<span style="color:var(--muted)">{cnt}</span></div>'
                    f'<div style="background:#0A1420;border-radius:3px;height:7px">'
                    f'<div style="width:{pct*100:.1f}%;height:100%;border-radius:3px;'
                    f'background:{color}"></div>'
                    f'</div></div>',
                    unsafe_allow_html=True,
                )

        # ── Latest incidents feed ──────────────────────────────
        with col_recent:
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.65rem;'
                'color:var(--muted);letter-spacing:2px;margin-bottom:12px">'
                '🕐  Most Recent Alerts</div>', unsafe_allow_html=True,
            )
            for inc in incidents[:10]:
                sev     = normalise_sev(inc)
                color   = SEV_COLORS.get(sev, "#3A607A")
                title   = (inc.get("title") or inc.get("name") or "Untitled")[:58]
                created = str(inc.get("created") or inc.get("createdDate") or "")[:16]
                st.markdown(
                    f'<div style="display:flex;align-items:center;gap:10px;'
                    f'padding:7px 12px;margin:3px 0;background:#060C16;'
                    f'border-radius:5px;border-left:3px solid {color}">'
                    f'<span style="font-family:var(--mono);font-size:0.58rem;'
                    f'color:{color};min-width:52px">{sev[:4]}</span>'
                    f'<span style="flex:1;font-size:0.8rem">{title}</span>'
                    f'<span style="font-family:var(--mono);font-size:0.58rem;'
                    f'color:var(--muted);white-space:nowrap">{created}</span>'
                    f'</div>',
                    unsafe_allow_html=True,
                )

        st.markdown("---")

        # ── Assignee breakdown ─────────────────────────────────
        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.65rem;'
            'color:var(--muted);letter-spacing:2px;margin-bottom:12px">'
            '👤  Team Workload</div>', unsafe_allow_html=True,
        )
        assignee_counts = Counter(
            str(i.get("assignee") or "Unassigned") for i in incidents
        )
        top_assignees = assignee_counts.most_common(6)
        if top_assignees:
            max_a  = top_assignees[0][1]
            cols_a = st.columns(len(top_assignees))
            for idx, (name, cnt) in enumerate(top_assignees):
                pct = cnt / max_a
                with cols_a[idx]:
                    st.markdown(
                        f'<div class="stat-mini">'
                        f'<div class="val">{cnt}</div>'
                        f'<div class="lbl">{name[:12]}</div>'
                        f'<div style="background:#0A1420;border-radius:2px;'
                        f'height:4px;margin-top:7px">'
                        f'<div style="width:{pct*100:.0f}%;height:100%;'
                        f'border-radius:2px;background:var(--accent)"></div>'
                        f'</div></div>',
                        unsafe_allow_html=True,
                    )


# ─────────────────────────────────────────────────────────────
# TAB 2 — INCIDENTS
# ─────────────────────────────────────────────────────────────
with tab_inc:
    st.markdown(
        '<div class="info-box"><div class="title">🚨 Security Alerts</div>'
        'This page lists all security incidents detected by NetWitness. '
        'Use the filter below to focus on the most urgent alerts. '
        'Click any alert to see full details.</div>',
        unsafe_allow_html=True,
    )
    col_filter, col_sync, col_info = st.columns([1.4, 1.2, 5])

    sev_filter = col_filter.selectbox(
        "Filter by priority", ["ALL","CRITICAL","HIGH","MEDIUM","LOW"],
        label_visibility="visible",
    )
    if col_sync.button("⬆️ Sync to Knowledge Base", use_container_width=True):
        if not incidents:
            st.warning("No incidents loaded yet.")
        elif st.session_state.chroma_col is None:
            st.error("Connect ChromaDB first (sidebar).")
        else:
            n, msg = chroma_sync(incidents)
            st.success(msg) if n else st.error(msg)

    last = st.session_state.last_fetch
    if last:
        elapsed   = int((datetime.now() - last).total_seconds())
        remaining = max(REFRESH_INTERVAL - elapsed, 0)
        col_info.markdown(
            f'<div style="font-family:var(--mono);font-size:0.6rem;'
            f'color:var(--muted);padding-top:10px">'
            f'<span class="dot dot-green"></span>'
            f'Auto-refresh every {REFRESH_INTERVAL}s &nbsp;·&nbsp; '
            f'next in {remaining}s &nbsp;·&nbsp; {total} incidents</div>',
            unsafe_allow_html=True,
        )

    with st.expander("🔧 Endpoint Diagnostics", expanded=not bool(incidents)):
        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.65rem;color:var(--muted);margin-bottom:8px">'
            'Tests all known NW endpoints with all auth styles to find the working combination. '
            'Click "✅ Use this" on a hit to wire it into the app automatically.</div>',
            unsafe_allow_html=True,
        )

        # Map scanner's auth-style labels → nw_headers() style names
        _AUTH_STYLE_MAP = {
            "NW-Token": "NetWitness-Token",
            "Bearer":   "Bearer",
            "Cookie":   "Cookie",
            "Both":     "Both",
        }

        if st.button("🔍 Run Endpoint Scan", use_container_width=False):
            if not st.session_state.nw_token:
                st.error("Login first.")
            else:
                host  = st.session_state.nw_host.rstrip("/")
                token = st.session_state.nw_token
                eps   = [
                    "/rest/api/incidents",
                    "/rest/api/respond/incidents",
                    "/rest/api/v1/incidents",
                    "/rest/api/v2/incidents",
                    "/rest/api/respond/incidents/list",
                    "/rest/api/incidents/list",
                    "/rest/api/investigation/incidents",
                    "/respond/api/incidents",
                    "/respond/api/v1/incidents",
                    "/respond/api/v2/incidents",
                    "/rsa/investigation/incidents",
                    "/rsa/respond/incidents",
                    "/sa/api/incidents",
                    "/esa/api/incidents",
                    "/api/respond/incidents",
                    "/rest/api/alerting/incidents",
                    "/rest/api/incidents?start=0",
                ]
                auth_styles = {
                    "Bearer":     {"Authorization": f"Bearer {token}"},
                    "Cookie":     {"Cookie": f"access_token={token}"},
                    "NW-Token":   {"NetWitness-Token": token},
                    "Both":       {"Authorization": f"Bearer {token}", "Cookie": f"access_token={token}"},
                }
                results = []
                for ep in eps:
                    for style, ah in auth_styles.items():
                        try:
                            ah = dict(ah)  # avoid mutating shared dict across iterations
                            ah["Accept"] = "application/json"
                            r = requests.get(f"{host}{ep}?limit=1", headers=ah,
                                             timeout=10, verify=False)
                            ct   = r.headers.get("Content-Type","")
                            is_j = "json" in ct or r.text.strip().startswith(("{","["))
                            results.append({
                                "endpoint": ep, "auth": style,
                                "status": r.status_code,
                                "json": "✅ JSON" if is_j else "❌ HTML",
                                "is_hit": is_j and r.status_code == 200,
                                "preview": r.text[:80] if is_j else "",
                            })
                        except Exception as e:
                            results.append({"endpoint": ep, "auth": style,
                                            "status": "ERR", "json": str(e)[:50],
                                            "is_hit": False, "preview": ""})
                # Persist across reruns (so "✅ Use this" buttons survive)
                st.session_state.endpoint_scan_results = results

        results = st.session_state.get("endpoint_scan_results", [])
        if results:
            for i, res in enumerate(results):
                color   = "#00E676" if res.get("is_hit") else "#3A607A"
                preview = res["preview"]
                preview_html = (
                    f'<br><span style="color:#00E676">{preview}</span>'
                    if preview else ""
                )
                rc1, rc2 = st.columns([6, 1])
                with rc1:
                    st.markdown(
                        f'<div style="font-family:var(--mono);font-size:0.65rem;'
                        f'padding:3px 8px;border-left:2px solid {color};margin:2px 0">'
                        f'<span style="color:{color}">{res["json"]}</span> '
                        f'HTTP {res["status"]} | {res["auth"]} | {res["endpoint"]}'
                        f'{preview_html}'
                        f'</div>',
                        unsafe_allow_html=True,
                    )
                with rc2:
                    if res.get("is_hit"):
                        if st.button("✅ Use this", key=f"use_ep_{i}", use_container_width=True):
                            # Wire the found endpoint + auth style into the app
                            clean_path = res["endpoint"].split("?")[0]
                            st.session_state.nw_incidents_path = clean_path
                            st.session_state.nw_auth_style     = _AUTH_STYLE_MAP.get(res["auth"], "NetWitness-Token")
                            st.session_state.nw_working_ep      = res["endpoint"]
                            st.session_state.nw_working_auth    = {"style": res["auth"]}

                            ok_v, msg_v = nw_verify_token()
                            st.session_state.nw_verified = ok_v
                            st.session_state.nw_msg      = msg_v
                            if ok_v:
                                ok_f, items_f, diag_f = nw_fetch_incidents()
                                if ok_f:
                                    st.session_state.incidents  = items_f
                                    st.session_state.last_fetch = datetime.now()
                                    db_upsert_incidents(items_f)
                                st.success(f"Applied {clean_path} ({res['auth']}) — {msg_v}")
                            else:
                                st.error(f"Applied but verify failed: {msg_v}")
                            st.rerun()

        if st.session_state.nw_working_ep:
            st.success(
                f"✅ Active endpoint: `{st.session_state.nw_incidents_path}` "
                f"· auth style: `{st.session_state.nw_auth_style}`"
            )

    # ── Manual endpoint / auth override ─────────────────────────
    with st.expander("⚙️ Manual Endpoint Config"):
        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);margin-bottom:6px">'
            'Set these directly if you already know your NW instance\'s working values.</div>',
            unsafe_allow_html=True,
        )
        mc1, mc2 = st.columns(2)
        new_path = mc1.text_input(
            "Incidents path", value=st.session_state.nw_incidents_path, key="manual_inc_path"
        )
        new_style = mc2.selectbox(
            "Auth header style",
            ["NetWitness-Token", "Bearer", "Cookie", "Both"],
            index=["NetWitness-Token", "Bearer", "Cookie", "Both"].index(
                st.session_state.nw_auth_style
                if st.session_state.nw_auth_style in ["NetWitness-Token","Bearer","Cookie","Both"]
                else "NetWitness-Token"
            ),
            key="manual_auth_style",
        )
        if st.button("💾 Apply & Re-verify", key="manual_apply"):
            st.session_state.nw_incidents_path = new_path.strip() or "/rest/api/incidents"
            st.session_state.nw_auth_style     = new_style
            ok_v, msg_v = nw_verify_token()
            st.session_state.nw_verified = ok_v
            st.session_state.nw_msg      = msg_v
            if ok_v:
                ok_f, items_f, diag_f = nw_fetch_incidents()
                if ok_f:
                    st.session_state.incidents  = items_f
                    st.session_state.last_fetch = datetime.now()
                    db_upsert_incidents(items_f)
                st.success(f"✅ {msg_v}")
            else:
                st.error(f"❌ {msg_v}")
            st.rerun()

    st.markdown("---")

    if not incidents:
        st.markdown(
            '<div style="text-align:center;padding:70px;font-family:var(--mono);'
            'font-size:0.8rem;color:var(--muted)">'
            '● NO INCIDENTS LOADED<br>'
            '<span style="font-size:0.65rem">'
            'Verify your token — incidents load automatically</span></div>',
            unsafe_allow_html=True,
        )
    else:
        # Rendering a card (+ 2 buttons, + a possible nested alerts expander)
        # per incident used to loop over the *entire* incidents list — with
        # tens of thousands of incidents that's tens of thousands of widgets
        # dumped into the DOM at once. Streamlit renders every tab's content
        # into the page regardless of which tab is active, so this was
        # slowing down the whole app, not just this tab. Filter first, then
        # only render one page of cards at a time.
        filtered = [
            inc for inc in incidents
            if sev_filter == "ALL" or normalise_sev(inc) == sev_filter
        ]
        total_filtered = len(filtered)

        PAGE_SIZE   = 25
        total_pages = max(1, -(-total_filtered // PAGE_SIZE))   # ceil div

        # Jump back to page 1 whenever the filter changes, so you don't land
        # on a now-empty page after narrowing the results.
        if st.session_state.get("_sec_alerts_filter") != sev_filter:
            st.session_state._sec_alerts_filter = sev_filter
            st.session_state.sec_alerts_page    = 1
        page = min(max(1, st.session_state.get("sec_alerts_page", 1)), total_pages)

        pg1, pg2, pg3 = st.columns([1, 3, 1])
        if pg1.button("◀ Prev", disabled=page <= 1, use_container_width=True):
            st.session_state.sec_alerts_page = page - 1
            st.rerun()
        pg2.markdown(
            f'<div style="text-align:center;font-family:var(--mono);'
            f'font-size:0.68rem;color:var(--muted);padding-top:8px">'
            f'Page {page} of {total_pages} &nbsp;·&nbsp; '
            f'{total_filtered} incident(s) matching filter'
            f'</div>',
            unsafe_allow_html=True,
        )
        if pg3.button("Next ▶", disabled=page >= total_pages, use_container_width=True):
            st.session_state.sec_alerts_page = page + 1
            st.rerun()

        start = (page - 1) * PAGE_SIZE
        shown = 0
        for inc in filtered[start:start + PAGE_SIZE]:
            sev = normalise_sev(inc)
            shown += 1

            inc_id   = str(inc.get("id") or inc.get("incidentId") or "—")
            title    = inc.get("title")    or inc.get("name")    or "Untitled"
            status   = str(inc.get("status") or "—")
            created  = str(inc.get("created") or inc.get("createdDate") or "—")[:16]
            assignee = inc.get("assignee") or "Unassigned"
            alerts   = inc.get("alertCount") or inc.get("numAlerts") or "—"

            st.markdown(
                f'<div class="card card-{sev.lower()}">'
                f'<div style="display:flex;align-items:center;gap:10px;flex-wrap:wrap">'
                f'<span class="badge badge-{sev.lower()}">{sev}</span>'
                f'<strong style="flex:1;font-size:0.9rem">{title}</strong>'
                f'<span style="color:var(--muted);font-size:0.62rem">Incident ID:</span>'
                f'<code style="color:var(--muted);font-size:0.68rem">{inc_id}</code>'
                f'</div>'
                f'<div style="margin-top:8px;font-size:0.75rem;color:var(--muted);'
                f'display:flex;gap:20px;flex-wrap:wrap">'
                f'<span>🕐 {created}</span><span>📌 {status}</span>'
                f'<span>👤 {assignee}</span><span>🚨 {alerts} alerts</span>'
                f'</div></div>',
                unsafe_allow_html=True,
            )
            b1, b2, _ = st.columns([0.9, 0.7, 6])
            if b1.button("🩺 Triage", key=f"chat_{inc_id}"):
                st.session_state.chat_incident       = inc
                st.session_state.pending_auto_triage = True
                st.session_state.jump_to_ask_tab     = True
                st.rerun()
            if b2.button("{ }", key=f"json_{inc_id}"):
                with st.expander(f"JSON — {inc_id}", expanded=True):
                    st.json(inc)

            # Associated Alerts / Logs
            alerts_list = inc.get("alerts")
            if alerts_list:
                with st.expander(f"🚨 Associated Alerts ({len(alerts_list)})", expanded=False):
                    for alert in alerts_list:
                        a_title = alert.get("title") or alert.get("name") or "Untitled Alert"
                        a_id = alert.get("id") or ""
                        a_source = alert.get("source") or "Unknown"
                        a_type = alert.get("type") or "Unknown"
                        a_created = alert.get("created") or alert.get("receivedTime") or ""
                        st.markdown(
                            f'<div style="background:#091624;padding:8px 12px;border-radius:4px;margin-bottom:6px;border-left:3px solid var(--accent)">'
                            f'<div style="display:flex;justify-content:between;align-items:center">'
                            f'<strong>{a_title}</strong>'
                            f'<code style="color:var(--muted);font-size:0.7rem;margin-left:auto">{a_id}</code>'
                            f'</div>'
                            f'<div style="font-size:0.72rem;color:var(--muted);margin-top:4px">'
                            f'Incident ID: <strong style="color:var(--accent)">{inc_id}</strong> &nbsp;·&nbsp; '
                            f'Source: {a_source} &nbsp;·&nbsp; Type: {a_type} &nbsp;·&nbsp; Time: {a_created}'
                            f'</div>'
                            f'</div>',
                            unsafe_allow_html=True
                        )
                        # Nested event details
                        events = alert.get("events")
                        if events:
                            for idx, ev in enumerate(events):
                                ev_src = ev.get("source", {})
                                ev_dst = ev.get("destination", {})
                                src_ip = ev_src.get("device", {}).get("ipAddress") or ev_src.get("ipAddress") or "—"
                                dst_ip = ev_dst.get("device", {}).get("ipAddress") or ev_dst.get("ipAddress") or "—"
                                user = ev_src.get("user", {}).get("username") or ev_src.get("username") or "—"
                                ev_proto = ev.get("ip_proto") or ev.get("protocol") or "—"
                                ev_port = ev_dst.get("device", {}).get("port") or ev_dst.get("port") or "—"
                                st.markdown(
                                    f'<div style="margin-left:15px;padding:4px 8px;font-family:var(--mono);font-size:0.7rem;color:var(--muted);border-left:1px dashed #1B4A62">'
                                    f'Event {idx+1}: 🧑‍💻 {user} | ➡️ {src_ip} → {dst_ip} (Port {ev_port}, Proto {ev_proto})'
                                    f'</div>',
                                    unsafe_allow_html=True
                                )

        if shown == 0:
            st.info(f"No incidents match filter: {sev_filter}")


# ─────────────────────────────────────────────────────────────
# TAB 3 — CHAT
# ─────────────────────────────────────────────────────────────
with tab_chat:
    st.markdown(
        '<div class="info-box"><div class="title">💬 Ask a Question</div>'
        'You can ask plain-language questions about your security alerts here. '
        'For example: <em>"What are the most critical incidents today?"</em> or '
        '<em>"Summarise the latest high-priority alerts."</em> '
        'You do not need any technical knowledge to use this.</div>',
        unsafe_allow_html=True,
    )

    # ── Global CSS for the spinning phase icon (injected once, never reset) ──
    st.markdown("""
    <style>
    @keyframes soc-phase-spin { to { transform: rotate(360deg); } }
    .soc-phase-spinner {
        display: inline-block;
        animation: soc-phase-spin 1.2s linear infinite;
        color: #E8623A;
        font-size: 0.95rem;
        line-height: 1;
        margin-right: 8px;
        vertical-align: middle;
    }
    </style>
    """, unsafe_allow_html=True)

    # ── Helper: parse uploaded file into an incident dict ──────
    def _parse_uploaded_file(uploaded_file) -> tuple[dict, str]:
        """
        Parse a Streamlit UploadedFile into an incident dict.
        Supports: .json, .csv, .txt, .log
        Returns (incident_dict, error_message).  error_message is "" on success.
        """
        import io
        name = uploaded_file.name
        ext  = name.rsplit(".", 1)[-1].lower()
        raw  = uploaded_file.read()

        try:
            if ext == "json":
                data = _json.loads(raw.decode("utf-8", errors="replace"))
                # Accept a list → take first item; accept a dict → use directly
                if isinstance(data, list):
                    if not data:
                        return {}, "JSON array is empty."
                    incident = data[0] if isinstance(data[0], dict) else {"raw": data[0]}
                elif isinstance(data, dict):
                    # NetWitness envelope: {"items": [...]}
                    if "items" in data and isinstance(data["items"], list) and data["items"]:
                        incident = data["items"][0]
                    else:
                        incident = data
                else:
                    return {}, "JSON root must be an object or array."

            elif ext == "csv":
                import csv, io as _io
                text    = raw.decode("utf-8", errors="replace")
                reader  = list(csv.DictReader(_io.StringIO(text)))
                if not reader:
                    return {}, "CSV has no data rows."
                # Use the first row as the incident dict
                incident = dict(reader[0])
                # Attach all rows as an alert list for richer context
                incident["_csv_alerts"] = reader

            elif ext in ("txt", "log"):
                text = raw.decode("utf-8", errors="replace")
                # Build a minimal incident dict wrapping the raw text
                incident = {
                    "id":          name,
                    "title":       f"Uploaded log — {name}",
                    "description": text[:4000],   # cap to avoid token overflow
                    "raw_log":     text,
                    "source":      "file_upload",
                }

            else:
                return {}, f"Unsupported file type '.{ext}'. Upload a .json, .csv, .txt, or .log file."

        except Exception as exc:
            return {}, f"Failed to parse file: {exc}"

        # Ensure there's always an id and title
        if "id" not in incident:
            incident["id"] = name
        if "title" not in incident:
            incident["title"] = name

        return incident, ""

    # ── Resolve which incident the agent will use ──────────────
    # Priority: NW incident (from Incidents tab) > uploaded file
    nw_inc       = st.session_state.chat_incident
    up_inc       = st.session_state.uploaded_incident
    active_inc   = nw_inc if nw_inc else up_inc   # what the agent receives

    # ── Context banner ─────────────────────────────────────────
    st.markdown(
        '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
        'letter-spacing:2px;margin-bottom:10px">■ INCIDENT CONTEXT</div>',
        unsafe_allow_html=True,
    )

    ctx_col, clr_col = st.columns([5, 1])

    with ctx_col:
        if nw_inc:
            sev = normalise_sev(nw_inc)
            st.markdown(
                f'<div style="background:#050F1A;border:1px solid var(--accent);'
                f'border-radius:6px;padding:10px 16px;font-family:var(--mono);'
                f'font-size:0.72rem">'
                f'<span class="badge badge-{sev.lower()}">{sev}</span>'
                f'&nbsp;'
                f'<strong>{nw_inc.get("id","?")}</strong> — {nw_inc.get("title","?")}'
                f'<span style="color:var(--muted);font-size:0.6rem;margin-left:10px">'
                f'● from NetWitness</span>'
                f'</div>',
                unsafe_allow_html=True,
            )
        elif up_inc:
            st.markdown(
                f'<div style="background:#050F1A;border:1px solid var(--warn);'
                f'border-radius:6px;padding:10px 16px;font-family:var(--mono);'
                f'font-size:0.72rem">'
                f'📎 <strong>{st.session_state.uploaded_filename}</strong>'
                f'<span style="color:var(--muted);font-size:0.6rem;margin-left:10px">'
                f'● from uploaded file</span>'
                f'</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                '<div style="background:#07080F;border:1px solid var(--border);'
                'border-radius:6px;padding:10px 16px;font-family:var(--mono);'
                'font-size:0.68rem;color:var(--muted)">'
                '⚠ No incident context — select one from the Incidents tab '
                'or upload a file below.'
                '</div>',
                unsafe_allow_html=True,
            )

    with clr_col:
        if nw_inc or up_inc:
            if st.button("✕ Clear", use_container_width=True, key="clear_ctx"):
                st.session_state.chat_incident    = None
                st.session_state.uploaded_incident = None
                st.session_state.uploaded_filename = ""
                st.rerun()

    st.markdown("---")

    # ── File uploader ──────────────────────────────────────────
    st.markdown(
        '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
        'letter-spacing:2px;margin-bottom:8px">■ UPLOAD INCIDENT FILE</div>',
        unsafe_allow_html=True,
    )

    uploaded_file = st.file_uploader(
        "Upload incident file",
        type=["json", "csv", "txt", "log"],
        label_visibility="collapsed",
        help="Upload a JSON export, CSV log, or plain-text log file to use as incident context.",
    )

    if uploaded_file is not None:
        # Only re-parse if it's a new file
        if uploaded_file.name != st.session_state.uploaded_filename:
            parsed_inc, err = _parse_uploaded_file(uploaded_file)
            if err:
                st.error(f"❌ {err}")
            else:
                st.session_state.uploaded_incident = parsed_inc
                st.session_state.uploaded_filename  = uploaded_file.name
                st.session_state.chat_incident      = None   # clear NW incident
                st.rerun()

    # Preview of parsed file fields
    if st.session_state.uploaded_incident and not nw_inc:
        preview_inc = st.session_state.uploaded_incident
        preview_keys = [k for k in preview_inc if not k.startswith("_")][:12]
        preview_html = " &nbsp;·&nbsp; ".join(
            f'<span style="color:var(--accent)">{k}</span>'
            f':<span style="color:var(--text)"> '
            f'{str(preview_inc[k])[:40]}</span>'
            for k in preview_keys
        )
        st.markdown(
            f'<div style="background:#060C16;border:1px solid var(--border);'
            f'border-radius:5px;padding:8px 12px;font-family:var(--mono);'
            f'font-size:0.62rem;margin:6px 0;line-height:1.8">'
            f'📋 Parsed fields: {preview_html}'
            f'</div>',
            unsafe_allow_html=True,
        )
        with st.expander("🔍 View full parsed incident JSON"):
            st.json({k: v for k, v in preview_inc.items() if k != "raw_log"})

    st.markdown("---")

    # ── Trigger hint ───────────────────────────────────────────
    if active_inc:
        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.6rem;color:var(--muted);'
            'margin-bottom:6px">💡 Type <strong style="color:var(--accent)">triage</strong>, '
            '<strong style="color:var(--accent)">ioc</strong>, '
            '<strong style="color:var(--accent)">classify</strong> or '
            '<strong style="color:var(--accent)">ticket</strong> to run the full triage pipeline.'
            '</div>',
            unsafe_allow_html=True,
        )

    # ── Chat input  (called first — always fixed at page bottom by Streamlit)
    user_input = st.chat_input("Ask the SOC agent…")

    # The "🩺 Triage" button sets this flag instead of the user typing a
    # message — synthesize the trigger word so it flows through the exact
    # same pipeline below, with no manual typing required.
    if not user_input and st.session_state.pending_auto_triage and active_inc:
        st.session_state.pending_auto_triage = False
        user_input = "Triage this incident"

    # Append user message immediately so it shows in the history render below
    if user_input:
        now = datetime.now().strftime("%H:%M:%S")
        st.session_state.chat_history.append(
            {"role": "user", "content": user_input, "ts": now}
        )

    # ── Chat history ───────────────────────────────────────────
    st.markdown(
        '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
        'letter-spacing:2px;margin-bottom:8px">■ CONVERSATION</div>',
        unsafe_allow_html=True,
    )

    if not st.session_state.chat_history:
        st.markdown(
            '<div style="text-align:center;padding:40px;font-family:var(--mono);'
            'font-size:0.78rem;color:var(--muted)">'
            '💬 SOC TRIAGE AGENT READY<br>'
            '<span style="font-size:0.62rem">'
            'Upload a file or select an incident, then type a message below</span></div>',
            unsafe_allow_html=True,
        )

    for msg in st.session_state.chat_history:
        ts = msg.get("ts", "")
        if msg["role"] == "user":
            st.markdown(
                f'<div class="bubble-user">'
                f'<div class="bubble-label" style="color:var(--muted)">YOU · {ts}</div>'
                f'{msg["content"]}</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                f'<div class="bubble-agent">'
                f'<div class="bubble-label">SOC AGENT · {ts}</div>'
                f'{msg["content"]}</div>',
                unsafe_allow_html=True,
            )

    # ── Thinking containers — in conversation flow, above fixed input ──────────
    # ── Agent execution ────────────────────────────────────────────────────────
    if user_input:

        reply = "⚠️ No response was generated."   # safe default

        # ── Triage trigger ─────────────────────────────────────────────────────
        if active_inc and _TRIAGE_TRIGGER.search(user_input):

            ALL_PHASES = [
                "IOC — Availability",
                "IOC — Confidentiality",
                "IOC — Integrity",
                "Risk Rating",
                "SOC Classification",
            ]
            PHASE_DESC = {
                "IOC — Availability":
                    "Scanning availability indicators — CPU spikes, reboots, buffer saturation",
                "IOC — Confidentiality":
                    "Scanning confidentiality indicators — exfil, geo anomalies, permission changes",
                "IOC — Integrity":
                    "Scanning integrity indicators — unknown binaries, hash anomalies, processes",
                "Risk Rating":
                    "Calculating risk across initiation, occurrence & adverse impact dimensions",
                "SOC Classification":
                    "Classifying incident severity and generating recommended actions",
            }

            # st.status() + st.write() is Streamlit's guaranteed progressive-update
            # pattern — st.write() calls inside the with-block appear live as each
            # phase fires; status.update(label=…) changes the header to the current phase.
            with st.status("⏳ Initialising triage pipeline…", expanded=True) as triage_status:

                thinking_panel = st.empty()   # token-by-token stream inside the status

                def on_progress(event: str, label: str, text: str = "") -> None:
                    key = next(
                        (p for p in ALL_PHASES
                         if label == p or label in p or p in label),
                        label,
                    )
                    if event == "phase_start":
                        desc = PHASE_DESC.get(key, label)
                        triage_status.update(label=f"🔍 {desc}…", expanded=True)
                        thinking_panel.empty()
                    elif event == "phase_complete":
                        result = f" — {text}" if text else ""
                        st.write(f"✅ **{key}**{result}")
                        thinking_panel.empty()

                try:
                    reply = soc_triage_chat_respond(
                        user_input,
                        incident           = active_inc,
                        llm_config         = get_cisco_cfg(),
                        progress_fn        = on_progress,
                        thinking_container = thinking_panel,
                    )
                    if not reply:
                        reply = "⚠️ Triage returned an empty response."
                    triage_status.update(
                        label="✅ Triage complete", state="complete", expanded=False
                    )
                except Exception as exc:
                    reply = f"❌ Triage error: {exc}"
                    triage_status.update(
                        label="❌ Triage failed", state="error", expanded=True
                    )

        # ── Plain Q&A fallback ─────────────────────────────────────────────────
        else:
            try:
                with st.spinner("Agent thinking…"):
                    reply = chat_respond(user_input, incident=active_inc)
                if not reply:
                    reply = "⚠️ Agent returned an empty response."
            except Exception as exc:
                reply = f"❌ Error: {exc}"

        # ── Pipeline auto-insert (fires on every triage trigger) ──────────
        if active_inc and _TRIAGE_TRIGGER.search(user_input):
            _inc_id  = str(active_inc.get("id") or active_inc.get("incidentId") or "")
            _title   = str(active_inc.get("title") or active_inc.get("name") or "Untitled")
            _sev     = normalise_sev(active_inc)
            _summary = reply[:500]

            pipeline_insert_full("alerts_to_triage", {
                "id": _inc_id or f"alert_{now.replace(':','')}",
                "incident_id": _inc_id, "title": _title,
                "severity": _sev, "summary": _summary})

            _needs = any(w in reply.lower() for w in [
                "investigate", "escalate", "critical", "high risk",
                "malicious", "confirmed", "requires action", "ioc confirmed"])
            if _needs:
                pipeline_insert_full("post_triage_investigate", {
                    "id": f"inv_{_inc_id}", "incident_id": _inc_id,
                    "title": _title, "severity": _sev, "summary": _summary})
            else:
                pipeline_insert_full("post_triage_no_investigate", {
                    "id": f"noinv_{_inc_id}", "incident_id": _inc_id,
                    "title": _title, "severity": _sev, "summary": _summary})

            _unc_m = re.search(r"#\d{5}[A-Z]+", reply)
            _unc   = _unc_m.group(0) if _unc_m else f"#TKT_{_inc_id[:6]}"
            pipeline_insert_full("initial_ticket", {
                "id": _unc, "incident_id": _inc_id,
                "title": f"Ticket {_unc} — {_title}",
                "severity": _sev, "summary": _summary})
            pipeline_insert_full("pending_ticket_report", {
                "id": f"pending_{_unc}", "incident_id": _inc_id,
                "title": f"[PENDING] {_title}", "severity": _sev,
                "summary": "Awaiting analyst sign-off before finalisation."})

        st.session_state.chat_history.append(
            {"role": "assistant", "content": reply,
             "ts": datetime.now().strftime("%H:%M:%S")}
        )
        st.rerun()

    if st.session_state.chat_history:
        if st.button("🗑️ Clear Chat", key="clear_chat"):
            st.session_state.chat_history = []
            st.rerun()


# ─────────────────────────────────────────────────────────────
# TAB 4 — CHROMADB
# ─────────────────────────────────────────────────────────────
with tab_chroma:
    st.markdown(
        '<div class="info-box"><div class="title">🧠 Knowledge Base</div>'
        'This is where the AI stores its understanding of your security alerts. '
        'Once synced, the AI can answer questions much more accurately. '
        'Use the search box below to find specific incidents by description.</div>',
        unsafe_allow_html=True,
    )
    if st.session_state.chroma_col is None:
        st.warning("⚠️ The Knowledge Base isn't connected yet. Connect it from the left panel under 🧠 Knowledge Base.")
    else:
        col = st.session_state.chroma_col
        st.markdown(f"### Knowledge Base — {col.count()} incidents indexed")
        st.markdown("---")

        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
            'letter-spacing:2px;margin-bottom:8px">■ SEMANTIC SEARCH</div>',
            unsafe_allow_html=True,
        )
        sq, sn = st.columns([5,1])
        query = sq.text_input("Query", placeholder="e.g. ransomware lateral movement C2",
                               label_visibility="collapsed")
        top_n = sn.number_input("N", 1, 20, 5, label_visibility="collapsed")

        if st.button("🔍 Search"):
            st.session_state.search_results = chroma_search(query, n=top_n)

        for r in st.session_state.search_results:
            m = r["meta"]
            st.markdown(
                f'<div style="background:#060C16;border:1px solid var(--border);'
                f'border-radius:5px;padding:10px 14px;margin:4px 0;'
                f'font-family:var(--mono);font-size:0.7rem">'
                f'<span class="badge badge-info">{r["score"]}% match</span>'
                f'&nbsp;<strong>{r["id"]}</strong>'
                f'&nbsp;<span style="color:var(--muted)">'
                f'sev:{m.get("severity","?")} · {m.get("status","?")}</span><br>'
                f'<span style="font-size:0.67rem;color:var(--text)">'
                f'{r["text"][:220]}…</span></div>',
                unsafe_allow_html=True,
            )

        st.markdown("---")
        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
            'letter-spacing:2px;margin-bottom:8px">■ ACTIONS</div>',
            unsafe_allow_html=True,
        )
        a1, a2 = st.columns(2)

        if a1.button("⬆️ Sync Incidents", use_container_width=True):
            if not incidents:
                st.warning("No incidents loaded yet.")
            else:
                n, msg = chroma_sync(incidents)
                st.success(msg) if n else st.error(msg)
                st.rerun()

        if a2.button("🗑️ Wipe & Reset", use_container_width=True):
            try:
                st.session_state.chroma_client.delete_collection("soc_incidents")
                chroma_connect("./chroma_db")
                st.success("Collection wiped and recreated.")
                st.rerun()
            except Exception as e:
                st.error(str(e))

        st.markdown("---")
        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
            'letter-spacing:2px;margin-bottom:8px">■ LANGCHAIN WIRING</div>',
            unsafe_allow_html=True,
        )
        st.code("""
# Replace chat_respond() body with your chain:

from langchain_community.vectorstores import Chroma
from langchain_anthropic import ChatAnthropic
from langchain.chains import RetrievalQA

vectorstore = Chroma(
    client=st.session_state.chroma_client,
    collection_name="soc_incidents",
    embedding_function=your_embedder,   # must match sync-time embedder
)
llm   = ChatAnthropic(model="claude-sonnet-4-20250514")
chain = RetrievalQA.from_chain_type(
    llm=llm,
    retriever=vectorstore.as_retriever(search_kwargs={"k": 5}),
)
return chain.invoke(user_msg)["result"]
        """, language="python")


# ─────────────────────────────────────────────────────────────
# TAB 5 — LOG HISTORY  (permanent SQLite store)
# ─────────────────────────────────────────────────────────────
with tab_log:
    st.markdown(
        '<div class="info-box"><div class="title">📁 History</div>'
        'All security alerts that have ever been loaded are saved here permanently, '
        'even after a restart. You can search, filter, export to Excel, or generate '
        'a report for any incident.</div>',
        unsafe_allow_html=True,
    )
    stats = db_stats()

    # ── Summary stats ──────────────────────────────────────────
    st.markdown(
        '<div style="font-size:0.75rem;font-weight:600;color:var(--muted);'
        'text-transform:uppercase;letter-spacing:1px;margin-bottom:12px">📊 Summary</div>',
        unsafe_allow_html=True,
    )
    ls1, ls2, ls3, ls4 = st.columns(4)
    ls1.metric("💾 Total Logged",    stats["total"])
    ls2.metric("🔄 Total Fetches",   stats["fetches"])
    ls3.metric("🚨 Critical (ever)", stats["by_sev"].get("CRITICAL", 0))
    ls4.metric("🟠 High (ever)",     stats["by_sev"].get("HIGH", 0))

    st.markdown("---")

    # ── Filters ────────────────────────────────────────────────
    st.markdown(
        '<div style="font-family:var(--mono);font-size:0.65rem;color:var(--muted);'
        'letter-spacing:2px;margin-bottom:10px">■ FILTER & SEARCH</div>',
        unsafe_allow_html=True,
    )
    hf1, hf2, hf3, hf4 = st.columns([1.2, 1.2, 2, 1])

    hist_sev    = hf1.selectbox("Severity", ["ALL","CRITICAL","HIGH","MEDIUM","LOW"],
                                 key="hist_sev", label_visibility="collapsed")
    # Collect unique statuses from DB for filter
    with db_connect() as _con:
        _statuses = ["ALL"] + [
            r[0] for r in _con.execute(
                "SELECT DISTINCT status FROM incidents ORDER BY status"
            ).fetchall() if r[0]
        ]
    hist_status = hf2.selectbox("Status", _statuses,
                                 key="hist_status", label_visibility="collapsed")
    hist_search = hf3.text_input("Search title / ID / assignee", "",
                                  placeholder="Type to filter…",
                                  label_visibility="collapsed")
    hist_limit  = hf4.number_input("Limit", 10, 5000, 200,
                                    label_visibility="collapsed")

    # ── Load rows ──────────────────────────────────────────────
    log_rows = db_load_incidents(
        severity=hist_sev,
        status=hist_status,
        search=hist_search,
        limit=hist_limit,
    )

    st.markdown(
        f'<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
        f'margin:8px 0">{len(log_rows)} records shown</div>',
        unsafe_allow_html=True,
    )

    # ── Export button ──────────────────────────────────────────
    csv_data = db_export_csv()
    if csv_data:
        st.download_button(
            label="📥 Export all as CSV",
            data=csv_data,
            file_name=f"soc_incidents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv",
        )

    st.markdown("---")

    # ── Incident rows ──────────────────────────────────────────
    if not log_rows:
        st.markdown(
            '<div style="text-align:center;padding:50px;font-family:var(--mono);'
            'font-size:0.78rem;color:var(--muted)">'
            '● NO RECORDS IN DATABASE<br>'
            '<span style="font-size:0.62rem">'
            'Incidents are saved automatically on every fetch</span></div>',
            unsafe_allow_html=True,
        )
    else:
        for row in log_rows:
            sev   = row.get("severity", "LOW")
            color = SEV_COLORS.get(sev, "#3A607A")
            inc_id    = row.get("id", "—")
            title     = row.get("title", "Untitled")
            status    = row.get("status", "—")
            assignee  = row.get("assignee", "—")
            alerts    = row.get("alert_count", "—")
            created   = str(row.get("created", ""))[:16]
            first_seen= str(row.get("first_seen",""))[:16]
            last_seen = str(row.get("last_seen", ""))[:16]

            st.markdown(
                f'<div class="card card-{sev.lower()}">'
                f'<div style="display:flex;align-items:center;gap:10px;flex-wrap:wrap">'
                f'<span class="badge badge-{sev.lower()}">{sev}</span>'
                f'<strong style="flex:1;font-size:0.88rem">{title}</strong>'
                f'<span style="color:var(--muted);font-size:0.62rem">Incident ID:</span>'
                f'<code style="color:var(--muted);font-size:0.68rem">{inc_id}</code>'
                f'</div>'
                f'<div style="margin-top:8px;font-size:0.73rem;color:var(--muted);'
                f'display:flex;gap:16px;flex-wrap:wrap">'
                f'<span>📌 {status}</span>'
                f'<span>👤 {assignee}</span>'
                f'<span>🚨 {alerts} alerts</span>'
                f'<span>📅 created {created}</span>'
                f'<span style="color:#1A4A62">first seen {first_seen}</span>'
                f'<span style="color:#1A4A62">last seen {last_seen}</span>'
                f'</div></div>',
                unsafe_allow_html=True,
            )

            bj, _ = st.columns([0.7, 8])
            if bj.button("{ }", key=f"log_json_{inc_id}"):
                try:
                    raw = _json.loads(row.get("raw_json") or "{}")
                except Exception:
                    raw = row
                with st.expander(f"Full JSON — {inc_id}", expanded=True):
                    st.json(raw)
            
            # Associated Alerts / Logs
            try:
                raw_dict = _json.loads(row.get("raw_json") or "{}")
            except Exception:
                raw_dict = {}
            alerts_list = raw_dict.get("alerts")
            if alerts_list:
                with st.expander(f"🚨 Associated Alerts ({len(alerts_list)})", expanded=False):
                    for alert in alerts_list:
                        a_title = alert.get("title") or alert.get("name") or "Untitled Alert"
                        a_id = alert.get("id") or ""
                        a_source = alert.get("source") or "Unknown"
                        a_type = alert.get("type") or "Unknown"
                        a_created = alert.get("created") or alert.get("receivedTime") or ""
                        st.markdown(
                            f'<div style="background:#091624;padding:8px 12px;border-radius:4px;margin-bottom:6px;border-left:3px solid var(--accent)">'
                            f'<div style="display:flex;justify-content:between;align-items:center">'
                            f'<strong>{a_title}</strong>'
                            f'<code style="color:var(--muted);font-size:0.7rem;margin-left:auto">{a_id}</code>'
                            f'</div>'
                            f'<div style="font-size:0.72rem;color:var(--muted);margin-top:4px">'
                            f'Incident ID: <strong style="color:var(--accent)">{inc_id}</strong> &nbsp;·&nbsp; '
                            f'Source: {a_source} &nbsp;·&nbsp; Type: {a_type} &nbsp;·&nbsp; Time: {a_created}'
                            f'</div>'
                            f'</div>',
                            unsafe_allow_html=True
                        )
                        # Nested event details
                        events = alert.get("events")
                        if events:
                            for idx, ev in enumerate(events):
                                ev_src = ev.get("source", {})
                                ev_dst = ev.get("destination", {})
                                src_ip = ev_src.get("device", {}).get("ipAddress") or ev_src.get("ipAddress") or "—"
                                dst_ip = ev_dst.get("device", {}).get("ipAddress") or ev_dst.get("ipAddress") or "—"
                                user = ev_src.get("user", {}).get("username") or ev_src.get("username") or "—"
                                ev_proto = ev.get("ip_proto") or ev.get("protocol") or "—"
                                ev_port = ev_dst.get("device", {}).get("port") or ev_dst.get("port") or "—"
                                st.markdown(
                                    f'<div style="margin-left:15px;padding:4px 8px;font-family:var(--mono);font-size:0.7rem;color:var(--muted);border-left:1px dashed #1B4A62">'
                                    f'Event {idx+1}: 🧑‍💻 {user} | ➡️ {src_ip} → {dst_ip} (Port {ev_port}, Proto {ev_proto})'
                                    f'</div>',
                                    unsafe_allow_html=True
                                )

# ═══════════════════════════════════════════════════════════════════════════════
# TAB 6 — PIPELINE DB
# Fully inline — single streamlit run app.py, no second process.
# st.stop() removed from tab_chroma so this tab always renders.
# ═══════════════════════════════════════════════════════════════════════════════
with tab_pipeline:
    st.markdown(
        '<div class="info-box"><div class="title">🔁 Data Pipeline</div>'
        'This tool helps IT staff manage how security data flows through the system — '
        'from collection, to review, to archiving. If you are not sure what this is for, '
        'you likely do not need to use it.</div>',
        unsafe_allow_html=True,
    )

    # ── session defaults ──────────────────────────────────────────────────────
    if "pl_stage"      not in st.session_state: st.session_state.pl_stage      = None
    if "pl_cv_results" not in st.session_state: st.session_state.pl_cv_results = []

    st.markdown(
        '<div style="font-family:var(--mono);font-size:0.65rem;color:var(--muted);'
        'letter-spacing:2px;margin-bottom:14px">'
        '■ SOC PIPELINE — CLICK SELECT UNDER ANY STAGE TO OPEN ITS DB VIEWER</div>',
        unsafe_allow_html=True,
    )

    # ══════════════════════════════════════════════════════════════════════════
    # STAGE CARDS  (always visible — one column per stage)
    # ══════════════════════════════════════════════════════════════════════════
    _pl_cols = st.columns(len(PIPELINE_STAGES))
    for _idx, _stage in enumerate(PIPELINE_STAGES):
        _cnt    = pipeline_count(_stage)
        _icon   = PIPELINE_ICONS[_stage]
        _label  = PIPELINE_LABELS[_stage]
        _color  = PIPELINE_COLORS[_stage]
        _active = st.session_state.pl_stage == _stage
        _border = f"3px solid {_color}" if _active else f"1px solid {_color}44"
        _bg     = f"{_color}18"          if _active else "#060C16"
        with _pl_cols[_idx]:
            st.markdown(
                f'<div style="background:{_bg};border:{_border};border-radius:7px;'
                f'padding:10px 8px;text-align:center;transition:all 0.15s">'
                f'<div style="font-size:1.3rem">{_icon}</div>'
                f'<div style="font-family:var(--mono);font-size:1.5rem;'
                f'color:{_color};margin:4px 0">{_cnt}</div>'
                f'<div style="font-family:var(--mono);font-size:0.48rem;'
                f'color:var(--muted);letter-spacing:1px;line-height:1.4">'
                f'{_label.upper()}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
            if st.button("Select", key=f"pl_sel_{_stage}", use_container_width=True):
                st.session_state.pl_stage      = None if _active else _stage
                st.session_state.pl_cv_results = []
                st.rerun()

    # ══════════════════════════════════════════════════════════════════════════
    # INLINE VIEWER  (expands below when a stage is selected)
    # ══════════════════════════════════════════════════════════════════════════
    _sel = st.session_state.pl_stage

    if not _sel:
        st.markdown(
            '<div style="text-align:center;padding:60px 20px;font-family:var(--mono);'
            'font-size:0.78rem;color:var(--muted)">'
            '↑ Click <strong style="color:var(--accent)">Select</strong> '
            'under any stage above to open its database viewer<br>'
            '<span style="font-size:0.62rem;margin-top:8px;display:block">'
            'Records auto-insert when the triage agent runs in the Chat tab.'
            '</span></div>',
            unsafe_allow_html=True,
        )
    else:
        _color = PIPELINE_COLORS[_sel]
        _icon  = PIPELINE_ICONS[_sel]
        _label = PIPELINE_LABELS[_sel]

        st.markdown("<br>", unsafe_allow_html=True)

        # ── header banner ─────────────────────────────────────────────────────
        st.markdown(
            f'<div style="background:#060C16;border:2px solid {_color};'
            f'border-left:8px solid {_color};border-radius:8px;'
            f'padding:18px 24px;margin-bottom:16px">'
            f'<div style="font-family:var(--mono);font-size:1.05rem;'
            f'color:{_color};letter-spacing:3px">{_icon} {_label.upper()}</div>'
            f'<div style="font-family:var(--mono);font-size:0.58rem;'
            f'color:var(--muted);margin-top:6px">'
            f'ChromaDB collection: <code style="color:var(--accent)">pipeline_{_sel}</code>'
            f'&nbsp;·&nbsp;SQLite table: <code style="color:var(--accent)">{_sel}</code>'
            f'</div></div>',
            unsafe_allow_html=True,
        )

        # ── metrics ───────────────────────────────────────────────────────────
        _sql_cnt = pipeline_count(_sel)
        _vec_cnt = pipeline_chroma_count(_sel)
        _vm1, _vm2, _vm3, _vm4 = st.columns(4)
        _vm1.metric("💾 SQLite Records",   _sql_cnt)
        _vm2.metric("🧠 ChromaDB Vectors", _vec_cnt)
        _vm3.metric("📋 Stage",            _label[:22])
        _vm4.metric("🗂️ Collection",       f"pipeline_{_sel}"[:24])

        # ── action strip ──────────────────────────────────────────────────────
        _ac1, _ac2, _ac3, _ac4, _ = st.columns([1, 1.4, 1, 1.1, 2.5])

        if _ac1.button("➕ Test Insert", key=f"pl_add_{_sel}"):
            import uuid as _uuid3
            pipeline_insert_full(_sel, {
                "id":          f"test_{str(_uuid3.uuid4())[:8]}",
                "incident_id": "DEMO-001",
                "title":       f"[Demo] Test record · {_label}",
                "severity":    "MEDIUM",
                "summary":     f"Auto-inserted test record into: {_label}",
            })
            st.success("Test record inserted!")
            st.rerun()

        if _sel == "pending_ticket_report":
            if _ac2.button("✅ Finalize All", key="pl_finalize_btn"):
                _pending = pipeline_load("pending_ticket_report")
                for _prow in _pending:
                    _fin = dict(_prow)
                    _fin["id"]      = f"final_{_prow['id']}"
                    _fin["title"]   = (_fin.get("title") or "").replace("[PENDING]","[FINAL]")
                    _fin["summary"] = "Finalized. " + (_prow.get("summary") or "")
                    pipeline_insert_full("finalized_report", _fin)
                st.success(f"Finalized {len(_pending)} records → Finalized Report.")
                st.rerun()

        if _ac3.button("🗑️ Clear Stage", key=f"pl_clear_{_sel}"):
            with _pl_con() as _c:
                _c.execute(f"DELETE FROM {_sel}")
                _c.commit()
            _wipe_col = _pl_chroma_col(_sel)
            if _wipe_col:
                try:
                    st.session_state.chroma_client.delete_collection(f"pipeline_{_sel}")
                except Exception:
                    pass
            st.success(f"Cleared {_label}.")
            st.rerun()

        if _ac4.button("✕ Close Viewer", key="pl_close_viewer"):
            st.session_state.pl_stage      = None
            st.session_state.pl_cv_results = []
            st.rerun()

        st.markdown("---")

        # ── sub-tabs ──────────────────────────────────────────────────────────
        _vtab_sql, _vtab_chroma = st.tabs(["💾  SQLITE RECORDS", "🧠  CHROMADB SEARCH"])

        # ════════════════════ SQLite Records ════════════════════
        with _vtab_sql:
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;'
                'color:var(--muted);letter-spacing:2px;margin-bottom:10px">'
                '■ ALL RECORDS IN THIS STAGE</div>',
                unsafe_allow_html=True,
            )
            _fc1, _fc2 = st.columns([5, 1])
            _sql_srch = _fc1.text_input(
                "filter", placeholder="Filter by title / ID / summary…",
                key=f"sql_srch_{_sel}", label_visibility="collapsed")
            _sql_lim = int(_fc2.number_input(
                "limit", 10, 2000, 200,
                key=f"sql_lim_{_sel}", label_visibility="collapsed"))

            _rows = pipeline_load(_sel, limit=_sql_lim)
            if _sql_srch.strip():
                _s = _sql_srch.lower()
                _rows = [r for r in _rows if (
                    _s in (r.get("title")   or "").lower() or
                    _s in (r.get("id")      or "").lower() or
                    _s in (r.get("summary") or "").lower())]

            if not _rows:
                st.markdown(
                    '<div style="text-align:center;padding:50px;font-family:var(--mono);'
                    'font-size:0.78rem;color:var(--muted)">● NO RECORDS FOUND<br>'
                    '<span style="font-size:0.62rem">'
                    'Run a triage in the Chat tab, or click ➕ Test Insert above.'
                    '</span></div>',
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(
                    f'<div style="font-family:var(--mono);font-size:0.6rem;'
                    f'color:var(--muted);margin-bottom:8px">{len(_rows)} records</div>',
                    unsafe_allow_html=True,
                )
                for _row in _rows:
                    _r_id  = _row.get("id", "—")
                    _r_ttl = _row.get("title", "Untitled")
                    _r_sev = str(_row.get("severity", "—")).upper()
                    _r_inc = _row.get("incident_id", "—")
                    _r_sum = (_row.get("summary") or "")[:200]
                    _r_ts  = str(_row.get("created_at", ""))[:16]
                    _sev_c = SEV_COLORS.get(_r_sev, "#3A607A")

                    # ── Stage-specific export badge ─────────────────────────
                    _is_csv_stage  = _sel in ("post_triage_investigate", "post_triage_no_investigate")
                    _is_docx_stage = _sel == "initial_ticket"
                    _export_badge  = ""
                    if _is_csv_stage:
                        _export_badge = (
                            f'<span style="background:#00403A;color:#0AF0A0;'
                            f'border:1px solid #0AF0A044;padding:2px 8px;'
                            f'border-radius:3px;font-family:var(--mono);font-size:0.56rem;'
                            f'margin-left:6px">📊 CSV</span>'
                        )
                    elif _is_docx_stage:
                        _export_badge = (
                            f'<span style="background:#1A0040;color:#A78BFA;'
                            f'border:1px solid #A78BFA44;padding:2px 8px;'
                            f'border-radius:3px;font-family:var(--mono);font-size:0.56rem;'
                            f'margin-left:6px">📄 DOCX</span>'
                        )

                    st.markdown(
                        f'<div style="background:#060C16;border:1px solid {_color}44;'
                        f'border-left:3px solid {_color};border-radius:6px;'
                        f'padding:10px 14px;margin:4px 0">'
                        f'<div style="display:flex;align-items:center;gap:10px;flex-wrap:wrap">'
                        f'<code style="font-family:var(--mono);font-size:0.65rem;color:{_color}">'
                        f'{_r_id}</code>'
                        f'<strong style="flex:1;font-size:0.85rem">{_r_ttl}</strong>'
                        f'{_export_badge}'
                        f'<span style="background:{_sev_c}22;color:{_sev_c};'
                        f'border:1px solid {_sev_c}44;padding:2px 8px;border-radius:3px;'
                        f'font-family:var(--mono);font-size:0.6rem">{_r_sev}</span>'
                        f'<span style="font-family:var(--mono);font-size:0.56rem;'
                        f'color:var(--muted)">{_r_ts}</span>'
                        f'</div>'
                        f'<div style="font-size:0.7rem;color:var(--muted);margin-top:6px">'
                        f'Incident: <code style="color:var(--accent)">{_r_inc}</code>'
                        f'{"&nbsp;·&nbsp;" + _r_sum if _r_sum else ""}'
                        f'</div></div>',
                        unsafe_allow_html=True,
                    )

                    # ── Button row: JSON | export | delete ──────────────────
                    if _is_csv_stage:
                        _bj1, _bj2, _bj3, _ = st.columns([0.6, 1.2, 0.28, 7])
                    elif _is_docx_stage:
                        _bj1, _bj2, _bj3, _ = st.columns([0.6, 1.3, 0.28, 7])
                    else:
                        _bj1, _bj3, _ = st.columns([0.6, 0.28, 9])
                        _bj2 = None

                    if _bj1.button("{ } JSON", key=f"pl_json_{_sel}_{_r_id}"):
                        try:
                            _raw_j = _json.loads(_row.get("raw_json") or "{}")
                        except Exception:
                            _raw_j = _row
                        with st.expander(f"JSON — {_r_id}", expanded=True):
                            st.json(_raw_j)

                    # ── CSV download for post-triage stages ─────────────────
                    if _is_csv_stage and _bj2 is not None:
                        _safe_id = re.sub(r"[^A-Za-z0-9_\-]", "_", _r_id)[:40]
                        _csv_bytes = _make_csv_bytes(_row)
                        _bj2.download_button(
                            label="📊 View as Sheet",
                            data=_csv_bytes,
                            file_name=f"triage_{_safe_id}.csv",
                            mime="text/csv",
                            key=f"pl_csv_{_sel}_{_r_id}",
                        )

                    # ── DOCX download for initial_ticket stage ──────────────
                    elif _is_docx_stage and _bj2 is not None:
                        _safe_id = re.sub(r"[^A-Za-z0-9_\-]", "_", _r_id)[:40]
                        _docx_bytes = _make_docx_bytes(_row)
                        # Detect if python-docx is available (bytes start with PK zip magic)
                        _is_real_docx = _docx_bytes[:2] == b'PK'
                        _bj2.download_button(
                            label="📄 View as Ticket",
                            data=_docx_bytes,
                            file_name=f"ticket_{_safe_id}.{'docx' if _is_real_docx else 'txt'}",
                            mime=(
                                "application/vnd.openxmlformats-officedocument"
                                ".wordprocessingml.document"
                                if _is_real_docx else "text/plain"
                            ),
                            key=f"pl_docx_{_sel}_{_r_id}",
                        )

                    if _bj3.button("✕", key=f"pl_del_{_sel}_{_r_id}"):
                        pipeline_delete(_sel, _r_id)
                        st.rerun()

        # ════════════════════ ChromaDB Search ════════════════════
        with _vtab_chroma:
            _chroma_ok = CHROMA_OK and bool(st.session_state.get("chroma_client"))

            if not _chroma_ok:
                st.markdown(
                    '<div style="background:#0A0608;border:1px solid #2A1010;'
                    'border-radius:6px;padding:12px 16px;font-family:var(--mono);'
                    'font-size:0.68rem;margin-bottom:12px">'
                    '<span style="color:var(--warn)">⚠ ChromaDB not connected</span><br>'
                    '<span style="color:var(--muted);font-size:0.6rem">'
                    'Connect ChromaDB from the sidebar. SQLite records persist regardless.'
                    '</span></div>',
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(
                    f'<div style="font-family:var(--mono);font-size:0.65rem;'
                    f'color:var(--muted);margin-bottom:10px">'
                    f'{pipeline_chroma_count(_sel)} vectors in '
                    f'<code style="color:var(--accent)">pipeline_{_sel}</code></div>',
                    unsafe_allow_html=True,
                )

            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
                'letter-spacing:2px;margin-bottom:8px">■ SEMANTIC SEARCH</div>',
                unsafe_allow_html=True,
            )
            _sq2, _sn2 = st.columns([5, 1])
            _cv_q = _sq2.text_input(
                "query", key=f"cv_q_{_sel}",
                placeholder="e.g. ransomware lateral movement C2",
                label_visibility="collapsed")
            _cv_n = int(_sn2.number_input(
                "topn", 1, 20, 5,
                key=f"cv_n_{_sel}", label_visibility="collapsed"))

            if st.button("🔍 Search", key=f"cv_srch_{_sel}"):
                if not _chroma_ok:
                    st.warning("Connect ChromaDB from the sidebar first.")
                elif not _cv_q.strip():
                    st.warning("Enter a query.")
                else:
                    st.session_state.pl_cv_results = pipeline_chroma_search(
                        _sel, _cv_q, n=_cv_n)

            for _r in st.session_state.pl_cv_results:
                _rm   = _r["meta"]
                _scol = PIPELINE_COLORS.get(_rm.get("stage", _sel), _color)
                st.markdown(
                    f'<div style="background:#060C16;border:1px solid {_scol}55;'
                    f'border-left:3px solid {_scol};border-radius:6px;'
                    f'padding:10px 14px;margin:4px 0">'
                    f'<div style="display:flex;align-items:center;gap:10px">'
                    f'<span style="background:{_scol}22;color:{_scol};'
                    f'border:1px solid {_scol}44;padding:2px 8px;border-radius:3px;'
                    f'font-family:var(--mono);font-size:0.6rem">{_r["score"]}%</span>'
                    f'<strong style="flex:1">{_r["id"]}</strong>'
                    f'<span style="font-family:var(--mono);font-size:0.58rem;'
                    f'color:var(--muted)">'
                    f'sev:{_rm.get("severity","?")} · {_rm.get("created","")[:10]}'
                    f'</span></div>'
                    f'<div style="font-size:0.72rem;color:var(--text);margin-top:6px">'
                    f'{str(_r["doc"])[:300]}</div></div>',
                    unsafe_allow_html=True,
                )

            st.markdown("---")
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
                'letter-spacing:2px;margin-bottom:8px">■ BROWSE ALL VECTORS</div>',
                unsafe_allow_html=True,
            )
            if st.button("📋 Load All Vectors", key=f"cv_all_{_sel}"):
                if not _chroma_ok:
                    st.warning("Connect ChromaDB from the sidebar first.")
                else:
                    _all_v = pipeline_chroma_all(_sel)
                    if not _all_v:
                        st.info("No vectors in this collection yet.")
                    for _v in _all_v:
                        _vm  = _v["meta"]
                        _vc  = PIPELINE_COLORS.get(_vm.get("stage", _sel), _color)
                        st.markdown(
                            f'<div style="background:#060C16;border:1px solid {_vc}44;'
                            f'border-left:3px solid {_vc};border-radius:6px;'
                            f'padding:8px 14px;margin:3px 0">'
                            f'<div style="display:flex;gap:10px;align-items:center">'
                            f'<code style="font-family:var(--mono);font-size:0.62rem;'
                            f'color:{_vc}">{_v["id"]}</code>'
                            f'<span style="font-family:var(--mono);font-size:0.56rem;'
                            f'color:var(--muted)">'
                            f'sev:{_vm.get("severity","?")} · {_vm.get("created","")[:10]}'
                            f'</span></div>'
                            f'<div style="font-size:0.7rem;color:var(--text);margin-top:5px">'
                            f'{str(_v["doc"])[:250]}</div></div>',
                            unsafe_allow_html=True,
                        )

            st.markdown("---")
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
                'letter-spacing:2px;margin-bottom:8px">■ COLLECTION ACTIONS</div>',
                unsafe_allow_html=True,
            )
            if st.button(f"🗑️ Wipe Chroma: pipeline_{_sel}", key=f"cv_wipe_{_sel}"):
                if not _chroma_ok:
                    st.warning("Connect ChromaDB from the sidebar first.")
                else:
                    try:
                        st.session_state.chroma_client.delete_collection(f"pipeline_{_sel}")
                        _pl_chroma_col(_sel)
                        st.success(f"Wiped pipeline_{_sel}.")
                        st.rerun()
                    except Exception as _we:
                        st.error(str(_we))