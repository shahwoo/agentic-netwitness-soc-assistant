from __future__ import annotations

import json
import os
import re
from backend.openai_client import extract_json_object, invoke_openai_text
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from jinja2 import Environment, FileSystemLoader, select_autoescape

from reporting.context_builder import build_context
from reporting.export_context_enhancer import enhance_export_context
from reporting.compact_renderer import count_placeholders, is_placeholder
from reporting.structured_report import (
    paragraph_contains_raw_pipe_table,
    parse_pipe_table,
    repair_pipe_tables_in_blocks,
)
from backend.export_cache import calculate_source_hash, is_cache_ready, mark_export_status

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except Exception:  # pragma: no cover
    Document = None


# Reporting exports must use the existing 4 templates in report_templates/ only.
REPORT_TEMPLATES: dict[str, dict[str, str]] = {
    "executive_summary": {
        "title": "Executive Summary",
        "template": "executive_summary_template.md.j2",
        "filename": "executive_summary",
    },
    "technical_findings": {
        "title": "Technical Findings",
        "template": "technical_findings_template.md.j2",
        "filename": "technical_findings",
    },
    "soc_analyst_review": {
        "title": "SOC Analyst Review",
        "template": "soc_analyst_review_template.md.j2",
        "filename": "soc_analyst_review",
    },
    "final_incident_report": {
        "title": "Final Incident Report",
        "template": "incident_report_template.md.j2",
        "filename": "final_incident_report",
    },
}

# Non-reporting agent exports use a separate folder so report_templates/ stays clean.
AGENT_TEMPLATES: dict[str, dict[str, str]] = {
    "parsing": {
        "title": "Parsing and Normalisation Report",
        "template": "parsing_normalisation_report_template.md.j2",
        "filename": "parsing_normalisation",
        "json_file": "parser_result.json",
    },
    "parsing_normalisation": {
        "title": "Parsing and Normalisation Report",
        "template": "parsing_normalisation_report_template.md.j2",
        "filename": "parsing_normalisation",
        "json_file": "parser_result.json",
    },
    "triage": {
        "title": "SOC Triage Assessment Report",
        "template": "triage_summary_template.md.j2",
        "filename": "triage_assessment",
        "json_file": "triage_result.json",
    },
    "threat_intel": {
        "title": "Threat Intelligence Enrichment Report",
        "template": "threat_intelligence_report_template.md.j2",
        "filename": "threat_intelligence_enrichment",
        "json_file": "threat_intel_result.json",
    },
    "threat_intelligence": {
        "title": "Threat Intelligence Enrichment Report",
        "template": "threat_intelligence_report_template.md.j2",
        "filename": "threat_intelligence_enrichment",
        "json_file": "threat_intel_result.json",
    },
    "triage_approval": {
        "title": "SOC Triage Approval Decision Report",
        "template": "approval_summary_template.md.j2",
        "filename": "triage_approval_decision",
        "json_file": "approval_result.json",
    },
    "approval": {
        "title": "SOC Approval Decision Report",
        "template": "approval_summary_template.md.j2",
        "filename": "approval_decision",
        "json_file": "approval_result.json",
    },
    "investigation": {
        "title": "SOC Investigation Findings Report",
        "template": "investigation_summary_template.md.j2",
        "filename": "investigation_findings",
        "json_file": "investigation_result.json",
    },
    "investigation_approval": {
        "title": "SOC Investigation Approval Decision Report",
        "template": "approval_summary_template.md.j2",
        "filename": "investigation_approval_decision",
        "json_file": "investigation_approval_result.json",
    },
    "soc_review": {
        "title": "SOC Analyst Report Review",
        "template": "soc_analyst_review_gate_template.md.j2",
        "filename": "soc_analyst_review_gate",
        "json_file": "soc_review_result.json",
    },
    "soc_analyst_review": {
        "title": "SOC Analyst Report Review",
        "template": "soc_analyst_review_gate_template.md.j2",
        "filename": "soc_analyst_review_gate",
        "json_file": "soc_review_result.json",
    },
}

FORBIDDEN_TEXT_PATTERNS = [
    "Template basis:",
    "Expected template sections:",
    "No report content was returned by the Reporting Agent",
    "No key points found in the saved output",
    "/Users/",
    "\\Users\\",
]

BRAND_BLUE = RGBColor(47, 106, 205)
BRAND_BLUE_DARK = RGBColor(28, 41, 60)
BRAND_GREY = RGBColor(92, 104, 121)
BORDER_BLUE = "2F6ACD"
LIGHT_BLUE = "EAF2FF"
LIGHT_GREY = "F4F7FB"
LIGHT_YELLOW = "FFF7DB"
LIGHT_GREEN = "E9F8EF"
LIGHT_RED = "FDEBEC"
WHITE = "FFFFFF"
TEXT_DARK = "1C293C"
BORDER_GREY = "C8D4E2"


def aegis_logo_path(project_root: Path) -> Path | None:
    """Return the report-safe Aegis logo path if it is available."""
    for candidate in [
        project_root / "report_assets" / "aegis-logo.png",
        project_root / "dashboard" / "assets" / "aegis-logo.png",
    ]:
        if candidate.exists() and candidate.stat().st_size > 0:
            return candidate
    return None


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def safe_filename(value: Any) -> str:
    text = str(value or "unknown")
    text = re.sub(r"[^A-Za-z0-9_.-]+", "_", text).strip("_")
    return text[:120] or "unknown"


def read_json(path: Path) -> dict[str, Any]:
    if not path.exists() or path.stat().st_size == 0:
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def scrub_template_context(value: Any, key: str = "") -> Any:
    """Remove local file paths and noisy artefacts before Jinja rendering."""
    noisy_keys = {
        "stdout", "stderr", "stdout_tail", "stderr_tail", "raw_stdout", "raw_stderr",
        "generated_reports", "report_manifest", "final_report_text", "output_files", "exports",
    }
    path_like_keys = {"path", "raw_report_path", "relative_path", "draft_path", "confirmed_path", "download_url"}
    if isinstance(value, dict):
        cleaned: dict[str, Any] = {}
        for k, v in value.items():
            lk = str(k).lower()
            if lk in noisy_keys:
                cleaned[k] = "Omitted from analyst-facing export. See system logs or JSON artefacts if raw execution details are required."
            elif lk in path_like_keys or lk.endswith("_path") or lk.endswith("_dir"):
                cleaned[k] = "Omitted from analyst-facing export."
            elif lk == "raw_thinking":
                cleaned[k] = "Omitted. Analyst export shows reasoning summary, not raw model scratchpad."
            else:
                cleaned[k] = scrub_template_context(v, lk)
        return cleaned
    if isinstance(value, list):
        return [scrub_template_context(v, key) for v in value]
    if isinstance(value, str):
        text = value
        text = re.sub(r"/Users/[^\s\"']+", "[local file path omitted]", text)
        text = re.sub(r"[A-Za-z]:\\\\[^\s\"']+", "[local file path omitted]", text)
        return text
    return value


def write_json(path: Path, payload: Any) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    return path


def first_present(*values: Any, default: Any = "Not Provided") -> Any:
    for value in values:
        if value not in (None, "", [], {}):
            return value
    return default


def to_list(value: Any) -> list[Any]:
    if value in (None, ""):
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, tuple):
        return list(value)
    return [value]


def get_nested(obj: dict[str, Any] | None, path: str, default: Any = None) -> Any:
    cur: Any = obj or {}
    for part in path.split("."):
        if isinstance(cur, dict) and part in cur:
            cur = cur[part]
        else:
            return default
    return cur


def load_outputs_context(project_root: Path, output_dir: Path, ticket: dict[str, Any] | None = None) -> tuple[dict[str, dict[str, Any]], list[str]]:
    input_dir = project_root / "inputs"
    outputs = {
        "processed_alert": read_json(output_dir / "processed_alert.json") or read_json(input_dir / "processed_alert.json"),
        "parser_result": read_json(output_dir / "parser_result.json") or read_json(input_dir / "parser_result.json"),
        "enriched_alert": read_json(output_dir / "enriched_alert.json") or read_json(input_dir / "enriched_alert.json"),
        "triage_result": read_json(output_dir / "triage_result.json") or read_json(input_dir / "triage_result.json"),
        "threat_intel_result": read_json(output_dir / "threat_intel_result.json") or read_json(input_dir / "threat_intel_result.json"),
        "investigation_result": read_json(output_dir / "investigation_result.json") or read_json(input_dir / "investigation_result.json"),
        "approval_result": read_json(output_dir / "approval_result.json") or read_json(input_dir / "approval_result.json"),
        "investigation_approval_result": read_json(output_dir / "investigation_approval_result.json") or read_json(input_dir / "investigation_approval_result.json"),
        "soc_review_result": read_json(output_dir / "soc_review_result.json") or read_json(input_dir / "soc_review_result.json"),
    }
    ticket = ticket or {}
    for src_key, dst_key in [
        ("parsing_result", "parser_result"),
        ("triage_result", "triage_result"),
        ("threat_intel_result", "threat_intel_result"),
        ("investigation_result", "investigation_result"),
        ("approval_result", "approval_result"),
        ("investigation_approval_result", "investigation_approval_result"),
        ("soc_review_result", "soc_review_result"),
    ]:
        if isinstance(ticket.get(src_key), dict) and ticket[src_key]:
            outputs[dst_key] = ticket[src_key]
    warnings = []
    for key in ["enriched_alert", "triage_result", "investigation_result"]:
        if not outputs.get(key):
            warnings.append(f"Missing {key}. Export will mark unavailable fields as Not Provided.")
    return outputs, warnings


def build_report_context(project_root: Path, output_dir: Path, ticket: dict[str, Any] | None = None) -> dict[str, Any]:
    inputs, warnings = load_outputs_context(project_root, output_dir, ticket)
    context = build_context(inputs, warnings, output_dir=output_dir)
    ticket = ticket or {}
    if ticket:
        context["ticket"] = ticket
        context["ticket_id"] = first_present(ticket.get("ticket_id"), context.get("ticket_id"), context.get("incident_id"))
        context["case_title"] = first_present(ticket.get("title"), context.get("case_title"), default="Not Provided")
        context["owner"] = first_present(ticket.get("owner"), default="Not Provided")
        context["current_stage"] = first_present(ticket.get("current_stage"), default=context.get("current_stage", "Not Provided"))
    context = enhance_export_context(context, ticket)
    context.setdefault("render_quality_checks", {})
    context["render_quality_checks"]["total_placeholders"] = count_placeholders(context)
    context["render_quality_checks"]["placeholders_by_section"] = {
        "evidence_register": count_placeholders(context.get("compact_evidence_register") or context.get("evidence") or []),
        "approval": count_placeholders(context.get("approval") or {}),
        "containment": count_placeholders(context.get("containment") or {}),
        "data_impact": count_placeholders(context.get("data_impact_summary")),
        "chain_of_custody": count_placeholders(context.get("chain_of_custody_note")),
    }
    context["render_quality_checks"]["columns_hidden"] = (context.get("compact_evidence_register") or {}).get("hidden_columns", [])
    compact_tables = context.get("compact_tables") or {}
    context["render_quality_checks"]["rows_hidden"] = sum(int((t or {}).get("hidden_rows") or 0) for t in compact_tables.values() if isinstance(t, dict)) + int((context.get("compact_evidence_register") or {}).get("hidden_rows") or 0)
    context["render_quality_checks"]["ai_narrative_sections"] = ["executive_summary", "technical_analysis", "business_impact_explanation", "attack_narrative", "conclusion", "analyst_friendly_explanation", "soc_analyst_review_checklist"] if context.get("llm_used") else []
    return scrub_template_context(context)


def build_agent_context(project_root: Path, output_dir: Path, ticket: dict[str, Any] | None, agent_key: str) -> dict[str, Any]:
    report_context = build_report_context(project_root, output_dir, ticket)
    template_cfg = AGENT_TEMPLATES[agent_key]
    output = read_json(output_dir / template_cfg["json_file"])
    ticket = ticket or {}
    if isinstance(ticket.get(f"{agent_key}_result"), dict) and ticket.get(f"{agent_key}_result"):
        output = ticket[f"{agent_key}_result"]
    llm_fields = build_agent_llm_fields(agent_key, report_context, output)
    agent_context = {
        **report_context,
        "agent_key": agent_key,
        "agent_title": template_cfg["title"],
        "agent_output": output,
        "triage": report_context.get("raw_inputs", {}).get("triage_result", {}) or read_json(output_dir / "triage_result.json"),
        "investigation": report_context.get("raw_inputs", {}).get("investigation_result", {}) or read_json(output_dir / "investigation_result.json"),
        "approval_raw": report_context.get("raw_inputs", {}).get("approval_result", {}) or read_json(output_dir / "approval_result.json"),
        "llm_agent": llm_fields,
        "generated_at": utc_now(),
    }
    agent_context.setdefault("render_quality_checks", {})
    agent_context["render_quality_checks"].setdefault("total_placeholders", count_placeholders(agent_context))
    return agent_context


def build_agent_llm_fields(agent_key: str, context: dict[str, Any], output: dict[str, Any]) -> dict[str, Any]:
    """Return grounded narrative fields for stage report templates.

    The LLM is used for analyst-readable explanation only. Factual values such as
    alert ID, severity, assets, IOCs, reputation scores, and timestamps remain
    deterministic fields from parser/threat-intel JSON. If OpenAI is unavailable,
    deterministic fallback text fills the same template variables.
    """
    key = str(agent_key or "").lower()

    def parser_fact_pack() -> dict[str, Any]:
        extracted = output.get("important_extracted_fields") or {}
        return {
            "stage": "Parsing and Normalisation",
            "ticket_id": context.get("ticket_id") or context.get("incident_id"),
            "parser_status": output.get("parser_status") or output.get("status"),
            "input_shape": output.get("input_shape") or output.get("detected_input_format"),
            "normalised_alert_count": output.get("normalised_alert_count"),
            "selected_alert": {
                "alert_id": output.get("selected_alert_id") or extracted.get("alert_id") or context.get("alert_id"),
                "incident_id": extracted.get("incident_id") or context.get("incident_id"),
                "alert_name": extracted.get("alert_name") or context.get("case_title"),
                "severity": extracted.get("severity") or get_nested(context, "severity.label", "Not Provided"),
                "risk_score": extracted.get("risk_score"),
                "alert_time": extracted.get("alert_time"),
                "detection_source": "NetWitness",
            },
            "important_extracted_fields": {
                "source_ips": extracted.get("source_ips") or [],
                "destination_ips": extracted.get("destination_ips") or [],
                "internal_ips": extracted.get("internal_ips") or [],
                "external_ips": extracted.get("external_ips") or [],
                "hosts": extracted.get("hosts") or [],
                "users": extracted.get("users") or [],
                "emails": extracted.get("emails") or [],
                "file_names": extracted.get("file_names") or [],
                "file_hashes": extracted.get("file_hashes") or [],
                "process_names": extracted.get("process_names") or [],
                "mitre_technique_ids": extracted.get("mitre_technique_ids") or [],
            },
            "data_quality": {
                "parser_confidence": output.get("parser_confidence"),
                "parser_confidence_score": output.get("parser_confidence_score"),
                "missing_important_fields": output.get("missing_important_fields") or [],
                "warnings": output.get("warnings") or [],
            },
            "downstream_target": "Triage Agent",
        }

    def threat_fact_pack() -> dict[str, Any]:
        ti = output.get("threat_intelligence") or {}
        iocs = output.get("iocs") or ti.get("iocs") or {}
        return {
            "stage": "Threat Intelligence Enrichment",
            "ticket_id": context.get("ticket_id") or context.get("incident_id"),
            "alert_id": context.get("alert_id"),
            "severity": get_nested(context, "severity.label", "Not Provided"),
            "confidence": get_nested(context, "confidence.label", "Not Provided"),
            "iocs": {
                "file_hash": iocs.get("file_hash"),
                "possible_file_name": iocs.get("possible_file_name"),
                "ip_indicators": iocs.get("ip_indicators") or [],
                "domain_indicators": iocs.get("domain_indicators") or [],
            },
            "virustotal": ti.get("virustotal") or {},
            "abuseipdb": ti.get("abuseipdb") or {},
            "alienvault_otx": ti.get("alienvault_otx") or {},
            "enrichment_risk_score": output.get("enrichment_risk_score"),
            "enrichment_risk_level": output.get("enrichment_risk_level"),
            "enrichment_risk_reasons": output.get("enrichment_risk_reasons") or [],
            "notes": output.get("notes") or ti.get("notes") or [],
            "downstream_target": "SOC Approval and Investigation Agent",
        }

    if key in {"parsing", "parsing_normalisation"}:
        payload = parser_fact_pack()
        required_keys = [
            "parser_executive_summary",
            "normalisation_quality_assessment",
            "analyst_interpretation",
            "downstream_readiness",
            "triage_handoff_note",
        ]
    elif key in {"threat_intel", "threat_intelligence"}:
        payload = threat_fact_pack()
        required_keys = [
            "threat_intel_executive_summary",
            "ioc_reputation_assessment",
            "risk_interpretation",
            "limitations",
            "recommended_follow_up",
            "investigation_handoff_note",
        ]
    else:
        payload = {
            "ticket_id": context.get("ticket_id") or context.get("incident_id"),
            "case_title": context.get("case_title"),
            "severity": get_nested(context, "severity.label", "Not Provided"),
            "confidence": get_nested(context, "confidence.label", "Not Provided"),
            "classification": context.get("classification"),
            "likely_scenario": context.get("likely_scenario"),
            "agent_output": output,
        }
        required_keys = ["summary", "evidence_interpretation", "rationale", "limitations", "analyst_actions"]

    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if api_key and not api_key.startswith("replace_"):
        try:
            model = os.getenv("OPENAI_MODEL", "gpt-5.4-mini")
            prompt = (
                "Generate grounded SOC stage report narrative variables. Use only the JSON context. "
                "Do not invent affected assets, users, IOCs, timestamps, approvals, containment actions, or external reputation results. "
                "If evidence is missing, write 'Not confirmed in the provided evidence.' "
                f"Return JSON only with exactly these keys: {', '.join(required_keys)}.\n\n"
                + json.dumps(payload, ensure_ascii=False)[:14000]
            )
            raw = invoke_openai_text(
                prompt,
                system="You fill SOC report variables with grounded analyst-readable text. Facts are locked by the provided JSON.",
                model=model,
                temperature=0.2,
                max_output_tokens=1400,
                timeout=60,
            )
            data = extract_json_object(raw)
            if not data:
                raise ValueError("OpenAI returned no parseable JSON for report variables")
            if isinstance(data, dict):
                data["provider"] = "openai"
                data["model"] = model
                data["llm_input"] = scrub_template_context(payload)
                # Keep old generic fields available for older templates.
                data.setdefault("summary", data.get("parser_executive_summary") or data.get("threat_intel_executive_summary") or data.get("analyst_interpretation") or data.get("risk_interpretation") or "Not confirmed in the provided evidence.")
                data.setdefault("evidence_interpretation", data.get("normalisation_quality_assessment") or data.get("ioc_reputation_assessment") or data.get("analyst_interpretation") or "Not confirmed in the provided evidence.")
                data.setdefault("rationale", data.get("downstream_readiness") or data.get("risk_interpretation") or "Grounded in the latest stage output.")
                data.setdefault("analyst_actions", data.get("triage_handoff_note") or data.get("recommended_follow_up") or data.get("investigation_handoff_note") or "Review the output before moving to the next stage.")
                return data
        except Exception:
            pass

    # Deterministic fallback fills the same deep template fields.
    if key in {"parsing", "parsing_normalisation"}:
        alert = payload.get("selected_alert", {})
        dq = payload.get("data_quality", {})
        extracted = payload.get("important_extracted_fields", {})
        missing = dq.get("missing_important_fields") or []
        summary = (
            f"The parser completed with status {payload.get('parser_status') or 'Not Provided'} and selected alert "
            f"{alert.get('alert_id') or 'Not confirmed in the provided evidence.'}. It extracted SOC-ready fields for downstream triage."
        )
        quality = (
            f"Parser confidence is {dq.get('parser_confidence') or 'Not Provided'} with score {dq.get('parser_confidence_score') or 'Not Provided'}. "
            + (f"Missing fields requiring analyst awareness: {', '.join(map(str, missing))}." if missing else "No major missing parser fields were recorded.")
        )
        interpretation = (
            f"Severity is {alert.get('severity') or 'Not Provided'}. Hosts: {', '.join(map(str, extracted.get('hosts') or [])) or 'Not confirmed in the provided evidence.'}. "
            f"File hashes: {', '.join(map(str, extracted.get('file_hashes') or [])) or 'Not confirmed in the provided evidence.'}."
        )
        handoff = "Use processed_alert.json as the source-of-truth input for the Triage Agent. Do not rely on PDF/Word exports for machine handoff."
        return {
            "provider": "deterministic_fallback",
            "model": "local_grounded_template_fill",
            "llm_input": scrub_template_context(payload),
            "parser_executive_summary": summary,
            "normalisation_quality_assessment": quality,
            "analyst_interpretation": interpretation,
            "downstream_readiness": "Parser JSON output is ready for the Triage Agent once processed_alert.json has been written.",
            "triage_handoff_note": handoff,
            "summary": summary,
            "evidence_interpretation": interpretation,
            "rationale": quality,
            "limitations": "; ".join(map(str, missing)) if missing else "No major export limitations were recorded.",
            "analyst_actions": handoff,
        }

    if key in {"threat_intel", "threat_intelligence"}:
        iocs = payload.get("iocs", {})
        notes = payload.get("notes") or []
        reasons = payload.get("enrichment_risk_reasons") or []
        summary = (
            f"Threat intelligence enrichment completed with {payload.get('enrichment_risk_level') or 'Not Provided'} risk "
            f"and score {payload.get('enrichment_risk_score') if payload.get('enrichment_risk_score') is not None else 'Not Provided'}."
        )
        reputation = (
            f"Primary file hash: {iocs.get('file_hash') or 'Not confirmed in the provided evidence.'}. "
            f"Public IP indicators: {', '.join(map(str, iocs.get('ip_indicators') or [])) or 'None found'}. "
            f"Domain indicators: {', '.join(map(str, iocs.get('domain_indicators') or [])) or 'None found'}."
        )
        limits = "; ".join(map(str, notes)) if notes else "No skipped lookup notes were recorded."
        follow_up = "Validate external reputation results against internal NetWitness and endpoint telemetry before Investigation Agent execution."
        return {
            "provider": "deterministic_fallback",
            "model": "local_grounded_template_fill",
            "llm_input": scrub_template_context(payload),
            "threat_intel_executive_summary": summary,
            "ioc_reputation_assessment": reputation,
            "risk_interpretation": "; ".join(map(str, reasons)) if reasons else "No confirmed malicious external intelligence was found, or no usable IOC was available.",
            "limitations": limits,
            "recommended_follow_up": follow_up,
            "investigation_handoff_note": "Use enriched_alert.json as the source-of-truth IOC reputation context for Investigation and Reporting.",
            "summary": summary,
            "evidence_interpretation": reputation,
            "rationale": "; ".join(map(str, reasons)) if reasons else "No enrichment risk reasons were recorded.",
            "analyst_actions": follow_up,
        }

    sev = payload.get("severity") or "Not Provided"
    conf = payload.get("confidence") or "Not Provided"
    scenario = payload.get("likely_scenario") or "Not confirmed in the provided evidence."
    status = output.get("status") or output.get("report_status") or "Not Provided"
    summary = output.get("summary") or output.get("analyst_summary") or output.get("classification") or "Not confirmed in the provided evidence."
    missing = to_list(output.get("missing_evidence") or output.get("missing_fields") or context.get("missing_required_fields"))
    return {
        "provider": "deterministic_fallback",
        "model": "local_grounded_template_fill",
        "llm_input": scrub_template_context(payload),
        "summary": f"{agent_key.title()} output indicates status '{status}' for ticket {payload.get('ticket_id')}. Severity is {sev} with {conf} confidence. {summary}",
        "evidence_interpretation": f"The available context points to: {scenario}. Facts not present in the supplied evidence are kept as not confirmed.",
        "rationale": output.get("reason") or output.get("rationale") or "The assessment is derived from saved ticket context and the latest agent output. No unsupported facts were added.",
        "limitations": "; ".join(str(x.get("gap") if isinstance(x, dict) else x) for x in missing) if missing else "No major export limitations were recorded in the saved output.",
        "analyst_actions": output.get("recommended_next_action") or output.get("next_action") or "Review the output, validate missing evidence, and record analyst decision before closure.",
    }

def template_environment(project_root: Path, template_folder: str) -> Environment:
    template_dir = project_root / template_folder
    env = Environment(
        loader=FileSystemLoader(str(template_dir)),
        autoescape=select_autoescape(enabled_extensions=()),
        trim_blocks=False,
        lstrip_blocks=False,
        keep_trailing_newline=True,
    )
    env.filters["join"] = lambda value, sep=", ": sep.join(str(v) for v in to_list(value))
    return env


def render_jinja_template(project_root: Path, template_name: str, context: dict[str, Any], *, template_folder: str) -> str:
    env = template_environment(project_root, template_folder)
    return env.get_template(template_name).render(**context)


# -------------------- Markdown normalisation and parsing --------------------

def clean_markdown_text(text: str) -> str:
    text = str(text or "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Remove raw code-fence language prefixes that sometimes appear inline in appendices.
    text = re.sub(r"```(?:json|text|markdown)?", "```", text, flags=re.IGNORECASE)
    text = cleanup_report_text(text)
    return text.strip() + "\n"


def cleanup_report_text(text: str) -> str:
    """Apply deterministic report polish without changing incident facts."""
    text = str(text or "")
    replacements = {
        "attempt. scenario": "attempt scenario",
        "if approved..": "if approved.",
    }
    for old, new in replacements.items():
        text = re.sub(re.escape(old), new, text, flags=re.IGNORECASE)
    text = re.sub(r"\.{2,}", ".", text)
    text = re.sub(r"\s+([,.;:])", r"\1", text)
    text = re.sub(r"([.!?])\s+\1+", r"\1", text)
    return text


def split_inline_numbered_list(line: str) -> list[str] | None:
    """Turn '1. A. 2. B. 3. C.' into separate numbered lines.

    Avoids section numbers such as 9.1 because the pattern requires an integer
    followed by a full stop and a space, usually followed by an uppercase letter.
    """
    markers = list(re.finditer(r"(?:(?<=^)|(?<=\s))(\d{1,2})\.\s+(?=[A-Z])", line))
    if len(markers) < 2:
        return None
    parts: list[str] = []
    for idx, marker in enumerate(markers):
        start = marker.start()
        end = markers[idx + 1].start() if idx + 1 < len(markers) else len(line)
        part = line[start:end].strip()
        if part:
            parts.append(part)
    return parts or None


def split_gap_sentence(line: str) -> list[str] | None:
    """Split compact evidence-gap strings into bullets for readability."""
    if not re.search(r"\b(Critical|High|Medium|Low|Informational):\s+", line):
        return None
    if sum(1 for _ in re.finditer(r"\b(Critical|High|Medium|Low|Informational):\s+", line)) < 2:
        return None
    intro = re.split(r"\b(?:Critical|High|Medium|Low|Informational):\s+", line, maxsplit=1)[0].strip()
    segments = re.findall(r"\b(Critical|High|Medium|Low|Informational):\s+([^\.]+)(?:\.|$)", line)
    if not segments:
        return None
    out = []
    if intro:
        out.append(intro.rstrip(":") + ":")
    for priority, item in segments:
        out.append(f"- **{priority}:** {item.strip()}")
    return out


def normalise_markdown_for_report(markdown: str) -> str:
    lines = clean_markdown_text(markdown).split("\n")
    out: list[str] = []
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            out.append("")
            continue
        gap_split = split_gap_sentence(line)
        if gap_split:
            out.extend(gap_split)
            continue
        list_split = split_inline_numbered_list(line)
        if list_split:
            out.extend(list_split)
            continue
        out.append(line)
    return "\n".join(out).strip() + "\n"


def parse_markdown_table(lines: list[str], start: int) -> tuple[dict[str, Any] | None, int]:
    return parse_pipe_table(lines, start)


def _block_contains_raw_markdown_table(block: dict[str, Any]) -> bool:
    text = str(block.get("text") or "")
    if block.get("type") == "paragraph":
        if re.search(r"\|\s*:?-{2,}:?\s*(?:\|\s*:?-{2,}:?\s*)+\|?", text):
            return True
        if re.search(r"^\s*\|.+\|\s*$", text, flags=re.MULTILINE):
            return True
    return False


def validate_no_raw_markdown_tables(blocks: list[dict[str, Any]], section_title: str = "Report") -> None:
    repaired = repair_pipe_tables_in_blocks(blocks)
    blocks[:] = repaired
    current_section = section_title
    for index, block in enumerate(repaired):
        if isinstance(block, dict) and block.get("type") == "heading":
            current_section = str(block.get("text") or current_section)
            continue
        if not isinstance(block, dict) or block.get("type") != "paragraph":
            continue
        text = str(block.get("text") or "")
        if paragraph_contains_raw_pipe_table(text):
            preview = re.sub(r"\s+", " ", text).strip()
            preview = re.sub(
                r"(?i)\b(api[_ -]?key|access[_ -]?token|token|secret|authorization)\b\s*[:=]\s*\S+",
                r"\1=[REDACTED]",
                preview,
            )
            preview = preview[:180] + ("..." if len(preview) > 180 else "")
            raise ValueError(
                f'Export validation failed: raw table syntax remained in section "{current_section}", '
                f"block index {index}: {preview}"
            )


def markdown_to_report_blocks(markdown: str) -> list[dict[str, Any]]:
    text = normalise_markdown_for_report(markdown)
    lines = text.split("\n")
    blocks: list[dict[str, Any]] = []
    paragraph_buffer: list[str] = []
    in_code = False
    code_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text_value = " ".join(part.strip() for part in paragraph_buffer if part.strip()).strip()
            if text_value:
                blocks.append({"type": "paragraph", "text": text_value})
            paragraph_buffer = []

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                blocks.append({"type": "code", "text": "\n".join(code_buffer).strip()})
                code_buffer = []
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue
        table_block, next_i = parse_markdown_table(lines, i)
        if table_block:
            flush_paragraph()
            blocks.append(table_block)
            i = next_i
            continue
        if not stripped:
            flush_paragraph()
            i += 1
            continue
        if stripped in {"---", "***", "___"}:
            flush_paragraph()
            i += 1
            continue
        if stripped.startswith("#"):
            flush_paragraph()
            match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
            if match:
                blocks.append({"type": "heading", "level": len(match.group(1)), "text": match.group(2).strip()})
                i += 1
                continue
        if stripped.startswith(">"):
            flush_paragraph()
            blocks.append({"type": "callout", "kind": "info", "text": stripped.lstrip("> ").strip()})
            i += 1
            continue
        bullet = re.match(r"^[-*•]\s+(.+)$", stripped)
        if bullet:
            flush_paragraph()
            items = []
            while i < len(lines):
                m = re.match(r"^[-*•]\s+(.+)$", lines[i].strip())
                if not m:
                    break
                items.append(m.group(1).strip())
                i += 1
            blocks.append({"type": "bullet_list", "items": items})
            continue
        numbered = re.match(r"^(\d{1,2})\.\s+(.+)$", stripped)
        if numbered:
            flush_paragraph()
            items = []
            while i < len(lines):
                m = re.match(r"^(\d{1,2})\.\s+(.+)$", lines[i].strip())
                if not m:
                    break
                items.append(m.group(2).strip())
                i += 1
            blocks.append({"type": "numbered_list", "items": items})
            continue
        paragraph_buffer.append(stripped)
        i += 1
    flush_paragraph()
    if code_buffer:
        blocks.append({"type": "code", "text": "\n".join(code_buffer).strip()})
    return repair_pipe_tables_in_blocks(blocks)


# -------------------- DOCX styling helpers --------------------

def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "6")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), kwargs[edge])


def set_paragraph_bottom_border(paragraph, color: str = BORDER_BLUE, size: str = "12") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_pct: int = 5000) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_pct))
    tbl_w.set(qn("w:type"), "pct")


def set_cell_text(cell, text: Any, bold: bool = False, fill: str | None = None, size: int = 9, color: RGBColor | None = None) -> None:
    cell.text = ""
    if fill:
        set_cell_shading(cell, fill)
    set_cell_margins(cell)
    set_cell_border(cell, top=BORDER_GREY, bottom=BORDER_GREY, left=BORDER_GREY, right=BORDER_GREY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    add_inline_markdown(p, str(text if text is not None else ""), base_size=size)
    for run in p.runs:
        run.bold = bool(bold or run.bold)
        if color:
            run.font.color.rgb = color
        run.font.name = "Aptos"


def add_inline_markdown(paragraph, text: str, *, base_size: int = 10) -> None:
    text = str(text or "")
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            run.font.size = Pt(base_size)
            run.font.name = "Aptos"
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.size = Pt(base_size)
            run.font.name = "Aptos"
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(max(8, base_size - 1))
            run.font.color.rgb = RGBColor(78, 90, 105)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(base_size)
        run.font.name = "Aptos"


def add_brand_header(doc, title: str, subtitle: str, logo_path: Path | None = None) -> None:
    if logo_path and logo_path.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_after = Pt(8)
        logo_run = logo_p.add_run()
        logo_run.add_picture(str(logo_path), width=Inches(2.65))
    else:
        label = doc.add_paragraph()
        label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label.paragraph_format.space_after = Pt(2)
        r = label.add_run("AEGIS")
        r.font.size = Pt(9)
        r.bold = True
        r.font.color.rgb = BRAND_BLUE
        r.font.name = "Aptos"

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    title_p.paragraph_format.line_spacing = 0.95
    title_run = title_p.add_run(title)
    title_run.font.size = Pt(23)
    title_run.bold = True
    title_run.font.color.rgb = BRAND_BLUE_DARK
    title_run.font.name = "Aptos Display"

    if subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.space_after = Pt(8)
        sr = sub.add_run(subtitle)
        sr.font.size = Pt(9)
        sr.font.color.rgb = BRAND_GREY
        sr.font.name = "Aptos"

    divider = doc.add_paragraph()
    divider.paragraph_format.space_after = Pt(16)
    set_paragraph_bottom_border(divider, BORDER_BLUE, "16")


def _strip_duplicate_leading_heading(blocks: list[dict[str, Any]], title: str) -> list[dict[str, Any]]:
    """Avoid repeating the document title after the branded title block."""
    if not blocks:
        return []
    normalised_title = re.sub(r"\s+", " ", str(title or "").strip().lower())
    cleaned = list(blocks)
    while cleaned:
        first = cleaned[0]
        if not isinstance(first, dict) or first.get("type") != "heading":
            break
        text = re.sub(r"\s+", " ", str(first.get("text") or "").strip().lower())
        if text in {normalised_title, "cybersecurity incident post-incident review report"}:
            cleaned.pop(0)
            continue
        break
    return cleaned


def add_heading_paragraph(doc, text: str, level: int) -> None:
    sizes = {1: 17, 2: 14, 3: 12, 4: 11, 5: 10, 6: 10}
    colors = {1: BRAND_BLUE, 2: BRAND_BLUE, 3: BRAND_BLUE_DARK, 4: BRAND_BLUE_DARK, 5: BRAND_BLUE_DARK, 6: BRAND_BLUE_DARK}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    add_inline_markdown(p, text, base_size=sizes.get(level, 10))
    for run in p.runs:
        run.bold = True
        run.font.color.rgb = colors.get(level, BRAND_BLUE_DARK)
        run.font.name = "Aptos Display" if level <= 2 else "Aptos"
    if level == 1:
        set_paragraph_bottom_border(p, "D9E5F4", "6")


def infer_callout_kind(text: str) -> str:
    lower = text.lower()
    if any(term in lower for term in ["critical", "high:", "missing required", "review required", "not confirmed", "not provided"]):
        return "warning"
    if any(term in lower for term in ["completed", "approved", "pass"]):
        return "success"
    return "info"


def add_callout(doc, text: str, kind: str = "info") -> None:
    fill = {"info": LIGHT_BLUE, "warning": LIGHT_YELLOW, "success": LIGHT_GREEN, "danger": LIGHT_RED}.get(kind, LIGHT_BLUE)
    accent = {"info": BORDER_BLUE, "warning": "D6A300", "success": "179C52", "danger": "C93434"}.get(kind, BORDER_BLUE)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_cell_border(cell, left=accent, top="D7E1EF", bottom="D7E1EF", right="D7E1EF")
    set_cell_margins(cell, top=110, start=150, bottom=110, end=120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline_markdown(p, text, base_size=9)
    for run in p.runs:
        run.font.name = "Aptos"
    doc.add_paragraph("")


def add_code_block(doc, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F6F8FA")
    set_cell_border(cell, top="D0D7DE", bottom="D0D7DE", left="D0D7DE", right="D0D7DE")
    set_cell_margins(cell, top=90, start=110, bottom=90, end=110)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text[:12000])
    run.font.name = "Consolas"
    run.font.size = Pt(7.8)
    run.font.color.rgb = RGBColor(50, 55, 65)
    doc.add_paragraph("")


def add_list(doc, items: list[str], *, numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.05)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.08
        add_inline_markdown(p, item, base_size=9.8)
        for run in p.runs:
            run.font.name = "Aptos"
            run.font.color.rgb = BRAND_BLUE_DARK


def add_metadata_table(doc, meta: list[tuple[str, Any]]) -> None:
    if not meta:
        return
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table)
    for label, value in meta:
        row = table.add_row().cells
        set_cell_text(row[0], label, bold=True, fill=LIGHT_GREY, size=8.8, color=BRAND_BLUE_DARK)
        set_cell_text(row[1], value, bold=False, fill=WHITE, size=8.8, color=BRAND_BLUE_DARK)
    doc.add_paragraph("")


def is_summary_paragraph(text: str) -> bool:
    return bool(re.match(r"^(Summary|Executive Summary|Decision|Conclusion|Final Assessment|Analyst Guidance|Review Focus):", text, flags=re.I))


def add_report_table(doc, columns: list[str], rows: list[list[str]]) -> None:
    if not columns:
        return
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    set_table_width(table)
    for idx, col in enumerate(columns):
        set_cell_text(table.rows[0].cells[idx], col, bold=True, fill=LIGHT_BLUE, size=8.6, color=BRAND_BLUE_DARK)
    for raw_row in rows:
        values = list(raw_row or []) + [""] * (len(columns) - len(raw_row or []))
        cells = table.add_row().cells
        for idx, value in enumerate(values[:len(columns)]):
            fill = WHITE
            val_text = str(value or "")
            lower = val_text.lower()
            if idx > 0 and any(word in lower for word in ["critical", "fail", "required", "not confirmed", "not provided", "review"]):
                fill = LIGHT_YELLOW
            elif idx > 0 and any(word in lower for word in ["pass", "completed", "approved"]):
                fill = LIGHT_GREEN
            set_cell_text(cells[idx], val_text, bold=(idx == 0 and len(columns) <= 3), fill=fill, size=8.3, color=BRAND_BLUE_DARK)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph("")


def create_docx_from_blocks(path: Path, *, title: str, subtitle: str, blocks: list[dict[str, Any]], meta: list[tuple[str, Any]], logo_path: Path | None = None) -> Path:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Run: python3 -m pip install python-docx")
    blocks = repair_pipe_tables_in_blocks(blocks)
    validate_no_raw_markdown_tables(blocks, title)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.8)

    add_brand_header(doc, title, subtitle, logo_path=logo_path)
    add_metadata_table(doc, meta)

    for block in _strip_duplicate_leading_heading(blocks, title):
        if not isinstance(block, dict):
            continue
        btype = block.get("type")
        if btype == "heading":
            add_heading_paragraph(doc, str(block.get("text") or ""), min(max(int(block.get("level") or 1), 1), 6))
        elif btype == "paragraph":
            text = str(block.get("text") or "").strip()
            if not text:
                continue
            if is_summary_paragraph(text):
                add_callout(doc, text, "info")
            elif len(text) < 230 and any(term in text.lower() for term in ["not confirmed", "review required", "missing required"]):
                add_callout(doc, text, infer_callout_kind(text))
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.12
                add_inline_markdown(p, text, base_size=9.8)
                for run in p.runs:
                    run.font.name = "Aptos"
                    run.font.color.rgb = BRAND_BLUE_DARK
        elif btype == "callout":
            add_callout(doc, str(block.get("text") or ""), block.get("kind") or "info")
        elif btype == "bullet_list":
            add_list(doc, [str(i or "") for i in (block.get("items") or [])], numbered=False)
        elif btype == "numbered_list":
            add_list(doc, [str(i or "") for i in (block.get("items") or [])], numbered=True)
        elif btype == "table":
            add_report_table(doc, [str(c or "") for c in (block.get("columns") or [])], block.get("rows") or [])
        elif btype == "code":
            add_code_block(doc, str(block.get("text") or ""))

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def libreoffice_binary() -> str | None:
    # Common macOS path is included because Flask may not inherit the user's shell PATH.
    candidates = [
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> Path:
    office = libreoffice_binary()
    if not office:
        raise RuntimeError("LibreOffice is required for PDF conversion from DOCX. Install LibreOffice or make soffice available in PATH.")
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="lo_profile_") as profile:
        cmd = [
            office,
            "--headless",
            f"-env:UserInstallation=file://{profile}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ]
        completed = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=int(os.getenv("PDF_CONVERSION_TIMEOUT_SECONDS", "45")))
        produced = pdf_path.parent / f"{docx_path.stem}.pdf"
        if completed.returncode != 0 or not produced.exists():
            raise RuntimeError(f"DOCX to PDF conversion failed: {completed.stderr or completed.stdout}")
        if produced != pdf_path:
            if pdf_path.exists():
                pdf_path.unlink()
            produced.rename(pdf_path)
    return pdf_path


def validate_export_text(text: str) -> None:
    for pattern in FORBIDDEN_TEXT_PATTERNS:
        if pattern in text:
            raise ValueError(f"Export validation failed: forbidden text found: {pattern}")
    if re.search(r"{{\s*[^}]+\s*}}", text):
        raise ValueError("Export validation failed: unresolved Jinja2 placeholder found.")


def report_export_dir(output_dir: Path, ticket_id: str, report_key: str) -> Path:
    return output_dir / "exports" / safe_filename(ticket_id) / "reporting" / safe_filename(report_key)


def agent_export_dir(output_dir: Path, ticket_id: str, agent_key: str) -> Path:
    return output_dir / "exports" / safe_filename(ticket_id) / "agents" / safe_filename(agent_key)



def _existing_source_files(project_root: Path, output_dir: Path, template_folder: str, template_name: str, json_file: str | None = None) -> list[Path]:
    candidates = [project_root / template_folder / template_name]
    if template_folder == "report_templates":
        logo = aegis_logo_path(project_root)
        if logo:
            candidates.append(logo)
    common_names = [
        "processed_alert.json",
        "parser_result.json",
        "enriched_alert.json",
        "threat_intel_result.json",
        "triage_result.json",
        "investigation_result.json",
        "approval_result.json",
        "investigation_approval_result.json",
        "soc_review_result.json",
        "final_report.json",
    ]
    if json_file and json_file not in common_names:
        common_names.insert(0, json_file)
    for name in common_names:
        candidates.append(output_dir / name)
        candidates.append(project_root / "inputs" / name)
    # Keep deterministic order and avoid duplicates.
    deduped: list[Path] = []
    seen = set()
    for path in candidates:
        key = str(path)
        if key not in seen:
            deduped.append(path)
            seen.add(key)
    return deduped


def _load_cached_json_blocks(json_path: Path, source_hash: str) -> tuple[dict[str, Any], list[dict[str, Any]]] | None:
    if not json_path.exists() or json_path.stat().st_size == 0:
        return None
    try:
        payload = json.loads(json_path.read_text(encoding="utf-8"))
    except Exception:
        return None
    if not isinstance(payload, dict):
        return None
    if payload.get("source_hash") != source_hash:
        return None
    blocks = payload.get("structured_blocks")
    if not isinstance(blocks, list):
        return None
    repaired = repair_pipe_tables_in_blocks(blocks)
    validate_no_raw_markdown_tables(repaired)
    if repaired != blocks:
        payload["structured_blocks"] = repaired
        json_path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    return payload, repaired


def _write_cached_export_json(json_path: Path, export_json: dict[str, Any], out_dir: Path, source_hash: str) -> None:
    export_json["source_hash"] = source_hash
    write_json(json_path, export_json)
    mark_export_status(out_dir, key="json", status="ready", source_hash=source_hash, file_path=json_path, message="JSON export cache is ready.")


def generate_reporting_export(project_root: Path, output_dir: Path, ticket: dict[str, Any] | None, report_key: str, file_type: str) -> Path:
    if report_key not in REPORT_TEMPLATES:
        raise KeyError(f"Unknown reporting template: {report_key}")
    cfg = REPORT_TEMPLATES[report_key]
    ticket = ticket or {}
    pre_ticket_id = str(first_present(ticket.get("ticket_id"), default="unknown"))
    source_hash = calculate_source_hash(
        source_files=_existing_source_files(project_root, output_dir, "report_templates", cfg["template"], "final_report.json"),
        extra_payload={"ticket": ticket, "report_key": report_key, "template": cfg},
    )

    out_dir = report_export_dir(output_dir, pre_ticket_id, report_key)
    base = f"{safe_filename(pre_ticket_id)}_{cfg['filename']}"
    json_path = out_dir / f"{base}.json"
    docx_path = out_dir / f"{base}.docx"
    pdf_path = out_dir / f"{base}.pdf"

    if file_type == "json" and is_cache_ready(out_dir, "json", json_path, source_hash):
        return json_path
    if file_type in {"docx", "word"} and is_cache_ready(out_dir, "docx", docx_path, source_hash):
        return docx_path
    if file_type == "pdf" and is_cache_ready(out_dir, "pdf", pdf_path, source_hash):
        return pdf_path
    if file_type == "pdf" and is_cache_ready(out_dir, "docx", docx_path, source_hash):
        mark_export_status(out_dir, key="pdf", status="preparing", source_hash=source_hash, file_path=pdf_path, message="Converting cached Word document to PDF.")
        try:
            converted = convert_docx_to_pdf(docx_path, pdf_path)
            mark_export_status(out_dir, key="pdf", status="ready", source_hash=source_hash, file_path=converted, message="PDF export cache is ready.", related_paths={"docx": str(docx_path)})
            return converted
        except Exception as exc:
            mark_export_status(out_dir, key="pdf", status="failed", source_hash=source_hash, file_path=pdf_path, message=str(exc))
            raise

    cached_json = _load_cached_json_blocks(json_path, source_hash)
    if cached_json:
        export_json, blocks = cached_json
        context = export_json.get("context") if isinstance(export_json.get("context"), dict) else {}
    else:
        mark_export_status(out_dir, key="json", status="preparing", source_hash=source_hash, file_path=json_path, message="Building template context and rendering report markdown.")
        context = build_report_context(project_root, output_dir, ticket)
        ticket_id_from_context = str(first_present(ticket.get("ticket_id"), context.get("incident_id"), default=pre_ticket_id))
        if ticket_id_from_context != pre_ticket_id:
            out_dir = report_export_dir(output_dir, ticket_id_from_context, report_key)
            base = f"{safe_filename(ticket_id_from_context)}_{cfg['filename']}"
            json_path = out_dir / f"{base}.json"
            docx_path = out_dir / f"{base}.docx"
            pdf_path = out_dir / f"{base}.pdf"
        rendered = render_jinja_template(project_root, cfg["template"], context, template_folder="report_templates")
        validate_export_text(rendered)
        context.setdefault("render_quality_checks", {})
        context["render_quality_checks"]["rendered_placeholder_count"] = count_placeholders(rendered)
        context["render_quality_checks"]["rendered_placeholder_mentions"] = {
            token: len(re.findall(re.escape(token), rendered, flags=re.IGNORECASE))
            for token in ["Not Provided", "To Be Validated", "To Be Assigned", "Pending", "Not linked", "Evidence link unavailable", "Unavailable from source telemetry"]
        }
        blocks = markdown_to_report_blocks(rendered)
        validate_no_raw_markdown_tables(blocks)
        export_json = {
            "success": True,
            "ticket_id": ticket_id_from_context,
            "report_key": report_key,
            "report_title": context.get("report_title") if report_key == "final_incident_report" else cfg["title"],
            "template_used": f"report_templates/{cfg['template']}",
            "generated_at": utc_now(),
            "cache_status": "ready",
            "source_hash": source_hash,
            "context": context,
            "rendered_markdown": rendered,
            "structured_blocks": blocks,
        }
        _write_cached_export_json(json_path, export_json, out_dir, source_hash)

    if file_type == "json":
        return json_path

    mark_export_status(out_dir, key="docx", status="preparing", source_hash=source_hash, file_path=docx_path, message="Generating Word document.")
    try:
        document_title = str(export_json.get("report_title") or cfg["title"])
        create_docx_from_blocks(
            docx_path,
            title=document_title,
            subtitle="Template-based SOC report generated from the Reporting Agent and converted to PDF from Word.",
            blocks=blocks,
            logo_path=aegis_logo_path(project_root),
            meta=[
                ("Ticket ID", export_json.get("ticket_id") or pre_ticket_id),
                ("Incident ID", context.get("incident_id", "Not Provided")),
                ("Alert ID", context.get("alert_id", "Not Provided")),
                ("Severity", get_nested(context, "severity.label", "Not Provided")),
                ("Confidence", get_nested(context, "confidence.label", "Not Provided")),
                ("Generated At", utc_now()),
            ],
        )
        mark_export_status(out_dir, key="docx", status="ready", source_hash=source_hash, file_path=docx_path, message="Word export cache is ready.", related_paths={"json": str(json_path)})
    except Exception as exc:
        mark_export_status(out_dir, key="docx", status="failed", source_hash=source_hash, file_path=docx_path, message=str(exc))
        raise

    if file_type in {"docx", "word"}:
        return docx_path
    if file_type == "pdf":
        mark_export_status(out_dir, key="pdf", status="preparing", source_hash=source_hash, file_path=pdf_path, message="Converting Word document to PDF.")
        try:
            converted = convert_docx_to_pdf(docx_path, pdf_path)
            mark_export_status(out_dir, key="pdf", status="ready", source_hash=source_hash, file_path=converted, message="PDF export cache is ready.", related_paths={"json": str(json_path), "docx": str(docx_path)})
            return converted
        except Exception as exc:
            mark_export_status(out_dir, key="pdf", status="failed", source_hash=source_hash, file_path=pdf_path, message=str(exc))
            raise
    raise KeyError(f"Unsupported export type: {file_type}")


def generate_agent_export(project_root: Path, output_dir: Path, ticket: dict[str, Any] | None, agent_key: str, file_type: str) -> Path:
    if agent_key not in AGENT_TEMPLATES:
        raise KeyError(f"Unknown agent export: {agent_key}")
    cfg = AGENT_TEMPLATES[agent_key]
    ticket = ticket or {}
    pre_ticket_id = str(first_present(ticket.get("ticket_id"), default="unknown"))
    source_hash = calculate_source_hash(
        source_files=_existing_source_files(project_root, output_dir, "agent_summary_templates", cfg["template"], cfg.get("json_file")),
        extra_payload={"ticket": ticket, "agent_key": agent_key, "template": cfg},
    )

    out_dir = agent_export_dir(output_dir, pre_ticket_id, agent_key)
    base = f"{safe_filename(pre_ticket_id)}_{cfg['filename']}"
    json_path = out_dir / f"{base}.json"
    docx_path = out_dir / f"{base}.docx"
    pdf_path = out_dir / f"{base}.pdf"

    if file_type == "json" and is_cache_ready(out_dir, "json", json_path, source_hash):
        return json_path
    if file_type in {"docx", "word"} and is_cache_ready(out_dir, "docx", docx_path, source_hash):
        return docx_path
    if file_type == "pdf" and is_cache_ready(out_dir, "pdf", pdf_path, source_hash):
        return pdf_path
    if file_type == "pdf" and is_cache_ready(out_dir, "docx", docx_path, source_hash):
        mark_export_status(out_dir, key="pdf", status="preparing", source_hash=source_hash, file_path=pdf_path, message="Converting cached Word document to PDF.")
        try:
            converted = convert_docx_to_pdf(docx_path, pdf_path)
            mark_export_status(out_dir, key="pdf", status="ready", source_hash=source_hash, file_path=converted, message="PDF export cache is ready.", related_paths={"docx": str(docx_path)})
            return converted
        except Exception as exc:
            mark_export_status(out_dir, key="pdf", status="failed", source_hash=source_hash, file_path=pdf_path, message=str(exc))
            raise

    cached_json = _load_cached_json_blocks(json_path, source_hash)
    if cached_json:
        export_json, blocks = cached_json
        context = export_json.get("context") if isinstance(export_json.get("context"), dict) else {}
    else:
        mark_export_status(out_dir, key="json", status="preparing", source_hash=source_hash, file_path=json_path, message="Building template context and rendering agent report markdown.")
        context = build_agent_context(project_root, output_dir, ticket, agent_key)
        ticket_id_from_context = str(first_present(ticket.get("ticket_id"), context.get("ticket_id"), context.get("incident_id"), default=pre_ticket_id))
        if ticket_id_from_context != pre_ticket_id:
            out_dir = agent_export_dir(output_dir, ticket_id_from_context, agent_key)
            base = f"{safe_filename(ticket_id_from_context)}_{cfg['filename']}"
            json_path = out_dir / f"{base}.json"
            docx_path = out_dir / f"{base}.docx"
            pdf_path = out_dir / f"{base}.pdf"
        rendered = render_jinja_template(project_root, cfg["template"], context, template_folder="agent_summary_templates")
        validate_export_text(rendered)
        blocks = markdown_to_report_blocks(rendered)
        validate_no_raw_markdown_tables(blocks)
        export_json = {
            "success": True,
            "ticket_id": ticket_id_from_context,
            "agent_key": agent_key,
            "title": cfg["title"],
            "template_used": f"agent_summary_templates/{cfg['template']}",
            "generated_at": utc_now(),
            "cache_status": "ready",
            "source_hash": source_hash,
            "context": context,
            "agent_output": context.get("agent_output"),
            "llm_agent": context.get("llm_agent"),
            "rendered_markdown": rendered,
            "structured_blocks": blocks,
        }
        _write_cached_export_json(json_path, export_json, out_dir, source_hash)

    if file_type == "json":
        return json_path

    mark_export_status(out_dir, key="docx", status="preparing", source_hash=source_hash, file_path=docx_path, message="Generating Word document.")
    try:
        create_docx_from_blocks(
            docx_path,
            title=cfg["title"],
            subtitle="Readable SOC analyst summary generated from the selected agent output.",
            blocks=blocks,
            meta=[
                ("Ticket ID", export_json.get("ticket_id") or pre_ticket_id),
                ("Incident ID", context.get("incident_id", "Not Provided")),
                ("Alert ID", context.get("alert_id", "Not Provided")),
                ("Severity", get_nested(context, "severity.label", "Not Provided")),
                ("Confidence", get_nested(context, "confidence.label", "Not Provided")),
                ("Generated At", utc_now()),
            ],
        )
        mark_export_status(out_dir, key="docx", status="ready", source_hash=source_hash, file_path=docx_path, message="Word export cache is ready.", related_paths={"json": str(json_path)})
    except Exception as exc:
        mark_export_status(out_dir, key="docx", status="failed", source_hash=source_hash, file_path=docx_path, message=str(exc))
        raise

    if file_type in {"docx", "word"}:
        return docx_path
    if file_type == "pdf":
        mark_export_status(out_dir, key="pdf", status="preparing", source_hash=source_hash, file_path=pdf_path, message="Converting Word document to PDF.")
        try:
            converted = convert_docx_to_pdf(docx_path, pdf_path)
            mark_export_status(out_dir, key="pdf", status="ready", source_hash=source_hash, file_path=converted, message="PDF export cache is ready.", related_paths={"json": str(json_path), "docx": str(docx_path)})
            return converted
        except Exception as exc:
            mark_export_status(out_dir, key="pdf", status="failed", source_hash=source_hash, file_path=pdf_path, message=str(exc))
            raise
    raise KeyError(f"Unsupported export type: {file_type}")
