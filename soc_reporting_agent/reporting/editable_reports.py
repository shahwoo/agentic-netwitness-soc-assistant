
from __future__ import annotations

import json
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from reporting.structured_report import (
    blocks_from_text,
    blocks_to_plain_text,
    load_blocks,
    markdown_to_blocks,
    paragraph_contains_raw_pipe_table,
    repair_pipe_tables_in_blocks,
    save_blocks,
)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:  # pragma: no cover
    Document = None
    Inches = None
    Pt = None
    RGBColor = None
    WD_ALIGN_PARAGRAPH = None
    WD_TABLE_ALIGNMENT = None
    WD_CELL_VERTICAL_ALIGNMENT = None
    OxmlElement = None
    qn = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
except Exception:  # pragma: no cover
    A4 = None
    getSampleStyleSheet = None
    inch = None
    Paragraph = None
    SimpleDocTemplate = None
    Image = None
    Spacer = None
    Table = None
    TableStyle = None
    colors = None

# Five templates are still available to the renderer. The dashboard exposes the
# four SOC-facing report artefacts requested by the analyst. The triage review
# template remains a supporting template that can be merged or referenced by
# the analyst review, but it is not shown as a fifth primary report card.
REPORT_SECTION_CONFIG: dict[str, dict[str, str]] = {
    "executive_summary": {
        "title": "Executive Summary",
        "template": "executive_summary_template.md.j2",
        "filename": "executive_summary.txt",
        "description": "Management-level incident summary for handover and leadership review.",
    },
    "technical_findings": {
        "title": "Technical Findings",
        "template": "technical_findings_template.md.j2",
        "filename": "technical_findings.txt",
        "description": "Evidence, IOCs, technical observations, and validation points.",
    },
    "soc_analyst_review": {
        "title": "SOC Analyst Review",
        "template": "soc_analyst_review_template.md.j2",
        "filename": "soc_analyst_review.txt",
        "description": "Analyst judgement, limitations, approval notes, and review checklist.",
    },
    "soc_triage_review": {
        "title": "SOC Triage Review",
        "template": "soc_triage_review_template.md.j2",
        "filename": "soc_triage_review.txt",
        "description": "Supporting triage decision notes used by the analyst review.",
    },
    "final_incident_report": {
        "title": "Final Incident Report",
        "template": "incident_report_template.md.j2",
        "filename": "final_incident_report.txt",
        "description": "Complete incident report for analyst confirmation and export.",
    },
}

CORE_REPORT_KEYS = ["executive_summary", "technical_findings", "soc_analyst_review", "final_incident_report"]
SUPPORT_REPORT_KEYS = ["soc_triage_review"]
EXPORT_SECTION_ORDER = CORE_REPORT_KEYS


PROJECT_ROOT = Path(__file__).resolve().parents[1]


def _aegis_logo_path() -> Path | None:
    for candidate in [
        PROJECT_ROOT / "report_assets" / "aegis-logo.png",
        PROJECT_ROOT / "dashboard" / "assets" / "aegis-logo.png",
    ]:
        if candidate.exists() and candidate.stat().st_size > 0:
            return candidate
    return None


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def markdown_to_plain_text(value: Any) -> str:
    text = str(value or "")
    if not text:
        return ""
    text = re.sub(r"```[a-zA-Z0-9_-]*\n?", "", text)
    text = text.replace("```", "")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"!\[([^\]]*)\]\([^\)]*\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]*\)", r"\1", text)
    text = re.sub(r"^\s{0,3}#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s{0,3}>\s?", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*+]\s+", "- ", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"(^|[^*])\*([^*\n]+)\*", r"\1\2", text)
    text = re.sub(r"(^|[^_])_([^_\n]+)_", r"\1\2", text)
    text = re.sub(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", "", text, flags=re.MULTILINE)

    def _table_row(match: re.Match) -> str:
        row = match.group(1)
        cells = [c.strip() for c in row.split("|") if c.strip()]
        if len(cells) >= 2:
            return f"{cells[0]}: " + " | ".join(cells[1:])
        return " | ".join(cells)

    text = re.sub(r"^\s*\|(.+)\|\s*$", _table_row, text, flags=re.MULTILINE)
    text = text.replace("---", "")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def incident_report_dir(output_dir: Path, incident_id: str) -> Path:
    return output_dir / incident_id / "reports"


def editable_dir(output_dir: Path, incident_id: str) -> Path:
    # Backwards-compatible renderer target. Files are copied into drafts by the manifest builder.
    return incident_report_dir(output_dir, incident_id) / "editable"


def drafts_dir(output_dir: Path, incident_id: str) -> Path:
    return incident_report_dir(output_dir, incident_id) / "drafts"


def confirmed_dir(output_dir: Path, incident_id: str) -> Path:
    return incident_report_dir(output_dir, incident_id) / "confirmed"


def exports_dir(output_dir: Path, incident_id: str) -> Path:
    return incident_report_dir(output_dir, incident_id) / "exports"


def draft_history_dir(output_dir: Path, incident_id: str, section_key: str) -> Path:
    return incident_report_dir(output_dir, incident_id) / "draft_history" / section_key


def final_dir(output_dir: Path, incident_id: str) -> Path:
    # Backwards-compatible old name.
    return exports_dir(output_dir, incident_id)


def manifest_path(output_dir: Path, incident_id: str) -> Path:
    return incident_report_dir(output_dir, incident_id) / "report_manifest.json"


def _rel(output_dir: Path, path: Path) -> str:
    try:
        return str(path.relative_to(output_dir.parent))
    except Exception:
        return str(path)


def _read(path: str | Path | None) -> str:
    if not path:
        return ""
    p = Path(path)
    return p.read_text(encoding="utf-8", errors="ignore") if p.exists() else ""


def _write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(str(text or ""), encoding="utf-8")


def _safe_resolve(path: Path) -> Path:
    try:
        return path.resolve()
    except Exception:
        return path.absolute()


def _draft_history_entry(output_dir: Path, incident_id: str, section_key: str, title: str, text: str, analyst: str) -> dict[str, Any]:
    hdir = draft_history_dir(output_dir, incident_id, section_key)
    hdir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%S%fZ")
    path = hdir / f"{stamp}.txt"
    _write(path, text)
    preview = markdown_to_plain_text(text)[:500]
    return {
        "created_at": utc_now(),
        "saved_by": analyst,
        "title": title,
        "path": str(path),
        "relative_path": _rel(output_dir, path),
        "preview": preview,
    }


def build_report_manifest(output_dir: Path, incident_id: str, generated_sections: dict[str, str], context: dict[str, Any] | None = None) -> dict[str, Any]:
    context = context or {}
    ddir = drafts_dir(output_dir, incident_id)
    cdir = confirmed_dir(output_dir, incident_id)
    edir = exports_dir(output_dir, incident_id)
    ddir.mkdir(parents=True, exist_ok=True)
    cdir.mkdir(parents=True, exist_ok=True)
    edir.mkdir(parents=True, exist_ok=True)

    # Use the supporting triage template by appending a short supporting section
    # into SOC Analyst Review when the support file exists. This keeps the user-facing
    # output at four reports while still using the triage template in the renderer.
    support_text = _read(generated_sections.get("soc_triage_review"))

    sections: dict[str, dict[str, Any]] = {}
    for key in CORE_REPORT_KEYS:
        cfg = REPORT_SECTION_CONFIG[key]
        src = Path(generated_sections.get(key) or editable_dir(output_dir, incident_id) / cfg["filename"])
        text = _read(src)
        structured_src = Path(generated_sections.get(f"{key}_structured") or editable_dir(output_dir, incident_id) / f"{key}.json")
        blocks = repair_pipe_tables_in_blocks(load_blocks(structured_src))
        if not blocks:
            blocks = blocks_from_text(text) if text else []
        if key == "soc_analyst_review" and support_text and "Supporting SOC Triage Review" not in text:
            support_blocks = load_blocks(generated_sections.get("soc_triage_review_structured")) or blocks_from_text(support_text)
            text = (text.rstrip() + "\n\nSupporting SOC Triage Review\n\n" + support_text.strip()).strip()
            blocks = blocks + [{"type": "heading", "level": 2, "text": "Supporting SOC Triage Review"}] + support_blocks
        draft_path = ddir / cfg["filename"]
        structured_draft_path = ddir / f"{key}.json"
        _write(draft_path, text)
        save_blocks(structured_draft_path, blocks)
        sections[key] = {
            "key": key,
            "title": cfg["title"],
            "description": cfg["description"],
            "template": cfg["template"],
            "filename": cfg["filename"],
            "status": "draft",
            "draft_path": str(draft_path),
            "draft_relative_path": _rel(output_dir, draft_path),
            "structured_draft_path": str(structured_draft_path),
            "structured_draft_relative_path": _rel(output_dir, structured_draft_path),
            "confirmed_path": None,
            "confirmed_relative_path": None,
            "structured_confirmed_path": None,
            "structured_confirmed_relative_path": None,
            "exports": {},
            "last_saved_at": utc_now(),
            "last_saved_by": "Reporting Agent",
            "confirmed_at": None,
            "confirmed_by": None,
            "draft_history": [],
        }
    support_sections = {}
    for key in SUPPORT_REPORT_KEYS:
        if key in generated_sections:
            cfg = REPORT_SECTION_CONFIG[key]
            support_sections[key] = {
                "key": key,
                "title": cfg["title"],
                "template": cfg["template"],
                "path": generated_sections[key],
                "relative_path": _rel(output_dir, Path(generated_sections[key])),
            }
    manifest = {
        "schema_version": "editable-report-manifest-v2",
        "incident_id": incident_id,
        "report_status": "draft_ready_for_analyst_review",
        "display_status": "Draft ready for analyst review",
        "sections": sections,
        "section_order": CORE_REPORT_KEYS,
        "support_sections": support_sections,
        "draft_reports": list(CORE_REPORT_KEYS),
        "confirmed_reports": [],
        "created_at": utc_now(),
        "updated_at": utc_now(),
        "confirmed_by": None,
        "confirmed_at": None,
        "exports": {},
        "source_context": {
            "report_status": context.get("report_status"),
            "validation_status": context.get("validation_status"),
            "data_consistency_status": context.get("data_consistency_status"),
        },
    }
    path = manifest_path(output_dir, incident_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return manifest


def load_manifest(output_dir: Path, incident_id: str | None = None) -> dict[str, Any]:
    if incident_id:
        path = manifest_path(output_dir, incident_id)
        if path.exists():
            return json.loads(path.read_text(encoding="utf-8"))
    manifests = sorted(output_dir.glob("*/reports/report_manifest.json"), key=lambda p: p.stat().st_mtime, reverse=True)
    if manifests:
        return json.loads(manifests[0].read_text(encoding="utf-8"))
    return {}


def save_manifest(output_dir: Path, manifest: dict[str, Any]) -> None:
    incident_id = manifest.get("incident_id") or "INC-0001"
    manifest["updated_at"] = utc_now()
    manifest["draft_reports"] = [k for k, s in (manifest.get("sections") or {}).items() if s.get("status") in {"draft", "draft_revision"} and s.get("draft_path")]
    manifest["confirmed_reports"] = [k for k, s in (manifest.get("sections") or {}).items() if s.get("status") in {"confirmed", "exported"} and s.get("confirmed_path")]
    if manifest["confirmed_reports"] and not manifest["draft_reports"]:
        manifest["report_status"] = "confirmed_by_analyst"
        manifest["display_status"] = "All reports confirmed by analyst"
    elif manifest["confirmed_reports"]:
        manifest["report_status"] = "partially_confirmed"
        manifest["display_status"] = "Partially confirmed"
    elif manifest["draft_reports"]:
        manifest["report_status"] = "analyst_editing"
        manifest["display_status"] = "Analyst editing"
    path = manifest_path(output_dir, incident_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")


def list_reports(output_dir: Path, incident_id: str | None = None) -> dict[str, Any]:
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    sections = manifest.get("sections") or {}
    return {
        "success": True,
        "manifest": manifest,
        "reports": sections,
        "draft_reports": [sections[k] for k in manifest.get("draft_reports", []) if k in sections],
        "confirmed_reports": [sections[k] for k in manifest.get("confirmed_reports", []) if k in sections],
        "section_order": manifest.get("section_order", CORE_REPORT_KEYS),
    }


def _section_text(section: dict[str, Any]) -> tuple[str, str]:
    if section.get("draft_path") and Path(section["draft_path"]).exists():
        return "draft", _read(section["draft_path"])
    if section.get("confirmed_path") and Path(section["confirmed_path"]).exists():
        return "confirmed", _read(section["confirmed_path"])
    return "missing", ""


def _section_blocks(section: dict[str, Any]) -> tuple[str, list[dict[str, Any]]]:
    if section.get("structured_draft_path") and Path(section["structured_draft_path"]).exists():
        return "draft", repair_pipe_tables_in_blocks(load_blocks(section["structured_draft_path"]))
    if section.get("structured_confirmed_path") and Path(section["structured_confirmed_path"]).exists():
        return "confirmed", repair_pipe_tables_in_blocks(load_blocks(section["structured_confirmed_path"]))
    source, text = _section_text(section)
    return source, blocks_from_text(text) if text else []


def read_section(output_dir: Path, section_key: str, incident_id: str | None = None) -> dict[str, Any]:
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    sections = manifest.get("sections") or {}
    if section_key not in sections:
        raise KeyError(f"Unknown report section: {section_key}")
    source, text = _section_text(sections[section_key])
    block_source, blocks = _section_blocks(sections[section_key])
    return {"manifest": manifest, "section": sections[section_key], "text": text, "blocks": blocks, "source": source, "block_source": block_source}


def save_section(output_dir: Path, section_key: str, text: str, analyst: str = "SOC Analyst", incident_id: str | None = None, blocks: list[dict[str, Any]] | None = None) -> dict[str, Any]:
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    section = manifest.get("sections", {}).get(section_key)
    if not section:
        raise KeyError(f"Unknown report section: {section_key}")
    incident_id = manifest.get("incident_id") or "INC-0001"
    draft_path = drafts_dir(output_dir, incident_id) / section["filename"]
    structured_draft_path = drafts_dir(output_dir, incident_id) / f"{section_key}.json"
    clean_blocks = repair_pipe_tables_in_blocks(blocks) if isinstance(blocks, list) else None
    if clean_blocks is None:
        clean_text_input = markdown_to_plain_text(text)
        clean_blocks = blocks_from_text(clean_text_input) if clean_text_input else []
    clean_text = blocks_to_plain_text(clean_blocks) if clean_blocks else markdown_to_plain_text(text)
    _write(draft_path, clean_text)
    save_blocks(structured_draft_path, clean_blocks)
    history = list(section.get("draft_history") or [])
    history.append(_draft_history_entry(output_dir, incident_id, section_key, section.get("title") or section_key, clean_text, analyst))
    section["draft_history"] = history[-20:]
    section["status"] = "draft_revision" if section.get("confirmed_path") else "draft"
    section["draft_path"] = str(draft_path)
    section["draft_relative_path"] = _rel(output_dir, draft_path)
    section["structured_draft_path"] = str(structured_draft_path)
    section["structured_draft_relative_path"] = _rel(output_dir, structured_draft_path)
    section["last_saved_at"] = utc_now()
    section["last_saved_by"] = analyst
    manifest["sections"][section_key] = section
    save_manifest(output_dir, manifest)
    return {"success": True, "manifest": manifest, "section": section, "text": _read(draft_path), "blocks": clean_blocks, "message": f"Draft saved for {section['title']}"}


def confirm_section(output_dir: Path, section_key: str, analyst: str = "SOC Analyst", incident_id: str | None = None) -> dict[str, Any]:
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    section = manifest.get("sections", {}).get(section_key)
    if not section:
        raise KeyError(f"Unknown report section: {section_key}")
    incident_id = manifest.get("incident_id") or "INC-0001"
    src = Path(section.get("draft_path") or section.get("confirmed_path") or "")
    if not src.exists():
        raise FileNotFoundError(f"No draft exists for {section.get('title') or section_key}. Save a draft first.")
    cdir = confirmed_dir(output_dir, incident_id)
    cdir.mkdir(parents=True, exist_ok=True)
    dst = cdir / section["filename"]
    src_is_dst = _safe_resolve(src) == _safe_resolve(dst)
    if not src_is_dst:
        shutil.copy2(src, dst)

    structured_src = Path(section.get("structured_draft_path") or section.get("structured_confirmed_path") or "")
    structured_dst = cdir / f"{section_key}.json"
    if structured_src.exists():
        confirmed_blocks = repair_pipe_tables_in_blocks(load_blocks(structured_src))
        save_blocks(structured_dst, confirmed_blocks)
    else:
        save_blocks(structured_dst, blocks_from_text(_read(dst)))

    # Remove active draft after confirmation to make the status clear.
    # If the source is already the confirmed file, do not unlink it.
    draft_value = section.get("draft_path")
    if draft_value:
        draft_path = Path(draft_value)
        if draft_path.exists() and _safe_resolve(draft_path) != _safe_resolve(dst):
            try:
                draft_path.unlink()
            except Exception:
                pass
    structured_draft_value = section.get("structured_draft_path")
    if structured_draft_value:
        structured_draft_path = Path(structured_draft_value)
        if structured_draft_path.exists() and _safe_resolve(structured_draft_path) != _safe_resolve(structured_dst):
            try:
                structured_draft_path.unlink()
            except Exception:
                pass
    section["status"] = "confirmed"
    section["draft_path"] = None
    section["draft_relative_path"] = None
    section["structured_draft_path"] = None
    section["structured_draft_relative_path"] = None
    section["confirmed_path"] = str(dst)
    section["confirmed_relative_path"] = _rel(output_dir, dst)
    section["structured_confirmed_path"] = str(structured_dst)
    section["structured_confirmed_relative_path"] = _rel(output_dir, structured_dst)
    section["confirmed_at"] = utc_now()
    section["confirmed_by"] = analyst
    manifest["sections"][section_key] = section
    manifest["confirmed_by"] = analyst
    manifest["confirmed_at"] = utc_now()
    save_manifest(output_dir, manifest)
    return {"success": True, "manifest": manifest, "section": section, "text": _read(dst), "blocks": load_blocks(structured_dst), "message": f"{section['title']} confirmed"}


def confirm_report(output_dir: Path, analyst: str = "SOC Analyst", incident_id: str | None = None) -> dict[str, Any]:
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    for key in list(manifest.get("section_order") or CORE_REPORT_KEYS):
        # Reload after each confirm.
        current = load_manifest(output_dir, incident_id)
        section = (current.get("sections") or {}).get(key)
        if section and (section.get("draft_path") or section.get("confirmed_path")):
            confirm_section(output_dir, key, analyst=analyst, incident_id=current.get("incident_id"))
    manifest = load_manifest(output_dir, incident_id)
    manifest["report_status"] = "confirmed_by_analyst"
    manifest["display_status"] = "All reports confirmed by analyst"
    manifest["confirmed_by"] = analyst
    manifest["confirmed_at"] = utc_now()
    save_manifest(output_dir, manifest)
    return {"success": True, "manifest": manifest, "message": "All report sections confirmed"}



def list_section_drafts(output_dir: Path, section_key: str, incident_id: str | None = None) -> dict[str, Any]:
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    section = (manifest.get("sections") or {}).get(section_key)
    if not section:
        raise KeyError(f"Unknown report section: {section_key}")
    drafts = []
    for entry in reversed(list(section.get("draft_history") or [])):
        path = Path(entry.get("path", ""))
        text = _read(path)
        item = dict(entry)
        item["text"] = text
        item["preview"] = markdown_to_plain_text(text)[:700]
        drafts.append(item)
    return {"success": True, "manifest": manifest, "section": section, "title": section.get("title") or section_key, "drafts": drafts}

def _confirmed_required(section: dict[str, Any]) -> None:
    if section.get("status") not in {"confirmed", "exported"} or not section.get("confirmed_path"):
        raise PermissionError("This report must be confirmed by the SOC analyst before export.")
    if not Path(section["confirmed_path"]).exists():
        raise FileNotFoundError("Confirmed report file not found.")


def _docx_write(path: Path, title: str, text: str, incident_id: str, manifest: dict[str, Any]) -> None:
    blocks = blocks_from_text(text)
    _docx_write_blocks(path, title, blocks, incident_id, manifest)



def _set_cell_shading(cell: Any, fill: str) -> None:
    if OxmlElement is None or qn is None:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def _apply_report_font(run: Any, name: str = "Georgia") -> None:
    try:
        run.font.name = name
        if qn is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
            run._element.rPr.rFonts.set(qn("w:ascii"), name)
            run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    except Exception:
        pass


def _set_cell_text(cell: Any, text: Any, *, bold: bool = False, font_size: int = 10, color: str = "1f2937") -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(str(text or ""))
    _apply_report_font(run)
    run.bold = bool(bold)
    if Pt is not None:
        run.font.size = Pt(font_size)
    if RGBColor is not None and color:
        try:
            run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
        except Exception:
            pass
    if WD_CELL_VERTICAL_ALIGNMENT is not None:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def _style_table(table: Any, *, header_fill: str = "EAF2FB", first_col_fill: str | None = None) -> None:
    table.style = "Table Grid"
    if WD_TABLE_ALIGNMENT is not None:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if r_idx == 0:
                _set_cell_shading(cell, header_fill)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        if RGBColor is not None:
                            run.font.color.rgb = RGBColor.from_string("1E3A8A")
            elif first_col_fill and c_idx == 0:
                _set_cell_shading(cell, first_col_fill)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
            for para in cell.paragraphs:
                if Pt is not None:
                    para.paragraph_format.space_after = Pt(0)


def _add_title_block(doc: Any, title: str, incident_id: str, manifest: dict[str, Any]) -> None:
    logo_path = _aegis_logo_path()
    if logo_path:
        logo = doc.add_paragraph()
        if WD_ALIGN_PARAGRAPH is not None:
            logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if Pt is not None:
            logo.paragraph_format.space_after = Pt(8)
        logo_run = logo.add_run()
        logo_run.add_picture(str(logo_path), width=Inches(2.65))
    else:
        label = doc.add_paragraph()
        if WD_ALIGN_PARAGRAPH is not None:
            label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label.add_run("AEGIS")
        _apply_report_font(label_run)
        label_run.bold = True
        if Pt is not None:
            label_run.font.size = Pt(10)
        if RGBColor is not None:
            label_run.font.color.rgb = RGBColor.from_string("2F66D0")

    heading = doc.add_paragraph()
    if WD_ALIGN_PARAGRAPH is not None:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run(title)
    _apply_report_font(heading_run)
    heading_run.bold = True
    if Pt is not None:
        heading_run.font.size = Pt(30)
    if RGBColor is not None:
        heading_run.font.color.rgb = RGBColor.from_string("172033")

    subtitle = doc.add_paragraph("Reviewed SOC report section exported from the analyst-confirmed dashboard draft and converted to PDF from Word.")
    if WD_ALIGN_PARAGRAPH is not None:
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _run in subtitle.runs:
        _apply_report_font(_run)
    if Pt is not None:
        subtitle.runs[0].font.size = Pt(11)
    if RGBColor is not None:
        subtitle.runs[0].font.color.rgb = RGBColor.from_string("526071")

    divider = doc.add_paragraph()
    divider_run = divider.add_run("━" * 80)
    _apply_report_font(divider_run)
    if RGBColor is not None:
        divider_run.font.color.rgb = RGBColor.from_string("2F66D0")
    if Pt is not None:
        divider_run.font.size = Pt(8)

    meta = doc.add_table(rows=1, cols=2)
    meta.style = "Table Grid"
    values = [
        ("Incident ID", incident_id),
        ("Report Status", "Confirmed by SOC Analyst"),
        ("Confirmed By", manifest.get("confirmed_by") or "SOC Analyst"),
        ("Exported At", utc_now()),
    ]
    for idx, (field, value) in enumerate(values):
        row = meta.rows[0] if idx == 0 else meta.add_row()
        _set_cell_text(row.cells[0], field, bold=True, font_size=10, color="172033")
        _set_cell_text(row.cells[1], value, font_size=10, color="172033")
        _set_cell_shading(row.cells[0], "F3F6FA")
    doc.add_paragraph("")

def _docx_write_blocks(path: Path, title: str, blocks: list[dict[str, Any]], incident_id: str, manifest: dict[str, Any]) -> None:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
    blocks = repair_pipe_tables_in_blocks(blocks)
    _validate_no_raw_markdown_tables(blocks, title)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    try:
        normal.font.name = "Georgia"
        if qn is not None:
            normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Georgia")
            normal._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
            normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    except Exception:
        pass
    if Pt is not None:
        normal.font.size = Pt(10.5)
    if RGBColor is not None:
        normal.font.color.rgb = RGBColor.from_string("172033")

    _add_title_block(doc, title, incident_id, manifest)

    for block in _strip_duplicate_leading_heading(blocks or [], title):
        btype = block.get("type") if isinstance(block, dict) else None
        if btype == "heading":
            level = int(block.get("level") or 2)
            para = doc.add_paragraph()
            run = para.add_run(str(block.get("text") or ""))
            _apply_report_font(run)
            run.bold = True
            if Pt is not None:
                run.font.size = Pt(22 if level <= 1 else 16 if level == 2 else 13)
            if RGBColor is not None:
                run.font.color.rgb = RGBColor.from_string("2F66D0" if level <= 2 else "1F2937")
            if Pt is not None:
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(6)
        elif btype == "paragraph":
            text = str(block.get("text") or "").strip()
            if text:
                para = doc.add_paragraph(text)
                for _run in para.runs:
                    _apply_report_font(_run)
                if Pt is not None:
                    para.paragraph_format.space_after = Pt(6)
        elif btype == "bullet_list":
            for item in block.get("items") or []:
                item_text = str(item or "").strip()
                if item_text:
                    bullet_para = doc.add_paragraph(item_text, style="List Bullet")
                    for _run in bullet_para.runs:
                        _apply_report_font(_run)
        elif btype == "table":
            columns = [str(c or "") for c in (block.get("columns") or [])]
            rows = block.get("rows") or []
            if not columns:
                continue
            table = doc.add_table(rows=1, cols=len(columns))
            for idx, col in enumerate(columns):
                _set_cell_text(table.rows[0].cells[idx], col, bold=True, font_size=10, color="1E3A8A")
            for row in rows:
                cells = table.add_row().cells
                row_values = list(row or []) + [""] * (len(columns) - len(row or []))
                for idx, value in enumerate(row_values[:len(columns)]):
                    _set_cell_text(cells[idx], value, bold=(idx == 0 and len(columns) == 2), font_size=9, color="172033")
            _style_table(table, header_fill="EAF2FB", first_col_fill="F3F6FA" if len(columns) == 2 else None)
            doc.add_paragraph("")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def _strip_duplicate_leading_heading(blocks: list[dict[str, Any]], title: str) -> list[dict[str, Any]]:
    if not blocks:
        return []
    normalised_title = re.sub(r"\s+", " ", str(title or "").strip().lower())
    cleaned = list(blocks)
    while cleaned:
        first = cleaned[0]
        if not isinstance(first, dict) or first.get("type") != "heading":
            break
        text = re.sub(r"\s+", " ", str(first.get("text") or "").strip().lower())
        if text in {normalised_title, "cybersecurity incident post-incident review report"}:
            cleaned.pop(0)
            continue
        break
    return cleaned

def _pdf_escape(text: str) -> str:
    return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pdf_para(text: str, style: Any) -> Any:
    return Paragraph(_pdf_escape(str(text or "")), style)


def _pdf_write(path: Path, title: str, text: str, incident_id: str, manifest: dict[str, Any]) -> None:
    blocks = blocks_from_text(text)
    _pdf_write_blocks(path, title, blocks, incident_id, manifest)


def _pdf_write_blocks(path: Path, title: str, blocks: list[dict[str, Any]], incident_id: str, manifest: dict[str, Any]) -> None:
    if SimpleDocTemplate is None or Table is None:
        raise RuntimeError("reportlab is not installed. Run: pip install reportlab")
    blocks = repair_pipe_tables_in_blocks(blocks)
    _validate_no_raw_markdown_tables(blocks, title)
    styles = getSampleStyleSheet()
    story = []
    logo_path = _aegis_logo_path()
    if Image is not None and logo_path:
        logo = Image(str(logo_path), width=2.65 * inch, height=0.86 * inch)
        logo.hAlign = "CENTER"
        story.extend([logo, Spacer(1, 0.12 * inch)])
    story.extend([_pdf_para(title, styles["Title"]), Spacer(1, 0.15 * inch)])
    meta_data = [[_pdf_para("Field", styles["Heading5"]), _pdf_para("Value", styles["Heading5"])],
                 [_pdf_para("Incident ID", styles["BodyText"]), _pdf_para(incident_id, styles["BodyText"])],
                 [_pdf_para("Confirmed by", styles["BodyText"]), _pdf_para(manifest.get("confirmed_by") or "SOC Analyst", styles["BodyText"])],
                 [_pdf_para("Generated at", styles["BodyText"]), _pdf_para(utc_now(), styles["BodyText"])] ]
    meta = Table(meta_data, repeatRows=1, hAlign="LEFT", colWidths=[1.6*inch, 4.8*inch])
    meta.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.4, colors.HexColor("#9AA4B2")),
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#E8EEF6")),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("LEFTPADDING", (0,0), (-1,-1), 6),
        ("RIGHTPADDING", (0,0), (-1,-1), 6),
    ]))
    story.extend([meta, Spacer(1, 0.2*inch)])

    for block in _strip_duplicate_leading_heading(blocks or [], title):
        if not isinstance(block, dict):
            continue
        btype = block.get("type")
        if btype == "heading":
            level = int(block.get("level") or 2)
            style_name = "Heading1" if level <= 1 else ("Heading2" if level == 2 else "Heading3")
            story.extend([Spacer(1, 0.08*inch), _pdf_para(block.get("text") or "", styles[style_name])])
        elif btype == "paragraph":
            text = str(block.get("text") or "").strip()
            if text:
                story.append(_pdf_para(text, styles["BodyText"]))
                story.append(Spacer(1, 0.06*inch))
        elif btype == "bullet_list":
            for item in block.get("items") or []:
                if str(item or "").strip():
                    story.append(_pdf_para("• " + str(item), styles["BodyText"]))
        elif btype == "table":
            columns = [str(c or "") for c in (block.get("columns") or [])]
            rows = block.get("rows") or []
            if not columns:
                continue
            data = [[_pdf_para(c, styles["Heading5"]) for c in columns]]
            for row in rows:
                values = list(row or []) + [""] * (len(columns) - len(row or []))
                data.append([_pdf_para(v, styles["BodyText"]) for v in values[:len(columns)]])
            usable_width = 7.0 * inch
            col_width = usable_width / max(1, len(columns))
            table = Table(data, repeatRows=1, hAlign="LEFT", colWidths=[col_width] * len(columns))
            table.setStyle(TableStyle([
                ("GRID", (0,0), (-1,-1), 0.35, colors.HexColor("#9AA4B2")),
                ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#E8EEF6")),
                ("VALIGN", (0,0), (-1,-1), "TOP"),
                ("LEFTPADDING", (0,0), (-1,-1), 5),
                ("RIGHTPADDING", (0,0), (-1,-1), 5),
                ("TOPPADDING", (0,0), (-1,-1), 4),
                ("BOTTOMPADDING", (0,0), (-1,-1), 4),
            ]))
            story.extend([table, Spacer(1, 0.15*inch)])
    path.parent.mkdir(parents=True, exist_ok=True)
    SimpleDocTemplate(str(path), pagesize=A4, rightMargin=0.45*inch, leftMargin=0.45*inch, topMargin=0.5*inch, bottomMargin=0.5*inch).build(story)


def _confirmed_blocks_or_text(section: dict[str, Any]) -> tuple[list[dict[str, Any]], str]:
    structured_path = section.get("structured_confirmed_path")
    original_blocks = load_blocks(structured_path)
    blocks = repair_pipe_tables_in_blocks(original_blocks)
    text = _read(section.get("confirmed_path"))
    if not blocks:
        blocks = blocks_from_text(text)
        blocks = repair_pipe_tables_in_blocks(blocks)
    _validate_no_raw_markdown_tables(blocks, section.get("title"))
    if structured_path and blocks != original_blocks:
        save_blocks(Path(structured_path), blocks)
    return blocks, text


def _safe_raw_table_preview(text: Any, limit: int = 180) -> str:
    preview = re.sub(r"\s+", " ", str(text or "")).strip()
    preview = re.sub(
        r"(?i)\b(api[_ -]?key|access[_ -]?token|token|secret|authorization)\b\s*[:=]\s*\S+",
        r"\1=[REDACTED]",
        preview,
    )
    return preview[:limit] + ("..." if len(preview) > limit else "")


def _validate_no_raw_markdown_tables(blocks: list[dict[str, Any]], section_title: str | None = None) -> None:
    repaired = repair_pipe_tables_in_blocks(blocks)
    if isinstance(blocks, list):
        blocks[:] = repaired
    current_section = section_title or "Report"
    for index, block in enumerate(repaired):
        if isinstance(block, dict) and block.get("type") == "heading":
            current_section = str(block.get("text") or current_section)
            continue
        if not isinstance(block, dict) or block.get("type") != "paragraph":
            continue
        text = str(block.get("text") or "")
        if paragraph_contains_raw_pipe_table(text):
            preview = _safe_raw_table_preview(text)
            raise ValueError(
                f'Export validation failed: raw table syntax remained in section "{current_section}", '
                f"block index {index}: {preview}"
            )


def export_section_docx(output_dir: Path, section_key: str, incident_id: str | None = None) -> dict[str, Any]:
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    section = (manifest.get("sections") or {}).get(section_key)
    if not section:
        raise KeyError(f"Unknown report section: {section_key}")
    _confirmed_required(section)
    incident_id = manifest.get("incident_id") or "INC-0001"
    blocks, text = _confirmed_blocks_or_text(section)
    path = exports_dir(output_dir, incident_id) / f"{section_key}.docx"
    _docx_write_blocks(path, section.get("title") or section_key, blocks, incident_id, manifest)
    section.setdefault("exports", {})["docx"] = {"path": str(path), "relative_path": _rel(output_dir, path), "created_at": utc_now()}
    section["status"] = "exported"
    manifest["sections"][section_key] = section
    save_manifest(output_dir, manifest)
    return {"success": True, "manifest": manifest, "section": section, "path": str(path), "download_url": f"/api/reports/{section_key}/download/docx", "message": "Word export ready"}


def export_section_pdf(output_dir: Path, section_key: str, incident_id: str | None = None) -> dict[str, Any]:
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    section = (manifest.get("sections") or {}).get(section_key)
    if not section:
        raise KeyError(f"Unknown report section: {section_key}")
    _confirmed_required(section)
    incident_id = manifest.get("incident_id") or "INC-0001"
    docx_result = export_section_docx(output_dir, section_key, incident_id=incident_id)
    docx_path = Path(docx_result["path"])
    path = exports_dir(output_dir, incident_id) / f"{section_key}.pdf"
    try:
        from reporting.template_document_exporter import convert_docx_to_pdf
        convert_docx_to_pdf(docx_path, path)
    except Exception:
        blocks, text = _confirmed_blocks_or_text(section)
        _pdf_write_blocks(path, section.get("title") or section_key, blocks, incident_id, manifest)
    manifest = load_manifest(output_dir, incident_id)
    section = (manifest.get("sections") or {}).get(section_key) or section
    section.setdefault("exports", {})["pdf"] = {"path": str(path), "relative_path": _rel(output_dir, path), "created_at": utc_now(), "source_docx": str(docx_path)}
    section["status"] = "exported"
    manifest["sections"][section_key] = section
    save_manifest(output_dir, manifest)
    return {"success": True, "manifest": manifest, "section": section, "path": str(path), "download_url": f"/api/reports/{section_key}/download/pdf", "message": "PDF export ready from confirmed Word document"}


def _combined_blocks(output_dir: Path, manifest: dict[str, Any]) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    for key in manifest.get("section_order") or CORE_REPORT_KEYS:
        section = (manifest.get("sections") or {}).get(key)
        if not section:
            continue
        title = section.get("title") or key.replace("_", " ").title()
        section_blocks = repair_pipe_tables_in_blocks(load_blocks(section.get("structured_confirmed_path") or section.get("structured_draft_path")))
        if not section_blocks:
            text = _read(section.get("confirmed_path") or section.get("draft_path"))
            section_blocks = blocks_from_text(text)
        blocks.append({"type": "heading", "level": 1, "text": title})
        blocks.extend(_strip_duplicate_leading_heading(section_blocks, title))
    return blocks


def export_docx(output_dir: Path, incident_id: str | None = None) -> dict[str, Any]:
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    if manifest.get("draft_reports"):
        raise PermissionError("All reports must be confirmed before combined DOCX export.")
    incident_id = manifest.get("incident_id") or "INC-0001"
    blocks = _combined_blocks(output_dir, manifest)
    path = exports_dir(output_dir, incident_id) / "combined_incident_report.docx"
    _docx_write_blocks(path, f"Combined Cybersecurity Incident Report - {incident_id}", blocks, incident_id, manifest)
    manifest.setdefault("exports", {})["docx"] = {"path": str(path), "relative_path": _rel(output_dir, path), "created_at": utc_now()}
    save_manifest(output_dir, manifest)
    return {"success": True, "path": str(path), "manifest": manifest, "download_url": "/api/reports/download/docx"}


def export_pdf(output_dir: Path, incident_id: str | None = None) -> dict[str, Any]:
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found. Run Reporting Agent first.")
    if manifest.get("draft_reports"):
        raise PermissionError("All reports must be confirmed before combined PDF export.")
    incident_id = manifest.get("incident_id") or "INC-0001"
    docx_result = export_docx(output_dir, incident_id=incident_id)
    docx_path = Path(docx_result["path"])
    path = exports_dir(output_dir, incident_id) / "combined_incident_report.pdf"
    try:
        from reporting.template_document_exporter import convert_docx_to_pdf
        convert_docx_to_pdf(docx_path, path)
    except Exception:
        blocks = _combined_blocks(output_dir, manifest)
        _pdf_write_blocks(path, f"Combined Cybersecurity Incident Report - {incident_id}", blocks, incident_id, manifest)
    manifest = load_manifest(output_dir, incident_id)
    manifest.setdefault("exports", {})["pdf"] = {"path": str(path), "relative_path": _rel(output_dir, path), "created_at": utc_now(), "source_docx": str(docx_path)}
    save_manifest(output_dir, manifest)
    return {"success": True, "path": str(path), "manifest": manifest, "download_url": "/api/reports/download/pdf"}


def download_path(output_dir: Path, section_key: str | None, file_type: str, incident_id: str | None = None) -> Path:
    manifest = load_manifest(output_dir, incident_id)
    if not manifest:
        raise FileNotFoundError("No report manifest found")
    if section_key:
        section = (manifest.get("sections") or {}).get(section_key)
        if not section:
            raise KeyError(f"Unknown report section: {section_key}")
        export = (section.get("exports") or {}).get(file_type)
    else:
        export = (manifest.get("exports") or {}).get(file_type)
    if not export:
        raise FileNotFoundError(f"No {file_type} export found")
    path = Path(export.get("path", ""))
    if not path.exists():
        raise FileNotFoundError("Export file not found")
    return path
