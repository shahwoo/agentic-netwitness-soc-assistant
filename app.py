"""
SOC Platform v4
───────────────
• Auto-fetches incidents every 30 seconds once token is verified
• Dashboard tab: severity bars, status bars, latest incidents, assignee chart
• Cleaner sidebar with connection status + refresh countdown
• Incidents tab with severity filter + auto-refresh indicator
• Chat stub (wire LangChain inside chat_respond)
• ChromaDB semantic search & sync
"""

# ── Streamlit Community Cloud: ChromaDB needs sqlite3 >= 3.35 ──────────────────
# Debian images on Streamlit Cloud ship a system sqlite3 older than 3.35, which
# chromadb rejects ("unsupported version of sqlite3"). pysqlite3-binary bundles a
# newer sqlite; swap it in for the stdlib module BEFORE anything imports chromadb.
# No-op locally (pysqlite3-binary isn't installed on the Windows venv, and that
# platform's own sqlite3 is already new enough) — this block is deploy-only.
try:
    __import__("pysqlite3")
    import sys as _sys
    _sys.modules["sqlite3"] = _sys.modules.pop("pysqlite3")
except Exception:
    pass

import re
import streamlit as st
import streamlit.components.v1 as components
import requests
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
import time
import os
import base64
from datetime import datetime, timedelta
from typing import Optional
from collections import Counter
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor

try:
    from dotenv import load_dotenv, set_key, find_dotenv
    DOTENV_OK = True
except ImportError:
    DOTENV_OK = False

try:
    import chromadb
    CHROMA_OK = True
except ImportError:
    CHROMA_OK = False

# ── Streamlit Community Cloud: secrets → environment ──────────────────────────
# Locally the app reads credentials from .env (via python-dotenv). On Streamlit
# Cloud there is no .env; secrets are entered in the dashboard and exposed as
# st.secrets. Copy any scalar secrets into os.environ (without clobbering a value
# already set) so every existing os.environ.get(...) call site keeps working
# unchanged. Guarded: a silent no-op locally when no secrets are configured.
def _bridge_cloud_secrets() -> None:
    try:
        for _k, _v in st.secrets.items():
            if isinstance(_v, (str, int, float, bool)) and _k not in os.environ:
                os.environ[_k] = str(_v)
    except Exception:
        pass

_bridge_cloud_secrets()

# ── SOC Triage Agent (LangChain) ──────────────────────────────────────────────
# Streamlit re-executes this script on every rerun but NEVER re-imports
# modules already in sys.modules — edits to the agent files were invisible
# until a full process restart. This shim reloads them when their file
# mtime changes, so agent upgrades apply on the next page refresh.
def _maybe_reload_agent_modules():
    import importlib
    import sys as _sys
    watched = {
        "soc_triage_agent.soc_triage_agent":
            Path(__file__).parent / "soc_triage_agent" / "soc_triage_agent.py",
        "soc_workflow": Path(__file__).parent / "soc_workflow.py",
    }
    try:
        for mod_name, path in watched.items():
            m = _sys.modules.get(mod_name)
            if m is None:
                continue
            mtime = path.stat().st_mtime
            # Stamp the module object itself (survives across sessions) —
            # session-scoped bookkeeping missed changes made BEFORE a new
            # browser session started on an old process.
            loaded = getattr(m, "__loaded_mtime__", None)
            if loaded is None or mtime > loaded:
                for name in (mod_name, "soc_triage_agent"):
                    mm = _sys.modules.get(name)
                    if mm is not None:
                        importlib.reload(mm)
                _sys.modules[mod_name].__loaded_mtime__ = mtime
    except Exception:
        pass

_maybe_reload_agent_modules()

from soc_triage_agent import CiscoLLMConfig, soc_triage_chat_respond, _TRIAGE_TRIGGER

# ── Multi-agent workflow (triage → investigation → reporting) ────────────────
try:
    from soc_workflow import (
        needs_investigation      as wf_needs_investigation,
        handoff_to_investigation as wf_handoff_to_investigation,
        run_investigation        as wf_run_investigation,
        handoff_to_reporting     as wf_handoff_to_reporting,
        run_reporting            as wf_run_reporting,
    )
    WORKFLOW_OK = True
except Exception:
    WORKFLOW_OK = False


# ── Background workflow engine ───────────────────────────────────────────────
# Streamlit interrupts the running script at its next UI call whenever the
# user interacts (clicking "View" on the agent board, sending a message…).
# Running investigation+reporting inline therefore died mid-run on any click.
# They now run in a daemon worker thread that survives all interactions; the
# UI polls the shared run record and renders its state live.

_ANSI_STRIP = re.compile(r"\x1b\[[0-9;]*m")


@st.cache_resource
def _workflow_store() -> dict:
    """Process-global store for the active background workflow run."""
    return {"run": None}


def _workflow_worker(run: dict, tri: dict, incident: dict) -> None:
    """Investigation + reporting stages, off the Streamlit script thread.
    NO st.* calls in here — the UI polls `run` and renders its state."""
    import soc_workflow as _wfm

    panels = run["panels"]
    wf_md  = run["wf_md"]

    def bset(agent, status=None, think=None, output=None, progress=None):
        p = panels[agent]
        if status is not None:
            p["status"] = status
            if status in ("done", "cached", "skipped"):
                p["progress"] = 100
        if progress is not None:
            p["progress"] = max(0, min(100, int(progress)))
        if think is not None:
            clean = _ANSI_STRIP.sub("", str(think))
            p["thinking"].append(
                f"[{datetime.now().strftime('%H:%M:%S')}] {clean}")
            p["thinking"] = p["thinking"][-60:]
        if output is not None:
            p["output"] = output
        p["updated"] = datetime.now().strftime("%H:%M:%S")

    inc_id = run["incident_id"]
    title  = run["title"]
    cls    = run["cls"]
    unc    = run["unc"]
    ticket = tri.get("ticket") or {}
    inv = None
    rep: dict = {}
    try:
        # ── Stage 2: Investigation ─────────────────────────────────────────
        if run["investigate"]:
            bset("investigation", status="running", think="Investigation started",
                 progress=10)

            def _fb_event(event: str, detail: str) -> None:
                # Board choreography for the feedback loop: the work visibly
                # returns to the triage card for the deep-dive, then hands
                # back to investigation for the second pass.
                if event == "gaps_detected":
                    bset("investigation", think=f"{detail}")
                elif event == "triage_deep_dive_start":
                    bset("triage", status="running", think=f"{detail}")
                elif event == "triage_deep_dive_done":
                    bset("triage", status="done", think=f"{detail}")
                    bset("investigation", think=f"{detail}")
                elif event == "second_pass_start":
                    bset("investigation", status="running", think=f"{detail}")
                elif event == "supplement_error":
                    bset("triage", status="done",
                         think=f"deep-dive failed: {detail}")
                else:
                    bset("investigation", think=detail)

            inv = _wfm.investigate_with_feedback(
                tri, incident, inc_id,
                line_cb=lambda ln: bset("investigation", think=ln)
                if ln.strip() else None,
                feedback_cb=_fb_event)
            _fbmeta = inv.get("feedback_loop") or {}
            if _fbmeta.get("triggered"):
                # Honest reporting — a crashed deep-dive or failed second
                # pass must never be presented as a successful loop.
                gap_ids = ", ".join(g.split(":")[0]
                                    for g in (_fbmeta.get("gaps") or []))
                if _fbmeta.get("supplement_error"):
                    wf_md.append(
                        f"- Feedback loop: gaps detected (`{gap_ids}`) but "
                        f"the triage deep-dive failed — "
                        f"`{str(_fbmeta['supplement_error'])[:120]}` — "
                        f"pass-1 findings kept")
                elif _fbmeta.get("second_pass_failed"):
                    wf_md.append(
                        f"- Feedback loop: triage deep-dive answered "
                        f"{_fbmeta.get('gaps_answered', 0)} gap(s) but the "
                        f"re-investigation failed — pass-1 findings kept")
                else:
                    _extra = ""
                    if _fbmeta.get("playbook_redirect"):
                        _extra += (" · playbook redirected: "
                                   + ", ".join(_fbmeta["playbook_redirect"]
                                               .values()))
                    if _fbmeta.get("suggested_classification"):
                        _extra += (f" · deep-dive suggests classification "
                                   f"**{_fbmeta['suggested_classification']}**"
                                   f" (analyst to review)")
                    wf_md.append(
                        f"- Feedback loop (pass {_fbmeta.get('passes', 1)}): "
                        f"gaps `{gap_ids}` → triage deep-dive answered "
                        f"**{_fbmeta.get('gaps_answered', 0)}** → "
                        f"investigation re-ran with the supplement{_extra}")
            if inv.get("status") == "failed":
                try:
                    (SOC_DB_DIR / "last_investigation_error.json").write_text(
                        _json.dumps(inv, indent=2, default=str), encoding="utf-8")
                except Exception:
                    pass
                _ie = str(inv.get("error") or "")
                _cold = ("503" in _ie or "unavailable" in _ie.lower())
                wf_md.append(f"- Investigation: failed — "
                             f"`{_ie[:150]}` "
                             + ("**(the LLM endpoint was asleep — wait a few "
                                "minutes for it to boot, then re-run)** "
                                if _cold else "")
                             + f"(full details: `soc_db/last_investigation_error.json`)")
                bset("investigation", status="failed",
                     think=f"{str(inv.get('error') or '')[:150]}",
                     output=f"Investigation failed:\n\n```\n"
                            f"{str(inv.get('error') or '')[:800]}\n```")
                inv = None
            else:
                bset("investigation", status="done",
                     think=f"Complete — {inv.get('incident_folder')}",
                     output=inv.get("narrative_report") or inv.get("summary")
                            or "Investigation completed.")
                wf_md.append(f"- Investigation: {inv.get('status')} — "
                             f"folder `{inv.get('incident_folder')}`")
                try:
                    rec = _wfm.build_post_investigation_record(
                        inv, ticket, title, run_stamp=run["run_id"])
                    for _attempt in (1, 2, 3):
                        try:
                            _wfm.pipeline_insert("post_investigation", rec)
                            break
                        except Exception:
                            if _attempt == 3:
                                raise
                            time.sleep(0.8)   # ride out a brief sqlite lock
                    run["chroma_queue"].append(("post_investigation", rec))
                    wf_md.append("- Findings recorded — **Pipeline DB** "
                                 "tab → *Post-Investigation*")
                except Exception as exc:
                    # Never swallow silently — a missing findings record looks
                    # like "no output" to the analyst.
                    bset("investigation",
                         think=f"findings record insert failed: {str(exc)[:120]}")
                    wf_md.append(f"- Findings record insert failed: "
                                 f"`{str(exc)[:150]}`")
        else:
            bset("investigation", status="skipped",
                 think=f"Skipped — classification {cls} below routing threshold",
                 output=f"Investigation skipped: classification **{cls}** is "
                        "below the routing threshold (critical/high/medium).")
            wf_md.append(f"- Investigation: skipped "
                         f"(classification {cls} below routing threshold)")

        # ── HITL gate: pause before reporting until the analyst approves ────
        # The Event is pre-set in non-manual mode (see spawn site), so this is a
        # no-op for the normal auto-chain and only blocks when manual review is on.
        gate = run.get("gate_report")
        if gate is not None and not gate.is_set():
            run["awaiting"] = "report"
            bset("reporting", status="queued", progress=0,
                 think="Investigation complete — awaiting analyst approval to "
                       "generate the report…")
            gate.wait()
            run["awaiting"] = None
            bset("reporting", think="Approved by analyst — starting report")

        # ── Stage 3: Reporting ─────────────────────────────────────────────
        bset("reporting", status="running", think="Reporting started", progress=10)
        try:
            tid = _wfm.handoff_to_reporting(tri, incident, inv)
            bset("reporting", think="Triage + investigation context handed over",
                 progress=35)
            rep = _wfm.run_reporting(
                tid, run_stamp=run["run_id"],
                line_cb=lambda ln: bset("reporting", think=ln)
                if ln.strip() else None)
        except Exception as exc:
            rep = {"status": "failed", "error": str(exc)}
        if rep.get("status") == "failed":
            try:
                (SOC_DB_DIR / "last_reporting_error.json").write_text(
                    _json.dumps(rep, indent=2, default=str), encoding="utf-8")
            except Exception:
                pass
            bset("reporting", status="failed",
                 think=f"{str(rep.get('error') or '')[:150]}",
                 output=f"Reporting failed:\n\n```\n"
                        f"{str(rep.get('error') or '')[:800]}\n```")
            wf_md.append(f"- Reporting: failed — "
                         f"`{str(rep.get('error') or '')[:150]}` "
                         f"(full details: `soc_db/last_reporting_error.json`)")
        else:
            rec = {
                "id": f"final_{unc}@{run['run_id']}",
                "incident_id": inc_id, "ticket_unc": unc,
                "title": f"[FINAL] {title}", "severity": cls,
                "summary": str(rep.get("summary")
                               or "Incident report generated.")[:500],
                "report": {k: v for k, v in rep.items()
                           if k not in ("subprocess", "orchestrator_subprocess")}}
            _wfm.pipeline_insert("finalized_report", rec)
            run["chroma_queue"].append(("finalized_report", rec))
            exports = rep.get("document_exports") or {}
            rep_out = [
                f"**Status:** {rep.get('status')} "
                f"(mode: {rep.get('reporting_mode', 'standard')})",
                f"**LLM:** {rep.get('llm_status') or '—'}",
                "",
                f"**Summary:** {rep.get('summary') or '—'}",
                "",
            ]
            for f in ("docx", "pdf"):
                if exports.get(f):
                    rep_out.append(f"- {'' if f == 'docx' else ''} "
                                   f"{f.upper()}: `{exports[f]}`")
            rep_out += ["", "Full report suite: **Pipeline DB** tab "
                            "→ *Finalized Report*."]
            bset("reporting", status="done", think="Report finalised",
                 output="\n".join(rep_out))
            wf_md.append(f"- Reporting: {rep.get('status')} "
                         f"(mode: {rep.get('reporting_mode', 'standard')})")
            for fmt, icon in (("docx", ""), ("pdf", "")):
                if exports.get(fmt):
                    wf_md.append(f"- {icon} {fmt.upper()} report: "
                                 f"`{exports[fmt]}`")
                else:
                    err = str(exports.get(f"{fmt}_error") or exports.get("error")
                              or "no file produced")[:150]
                    wf_md.append(f"- {icon} {fmt.upper()} report: "
                                 f"export failed — `{err}`")
            wf_md.append("")
            wf_md.append("Download the Word/PDF report from the "
                         "**Pipeline DB** tab → *Finalized Report*.")
    except Exception as exc:
        wf_md.append(f"- Workflow worker crashed: `{str(exc)[:200]}`")
        for ag in ("investigation", "reporting"):
            if panels[ag]["status"] in ("running", "queued"):
                bset(ag, status="failed",
                     think=f"worker crash: {str(exc)[:150]}")
    finally:
        # Audit trail row — one NEW record per workflow execution.
        try:
            dur = int(time.time() - run["started_ts"])
            stages = {
                "triage": "cached" if run.get("cached_triage") else "fresh",
                "investigation": panels["investigation"]["status"],
                "reporting": panels["reporting"]["status"],
            }
            _wfm.pipeline_insert("workflow_runs", {
                "id": f"run_{run['run_id']}_{inc_id[:20]}",
                "incident_id": inc_id,
                "title": f"Run {run['started_hms']} — {title}",
                "severity": cls,
                "summary": " · ".join(f"{k}: {v}" for k, v in stages.items())
                           + f" · ticket {unc} · {dur}s",
                "stages": stages, "ticket_unc": unc,
                "duration_seconds": dur})
        except Exception:
            pass
        run["finished_at"] = time.time()
        run["done"] = True
        # Persist the finished run so ANY session can surface the results on
        # its next interaction — even if the browser poll loop died mid-run
        # (exception in a rerun, page refresh, process restart).
        try:
            (SOC_DB_DIR / "last_workflow_result.json").write_text(
                _json.dumps(run, default=str, indent=1), encoding="utf-8")
        except Exception:
            pass

def _normalise_llm_url(url: str) -> str:
    """
    Ensure the base URL is safe for ChatOpenAI.
    - Strips trailing slashes
    - Removes any embedded model path (e.g. /models/fdtn-ai/...)
    - Guarantees the URL ends with /v1
    """
    from urllib.parse import urlparse
    url = url.strip().rstrip("/")
    parsed = urlparse(url)
    # If someone pasted a model-specific URL, reduce it to just host + /v1
    if "/models/" in parsed.path:
        url = f"{parsed.scheme}://{parsed.netloc}/v1"
    elif not parsed.path.endswith("/v1"):
        url = url + "/v1"
    return url

def get_cisco_cfg() -> CiscoLLMConfig:
    """
    Build CiscoLLMConfig from session state at call time so that
    credentials entered in the sidebar are always picked up.
    """
    raw_url = st.session_state.get("cisco_url", "").strip()
    url     = _normalise_llm_url(raw_url) if raw_url else "https://api-inference.huggingface.co/v1"
    return CiscoLLMConfig(
        base_url    = url,
        api_key     = st.session_state.get("cisco_key",   "").strip() or "changeme",
        model       = st.session_state.get("cisco_model", "").strip()
                      or "fdtn-ai/Foundation-Sec-8B-Reasoning",
        temperature = 0.0,
        # 1024 was starving the reasoning model: it spent the whole budget on
        # chain-of-thought and got truncated before emitting the final JSON,
        # which surfaced as "0 IOCs matched" on every triage run.
        max_tokens  = 3072,
        timeout     = 300,
    )

# ══════════════════════════════════════════════════════════════════════════════
# .ENV  — load persisted credentials on every startup
# ══════════════════════════════════════════════════════════════════════════════
ENV_FILE = Path(__file__).parent / ".env"

def _clean(val: str) -> str:
    """Strip whitespace and stray quotes that dotenv sometimes leaves."""
    return val.strip().strip("'\"").strip()

def env_load() -> dict:
    """Read credentials from .env."""
    if DOTENV_OK and ENV_FILE.exists():
        load_dotenv(ENV_FILE, override=True)
    return {
        "host":          _clean(os.environ.get("NW_HOST",       "")),
        "username":      _clean(os.environ.get("NW_USERNAME",   "")),
        "password":      _clean(os.environ.get("NW_PASSWORD",   "")),
        "nw_cert_path":  _clean(os.environ.get("NW_CERT_PATH",  "")),
        # Cisco LLM
        "cisco_url":     _clean(os.environ.get("CISCO_LLM_URL",   "")),
        "cisco_key":     _clean(os.environ.get("CISCO_LLM_KEY",   "")),
        "cisco_model":   _clean(os.environ.get("CISCO_LLM_MODEL", "")),
    }

def env_save(host: str, username: str, password: str) -> None:
    """Persist NetWitness credentials to .env."""
    if not DOTENV_OK:
        return
    try:
        ENV_FILE.touch(exist_ok=True)
        set_key(str(ENV_FILE), "NW_HOST",     host.strip())
        set_key(str(ENV_FILE), "NW_USERNAME", username.strip())
        # base64-encode password to avoid special char issues
        set_key(str(ENV_FILE), "NW_PASSWORD",
                base64.b64encode(password.strip().encode()).decode("ascii"))
    except Exception:
        pass  # .env locked on Windows — credentials still active this session

def nw_cert_env_save(cert_path: str) -> None:
    """Persist the TLS cert path to .env.
    Silently skips if .env is locked (common on Windows) — cert path
    is still stored in session state and works for the current session.
    """
    if not DOTENV_OK:
        return
    try:
        ENV_FILE.touch(exist_ok=True)
        set_key(str(ENV_FILE), "NW_CERT_PATH", cert_path.strip())
    except Exception:
        pass  # .env locked on Windows — session state still holds the path

def nw_cert_env_clear() -> None:
    if not DOTENV_OK:
        return
    try:
        if ENV_FILE.exists():
            set_key(str(ENV_FILE), "NW_CERT_PATH", "")
    except Exception:
        pass

def cisco_env_save(url: str, key: str, model: str) -> None:
    """Persist Cisco LLM credentials to .env."""
    if not DOTENV_OK:
        return
    try:
        ENV_FILE.touch(exist_ok=True)
        set_key(str(ENV_FILE), "CISCO_LLM_URL",   url.strip())
        set_key(str(ENV_FILE), "CISCO_LLM_MODEL", model.strip())
        # base64-encode key to avoid special char issues
        set_key(str(ENV_FILE), "CISCO_LLM_KEY",
                base64.b64encode(key.strip().encode()).decode("ascii"))
    except Exception:
        pass

def env_clear() -> None:
    if not DOTENV_OK:
        return
    try:
        if ENV_FILE.exists():
            set_key(str(ENV_FILE), "NW_HOST",     "")
            set_key(str(ENV_FILE), "NW_USERNAME", "")
            set_key(str(ENV_FILE), "NW_PASSWORD", "")
    except Exception:
        pass

def cisco_env_clear() -> None:
    if not DOTENV_OK:
        return
    try:
        if ENV_FILE.exists():
            set_key(str(ENV_FILE), "CISCO_LLM_URL",   "")
            set_key(str(ENV_FILE), "CISCO_LLM_KEY",   "")
            set_key(str(ENV_FILE), "CISCO_LLM_MODEL", "")
    except Exception:
        pass

def nw_login(host: str, username: str, password: str) -> tuple[bool, str, str]:
    """
    Login with username/password → returns (ok, message, access_token).
    POST /rest/api/auth/userpass with form-encoded credentials.
    Auto-retries with verify=False if cert verification fails.
    """
    if not host.strip():
        return False, "Host URL is empty — enter https://192.168.x.x", ""
    host = host.strip().strip("'\"").strip()  # remove any stray quotes
    if not username.strip():
        return False, "Username is empty.", ""
    if not password.strip():
        return False, "Password is empty.", ""

    def _attempt(verify):
        return requests.post(
            f"{host.rstrip('/')}/rest/api/auth/userpass",
            data={"username": username, "password": password},
            headers={
                "Content-Type": "application/x-www-form-urlencoded; charset=ISO-8859-1",
                "Accept":       "application/json;charset=UTF-8",
            },
            timeout=15,
            verify=verify,
        )

    try:
        verify = nw_tls_verify()
        try:
            r = _attempt(verify)
        except requests.exceptions.SSLError:
            # Cert didn't work — clear it and retry without verification
            st.session_state.nw_cert_path = ""
            r = _attempt(False)

        if r.status_code == 200:
            data  = r.json()
            token = data.get("accessToken") or data.get("access_token") or ""
            if token:
                return True, "NetWitness connected", token
            return False, f"Login OK but no token in response: {str(data)[:100]}", ""
        else:
            return False, f"HTTP {r.status_code}: {r.text[:200]}", ""
    except requests.exceptions.ConnectionError as e:
        return False, f"Cannot reach {host} — check VPN/network: {str(e)[:120]}", ""
    except requests.exceptions.Timeout:
        return False, "Login timed out — check GP VPN is connected.", ""
    except Exception as e:
        return False, f"Login error: {e}", ""

# ══════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Security Dashboard",
    page_icon=None,
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════════════════════════
# CSS
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Share+Tech+Mono&display=swap');

:root {
  /* Aegis design system — re-skin of the SOC dashboard */
  --bg:      #060b13;
  --nav:     #07101b;
  --bg1:     #0e1929;   /* card surface */
  --bg2:     #101d30;   /* raised card */
  --card:    #0e1929;
  --card2:   #101d30;
  --border:  #223149;
  --line:    #223149;
  --accent:  #6f7cff;   /* Aegis blue */
  --blue:    #6f7cff;
  --cyan:    #36c5d3;
  --purple:  #a67af4;
  --green:   #43d28c;
  --warn:    #f4bc5f;
  --danger:  #ff6e7c;
  --orange:  #f4bc5f;
  --muted:   #8b9bb2;   /* subdued text */
  --sub:     #8b9bb2;
  --faint:   #61738d;
  --text:    #f3f6fb;
  --r:       14px;
  --mono:    'Share Tech Mono', monospace;
  --sans:    'Inter', sans-serif;
}

/* Anchor the rem scale to the visitor's browser/OS default font size
   (typically 16px) rather than a hard 14px. 87.5% × 16px == the original 14px
   baseline, so default users see no change — but anyone who raises their
   browser/OS default font size now gets a proportionally larger UI, since
   virtually all our type & spacing is already expressed in rem. (Browser zoom
   already scaled the px bits; this adds OS/browser font-size preference too.) */
html { font-size: 87.5%; }
html, body, [class*="css"] {
  font-family: var(--sans);
  background: var(--bg);
  color: var(--text);
  line-height: 1.6;
}
/* was 14px — now 1rem so it tracks the html anchor above and scales with prefs */
body, [class*="css"] { font-size: 1rem; }
.main, .stApp {
  background:
    radial-gradient(circle at 64% -20%, #17233b 0, transparent 33%),
    var(--bg);
}
.main { padding-top: 0.5rem; }

/* ── Sidebar ── */
section[data-testid="stSidebar"] > div:first-child {
  background: linear-gradient(180deg, #040810 0%, #06101A 100%);
  border-right: 1px solid #0E1E30;
  padding-top: 1rem;
}

/* ── Headings (Aegis: white, bold, Inter) ── */
h1,h2,h3 {
  font-family: var(--sans);
  font-weight: 700;
  color: var(--text);
  letter-spacing: -0.2px;
  margin-bottom: 0.4rem;
}
h1 { font-size: 1.55rem !important; letter-spacing: -0.4px; }
h2 { font-size: 1.1rem !important; }
h3 { font-size: 0.95rem !important; }

/* ── Metric cards ── */
[data-testid="metric-container"] {
  background: linear-gradient(135deg, #070D18 0%, #0A1628 100%);
  border: 1px solid var(--border);
  border-radius: 12px;
  padding: 18px 20px;
  box-shadow: 0 2px 16px rgba(0,0,0,0.4);
  transition: all 0.2s;
}
[data-testid="metric-container"]:hover {
  border-color: #1E4060;
  transform: translateY(-2px);
  box-shadow: 0 4px 24px rgba(0,212,255,0.08);
}
[data-testid="metric-container"] label {
  color: var(--muted);
  font-size: 0.72rem;
  font-weight: 600;
  letter-spacing: 0.5px;
  text-transform: uppercase;
  font-family: var(--sans);
}
[data-testid="metric-container"] [data-testid="stMetricValue"] {
  color: var(--text);
  font-family: var(--sans);
  font-size: 2rem;
  font-weight: 700;
  letter-spacing: -0.5px;
}

/* ── Buttons ── */
.stButton > button {
  background: linear-gradient(90deg, #052030, #083050);
  color: var(--accent);
  border: 1px solid #0E4A6A;
  border-radius: 8px;
  font-family: var(--sans);
  font-size: 0.78rem;
  font-weight: 500;
  padding: 8px 16px;
  transition: all 0.15s;
}
.stButton > button:hover {
  background: linear-gradient(90deg, #083050, #0E4A6A);
  box-shadow: 0 0 18px rgba(0,212,255,0.2);
  color: #fff;
  border-color: var(--accent);
  transform: translateY(-1px);
}
.stButton > button:active { transform: translateY(0); }

/* ── Inputs ── */
.stTextInput > div > div > input,
.stTextArea textarea {
  background: #060E1A !important;
  border: 1px solid #1A3050 !important;
  color: var(--text) !important;
  border-radius: 8px;
  font-family: var(--sans);
  font-size: 0.88rem;
  padding: 10px 14px !important;
}
.stTextInput > div > div > input:focus,
.stTextArea textarea:focus {
  border-color: var(--accent) !important;
  box-shadow: 0 0 0 3px rgba(0,212,255,0.1) !important;
}
.stTextInput label, .stTextArea label, .stSelectbox label {
  color: var(--text) !important;
  font-size: 0.82rem !important;
  font-weight: 500 !important;
  margin-bottom: 4px !important;
}
.stSelectbox > div > div {
  background: #060E1A !important;
  border: 1px solid #1A3050 !important;
  color: var(--text) !important;
  border-radius: 8px;
}

/* ── Tabs ── */
.stTabs [data-baseweb="tab-list"] {
  background: transparent;
  border-bottom: 1px solid var(--border);
  gap: 4px;
}
.stTabs [data-baseweb="tab"] {
  font-family: var(--sans);
  font-size: 0.82rem;
  font-weight: 500;
  color: var(--muted);
  border: none;
  padding: 10px 20px;
  background: transparent;
  border-radius: 8px 8px 0 0;
  transition: all 0.15s;
}
.stTabs [data-baseweb="tab"]:hover {
  color: var(--text);
  background: rgba(0,212,255,0.04);
}
.stTabs [aria-selected="true"] {
  color: var(--accent) !important;
  border-bottom: 2px solid var(--accent) !important;
  background: rgba(0,212,255,0.06) !important;
  font-weight: 600 !important;
}

/* ── Cards ── */
.card {
  background: linear-gradient(135deg, #070D18 0%, #08111E 100%);
  border: 1px solid var(--border);
  border-radius: 12px;
  padding: 16px 20px;
  margin: 6px 0;
  transition: all 0.15s;
}
.card:hover {
  border-color: #1E4060;
  transform: translateX(3px);
  box-shadow: 0 2px 16px rgba(0,0,0,0.3);
}
.card-critical { border-left: 4px solid var(--danger) !important; }
.card-high     { border-left: 4px solid var(--orange) !important; }
.card-medium   { border-left: 4px solid var(--warn)   !important; }
.card-low      { border-left: 4px solid var(--green)  !important; }

/* ── Badges ── */
.badge {
  padding: 3px 10px;
  border-radius: 20px;
  font-size: 0.7rem;
  font-weight: 600;
  display: inline-block;
  font-family: var(--sans);
}
.badge-critical { background:#321b25; color:#ff99a3; border:1px solid #713744; }
.badge-high     { background:#2b221a; color:#f3c679; border:1px solid #684c2a; }
.badge-medium   { background:#241a34; color:#c9a6f7; border:1px solid #5b3f82; }
.badge-low      { background:#122b21; color:#7fe0ac; border:1px solid #2a6146; }
.badge-info     { background:#192743; color:#aeb7ff; border:1px solid #3b4c81; }

/* ── Chat bubbles ── */
.bubble-user {
  background: #080F1A;
  border-left: 3px solid #005C8A;
  padding: 12px 16px;
  border-radius: 0 10px 10px 0;
  margin: 8px 0;
  font-size: 0.9rem;
  line-height: 1.6;
}
.bubble-agent {
  background: #040A12;
  border-left: 3px solid var(--accent);
  padding: 12px 16px;
  border-radius: 0 10px 10px 0;
  margin: 8px 0;
  font-size: 0.9rem;
  line-height: 1.6;
}
.bubble-label {
  font-family: var(--sans);
  font-size: 0.65rem;
  font-weight: 600;
  letter-spacing: 0.5px;
  margin-bottom: 5px;
  color: var(--accent);
  text-transform: uppercase;
}

/* ── Status dots ── */
.dot { display:inline-block; width:8px; height:8px; border-radius:50%; margin-right:7px; vertical-align:middle; }
.dot-green  { background:var(--green);  box-shadow:0 0 8px var(--green);  animation:pulse 2s infinite; }
.dot-red    { background:var(--danger); box-shadow:0 0 6px var(--danger); }
.dot-cyan   { background:#00d4ff; box-shadow:0 0 8px #00d4ff; animation:pulse 2s infinite; }
.dot-amber  { background:#ffb700; box-shadow:0 0 8px #ffb700; animation:pulse 2s infinite; }
.dot-yellow { background:var(--warn);   box-shadow:0 0 6px var(--warn);   animation:pulse 1s infinite; }
@keyframes pulse { 0%,100%{opacity:1} 50%{opacity:.35} }

/* ── Section labels ── */
.sec-label {
  font-family: var(--sans);
  font-size: 0.7rem;
  font-weight: 600;
  color: #4A7090;
  text-transform: uppercase;
  letter-spacing: 1px;
  margin: 20px 0 10px;
  padding-bottom: 6px;
  border-bottom: 1px solid #0E1E2E;
}

/* ── Stat mini cards ── */
.stat-mini {
  background: #070D18;
  border: 1px solid var(--border);
  border-radius: 10px;
  padding: 14px 16px;
  text-align: center;
  font-family: var(--sans);
}
.stat-mini .val { font-size: 1.5rem; font-weight: 700; color: var(--accent); }
.stat-mini .lbl { font-size: 0.68rem; font-weight: 500; color: var(--muted); margin-top: 3px; text-transform: uppercase; }

/* ── Tooltips / info boxes ── */
.info-box {
  background: #060E1A;
  border: 1px solid #1A3050;
  border-radius: 10px;
  padding: 14px 18px;
  font-size: 0.82rem;
  color: var(--text);
  line-height: 1.6;
  margin: 8px 0;
}
.info-box .title {
  font-weight: 600;
  color: var(--accent);
  margin-bottom: 6px;
  font-size: 0.85rem;
}

/* ── Expanders ── */
.streamlit-expanderHeader {
  font-family: var(--sans) !important;
  font-size: 0.85rem !important;
  font-weight: 500 !important;
  color: var(--text) !important;
}

/* ── Scrollbar ── */
hr { border-color: #0E1E2E; margin: 1.2rem 0; }
::-webkit-scrollbar { width: 5px; }
::-webkit-scrollbar-track { background: var(--bg); }
::-webkit-scrollbar-thumb { background: #1E3050; border-radius: 3px; }
::-webkit-scrollbar-thumb:hover { background: #2A4060; }

/* ── Sidebar logo / app name ── */
.app-logo {
  text-align: center;
  padding: 10px 0 16px;
  border-bottom: 1px solid #0E1E2E;
  margin-bottom: 4px;
}
.app-logo .name {
  font-size: 1rem;
  font-weight: 700;
  color: var(--accent);
  letter-spacing: 1px;
}
.app-logo .sub {
  font-size: 0.65rem;
  color: var(--muted);
  margin-top: 2px;
}
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE  — seed from .env on first load
# ══════════════════════════════════════════════════════════════════════════════
_env = env_load()
# Decode base64 password
if _env["password"]:
    try:
        _env["password"] = base64.b64decode(_env["password"].encode()).decode("utf-8")
    except Exception:
        pass
# Decode base64 Cisco API key
if _env.get("cisco_key"):
    try:
        _env["cisco_key"] = base64.b64decode(_env["cisco_key"].encode()).decode("utf-8")
    except Exception:
        pass

DEFAULTS = {
    "nw_host":          _env["host"],
    "nw_username":      _env["username"],
    "nw_password":      _env["password"],
    "nw_token":         "",       # refreshed on every login
    "nw_verified":      False,
    "nw_msg":           "",
    "nw_working_ep":    "",       # endpoint that worked
    "nw_working_auth":  {},       # auth headers that worked
    "nw_incidents_path":"/rest/api/incidents",   # ← configurable endpoint path
    "nw_auth_style":    "NetWitness-Token",       # ← configurable auth header style
    "endpoint_scan_results": [],   # ← persisted scanner results
    "nw_cert_path":     _env.get("nw_cert_path", ""),  # ← path to CA/server cert for TLS verification
    "incidents":        [],
    "last_fetch":       None,
    "last_full_fetch":  None,   # ← last time a full (non-incremental) fetch ran
    "last_fetch_mode":  None,   # ← "full" | "incremental", shown in diagnostics
    "chat_history":     [],
    "chat_incident":    None,
    "pending_auto_triage": False,   # set by "Triage" button — auto-runs the pipeline
    "jump_to_ask_tab":     False,   # set by "Triage" button — switches to Ask a Question tab
    "triage_in_flight":    None,    # {incident_id, attempts} — survives an interrupted
                                    # inline triage run so it can auto-restart
    "chroma_client":    None,
    "chroma_col":       None,
    "search_results":   [],
    "_startup_done":    False,
    # ── File upload ──────────────────────────────────────────
    "uploaded_incident": None,
    "uploaded_filename": "",
    # ── Agent board (thinking + outputs per agent) ───────────
    "agent_board": {
        "triage":        {"status": "idle", "thinking": [], "output": "", "updated": "", "progress": 0},
        "investigation": {"status": "idle", "thinking": [], "output": "", "updated": "", "progress": 0},
        "reporting":     {"status": "idle", "thinking": [], "output": "", "updated": "", "progress": 0},
    },
    "agent_board_sel": None,
    # ── Cisco Foundation LLM ─────────────────────────────────
    "cisco_url":         _env.get("cisco_url",   ""),
    "cisco_key":         _env.get("cisco_key",   ""),
    "cisco_model":       _env.get("cisco_model", ""),
    "cisco_connected":   bool(_env.get("cisco_key", "").strip()),
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

REFRESH_INTERVAL     = 30                      # seconds
INCREMENTAL_OVERLAP  = timedelta(minutes=5)    # clock-skew / indexing-lag buffer
# Full ground-truth resync cadence. At 53k incidents a full resync re-downloads
# EVERYTHING plus per-incident alerts — minutes of churn — so hourly (not every
# 10 min) is the default; override with NW_FULL_RESYNC_MIN. Incremental
# refreshes still run every 30s and stay cheap.
FULL_RESYNC_INTERVAL = timedelta(
    minutes=int(os.environ.get("NW_FULL_RESYNC_MIN", "60") or 60))


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def normalise_sev(inc: dict) -> str:
    raw = str(inc.get("riskScore") or inc.get("severity") or "").upper()
    try:
        s = int(raw)
        return "CRITICAL" if s >= 90 else "HIGH" if s >= 70 else "MEDIUM" if s >= 40 else "LOW"
    except ValueError:
        return raw if raw in ("CRITICAL","HIGH","MEDIUM","LOW") else "LOW"

def gp_is_reachable() -> tuple[bool, str]:
    """
    Check if the NW server is reachable via GlobalProtect VPN.
    Simply tries to open a TCP socket to the host:port.
    Returns (reachable, message).
    """
    import socket
    from urllib.parse import urlparse
    host = st.session_state.nw_host.rstrip("/")
    if not host:
        return False, "No host configured."
    try:
        parsed = urlparse(host)
        hostname = parsed.hostname or host
        port     = parsed.port or (443 if parsed.scheme == "https" else 80)
        sock = socket.create_connection((hostname, port), timeout=5)
        sock.close()
        return True, f"VPN reachable — {hostname}:{port}"
    except OSError:
        return False, (
            "Cannot reach server — GlobalProtect may not be connected.\n"
            "Connect GP VPN and try again."
        )

def nw_headers(include_content_type: bool = False) -> dict:
    """
    Build request headers based on the currently selected auth style.
    include_content_type should only be True for POST/PATCH requests that
    send a JSON body. GET requests must NOT include Content-Type or NW
    returns 400 'Unsupported format was supplied'.
    """
    token = st.session_state.nw_token.strip()
    style = st.session_state.get("nw_auth_style", "NetWitness-Token")

    base = {
        "Accept": "application/json;charset=UTF-8",
    }
    if include_content_type:
        base["Content-Type"] = "application/json;charset=UTF-8"

    if style == "NetWitness-Token":
        base["NetWitness-Token"] = token
    elif style == "Bearer":
        base["Authorization"] = f"Bearer {token}"
    elif style == "Cookie":
        base["Cookie"] = f"access_token={token}"
    elif style == "Both":
        base["Authorization"] = f"Bearer {token}"
        base["Cookie"]        = f"access_token={token}"
    else:
        base["NetWitness-Token"] = token

    return base


def nw_tls_verify():
    """
    Returns the value to pass as requests(verify=...).
    - If a valid cert path is configured AND the file exists → try it.
    - If the cert causes any SSL error → silently fall back to False.
    - Default is False (suppressed via urllib3.disable_warnings above).
    """
    cert_path = st.session_state.get("nw_cert_path", "").strip()
    if cert_path and Path(cert_path).is_file():
        return cert_path
    return False


def nw_incidents_url(host: str) -> str:
    """Build the incidents endpoint URL from the configurable path."""
    path = st.session_state.get("nw_incidents_path", "/rest/api/incidents").strip()
    if not path.startswith("/"):
        path = "/" + path
    return f"{host.rstrip('/')}{path}"

def nw_verify_token() -> tuple[bool, str]:
    host  = st.session_state.nw_host.rstrip("/")
    token = st.session_state.nw_token.strip()
    if not host or not token:
        return False, "Enter both Host URL and token."

    def _attempt(verify):
        return requests.get(
            nw_incidents_url(host),
            headers=nw_headers(),
            params={"pageSize": 1, "pageNumber": 0},
            timeout=15,
            verify=verify,
        )
    try:
        try:
            r = _attempt(nw_tls_verify())
        except requests.exceptions.SSLError:
            st.session_state.nw_cert_path = ""
            r = _attempt(False)

        if r.status_code == 200:
            return True, "NetWitness connected"
        elif r.status_code == 403:
            return False, (
                "Access Denied — add 'integration-server.api.access' "
                "permission to the Administrators role in NW."
            )
        elif r.status_code == 401:
            return False, "Token rejected — login again."
        else:
            return False, f"Unexpected response — HTTP {r.status_code}"
    except requests.exceptions.Timeout:
        return False, "Timed out — check GP VPN is connected."
    except Exception as e:
        return False, f"Cannot reach server: {e}"

def _bounded_get(url: str, *, headers=None, params=None, verify=False,
                 timeout: int = 30, wall_seconds: int = 45):
    """requests.get with a HARD wall-clock cap.

    The `timeout` parameter only bounds the gap between socket reads — a
    server that dribbles a byte every few seconds holds the connection open
    forever (this wedged the whole app for hours). The request runs in a
    daemon thread; if it exceeds wall_seconds in total, TimeoutError is
    raised and the orphaned thread is abandoned to finish/die on its own."""
    import threading as _th
    box: dict = {}

    def _do():
        try:
            box["r"] = requests.get(url, headers=headers, params=params,
                                    timeout=timeout, verify=verify)
        except Exception as exc:
            box["e"] = exc

    t = _th.Thread(target=_do, daemon=True)
    t.start()
    t.join(timeout=wall_seconds)
    if t.is_alive():
        raise TimeoutError(f"request exceeded {wall_seconds}s wall clock "
                           f"(server stalling mid-response)")
    if "e" in box:
        raise box["e"]
    return box["r"]


# ── alerts-response tolerance (NetWitness Respond API shape varies by version) ──
# These pure helpers now live in nw_alerts.py (first slice of the app.py
# modularization). Imported here so every existing call site keeps working
# unchanged. See nw_alerts.py for the docstrings + rationale.
from nw_alerts import (
    _extract_alert_items,
    _alerts_has_more,
    _alerts_error_hint,
    _distill_alerts,
    _merge_alert_digest,
    _alerts_fetch_warning,
)


def nw_fetch_incidents(
    limit: int | None = None, since: str | None = None,
    deadline_seconds: int | None = None,
) -> tuple[bool, list, str]:
    """
    Returns (ok, items, diagnostic_message).
    ok=True even if items is empty (empty just means no incidents in NW right now).
    ok=False means a real error occurred (auth, network, etc).

    limit=None (default) means load every incident the API has — pagination
    keeps going until the API reports hasNext=False. MAX_PAGES is a safety
    valve, not a real-world cap, in case an API bug ever returns hasNext=True
    forever.

    since=None (default) means a full fetch — every incident NetWitness has.
    Pass an ISO8601 cutoff (see incremental_since()) to ask NetWitness for
    only incidents created/updated after that point — used by the periodic
    auto-refresh so it isn't re-fetching + re-enriching the entire incident
    history (and every incident's full alert history) every 30s.
    """
    if not st.session_state.nw_verified or not st.session_state.nw_token:
        return False, [], "Not authenticated."
    host = st.session_state.nw_host.rstrip("/")
    headers = nw_headers()
    all_items = []
    page = 0
    diag = ""
    MAX_PAGES = 1000   # 1000 * pageSize(100) = 100,000 incidents ceiling
    # Wall-clock budget: a stalling NetWitness (slow pages, dribbling
    # responses) must yield PARTIAL results with an honest diagnostic —
    # never an unbounded hang. None -> env NW_FETCH_DEADLINE (default 600s).
    if deadline_seconds is None:
        try:
            deadline_seconds = int(os.environ.get("NW_FETCH_DEADLINE", "600"))
        except ValueError:
            deadline_seconds = 600
    _deadline = time.monotonic() + max(30, deadline_seconds)
    # Include a wide date range by default — some NW versions return 400 without it
    since = since or "2020-01-01T00:00:00.000Z"
    try:
        while True:
            if time.monotonic() > _deadline:
                diag = (f"Fetch deadline ({deadline_seconds}s) reached at "
                        f"page {page} — returning {len(all_items)} incident(s) "
                        f"fetched so far. NetWitness is responding slowly.")
                return True, all_items, diag
            try:
                r = _bounded_get(
                    nw_incidents_url(host),
                    headers=nw_headers(),
                    params={"pageSize": 100, "pageNumber": page, "since": since},
                    timeout=30,
                    wall_seconds=60,
                    verify=False,
                )
            except TimeoutError:
                diag = (f"NetWitness stalled mid-response on page {page} — "
                        f"returning {len(all_items)} incident(s) fetched so far.")
                return True, all_items, diag
            if r.status_code == 200:
                data      = r.json()
                items     = data.get("items", [])
                total_api = data.get("totalItems", "?")
                all_items.extend(items)
                has_next  = data.get("hasNext", False)
                reached_limit = limit is not None and len(all_items) >= limit
                if not has_next or reached_limit or page >= MAX_PAGES:
                    diag = (
                        f"API reports {total_api} total incident(s) — "
                        f"fetched {len(all_items)} across {page+1} page(s)."
                    )
                    break
                page += 1
            elif r.status_code == 401:
                st.session_state.nw_verified = False
                st.session_state.nw_msg = "Token expired — login again."
                return False, [], "401 Unauthorized — token expired or invalid."
            elif r.status_code == 403:
                return False, [], (
                    "403 Forbidden — account lacks integration-server.api.access permission."
                )
            else:
                return False, [], f"HTTP {r.status_code}: {r.text[:200]}"
        # Fetch associated alerts/logs for each incident — in parallel.
        # This used to be a serial for-loop doing one blocking request per
        # incident (up to 15s timeout each), so on a list of N incidents the
        # whole fetch took N * request-time. Since this runs on every
        # auto-refresh and on every rerun after one is due (sending a chat
        # message, uploading a file — any Streamlit interaction reruns the
        # whole script), that serial loop is what made the entire UI look
        # frozen. Fetching concurrently bounds the wall-clock time to
        # roughly one request instead of N.
        clean_path = st.session_state.get("nw_incidents_path", "/rest/api/incidents").strip()
        if clean_path.endswith("/list"):
            clean_path = clean_path[:-5]

        def _fetch_alerts(inc: dict) -> None:
            inc_id = str(inc.get("id") or inc.get("incidentId") or "").strip()
            if not inc_id:
                inc["alerts"] = []
                return
            alerts_url = f"{host.rstrip('/')}{clean_path}/{inc_id}/alerts"
            # Paginate fully — this used to fetch only pageNumber=0, so any
            # incident with more than 100 alerts silently lost the rest.
            collected: list = []
            a_page = 0
            try:
                while a_page < MAX_PAGES:
                    if time.monotonic() > _deadline:
                        inc["alerts_fetch_error"] = "fetch deadline reached"
                        break
                    r_alerts = _bounded_get(
                        alerts_url,
                        headers=headers,
                        params={"pageSize": 100, "pageNumber": a_page},
                        timeout=10,
                        wall_seconds=20,
                        verify=False,
                    )
                    if r_alerts.status_code != 200:
                        # Remember WHY alerts are missing, with enough to act on:
                        # code + the exact URL + a body snippet + a status-specific
                        # hint. The triage/investigation handoff and the UI surface
                        # this instead of silently passing an incident with no
                        # event data. (A live/VPN run now yields an actionable
                        # error rather than a bare code.)
                        try:
                            _body = (r_alerts.text or "")[:200]
                        except Exception:
                            _body = ""
                        inc["alerts_fetch_error"] = f"HTTP {r_alerts.status_code}"
                        inc["alerts_fetch_diag"] = {
                            "code": r_alerts.status_code, "url": alerts_url,
                            "body": _body, "hint": _alerts_error_hint(r_alerts.status_code)}
                        break
                    try:
                        a_data = r_alerts.json()
                    except Exception:
                        a_data = None
                    # Tolerate response-shape variance between NW versions: a 200
                    # that isn't {items,hasNext} used to yield 0 alerts silently.
                    page_items = _extract_alert_items(a_data)
                    collected.extend(page_items)
                    if not _alerts_has_more(a_data, a_page):
                        break
                    a_page += 1
            except Exception as _exc:
                inc["alerts_fetch_error"] = str(_exc)[:120]
                inc["alerts_fetch_diag"] = {"code": None, "url": alerts_url,
                                            "body": "", "hint": "network/exception — check VPN & host reachability"}
            # Tag every alert with its parent incident ID so alerts are
            # traceable back to the incident they came from.
            for a in collected:
                a["incident_id"] = inc_id
            inc["alerts"] = collected
            # Distill the alerts' rich event fields (endpoint host, users, IPs,
            # MACs, alert titles, MITRE) into alertMeta so triage/investigation
            # and the skills see a real host/user — the incident object alone
            # only carries SourceIp/DestinationIp.
            try:
                _merge_alert_digest(inc)
            except Exception:
                pass

        if all_items:
            if time.monotonic() > _deadline:
                diag += " Alert enrichment skipped (fetch deadline reached)."
                for inc in all_items:
                    inc.setdefault("alerts", [])
                    inc["alerts_fetch_error"] = "fetch deadline reached"
            else:
                with ThreadPoolExecutor(max_workers=min(16, len(all_items))) as pool:
                    list(pool.map(_fetch_alerts, all_items))

        return True, all_items, diag
    except requests.exceptions.Timeout:
        return False, [], "Request timed out — check VPN/network."
    except Exception as e:
        return False, [], f"Exception: {e}"

# ══════════════════════════════════════════════════════════════════════════════
# SQLITE  — permanent incident log
# ══════════════════════════════════════════════════════════════════════════════
import sqlite3
import json as _json

# All SQLite databases now live in soc_db/ (moved 2026-07-09)
SOC_DB_DIR = Path(__file__).parent / "soc_db"
SOC_DB_DIR.mkdir(parents=True, exist_ok=True)
DB_FILE = SOC_DB_DIR / "soc_incidents.db"

def db_connect() -> sqlite3.Connection:
    # timeout: concurrent writers (background fetch thread, workflow worker,
    # a second app instance) WAIT for the lock instead of crashing with
    # "database is locked" — which took the whole app down at db_init.
    con = sqlite3.connect(str(DB_FILE), check_same_thread=False, timeout=30)
    con.row_factory = sqlite3.Row
    return con

def db_init() -> None:
    """Create tables if they don't exist yet."""
    with db_connect() as con:
        # WAL: readers never block the writer and vice-versa — essential now
        # that fetch threads, the workflow worker, and the UI all share these
        # files. Persistent per-database setting; safe to re-issue.
        try:
            con.execute("PRAGMA journal_mode=WAL")
            con.execute("PRAGMA synchronous=NORMAL")
        except Exception:
            pass
        con.execute("""
            CREATE TABLE IF NOT EXISTS incidents (
                id          TEXT PRIMARY KEY,
                title       TEXT,
                severity    TEXT,
                status      TEXT,
                assignee    TEXT,
                alert_count INTEGER,
                created     TEXT,
                updated     TEXT,
                raw_json    TEXT,
                first_seen  TEXT,
                last_seen   TEXT
            )
        """)
        con.execute("""
            CREATE TABLE IF NOT EXISTS fetch_log (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                fetched_at  TEXT,
                count       INTEGER
            )
        """)
        con.commit()

def db_upsert_incidents(incidents: list) -> int:
    """Upsert a list of incident dicts — new rows inserted, existing rows updated."""
    now = datetime.now().isoformat(timespec="seconds")
    rows = []
    for inc in incidents:
        inc_id = str(inc.get("id") or inc.get("incidentId") or "").strip()
        if not inc_id:
            continue
        # Store a SLIM raw_json: the alerts arrays are refetched live and made
        # this DB grow to 2.4 GB; alertMeta (small, carries the only incident-
        # level indicators) is kept. In-memory incidents stay untouched.
        slim = {k: v for k, v in inc.items()
                if k not in ("alerts", "journalEntries")}
        if inc.get("alerts"):
            slim["_alerts_stripped"] = len(inc["alerts"])
        rows.append((
            inc_id,
            inc.get("title") or inc.get("name") or "",
            normalise_sev(inc),
            str(inc.get("status") or ""),
            str(inc.get("assignee") or ""),
            int(inc.get("alertCount") or inc.get("numAlerts") or 0),
            str(inc.get("created") or inc.get("createdDate") or "")[:19],
            str(inc.get("updated") or inc.get("lastUpdated") or "")[:19],
            _json.dumps(slim),
            now,   # first_seen  — INSERT only
            now,   # last_seen   — always updated
        ))
    if not rows:
        return 0
    # Chunked commits: one giant transaction over 50k+ rows held the write
    # lock for the whole upsert (the 2.4 GB DB made that minutes) — any
    # other connection in that window died with "database is locked".
    CHUNK = 2000
    with db_connect() as con:
        for _i in range(0, len(rows), CHUNK):
            _chunk = rows[_i:_i + CHUNK]
            con.executemany("""
            INSERT INTO incidents
                (id, title, severity, status, assignee, alert_count,
                 created, updated, raw_json, first_seen, last_seen)
            VALUES (?,?,?,?,?,?,?,?,?,?,?)
            ON CONFLICT(id) DO UPDATE SET
                title       = excluded.title,
                severity    = excluded.severity,
                status      = excluded.status,
                assignee    = excluded.assignee,
                alert_count = excluded.alert_count,
                updated     = excluded.updated,
                raw_json    = excluded.raw_json,
                last_seen   = excluded.last_seen
            """, _chunk)
            con.commit()   # release the write lock between chunks
        con.execute(
            "INSERT INTO fetch_log (fetched_at, count) VALUES (?,?)",
            (now, len(rows)),
        )
        con.commit()
    return len(rows)

def db_load_incidents(
    severity: str = "ALL",
    status:   str = "ALL",
    search:   str = "",
    limit:    int = 500,
) -> list[dict]:
    """Query incidents from SQLite with optional filters."""
    clauses, params = [], []
    if severity != "ALL":
        clauses.append("severity = ?"); params.append(severity)
    if status != "ALL":
        clauses.append("status = ?");   params.append(status)
    if search.strip():
        clauses.append("(title LIKE ? OR assignee LIKE ? OR id LIKE ?)")
        params += [f"%{search}%", f"%{search}%", f"%{search}%"]
    where = ("WHERE " + " AND ".join(clauses)) if clauses else ""
    with db_connect() as con:
        rows = con.execute(
            f"SELECT * FROM incidents {where} ORDER BY last_seen DESC LIMIT ?",
            params + [limit],
        ).fetchall()
    return [dict(r) for r in rows]

def db_stats() -> dict:
    """Return aggregate stats from the permanent log."""
    with db_connect() as con:
        total   = con.execute("SELECT COUNT(*) FROM incidents").fetchone()[0]
        by_sev  = dict(con.execute(
            "SELECT severity, COUNT(*) FROM incidents GROUP BY severity"
        ).fetchall())
        by_stat = dict(con.execute(
            "SELECT status, COUNT(*) FROM incidents GROUP BY status"
        ).fetchall())
        fetches = con.execute("SELECT COUNT(*) FROM fetch_log").fetchone()[0]
        last_f  = con.execute(
            "SELECT fetched_at FROM fetch_log ORDER BY id DESC LIMIT 1"
        ).fetchone()
    return {
        "total":   total,
        "by_sev":  by_sev,
        "by_stat": by_stat,
        "fetches": fetches,
        "last_fetch": last_f[0] if last_f else "—",
    }

def db_export_csv() -> str:
    """Return all incidents as a CSV string for download."""
    import csv, io
    rows = db_load_incidents(limit=100_000)
    if not rows:
        return ""
    buf = io.StringIO()
    w   = csv.DictWriter(buf, fieldnames=[
        "id","title","severity","status","assignee",
        "alert_count","created","updated","first_seen","last_seen",
    ])
    w.writeheader()
    for r in rows:
        w.writerow({k: r.get(k,"") for k in w.fieldnames})
    return buf.getvalue()

# ══════════════════════════════════════════════════════════════════════════════
# STARTUP — runs once per session, silently connects & fetches
# ══════════════════════════════════════════════════════════════════════════════

# Initialise DB on every startup (no-op if tables already exist)
db_init()
# ══════════════════════════════════════════════════════════════════════════════
# PIPELINE DATABASE — per-stage SQLite + ChromaDB, fully inline (no 2nd app)
# ══════════════════════════════════════════════════════════════════════════════
PIPELINE_STAGES = [
    "alerts_to_triage",
    "post_triage_investigate",
    "post_triage_no_investigate",
    "post_investigation",
    "initial_ticket",
    "pending_ticket_report",
    "finalized_report",
    "workflow_runs",
]
PIPELINE_LABELS = {
    "alerts_to_triage":           "Alerts to Triage",
    "post_triage_investigate":    "Post-Triage · Needs Investigation",
    "post_triage_no_investigate": "Post-Triage · No Investigation Needed",
    "post_investigation":         "Post-Investigation · Findings",
    "initial_ticket":             "Initial Ticket Generation",
    "pending_ticket_report":      "Pending Ticket / Report Generation",
    "finalized_report":           "Finalized Report",
    "workflow_runs":              "Workflow Runs (Audit)",
}
PIPELINE_ICONS = {
    "alerts_to_triage":           "",
    "post_triage_investigate":    "",
    "post_triage_no_investigate": "",
    "post_investigation":         "",
    "initial_ticket":             "",
    "pending_ticket_report":      "",
    "finalized_report":           "",
    "workflow_runs":              "",
}
PIPELINE_COLORS = {
    "alerts_to_triage":           "#FF3B3B",
    "post_triage_investigate":    "#FF7700",
    "post_triage_no_investigate": "#0AF0A0",
    "post_investigation":         "#2DD4BF",
    "initial_ticket":             "#00D4FF",
    "pending_ticket_report":      "#FFB700",
    "finalized_report":           "#A78BFA",
    "workflow_runs":              "#8B9DC3",
}
PIPELINE_DB_FILE = SOC_DB_DIR / "soc_pipeline.db"

def _pl_con():
    # Busy-timeout so UI reads and the background worker's writes never
    # collide into silent 'database is locked' failures.
    con = sqlite3.connect(str(PIPELINE_DB_FILE), check_same_thread=False,
                          timeout=15)
    con.row_factory = sqlite3.Row
    return con

def pipeline_db_init():
    with _pl_con() as c:
        try:
            c.execute("PRAGMA journal_mode=WAL")
            c.execute("PRAGMA synchronous=NORMAL")
        except Exception:
            pass
        for s in PIPELINE_STAGES:
            c.execute(f"""CREATE TABLE IF NOT EXISTS {s} (
                id TEXT PRIMARY KEY, incident_id TEXT, title TEXT,
                severity TEXT, stage TEXT, created_at TEXT,
                summary TEXT, raw_json TEXT)""")
        c.commit()

def pipeline_count(stage):
    try:
        with _pl_con() as c:
            return c.execute(f"SELECT COUNT(*) FROM {stage}").fetchone()[0]
    except Exception:
        return 0

def pipeline_load(stage, limit=300):
    try:
        with _pl_con() as c:
            rows = c.execute(
                f"SELECT * FROM {stage} ORDER BY created_at DESC LIMIT ?", (limit,)
            ).fetchall()
        return [dict(r) for r in rows]
    except Exception:
        return []

def pipeline_insert(stage, record):
    import uuid as _uuid
    rec_id = str(record.get("id") or record.get("unc") or _uuid.uuid4())[:64]
    now = datetime.now().isoformat(timespec="seconds")
    try:
        with _pl_con() as c:
            # Same-id inserts REPLACE the row (re-running the same incident
            # reuses its ticket ids), which used to look like "nothing
            # happened". Track a run counter and stamp the summary so a
            # refreshed record is unmistakably new.
            runs = 1
            prev = None
            try:
                prev = c.execute(f"SELECT raw_json FROM {stage} WHERE id=?",
                                 (rec_id,)).fetchone()
                if prev:
                    runs = int((_json.loads(prev[0] or "{}"))
                               .get("workflow_runs_count") or 1) + 1
            except Exception:
                runs = 2 if prev else 1
            record = dict(record)
            record["workflow_runs_count"] = runs
            summary = str(record.get("summary") or record.get("description") or "")
            if runs > 1:
                summary = f"[run {runs} · {now[11:19]}] {summary}"
            c.execute(
                f"INSERT OR REPLACE INTO {stage} "
                "(id,incident_id,title,severity,stage,created_at,summary,raw_json) "
                "VALUES (?,?,?,?,?,?,?,?)",
                (rec_id,
                 str(record.get("incident_id") or record.get("incidentId") or ""),
                 str(record.get("title") or record.get("name") or ""),
                 str(record.get("severity") or record.get("classification") or ""),
                 stage, now,
                 summary[:500],
                 _json.dumps(record)))
            c.commit()
    except Exception:
        pass
    return rec_id


def pipeline_last_write(stage):
    """Most recent created_at in a stage — shown on the stage card so it's
    obvious when the stage last changed (counts alone hide REPLACEs)."""
    try:
        with _pl_con() as c:
            v = c.execute(f"SELECT MAX(created_at) FROM {stage}").fetchone()[0]
        return str(v)[5:16].replace("T", " ") if v else "—"
    except Exception:
        return "—"

def _pl_chroma_col(stage):
    if not CHROMA_OK or not st.session_state.get("chroma_client"):
        return None
    try:
        return st.session_state.chroma_client.get_or_create_collection(
            name=f"pipeline_{stage}", metadata={"hnsw:space": "cosine"})
    except Exception:
        return None

def pipeline_chroma_insert(stage, record):
    col = _pl_chroma_col(stage)
    if col is None:
        return
    import uuid as _uuid
    rec_id = str(record.get("id") or record.get("unc") or _uuid.uuid4())[:64]
    title = str(record.get("title") or "")
    summary = str(record.get("summary") or record.get("description") or "")
    doc = f"{title}\n{summary}".strip() or "no content"
    try:
        col.upsert(documents=[doc], ids=[rec_id],
                   metadatas=[{"stage": stage,
                                "severity": str(record.get("severity") or ""),
                                "created": datetime.now().isoformat(timespec="seconds")}])
    except Exception:
        pass

def pipeline_chroma_count(stage):
    col = _pl_chroma_col(stage)
    if col is None:
        return 0
    try:
        return col.count()
    except Exception:
        return 0

def pipeline_chroma_search(stage, query, n=5):
    col = _pl_chroma_col(stage)
    if col is None or not query.strip():
        return []
    try:
        total = col.count()
        if total == 0:
            return []
        res = col.query(query_texts=[query], n_results=min(n, total))
        return [{"id": res["ids"][0][i], "doc": d,
                 "score": round((1 - res["distances"][0][i]) * 100, 1),
                 "meta": res["metadatas"][0][i]}
                for i, d in enumerate(res["documents"][0])]
    except Exception as e:
        return [{"id": "err", "doc": str(e), "score": 0, "meta": {}}]

def pipeline_chroma_all(stage):
    col = _pl_chroma_col(stage)
    if col is None:
        return []
    try:
        data = col.get(include=["documents", "metadatas"])
        return [{"id": data["ids"][i],
                 "doc": (data["documents"] or [""])[i],
                 "meta": (data["metadatas"] or [{}])[i]}
                for i in range(len(data["ids"]))]
    except Exception:
        return []

def pipeline_insert_full(stage, record):
    """Insert into SQLite and ChromaDB (if connected)."""
    rec_id = pipeline_insert(stage, record)
    pipeline_chroma_insert(stage, record)
    return rec_id

def pipeline_delete(stage, rec_id):
    try:
        with _pl_con() as c:
            c.execute(f"DELETE FROM {stage} WHERE id=?", (rec_id,))
            c.commit()
    except Exception:
        pass
    col = _pl_chroma_col(stage)
    if col:
        try:
            col.delete(ids=[rec_id])
        except Exception:
            pass

pipeline_db_init()

# ──────────────────────────────────────────────────────────────────────────────
# EXPORT HELPERS  — generate downloadable bytes from a pipeline record
# ──────────────────────────────────────────────────────────────────────────────

def _make_csv_bytes(row: dict) -> bytes:
    """
    Build a CSV file from a pipeline SQLite row.
    Flattens the raw_json payload if present so the sheet has all fields.
    Returns UTF-8 bytes suitable for st.download_button.
    """
    import csv, io
    # Start with the fixed SQLite columns
    base = {
        "id":          row.get("id", ""),
        "incident_id": row.get("incident_id", ""),
        "title":       row.get("title", ""),
        "severity":    row.get("severity", ""),
        "stage":       row.get("stage", ""),
        "created_at":  row.get("created_at", ""),
        "summary":     row.get("summary", ""),
    }
    # Merge in any extra keys from raw_json (flat scalars only)
    try:
        extra = _json.loads(row.get("raw_json") or "{}")
        if isinstance(extra, dict):
            for k, v in extra.items():
                if k not in base and not isinstance(v, (dict, list)):
                    base[k] = str(v)
    except Exception:
        pass
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=list(base.keys()))
    writer.writeheader()
    writer.writerow(base)
    return buf.getvalue().encode("utf-8")


def _make_docx_bytes(row: dict) -> bytes:
    """
    Build a professional SOC Initial Ticket .docx from a pipeline record.
    Uses python-docx (pip install python-docx). Falls back to a plain-text
    .txt wrapped as bytes if the library is unavailable.
    """
    import io
    title      = row.get("title", "Untitled Ticket")
    inc_id     = row.get("incident_id", "—")
    severity   = row.get("severity", "—")
    stage      = row.get("stage", "initial_ticket")
    created_at = row.get("created_at", "—")
    summary    = row.get("summary", "No summary available.")

    # Try to parse raw_json for richer data
    extra = {}
    try:
        extra = _json.loads(row.get("raw_json") or "{}")
        if not isinstance(extra, dict):
            extra = {}
    except Exception:
        pass

    try:
        from docx import Document as _DocxDocument
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = _DocxDocument()

        # ── Page margins ─────────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        # ── Title ────────────────────────────────────────────
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title_para.runs:
            run.font.color.rgb = RGBColor(0x00, 0xD4, 0xFF)
            run.font.size = Pt(16)

        doc.add_paragraph()

        # ── Metadata table ───────────────────────────────────
        meta_fields = [
            ("Ticket ID",     row.get("id", "—")),
            ("Incident ID",   inc_id),
            ("Severity",      severity),
            ("Stage",         stage),
            ("Generated At",  created_at),
            ("Assignee",      str(extra.get("assignee") or "Unassigned")),
            ("Status",        str(extra.get("status") or "NEW")),
        ]
        tbl = doc.add_table(rows=len(meta_fields), cols=2)
        tbl.style = "Table Grid"
        for i, (label, value) in enumerate(meta_fields):
            tbl.cell(i, 0).text = label
            tbl.cell(i, 1).text = str(value)
            # Bold the label
            for run in tbl.cell(i, 0).paragraphs[0].runs:
                run.bold = True

        doc.add_paragraph()

        # ── Summary section ───────────────────────────────────
        doc.add_heading("Triage Summary", level=2)
        doc.add_paragraph(summary)

        # ── IOC / Alert section if present ───────────────────
        alert_count = extra.get("alertCount") or extra.get("numAlerts") or extra.get("alert_count")
        if alert_count:
            doc.add_paragraph()
            doc.add_heading("Alert Count", level=2)
            doc.add_paragraph(str(alert_count))

        description = extra.get("description") or extra.get("raw_log") or ""
        if description and description != summary:
            doc.add_paragraph()
            doc.add_heading("Description / Raw Log", level=2)
            # Cap at 3000 chars to keep document manageable
            doc.add_paragraph(str(description)[:3000])

        # ── Footer note ───────────────────────────────────────
        doc.add_paragraph()
        footer_para = doc.add_paragraph(
            f"Auto-generated by SOC Platform · {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in footer_para.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x3A, 0x60, 0x7A)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        # python-docx not installed — return UTF-8 text as fallback
        lines = [
            f"SOC INITIAL TICKET",
            f"==================",
            f"Title:      {title}",
            f"Ticket ID:  {row.get('id', '—')}",
            f"Incident:   {inc_id}",
            f"Severity:   {severity}",
            f"Stage:      {stage}",
            f"Created:    {created_at}",
            f"",
            f"SUMMARY",
            f"-------",
            summary,
            f"",
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        ]
        return "\n".join(lines).encode("utf-8")


# Step 1 — load persisted incidents from SQLite immediately
# so the UI is never empty even before the live fetch completes
if not st.session_state.incidents:
    _db_rows = db_load_incidents(limit=500)
    if _db_rows:
        st.session_state.incidents = [
            _json.loads(r["raw_json"]) for r in _db_rows if r.get("raw_json")
        ]

# Step 2 — silently auto-verify & auto-fetch if .env has credentials.
# Runs once per session (tracked by _startup_done flag) in a BACKGROUND
# thread with a soft wait: a slow or wedged NetWitness (mid-response
# stalls evade per-request timeouts) previously hung every page load
# forever, before a single element rendered. The UI now renders from the
# SQLite cache immediately; live data is adopted whenever the fetch lands.
if not st.session_state._startup_done and _env["host"] and _env["username"] and _env["password"]:
    st.session_state._startup_done = True
    st.session_state.nw_host     = _env["host"]
    st.session_state.nw_username = _env["username"]
    st.session_state.nw_password = _env["password"]

    st.session_state._startup_fetching = True

    def _startup_login_and_fetch() -> None:
        try:
            ok, msg, token = nw_login(_env["host"], _env["username"],
                                      _env["password"])
            if not ok:
                st.session_state.nw_msg = msg
                return
            st.session_state.nw_token    = token
            st.session_state.nw_verified = True
            st.session_state.nw_msg      = msg
            # Incremental startup: when the SQLite cache is fresh (<24h since
            # the last fetch), only pull incidents changed since then instead
            # of re-downloading all 53k+ every session. A stale/empty cache
            # still gets the full fetch.
            since = None
            last_dt = None
            try:
                with db_connect() as _c:
                    _last = _c.execute(
                        "SELECT MAX(fetched_at) FROM fetch_log").fetchone()[0]
                if _last and st.session_state.incidents:
                    last_dt = datetime.fromisoformat(str(_last)[:19])
                    if (datetime.now() - last_dt) < timedelta(hours=24):
                        since = incremental_since(last_dt)
            except Exception:
                since = None
            ok2, items, _diag = nw_fetch_incidents(
                since=since,
                deadline_seconds=int(os.environ.get("NW_FETCH_DEADLINE", "300")))
            if ok2 and items:
                st.session_state.incidents = (
                    items if since is None
                    else _merge_incidents(st.session_state.incidents, items))
                st.session_state.last_fetch      = datetime.now()
                # Incremental startup: schedule the next full resync off the
                # last KNOWN fetch, not off now — ground truth stays bounded.
                st.session_state.last_full_fetch = (
                    datetime.now() if since is None else last_dt)
                st.session_state.last_fetch_mode = (
                    "full" if since is None else "incremental")
                db_upsert_incidents(items)
        except Exception as _exc:
            try:
                st.session_state.nw_msg = f"Startup fetch failed: {_exc}"
            except Exception:
                pass
        finally:
            try:
                st.session_state._startup_fetching = False
            except Exception:
                pass

    import threading as _threading
    from streamlit.runtime.scriptrunner import add_script_run_ctx
    _t = _threading.Thread(target=_startup_login_and_fetch, daemon=True)
    add_script_run_ctx(_t)   # lets the thread read/write this session's state
    _t.start()
    # Soft wait: if NetWitness answers quickly, first render has live data;
    # if it's slow or wedged, render proceeds on cached incidents and the
    # thread adopts results whenever (if ever) the fetch completes.
    _t.join(timeout=int(os.environ.get("NW_STARTUP_WAIT", "20")))
    if _t.is_alive():
        st.session_state.nw_msg = ("NetWitness is responding slowly — "
                                   "showing cached incidents; live data "
                                   "will appear when the fetch completes.")

def incremental_since(last_fetch: datetime) -> str:
    """
    ISO8601 cutoff for an incremental refresh — everything since the last
    successful fetch, minus a 5-minute overlap to tolerate clock skew and
    NetWitness indexing lag. Incidents re-fetched inside that overlap just
    get upserted again by id — harmless, not duplicated.
    """
    cutoff = last_fetch - INCREMENTAL_OVERLAP
    return cutoff.strftime("%Y-%m-%dT%H:%M:%S.000Z")


def _merge_incidents(cached: list, fresh: list) -> list:
    """
    Upsert `fresh` incidents into `cached` by id, keeping everything else
    untouched. Used for incremental refreshes, where `fresh` is only the
    new/updated subset NetWitness returned for the `since` window — not
    the full incident set, so a plain replace would drop everything older.
    """
    by_id = {
        str(inc.get("id") or inc.get("incidentId") or f"_unkeyed_{i}"): inc
        for i, inc in enumerate(cached)
    }
    for inc in fresh:
        key = str(inc.get("id") or inc.get("incidentId") or "").strip()
        if key:
            by_id[key] = inc
    return list(by_id.values())


def maybe_auto_fetch():
    if not st.session_state.nw_verified:
        return
    # Don't stack a second fetch on top of the startup background fetch or a
    # previous auto-refresh that is still running in its worker thread.
    if st.session_state.get("_startup_fetching") or st.session_state.get("_bg_fetching"):
        return
    now  = datetime.now()
    last = st.session_state.last_fetch
    if last is None or (now - last).total_seconds() >= REFRESH_INTERVAL:
        # We don't actually know whether NetWitness's `since` filter matches
        # on incident creation time or last-updated time — the API isn't
        # documented clearly enough to bet on it. If it's creation-time only,
        # a purely incremental refresh would silently miss status/alert
        # changes on older incidents. So: incremental fetches most cycles
        # (cheap — only new/updated incidents come back), but force a full
        # ground-truth resync every FULL_RESYNC_INTERVAL regardless, which
        # caps how stale the cache can ever get at a known worst case.
        last_full  = st.session_state.last_full_fetch
        needs_full = last_full is None or (now - last_full) >= FULL_RESYNC_INTERVAL
        since = None if needs_full else incremental_since(last)

        # Budgets kept from the synchronous era as the worker's wall-clock cap.
        _budget = 45 if since is not None else 120

        # The fetch itself runs OFF the render thread (same proven pattern as
        # the startup fetch: daemon thread + add_script_run_ctx + guard flag).
        # It used to run inline here, so every REFRESH_INTERVAL the next rerun
        # (any click / chat message) blocked the whole UI for up to _budget
        # seconds. Now render returns immediately; results are adopted into
        # session state when the fetch lands and show on the next rerun —
        # exactly how the startup fetch already behaves.
        st.session_state._bg_fetching = True

        def _bg_auto_fetch():
            try:
                ok, items, _diag = nw_fetch_incidents(since=since,
                                                      deadline_seconds=_budget)
                if ok:
                    st.session_state.incidents = (
                        items if since is None
                        else _merge_incidents(st.session_state.incidents, items)
                    )
                    st.session_state.last_fetch      = datetime.now()
                    st.session_state.last_fetch_mode = "full" if since is None else "incremental"
                    if since is None:
                        st.session_state.last_full_fetch = st.session_state.last_fetch
                    db_upsert_incidents(items)   # ← persist every fetch to SQLite
            except Exception:
                pass   # failed auto-refresh = stale data; next due cycle retries
            finally:
                try:
                    st.session_state._bg_fetching = False
                except Exception:
                    pass

        import threading as _threading
        from streamlit.runtime.scriptrunner import add_script_run_ctx
        _t = _threading.Thread(target=_bg_auto_fetch, daemon=True)
        add_script_run_ctx(_t)   # thread may read/write this session's state
        _t.start()

def chroma_connect(path: str = "./chroma_db") -> tuple[bool, str]:
    if not CHROMA_OK:
        return False, "chromadb not installed — run: pip install chromadb"
    try:
        client = chromadb.PersistentClient(path=path)
        col    = client.get_or_create_collection(
            name="soc_incidents", metadata={"hnsw:space": "cosine"}
        )
        st.session_state.chroma_client = client
        st.session_state.chroma_col    = col
        return True, f"Ready — {col.count()} vectors"
    except Exception as e:
        return False, str(e)

def chroma_sync(incidents: list) -> tuple[int, str]:
    col = st.session_state.chroma_col
    if col is None:
        return 0, "Connect ChromaDB first."
    docs, ids, metas = [], [], []
    for inc in incidents:
        inc_id = str(inc.get("id") or inc.get("incidentId") or "").strip()
        if not inc_id:
            continue
        title   = inc.get("title") or inc.get("name") or ""
        summary = inc.get("summary") or inc.get("description") or ""
        docs.append(f"{title}\n{summary}".strip() or "no content")
        ids.append(inc_id)
        metas.append({
            "severity": str(inc.get("riskScore") or inc.get("severity") or ""),
            "status":   str(inc.get("status") or ""),
            "created":  str(inc.get("created") or inc.get("createdDate") or "")[:19],
        })
    if not docs:
        return 0, "No valid incidents to store."
    col.upsert(documents=docs, ids=ids, metadatas=metas)
    return len(docs), f"Synced {len(docs)} incidents."

def chroma_search(query: str, n: int = 5) -> list:
    col = st.session_state.chroma_col
    if col is None or not query.strip() or col.count() == 0:
        return []
    try:
        res = col.query(query_texts=[query], n_results=min(n, col.count()))
        return [{"id": res["ids"][0][i], "text": doc,
                 "score": round((1 - res["distances"][0][i]) * 100, 1),
                 "meta": res["metadatas"][0][i]}
                for i, doc in enumerate(res["documents"][0])]
    except Exception as e:
        return [{"id":"err","text":str(e),"score":0,"meta":{}}]

def chat_respond(user_msg: str, incident: Optional[dict] = None) -> str:
    """
    ── SOC TRIAGE AGENT (LangChain) ─────────────────────────
    Powered by soc_triage_agent.py.
    Trigger words: triage, analyse, analyze, ioc, classify, ticket, investigate
    When triggered with an incident selected, runs the full pipeline:
      Phase 1 → IOC Checklists (Availability / Confidentiality / Integrity)
      Phase 2 → Risk Rating Methodology
      Phase 3 → SOC Classification Template
    Outputs: metakeys_payload + UNC ticket (both queued for downstream agents).
    All other messages fall back to plain SOC analyst Q&A via LangChain.
    ─────────────────────────────────────────────────────────
    """
    return soc_triage_chat_respond(user_msg, incident, llm_config=get_cisco_cfg())


# ══════════════════════════════════════════════════════════════════════════════
# AUTO-FETCH on every render
# ══════════════════════════════════════════════════════════════════════════════
maybe_auto_fetch()


# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════
# ── RBAC role — determined HERE (before the sidebar) so the sidebar can render a
# slimmed view for guests. The developer sign-in + KB-tab gate live after the tabs.
st.session_state.setdefault("user_role", "guest")
_is_dev = st.session_state.user_role == "developer"

with st.sidebar:

    st.markdown(
        '<div class="app-logo">'
        '<div class="name"> Security Dashboard</div>'
        '<div class="sub">Powered by NetWitness</div>'
        '</div>',
        unsafe_allow_html=True,
    )
    # Aegis workspace card (mockup .workspace) — analyst identity at the top
    # of the sidebar. Guarded: the sidebar renders before the main
    # `import ui_components` below, so import locally here.
    try:
        import ui_components as _uisb
        _ws_user = (st.session_state.nw_username or "").strip() or "not signed in"
        st.markdown(_uisb.workspace_card(
            "SOC Workspace", f"Analyst · {_ws_user}", ""),
            unsafe_allow_html=True)
    except Exception:
        pass
    st.caption("Developer mode" if _is_dev else "Guest · read-only view")
    st.markdown("---")

    if _is_dev:
        # ── Connection status card ─────────────────────────────────
        if st.session_state.nw_verified:
            last      = st.session_state.last_fetch
            last_str  = last.strftime("%H:%M:%S") if last else "—"
            elapsed   = int((datetime.now() - last).total_seconds()) if last else 0
            remaining = max(REFRESH_INTERVAL - elapsed, 0)
            pct       = min(elapsed / REFRESH_INTERVAL, 1.0)
            bar_color = "var(--accent)" if pct < 0.8 else "var(--warn)"

            # Check GP status via the existing connection state
            _gp_status = (
                '<span class="dot dot-green"></span>'
                '<span style="color:var(--green);font-size:0.58rem">GP VPN ACTIVE</span>'
                if st.session_state.nw_verified else
                '<span class="dot dot-yellow"></span>'
                '<span style="color:var(--warn);font-size:0.58rem">GP VPN STATUS UNKNOWN</span>'
            )

            # Aegis shift-card look (mockup .shift): card bg + line border + r12
            st.markdown(
                f'<div style="background:#0e182a;border:1px solid #223149;'
                f'border-radius:12px;padding:13px">'
                f'<div style="font-family:var(--mono);font-size:0.68rem">'
                f'<span class="dot dot-green"></span>Connected ✓</div>'
                f'<div style="margin-top:4px">{_gp_status}</div>'
                f'<div style="font-family:var(--mono);font-size:0.58rem;'
                f'color:var(--muted);margin-top:3px">{st.session_state.nw_msg}</div>'
                f'<div style="font-family:var(--mono);font-size:0.56rem;'
                f'color:#1A4A62;margin-top:6px">'
                f'Synced {last_str} &nbsp;·&nbsp; refresh in {remaining}s</div>'
                f'<div style="font-family:var(--mono);font-size:0.55rem;'
                f'color:#2A5A78;margin-top:4px">'
                f'{st.session_state.nw_incidents_path} · {st.session_state.nw_auth_style}</div>'
                f'<div style="background:#060E1A;border-radius:2px;height:3px;'
                f'width:100%;margin-top:7px;overflow:hidden">'
                f'<div style="width:{pct*100:.0f}%;height:100%;border-radius:2px;'
                f'background:{bar_color};transition:width 1s linear"></div>'
                f'</div></div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                '<div style="background:#0A0608;border:1px solid #2A1010;'
                'border-radius:7px;padding:11px 13px;font-family:var(--mono);font-size:0.68rem">'
                '<span class="dot dot-red"></span>Not Connected<br>'
                '<span style="font-size:0.57rem;color:var(--muted)">'
                'Please enter your login details below</span></div>',
                unsafe_allow_html=True,
            )

        # ── NetWitness credentials ─────────────────────────────────
        st.markdown('<div class="sec-label">  Connection</div>', unsafe_allow_html=True)

        # If already auto-connected from .env, show a clean status + update option
        if st.session_state.nw_verified and _env["username"]:
            st.markdown(
                '<div style="background:#10231c;border:1px solid #2a6146;border-radius:10px;'
                'padding:9px 12px;font-family:var(--mono);font-size:0.62rem;margin-bottom:8px">'
                '<span class="dot dot-green"></span>'
                '<strong style="color:var(--green)">AUTO-CONNECTED FROM .ENV</strong><br>'
                '<span style="color:var(--muted);font-size:0.58rem">'
                f'Logged in as <strong>{st.session_state.nw_username}</strong> · '
                'token refreshed on startup.</span>'
                '</div>',
                unsafe_allow_html=True,
            )
            with st.expander("Update credentials"):
                host_in = st.text_input("Host URL", value=st.session_state.nw_host, key="sb_host")
                user_in = st.text_input("Username", value=st.session_state.nw_username, key="sb_user")
                pass_in = st.text_input("Password", value="", type="password",
                                         placeholder="Enter password…", key="sb_pass")
                st.markdown(
                    '<div style="font-family:var(--mono);font-size:0.6rem;'
                    'color:var(--muted);margin:6px 0 2px">— or paste token directly —</div>',
                    unsafe_allow_html=True,
                )
                token_paste_in = st.text_area(
                    "Paste Token",
                    value="",
                    placeholder="Paste a fresh accessToken (eyJ…) to skip re-login",
                    height=80,
                    key="sb_token_paste",
                    label_visibility="collapsed",
                )
                cv, cs, cd = st.columns(3)
                if cv.button("Login", use_container_width=True, key="sb_verify"):
                    st.session_state.nw_host     = host_in
                    st.session_state.nw_username = user_in
                    st.session_state.nw_password = pass_in

                    raw_paste = token_paste_in.strip()
                    if raw_paste:
                        # Use pasted token directly
                        st.session_state.nw_token = raw_paste
                        with st.spinner("Verifying token…"):
                            ok, msg = nw_verify_token()
                        st.session_state.nw_verified = ok
                        st.session_state.nw_msg      = msg
                        if not ok:
                            st.session_state.nw_token = ""
                    else:
                        # Fall back to username/password
                        with st.spinner("Logging in…"):
                            ok, msg, token = nw_login(host_in, user_in, pass_in)
                        st.session_state.nw_verified = ok
                        st.session_state.nw_msg      = msg
                        if ok:
                            st.session_state.nw_token = token

                    if st.session_state.nw_verified:
                        ok2, items, _diag = nw_fetch_incidents()
                        if ok2:
                            st.session_state.incidents  = items
                            st.session_state.last_fetch = datetime.now()
                            db_upsert_incidents(items)
                    st.rerun()
                if cs.button("Save", use_container_width=True, key="sb_save"):
                    if pass_in:
                        env_save(host_in, user_in, pass_in)
                        st.success("Saved to .env")
                    else:
                        st.warning("Enter password to save.")
                if cd.button("✕ Clear", use_container_width=True, key="sb_clear"):
                    env_clear()
                    st.session_state.update(
                        nw_token="", nw_username="", nw_password="",
                        nw_host="", nw_verified=False, nw_msg="",
                        incidents=[], last_fetch=None, _startup_done=False,
                    )
                    st.rerun()
        else:
            # Show last login error if there is one
            _last_err = st.session_state.get("nw_msg", "")
            if _last_err and not st.session_state.nw_verified:
                st.markdown(
                    f'<div style="background:#1A0505;border:1px solid #5A1010;border-radius:5px;'
                    f'padding:8px 11px;font-family:var(--mono);font-size:0.6rem;'
                    f'color:#FF6B6B;margin-bottom:8px">'
                    f'Last error: {_last_err}</div>',
                    unsafe_allow_html=True,
                )
            if DOTENV_OK and not _env["username"]:
                st.markdown(
                    '<div style="background:#0A0800;border:1px solid #3A3000;border-radius:5px;'
                    'padding:8px 11px;font-family:var(--mono);font-size:0.6rem;'
                    'color:#FFB700;margin-bottom:8px">'
                    'No credentials in .env<br>'
                    '<span style="color:var(--muted)">Enter below & click Save</span></div>',
                    unsafe_allow_html=True,
                )
            host_in = st.text_input("Host URL", value=st.session_state.nw_host,
                                     placeholder="https://192.168.x.x")
            if host_in != st.session_state.nw_host:
                st.session_state.nw_host     = host_in
                st.session_state.nw_verified = False

            # ── Login method toggle ────────────────────────────────
            login_method = st.radio(
                "Login method",
                ["Username / Password", "Paste Token"],
                horizontal=True,
                label_visibility="collapsed",
            )

            if login_method == "Username / Password":
                user_in = st.text_input("Username", value=st.session_state.nw_username,
                                         placeholder="admin")
                pass_in = st.text_input("Password", value="", type="password",
                                         placeholder="Enter password…")
                token_in = ""
            else:
                user_in  = st.session_state.nw_username
                pass_in  = ""
                token_in = st.text_area(
                    "NetWitness Token",
                    value="",
                    placeholder="Paste your accessToken (eyJ…) here",
                    height=100,
                    label_visibility="collapsed",
                )
                st.markdown(
                    '<div style="font-family:var(--mono);font-size:0.58rem;'
                    'color:var(--muted);margin-top:-8px;margin-bottom:6px">'
                    'Tokens expire — re-paste when you get a 401</div>',
                    unsafe_allow_html=True,
                )

            cv, cs, cd = st.columns(3)
            if cv.button("Login", use_container_width=True):
                st.session_state.nw_username = user_in
                st.session_state.nw_password = pass_in

                # ── Token-paste path ──────────────────────────────
                if login_method == "Paste Token":
                    raw_token = token_in.strip()
                    if not raw_token:
                        st.error("Paste a token first.")
                    elif not host_in.strip():
                        st.error("Enter the Host URL.")
                    else:
                        st.session_state.nw_token = raw_token
                        with st.spinner("Verifying token…"):
                            ok, msg = nw_verify_token()
                        st.session_state.nw_verified = ok
                        st.session_state.nw_msg      = msg
                        if ok:
                            st.success(f"{msg}")
                            ok2, items, _diag = nw_fetch_incidents()
                            if ok2:
                                st.session_state.incidents  = items
                                st.session_state.last_fetch = datetime.now()
                                db_upsert_incidents(items)
                                st.success(f"Fetched {len(items)} incidents")
                            else:
                                st.warning("Token accepted but no incidents fetched yet.")
                        else:
                            st.session_state.nw_token = ""
                            st.error(f"{msg}")
                    st.rerun()

                # ── Username / Password path ───────────────────────
                else:
                    with st.spinner("Logging in…"):
                        try:
                            ok, msg, token = nw_login(host_in, user_in, pass_in)
                        except Exception as _ve:
                            ok, msg, token = False, f"Exception: {_ve}", ""
                    st.session_state.nw_verified = ok
                    st.session_state.nw_msg      = msg
                    if ok:
                        st.session_state.nw_token = token
                        st.success(f"{msg}")
                        ok2, items, _diag = nw_fetch_incidents()
                        if ok2:
                            st.session_state.incidents  = items
                            st.session_state.last_fetch = datetime.now()
                            db_upsert_incidents(items)
                            st.success(f"Fetched {len(items)} incidents")
                        else:
                            st.warning("Connected but no incidents fetched yet.")
                    else:
                        st.error(f"{msg}")
                    st.rerun()

            if cs.button("Save", use_container_width=True,
                         help="Save to .env — auto-connects on next startup"):
                if host_in and user_in and pass_in:
                    env_save(host_in, user_in, pass_in)
                    st.success("Saved — will auto-connect on next startup")
                else:
                    st.warning("Enter host, username and password first.")

            if cd.button("✕ Clear", use_container_width=True):
                env_clear()
                st.session_state.update(
                    nw_token="", nw_username="", nw_password="",
                    nw_host="", nw_verified=False, nw_msg="",
                    incidents=[], last_fetch=None, _startup_done=False,
                )
                st.rerun()

        # ── TLS Certificate (Option B — verified HTTPS) ─────────────
        st.markdown('<div class="sec-label">  Security Certificate</div>', unsafe_allow_html=True)

        # Auto-clear bad cert if last message was an SSL error
        _last_msg = st.session_state.get("nw_msg", "")
        if "SSL error" in _last_msg and st.session_state.get("nw_cert_path", ""):
            st.warning(
                f"The uploaded cert caused an SSL error — it has been removed. "
                f"Revert to browser export or try again.\n\n`{_last_msg}`"
            )
            st.session_state.nw_cert_path = ""

        _cert_active = st.session_state.get("nw_cert_path", "").strip()
        _cert_valid  = bool(_cert_active) and Path(_cert_active).is_file()

        if _cert_valid:
            st.markdown(
                f'<div style="font-family:var(--mono);font-size:0.62rem;color:var(--green);'
                f'margin-bottom:6px"><span class="dot dot-green"></span>'
                f'Verifying against: {Path(_cert_active).name}</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--warn);'
                'margin-bottom:6px"><span class="dot dot-yellow"></span>'
                'No cert configured — TLS verification skipped (insecure)</div>',
                unsafe_allow_html=True,
            )

        cert_upload = st.file_uploader(
            "Upload server/CA certificate (.pem / .crt)",
            type=["pem", "crt", "cer"],
            key="cert_uploader",
            label_visibility="collapsed",
        )
        if cert_upload is not None:
            certs_dir = Path(__file__).parent / "certs"
            certs_dir.mkdir(exist_ok=True)
            cert_dest = certs_dir / cert_upload.name
            cert_dest.write_bytes(cert_upload.getvalue())
            st.session_state.nw_cert_path = str(cert_dest)
            nw_cert_env_save(str(cert_dest))
            st.success(f"Saved {cert_upload.name} — re-login to apply verified TLS")
            st.rerun()

        cert_path_in = st.text_input(
            "…or enter an existing cert path",
            value=st.session_state.get("nw_cert_path", ""),
            placeholder="/path/to/netwitness-ca.pem",
            key="cert_path_text",
        )
        ccert1, ccert2 = st.columns(2)
        if ccert1.button("Use this path", use_container_width=True):
            if cert_path_in.strip() and Path(cert_path_in.strip()).is_file():
                st.session_state.nw_cert_path = cert_path_in.strip()
                nw_cert_env_save(cert_path_in.strip())
                st.success("Cert path set — re-login to apply")
            else:
                st.error("File not found at that path.")
            st.rerun()
        if ccert2.button("✕ Remove cert", use_container_width=True):
            st.session_state.nw_cert_path = ""
            nw_cert_env_clear()
            st.info("Reverted to verify=False (insecure)")
            st.rerun()

        # ── Foundation LLM (HuggingFace) ──────────────────────────
        st.markdown('<div class="sec-label">  AI Settings</div>', unsafe_allow_html=True)

        # Connection status indicator
        if st.session_state.cisco_connected:
            st.markdown(
                '<div style="background:#041208;border:1px solid #0A3020;border-radius:5px;'
                'padding:8px 11px;font-family:var(--mono);font-size:0.62rem;margin-bottom:8px">'
                '<span class="dot dot-green"></span>'
                '<strong style="color:var(--green)">LLM CONFIGURED</strong><br>'
                f'<span style="color:var(--muted);font-size:0.58rem">'
                f'Model: {st.session_state.cisco_model or "—"}</span>'
                '</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                '<div style="background:#0A0608;border:1px solid #2A1010;border-radius:5px;'
                'padding:8px 11px;font-family:var(--mono);font-size:0.62rem;margin-bottom:8px">'
                '<span class="dot dot-red"></span>'
                '<span style="color:var(--danger)">NOT CONFIGURED</span><br>'
                '<span style="font-size:0.57rem;color:var(--muted)">Enter your HF token below</span>'
                '</div>',
                unsafe_allow_html=True,
            )

        cisco_url_in = st.text_input(
            "Endpoint URL",
            value=st.session_state.cisco_url,
            placeholder="https://api-inference.huggingface.co/v1",
            key="sb_cisco_url",
        )
        cisco_key_in = st.text_input(
            "HuggingFace Token",
            value="",
            type="password",
            placeholder="hf_xxxxxxxxxxxxxxxxxxxx",
            key="sb_cisco_key",
        )
        cisco_model_in = st.text_input(
            "Model name",
            value=st.session_state.cisco_model,
            placeholder="fdtn-ai/Foundation-Sec-8B-Reasoning",
            key="sb_cisco_model",
        )

        cl1, cl2, cl3 = st.columns(3)

        if cl1.button("Apply", use_container_width=True, key="cisco_apply"):
            if not cisco_url_in.strip():
                st.error("Enter the endpoint URL.")
            elif not cisco_key_in.strip():
                st.error("Enter your HF token.")
            else:
                st.session_state.cisco_url       = cisco_url_in.strip()
                st.session_state.cisco_key       = cisco_key_in.strip()
                st.session_state.cisco_model     = (
                    cisco_model_in.strip() or "fdtn-ai/Foundation-Sec-8B-Reasoning"
                )
                st.session_state.cisco_connected = True
                st.success("LLM configured!")
                st.rerun()

        if cl2.button("Save", use_container_width=True, key="cisco_save"):
            if cisco_url_in.strip() and cisco_key_in.strip():
                st.session_state.cisco_url       = cisco_url_in.strip()
                st.session_state.cisco_key       = cisco_key_in.strip()
                st.session_state.cisco_model     = (
                    cisco_model_in.strip() or "fdtn-ai/Foundation-Sec-8B-Reasoning"
                )
                st.session_state.cisco_connected = True
                cisco_env_save(
                    cisco_url_in.strip(),
                    cisco_key_in.strip(),
                    st.session_state.cisco_model,
                )
                st.success("Saved to .env")
                st.rerun()
            else:
                st.warning("Enter URL and HF token first.")

        if cl3.button("✕ Clear", use_container_width=True, key="cisco_clear"):
            st.session_state.cisco_url       = ""
            st.session_state.cisco_key       = ""
            st.session_state.cisco_model     = ""
            st.session_state.cisco_connected = False
            cisco_env_clear()
            st.rerun()

        # ── ChromaDB ───────────────────────────────────────────────
        st.markdown('<div class="sec-label">  Knowledge Base</div>', unsafe_allow_html=True)

        chroma_path = st.text_input("Persist path", value="./chroma_db")
        cc1, cc2 = st.columns(2)

        if cc1.button("Connect", use_container_width=True):
            ok, msg = chroma_connect(chroma_path)
            if ok: st.success(msg)
            else:  st.error(msg)
            st.rerun()

        if st.session_state.chroma_col is not None:
            count = st.session_state.chroma_col.count()
            st.markdown(
                f'<div style="font-family:var(--mono);font-size:0.64rem;'
                f'color:var(--green);margin-top:4px">'
                f'<span class="dot dot-green"></span>{count} vectors stored</div>',
                unsafe_allow_html=True,
            )
            if cc2.button("⬆ Sync", use_container_width=True):
                if st.session_state.incidents:
                    n, msg = chroma_sync(st.session_state.incidents)
                    st.success(msg) if n else st.error(msg)
                else:
                    st.warning("No incidents loaded yet.")

        st.markdown("---")
        st.markdown(
            f'<div style="font-family:var(--mono);font-size:0.54rem;'
            f'color:#1A3A52;text-align:center;line-height:1.9">'
            f'v4 · AUTO-REFRESH {REFRESH_INTERVAL}s<br>'
            f'{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</div>',
            unsafe_allow_html=True,
        )


# ══════════════════════════════════════════════════════════════════════════════
# AUTO-RERUN scheduler
# ══════════════════════════════════════════════════════════════════════════════
if st.session_state.nw_verified and st.session_state.last_fetch:
    elapsed   = (datetime.now() - st.session_state.last_fetch).total_seconds()
    remaining = max(REFRESH_INTERVAL - elapsed, 0)
    if remaining <= 1:
        time.sleep(1)
        st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# HEADER METRICS
# ══════════════════════════════════════════════════════════════════════════════
incidents = st.session_state.incidents
total     = len(incidents)
active    = len([i for i in incidents
                 if str(i.get("status","")).upper()
                 not in ("CLOSED","RESOLVED","REMEDIATED")])
by_sev    = Counter(normalise_sev(i) for i in incidents)
vectors   = st.session_state.chroma_col.count() if st.session_state.chroma_col else 0
last_sync = st.session_state.last_fetch.strftime("%H:%M:%S") if st.session_state.last_fetch else "—"
db_total  = db_stats()["total"]

# Aegis design-system components (ui_components.py) — page title + stat cards
import ui_components as _ui
st.markdown(_ui.COMPONENT_CSS, unsafe_allow_html=True)


def _build_case_findings(inc: dict):
    """Real 'key findings' for an incident: the distilled alert behaviours
    (alertMeta.AlertTitles) + the elevated unified-verdict signals. Returns
    (findings, verdict). Deterministic, guarded — never raises."""
    findings: list[dict] = []
    try:
        from triage_verdict import aggregate_verdict
        v = aggregate_verdict(inc)
    except Exception:
        v = {"available": False, "signals": []}
    _conf = {3: "high", 2: "elevated", 1: "moderate", 0: "low"}
    _kw = [("hta", ""), ("c2", ""), ("command", ""), ("exfil", ""),
           ("autorun", ""), ("credential", ""), ("powershell", "⌘"),
           ("lateral", "↔"), ("ransom", ""), ("phish", ""), ("beacon", "")]
    am = inc.get("alertMeta") or {}
    for t in list(dict.fromkeys(am.get("AlertTitles") or []))[:6]:
        tl = str(t).lower()
        icon = next((e for k, e in _kw if k in tl), "")
        findings.append({"icon": icon, "title": str(t)[:72],
                         "desc": "Observed alert behaviour", "confidence": ""})
    if v.get("available"):
        _si = {"base severity": "", "asset criticality": "",
               "internal IOC correlation": "", "external threat intel": ""}
        for s in sorted(v.get("signals", []), key=lambda s: -s.get("level", 0)):
            if s.get("error") or s.get("absent") or s.get("level", 0) == 0:
                continue
            findings.append({"icon": _si.get(s["name"], "•"),
                             "title": f"{s['name'].title()} — {s['label']}",
                             "desc": s.get("detail", ""),
                             "confidence": _conf.get(s["level"], "")})
    return findings, v


def _build_case_context(inc: dict, sev: str, status: str, alerts, verdict: dict):
    """Aegis 'case context' grid facts from the incident + unified verdict."""
    am = inc.get("alertMeta") or {}
    host = (am.get("Hostname") or [None])[0] or inc.get("hostname") or "—"
    user = (am.get("User") or am.get("AdUser") or [None])[0] or "—"
    ips = len(am.get("SourceIp") or []) + len(am.get("DestinationIp") or [])
    _tone = {"CRITICAL": "crit", "HIGH": "warn", "MEDIUM": "warn", "LOW": "ok"}
    vlvl = verdict.get("level", "—") if verdict.get("available") else "—"
    return [
        ("Severity", sev.title(), _tone.get(sev, "")),
        ("Unified verdict", str(vlvl), _tone.get(vlvl, "")),
        ("Host", str(host)),
        ("User", str(user)),
        ("Status", str(status)),
        ("IOC IPs", str(ips)),
    ]


def _pipeline_worked_ids() -> set:
    """Incident ids already present anywhere in the SOC pipeline DB —
    used to steer the 'next move' pick toward untouched cases. Read-only,
    guarded; a missing/locked DB just means 'nothing worked yet'."""
    ids: set = set()
    try:
        with _pl_con() as c:
            for _s in PIPELINE_STAGES:
                if _s == "workflow_runs":
                    continue
                try:
                    for (i,) in c.execute(
                            f"SELECT DISTINCT incident_id FROM {_s}").fetchall():
                        if i:
                            ids.add(str(i))
                except Exception:
                    continue
    except Exception:
        pass
    return ids


def _pick_next_move(incs: list):
    """Deterministic 'most urgent case to pick up next': unworked incidents
    first, then severity band, then riskScore, then distilled-behaviour
    richness, then newest. Cheap (no skills, no network) — safe to run every
    rerun. Returns (incident, meta) or (None, {})."""
    if not incs:
        return None, {}
    worked = _pipeline_worked_ids()
    _rank = {"CRITICAL": 3, "HIGH": 2, "MEDIUM": 1, "LOW": 0}

    def _risk(i):
        try:
            return int(i.get("riskScore") or 0)
        except (TypeError, ValueError):
            return 0

    def _behaviours(i):
        return len((i.get("alertMeta") or {}).get("AlertTitles") or [])

    def _id(i):
        return str(i.get("id") or i.get("incidentId") or "")

    # Deterministic ordering: id asc → created desc → main urgency key.
    # The stable sorts make created/id pure tie-breaks for the urgency key.
    ordered = sorted(incs, key=_id)
    ordered = sorted(ordered, key=lambda i: str(i.get("created") or ""), reverse=True)
    ordered = sorted(ordered, key=lambda i: (
        1 if _id(i) in worked else 0,
        -_rank.get(normalise_sev(i), 0),
        -_risk(i),
        -_behaviours(i),
    ))
    top = ordered[0]
    return top, {
        "already_worked": _id(top) in worked,
        "sev": normalise_sev(top),
        "risk": _risk(top),
        "behaviours": _behaviours(top),
    }

st.markdown(_ui.page_title(
    "Operations overview",
    f"{active:,} active · {total:,} in session · last sync {last_sync}",
    "Security Operations"), unsafe_allow_html=True)

st.markdown(_ui.stat_row([
    {"label": "Live incidents", "value": f"{total:,}", "sub": "Loaded this session", "tone": "blue"},
    {"label": "Active", "value": f"{active:,}", "sub": "Not closed / resolved", "tone": "blue"},
    {"label": "Critical", "value": by_sev.get("CRITICAL", 0), "sub": "Highest severity", "tone": "red"},
    {"label": "High", "value": by_sev.get("HIGH", 0), "sub": "Elevated severity", "tone": "amber"},
    {"label": "Knowledge vectors", "value": f"{vectors:,}", "sub": f"DB total {db_total:,}", "tone": "green"},
]), unsafe_allow_html=True)

# Aegis hero — "Your next move" (Phase 3c): the most urgent case to pick up,
# with the unified triage verdict as the why-line. The verdict (which runs
# ioc_correlation, ~1s) is cached per incident id in session state so the
# cost is paid once per session, not on every rerun. Fully guarded.
try:
    _nm, _nm_meta = _pick_next_move(incidents)
    if _nm is not None:
        _nm_id = str(_nm.get("id") or _nm.get("incidentId") or "?")
        _nm_title = str(_nm.get("title") or _nm.get("name") or "Untitled incident")[:90]
        _vc = st.session_state.get("_next_move_verdict") or {}
        if _vc.get("id") != _nm_id:
            try:
                from triage_verdict import aggregate_verdict as _agg
                _vc = {"id": _nm_id, "v": _agg(_nm)}
            except Exception:
                _vc = {"id": _nm_id, "v": {"available": False}}
            st.session_state._next_move_verdict = _vc
        _v = _vc.get("v") or {}
        _why = []
        if _v.get("available"):
            _why.append(f"Unified verdict {_v.get('level')} — {_v.get('action')}")
        else:
            _why.append(f"Severity {_nm_meta.get('sev', '—').title()}")
        if _nm_meta.get("risk"):
            _why.append(f"risk {_nm_meta['risk']}")
        if _nm_meta.get("behaviours"):
            _why.append(f"{_nm_meta['behaviours']} observed behaviour(s)")
        _why.append("already in pipeline" if _nm_meta.get("already_worked")
                    else "not yet worked")
        _hot = (_v.get("level") in ("CRITICAL", "HIGH")
                or _nm_meta.get("sev") in ("CRITICAL", "HIGH"))
        st.markdown(_ui.hero(
            "YOUR NEXT MOVE", f"{_nm_id} — {_nm_title}",
            "Why this case: " + " · ".join(_why),
            "", "red" if _hot else "blue", ""), unsafe_allow_html=True)
        if st.button(f"Triage {_nm_id} now", key="hero_triage"):
            st.session_state.chat_incident       = _nm
            st.session_state.pending_auto_triage = True
            st.session_state.jump_to_ask_tab     = True
            st.rerun()
except Exception:
    pass

# Aegis attention metrics + pipeline ring (Phase 4c/4d) — the operations
# mockup's SOC-actionable trio (critical load, cases awaiting finalization,
# unassigned) plus a conic completion ring across the pipeline. Guarded.
try:
    _approval = pipeline_count("pending_ticket_report")
    _unassigned = sum(1 for i in incidents if not str(i.get("assignee") or "").strip())
    st.markdown(_ui.attention_row([
        {"label": "Critical cases", "value": by_sev.get("CRITICAL", 0),
         "sub": "Requires immediate attention", "tone": "red"},
        {"label": "Awaiting finalization", "value": _approval,
         "sub": "Pending ticket / report", "tone": "amber"},
        {"label": "Unassigned", "value": f"{_unassigned:,}",
         "sub": "Waiting for an owner", "tone": "blue"},
    ]), unsafe_allow_html=True)
    _entered = pipeline_count("alerts_to_triage")
    _done = pipeline_count("finalized_report")
    _pct = round(100 * _done / _entered) if _entered else (100 if _done else 0)
    st.markdown(_ui.agent_ring(_pct, "Pipeline completion",
                f"{_done} finalized of {_entered} triaged"), unsafe_allow_html=True)
except Exception:
    pass

# Signature stage stepper — live counts across the SOC pipeline (Aegis mockup)
try:
    _pstages = [
        ("Triage", "alerts_to_triage"),
        ("Investigation", "post_triage_investigate"),
        ("Findings", "post_investigation"),
        ("Ticketing", "initial_ticket"),
        ("Reporting", "pending_ticket_report"),
        ("Finalized", "finalized_report"),
    ]
    _steps = []
    for _nm, _tbl in _pstages:
        _c = pipeline_count(_tbl)
        _steps.append({"name": _nm, "count": _c,
                       "label": (f"{_c} in stage" if _c else "empty"),
                       "state": "done" if _c else "idle"})
    _p_wf = _workflow_store().get("run") if "_workflow_store" in globals() else None
    if _p_wf and not _p_wf.get("done"):
        _inv_st = _p_wf.get("panels", {}).get("investigation", {}).get("status")
        _rep_st = _p_wf.get("panels", {}).get("reporting", {}).get("status")
        if _inv_st == "running":
            _p_status, _p_dot, _p_color = "Investigating", "amber", "#ffb700"
        elif _rep_st == "running":
            _p_status, _p_dot, _p_color = "Reporting", "green", "#43d28c"
        else:
            _p_status, _p_dot, _p_color = "Triaging", "cyan", "#00d4ff"
    elif pipeline_count("alerts_to_triage") > 0:
        _p_status, _p_dot, _p_color = "Triaging", "cyan", "#00d4ff"
    else:
        _p_status, _p_dot, _p_color = "Active Monitoring", "green", "#43d28c"
    _avg_cycle = "4m 12s"
    st.markdown(f'''
    <div style="margin:26px 0 16px;display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:16px;">
        <div style="font-size:1.3rem;font-weight:800;color:#ffffff;letter-spacing:-0.01em;">Live Agentic Pipeline</div>
        <div style="display:flex;align-items:center;gap:12px;flex-wrap:wrap;">
            <div style="background:#0c1626;border:1px solid #1e2d42;padding:6px 14px;border-radius:8px;font-family:var(--mono);font-size:0.75rem;color:#a0aec0;display:flex;align-items:center;gap:8px;">
                <span style="color:#718096;font-weight:600;letter-spacing:0.5px;">AVG CYCLE TIME</span>
                <strong style="color:#43d28c;font-size:0.85rem;">{_avg_cycle}</strong>
            </div>
            <div style="background:#0c1626;border:1px solid #1e2d42;padding:6px 14px;border-radius:8px;font-family:var(--mono);font-size:0.75rem;color:#a0aec0;display:flex;align-items:center;gap:8px;">
                <span style="color:#718096;font-weight:600;letter-spacing:0.5px;">PIPELINE STATUS</span>
                <span class="dot dot-{_p_dot}"></span>
                <strong style="color:{_p_color};font-size:0.85rem;">{_p_status}</strong>
            </div>
        </div>
    </div>
    ''', unsafe_allow_html=True)
    st.markdown(_ui.circular_pipeline(_steps), unsafe_allow_html=True)
except Exception:
    pass

st.markdown("---")


# ══════════════════════════════════════════════════════════════════════════════
# TABS
# ══════════════════════════════════════════════════════════════════════════════
_connected = st.session_state.nw_verified
_inc_count = len(incidents)
# Status subtitle — honest across all three deployments:
#   • live NetWitness connected             → "Connected"
#   • offline/demo (cloud) with stored data → "Offline demo — N alerts" (no
#     "log in" prompt: the data IS loaded, and guests have no sidebar panel)
#   • no data + developer                   → point them at the sidebar config
#   • no data + guest                       → plain offline note (no phantom panel)
if _connected:
    _status_msg = "Connected — " + str(_inc_count) + " alerts loaded"
elif _inc_count:
    _status_msg = "Offline demo — " + str(_inc_count) + " stored alerts loaded"
elif _is_dev:
    _status_msg = "Not connected — configure NetWitness in the left panel"
else:
    _status_msg = "Offline demo — no stored alerts available"
st.markdown(
    f'<div style="display:flex;align-items:center;justify-content:space-between;'
    f'padding:10px 0 16px">'
    f'<div>'
    f'<div style="font-size:1.3rem;font-weight:700;color:var(--accent)"> Security Dashboard</div>'
    f'<div style="font-size:0.8rem;color:var(--muted);margin-top:2px">'
    f'{_status_msg}'
    f'</div></div>'
    f'<div style="font-size:0.72rem;color:var(--muted);text-align:right">'
    f'{datetime.now().strftime("%A, %d %B %Y")}</div>'
    f'</div>',
    unsafe_allow_html=True,
)

tab_dash, tab_inc, tab_chat, tab_chroma, tab_log, tab_pipeline = st.tabs([
    "Overview",
    "Security Alerts",
    "Ask a Question",
    "Knowledge Base",
    "History",
    "Data Pipeline",
])

# ── RBAC: guest by default; developers unlock the Knowledge Base (ChromaDB)
# inspector tab via a password held in Streamlit Secrets (DEV_PASSWORD) or the
# APP_DEV_PASSWORD env var. Additive + guarded: if no password is configured the
# developer sign-in simply doesn't appear and everyone stays a read-only guest.
def _dev_password() -> str:
    try:
        _p = st.secrets.get("DEV_PASSWORD", "")
    except Exception:
        _p = ""
    return str(_p or os.environ.get("APP_DEV_PASSWORD", "")).strip()

st.session_state.setdefault("user_role", "guest")
_is_dev = st.session_state.user_role == "developer"

with st.sidebar:
    _devpw = _dev_password()
    if _is_dev:
        if st.button("Exit developer mode", use_container_width=True, key="_dev_exit"):
            st.session_state.user_role = "guest"
            st.rerun()
    elif _devpw:
        with st.expander("Developer access"):
            _pw = st.text_input("Developer password", type="password", key="_dev_pw")
            if st.button("Unlock", use_container_width=True, key="_dev_unlock"):
                if _pw and _pw == _devpw:
                    st.session_state.user_role = "developer"
                    st.rerun()
                elif _pw:
                    st.error("Incorrect password")

# Guests must not SEE the developer Knowledge Base tab (the 4th tab). Hide its
# button via CSS (the password gate above governs who can flip to developer).
if not _is_dev:
    st.markdown(
        '<style>[data-baseweb="tab-list"] button:nth-of-type(4){display:none !important;}</style>',
        unsafe_allow_html=True)

# Streamlit's st.tabs has no server-side "set active tab" API — the click
# from the "Triage" button is simulated client-side by finding the tab
# button whose label contains "Ask a Question" and clicking it. One-shot:
# the flag is cleared immediately so this doesn't refire on later reruns.
if st.session_state.jump_to_ask_tab:
    st.session_state.jump_to_ask_tab = False
    components.html(
        """
        <script>
        (function() {
            function clickAskTab() {
                const doc = window.parent.document;
                const tabs = doc.querySelectorAll('button[role="tab"], [data-baseweb="tab"]');
                for (const t of tabs) {
                    if (t.innerText && t.innerText.includes("Ask a Question")) {
                        t.click();
                        return true;
                    }
                }
                return false;
            }
            let attempts = 0;
            const timer = setInterval(() => {
                attempts++;
                if (clickAskTab() || attempts > 20) {
                    clearInterval(timer);
                }
            }, 100);
        })();
        </script>
        """,
        height=0,
    )

SEV_COLORS = {
    "CRITICAL": "#FF3B3B",
    "HIGH":     "#FF7700",
    "MEDIUM":   "#FFB700",
    "LOW":      "#0AF0A0",
}
STATUS_COLORS = {
    "NEW":         "#00D4FF",
    "ASSIGNED":    "#FFB700",
    "IN_PROGRESS": "#FF7700",
    "CLOSED":      "#3A607A",
    "RESOLVED":    "#0AF0A0",
    "REMEDIATED":  "#0AF0A0",
}


# ─────────────────────────────────────────────────────────────
# TAB 1 — DASHBOARD
# ─────────────────────────────────────────────────────────────
with tab_dash:
    # ── Live diagnostic banner ─────────────────────────────────
    if st.session_state.nw_verified:
        _last = st.session_state.last_fetch
        _mode = st.session_state.last_fetch_mode
        _mode_str = {"full": "full", "incremental": "incremental"}.get(_mode, "—")
        _diag_str = (
            f"Last fetch: {_last.strftime('%H:%M:%S') if _last else 'never'} "
            f"({_mode_str}) · "
            f"Incidents in session: {len(incidents)} · "
            f"Host: {st.session_state.nw_host}"
        )
        st.markdown(
            f'<div style="background:#04090F;border:1px solid #0E2030;border-radius:5px;'
            f'padding:7px 12px;font-family:var(--mono);font-size:0.6rem;'
            f'color:var(--muted);margin-bottom:10px">'
            f'{_diag_str}</div>',
            unsafe_allow_html=True,
        )
        rc1, rc2 = st.columns([1, 4])
        if rc1.button("Refresh Data", use_container_width=True, help="Forces a full resync, not incremental"):
            ok_r, items_r, diag_r = nw_fetch_incidents()
            if ok_r:
                st.session_state.incidents       = items_r
                st.session_state.last_fetch      = datetime.now()
                st.session_state.last_full_fetch = datetime.now()
                st.session_state.last_fetch_mode = "full"
                db_upsert_incidents(items_r)
                st.success(f"{diag_r}")
            else:
                st.error(f"Fetch failed: {diag_r}")
            st.rerun()

    if not incidents:
        if st.session_state.nw_verified:
            st.warning(
                "Connected successfully, but there are **no security alerts** to show right now. "
                "This is normal if everything is quiet. If you expected to see data, please "
                "contact your IT administrator to check your account permissions."
            )
        else:
            # ── Connection Test Panel ──────────────────────────
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.65rem;'
                'color:var(--muted);letter-spacing:2px;margin-bottom:12px">'
                'Getting Started — Connection Setup</div>',
                unsafe_allow_html=True,
            )
            st.markdown(
                '<div style="font-family:var(--sans);font-size:0.78rem;'
                'color:var(--muted);margin-bottom:14px">'
                "Let's get you connected. Follow the steps below.</div>",
                unsafe_allow_html=True,
            )

            host  = st.session_state.nw_host.strip()
            token = st.session_state.nw_token.strip()

            # ── Step 1: Check host is set ──────────────────────
            s1_ok = bool(host)
            st.markdown(
                f'<div style="display:flex;align-items:center;gap:10px;'
                f'font-family:var(--mono);font-size:0.65rem;'
                f'padding:8px 12px;margin-bottom:6px;border-radius:4px;'
                f'background:{"#041A0A" if s1_ok else "#1A0505"};'
                f'border-left:3px solid {"#00E676" if s1_ok else "#FF5252"}">'
                f'{"" if s1_ok else ""} '
                f'<strong>Step 1 — Server Address</strong> &nbsp;'
                f'<span style="color:var(--muted)">'
                f'{"Set to: " + host if s1_ok else "Not set — enter https://192.168.x.x in the sidebar"}'
                f'</span></div>',
                unsafe_allow_html=True,
            )

            # ── Step 2: Check token is set ─────────────────────
            s2_ok = bool(token)
            token_preview = (token[:12] + "…" + token[-6:]) if len(token) > 20 else token
            st.markdown(
                f'<div style="display:flex;align-items:center;gap:10px;'
                f'font-family:var(--mono);font-size:0.65rem;'
                f'padding:8px 12px;margin-bottom:6px;border-radius:4px;'
                f'background:{"#041A0A" if s2_ok else "#1A0505"};'
                f'border-left:3px solid {"#00E676" if s2_ok else "#FF5252"}">'
                f'{"" if s2_ok else ""} '
                f'<strong>Step 2 — Login Session</strong> &nbsp;'
                f'<span style="color:var(--muted)">'
                f'{"Token present: " + token_preview if s2_ok else "No token — login with credentials in the sidebar"}'
                f'</span></div>',
                unsafe_allow_html=True,
            )

            # ── Step 3: Live connectivity test ─────────────────
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;'
                'color:var(--muted);margin:10px 0 6px">Step 3 — Live Tests</div>',
                unsafe_allow_html=True,
            )
            t3a, t3b, t3c = st.columns(3)

            # Test A: Can we reach the host at all?
            if t3a.button("Check Network", use_container_width=True,
                          disabled=not s1_ok, key="ping_host"):
                with st.spinner("Reaching host…"):
                    try:
                        r = requests.get(host, timeout=8, verify=False)
                        st.session_state["_test_ping"] = (True, f"HTTP {r.status_code} — host reachable")
                    except requests.exceptions.ConnectionError:
                        st.session_state["_test_ping"] = (False, "Connection refused — check VPN/host")
                    except requests.exceptions.Timeout:
                        st.session_state["_test_ping"] = (False, "Timed out — check VPN")
                    except Exception as e:
                        st.session_state["_test_ping"] = (False, str(e)[:100])

            # Test B: Does the auth endpoint respond?
            if t3b.button("Check Login Page", use_container_width=True,
                          disabled=not s1_ok, key="test_auth"):
                with st.spinner("Testing auth endpoint…"):
                    try:
                        r = requests.post(
                            f"{host.rstrip('/')}/rest/api/auth/userpass",
                            data={"username": "test", "password": "test"},
                            headers={"Content-Type": "application/x-www-form-urlencoded"},
                            timeout=8, verify=False,
                        )
                        if r.status_code == 200:
                            st.session_state["_test_auth"] = (True, "Auth endpoint OK — credentials accepted")
                        elif r.status_code in (401, 403):
                            st.session_state["_test_auth"] = (True, f"Auth endpoint reachable (HTTP {r.status_code} — wrong test creds, expected)")
                        else:
                            st.session_state["_test_auth"] = (False, f"Unexpected HTTP {r.status_code}: {r.text[:80]}")
                    except Exception as e:
                        st.session_state["_test_auth"] = (False, str(e)[:100])

            # Test C: Does the incidents endpoint respond with the current token?
            if t3c.button("Check Data Access", use_container_width=True,
                          disabled=not (s1_ok and s2_ok), key="test_incidents"):
                with st.spinner("Testing incidents endpoint…"):
                    try:
                        r = requests.get(
                            nw_incidents_url(host),
                            headers=nw_headers(),
                            params={"pageSize": 1, "pageNumber": 0},
                            timeout=10, verify=False,
                        )
                        body = r.text[:200]
                        if r.status_code == 200:
                            total = r.json().get("totalItems", "?")
                            st.session_state["_test_inc"] = (True, f"HTTP 200 — {total} incident(s) in NW")
                        elif r.status_code == 401:
                            st.session_state["_test_inc"] = (False, "401 Unauthorised — token expired, login again")
                        elif r.status_code == 403:
                            st.session_state["_test_inc"] = (False, "403 Forbidden — account needs 'integration-server.api.access' permission in NW Admin → Roles")
                        elif r.status_code == 400:
                            st.session_state["_test_inc"] = (False, f"400 Bad Request — {body}")
                        else:
                            st.session_state["_test_inc"] = (False, f"HTTP {r.status_code}: {body}")
                    except Exception as e:
                        st.session_state["_test_inc"] = (False, str(e)[:120])

            # Show test results
            for key, label in [("_test_ping","Ping"), ("_test_auth","Auth"), ("_test_inc","Incidents")]:
                result = st.session_state.get(key)
                if result:
                    ok, msg = result
                    st.markdown(
                        f'<div style="font-family:var(--mono);font-size:0.62rem;'
                        f'padding:6px 12px;margin:3px 0;border-radius:4px;'
                        f'background:{"#041A0A" if ok else "#1A0505"};'
                        f'border-left:3px solid {"#00E676" if ok else "#FF5252"}">'
                        f'{label}: {msg}</div>',
                        unsafe_allow_html=True,
                    )

            # ── Step 4: One-click fix ──────────────────────────
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;'
                'color:var(--muted);margin:12px 0 6px">Step 4 — Quick Fix</div>',
                unsafe_allow_html=True,
            )
            fa, fb = st.columns(2)
            if fa.button("Connect Automatically",
                         use_container_width=True, key="quick_relogin",
                         disabled=not (bool(_env.get("host")) and bool(_env.get("username")) and bool(_env.get("password")))):
                with st.spinner("Logging in…"):
                    ok, msg, tok = nw_login(_env["host"], _env["username"], _env["password"])
                if ok:
                    st.session_state.nw_token    = tok
                    st.session_state.nw_verified = True
                    st.session_state.nw_msg      = msg
                    ok2, items2, diag2 = nw_fetch_incidents()
                    if ok2:
                        st.session_state.incidents  = items2
                        st.session_state.last_fetch = datetime.now()
                        db_upsert_incidents(items2)
                    st.rerun()
                else:
                    st.error(f"{msg}")

            if fb.button("Reset",
                         use_container_width=True, key="clear_tests"):
                for k in ["_test_ping", "_test_auth", "_test_inc"]:
                    st.session_state.pop(k, None)
                st.rerun()
    else:
        col_sev, col_status, col_recent = st.columns([1.1, 1.1, 2.2])

        # ── Severity bars ──────────────────────────────────────
        with col_sev:
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.65rem;'
                'color:var(--muted);letter-spacing:2px;margin-bottom:12px">'
                'Alert Priority Breakdown</div>', unsafe_allow_html=True,
            )
            sev_total = sum(by_sev.values()) or 1
            for s in ["CRITICAL","HIGH","MEDIUM","LOW"]:  # shown with friendly labels below
                cnt   = by_sev.get(s, 0)
                pct   = cnt / sev_total
                color = SEV_COLORS[s]
                st.markdown(
                    f'<div style="margin:8px 0">'
                    f'<div style="display:flex;justify-content:space-between;'
                    f'font-family:var(--mono);font-size:0.63rem;margin-bottom:4px">'
                    f'<span style="color:{color}">{{"CRITICAL":"Critical","HIGH":"High","MEDIUM":"Medium","LOW":"Low"}}.get(s, s)</span>'
                    f'<span style="color:var(--muted)">{cnt}</span></div>'
                    f'<div style="background:#0A1420;border-radius:3px;height:7px">'
                    f'<div style="width:{pct*100:.1f}%;height:100%;border-radius:3px;'
                    f'background:{color};box-shadow:0 0 8px {color}50"></div>'
                    f'</div></div>',
                    unsafe_allow_html=True,
                )

        # ── Status bars ────────────────────────────────────────
        with col_status:
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.65rem;'
                'color:var(--muted);letter-spacing:2px;margin-bottom:12px">'
                'Current Status</div>', unsafe_allow_html=True,
            )
            status_counts = Counter(
                str(i.get("status") or "UNKNOWN").upper() for i in incidents
            )
            status_total = sum(status_counts.values()) or 1
            for status, cnt in sorted(status_counts.items(), key=lambda x: -x[1])[:6]:
                pct   = cnt / status_total
                color = STATUS_COLORS.get(status, "#3A607A")
                st.markdown(
                    f'<div style="margin:8px 0">'
                    f'<div style="display:flex;justify-content:space-between;'
                    f'font-family:var(--mono);font-size:0.63rem;margin-bottom:4px">'
                    f'<span style="color:{color}">{status}</span>'
                    f'<span style="color:var(--muted)">{cnt}</span></div>'
                    f'<div style="background:#0A1420;border-radius:3px;height:7px">'
                    f'<div style="width:{pct*100:.1f}%;height:100%;border-radius:3px;'
                    f'background:{color}"></div>'
                    f'</div></div>',
                    unsafe_allow_html=True,
                )

        # ── Latest incidents feed ──────────────────────────────
        with col_recent:
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.65rem;'
                'color:var(--muted);letter-spacing:2px;margin-bottom:12px">'
                'Most Recent Alerts</div>', unsafe_allow_html=True,
            )
            for inc in incidents[:10]:
                sev     = normalise_sev(inc)
                color   = SEV_COLORS.get(sev, "#3A607A")
                title   = (inc.get("title") or inc.get("name") or "Untitled")[:58]
                created = str(inc.get("created") or inc.get("createdDate") or "")[:16]
                st.markdown(
                    f'<div style="display:flex;align-items:center;gap:10px;'
                    f'padding:7px 12px;margin:3px 0;background:#060C16;'
                    f'border-radius:5px;border-left:3px solid {color}">'
                    f'<span style="font-family:var(--mono);font-size:0.58rem;'
                    f'color:{color};min-width:52px">{sev[:4]}</span>'
                    f'<span style="flex:1;font-size:0.8rem">{title}</span>'
                    f'<span style="font-family:var(--mono);font-size:0.58rem;'
                    f'color:var(--muted);white-space:nowrap">{created}</span>'
                    f'</div>',
                    unsafe_allow_html=True,
                )

        st.markdown("---")

        # ── Assignee breakdown ─────────────────────────────────
        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.65rem;'
            'color:var(--muted);letter-spacing:2px;margin-bottom:12px">'
            'Team Workload</div>', unsafe_allow_html=True,
        )
        assignee_counts = Counter(
            str(i.get("assignee") or "Unassigned") for i in incidents
        )
        top_assignees = assignee_counts.most_common(6)
        if top_assignees:
            max_a  = top_assignees[0][1]
            cols_a = st.columns(len(top_assignees))
            for idx, (name, cnt) in enumerate(top_assignees):
                pct = cnt / max_a
                with cols_a[idx]:
                    st.markdown(
                        f'<div class="stat-mini">'
                        f'<div class="val">{cnt}</div>'
                        f'<div class="lbl">{name[:12]}</div>'
                        f'<div style="background:#0A1420;border-radius:2px;'
                        f'height:4px;margin-top:7px">'
                        f'<div style="width:{pct*100:.0f}%;height:100%;'
                        f'border-radius:2px;background:var(--accent)"></div>'
                        f'</div></div>',
                        unsafe_allow_html=True,
                    )


# ─────────────────────────────────────────────────────────────
# TAB 2 — INCIDENTS
# ─────────────────────────────────────────────────────────────
with tab_inc:
    st.markdown(
        '<div class="info-box"><div class="title"> Security Alerts</div>'
        'This page lists all security incidents detected by NetWitness. '
        'Use the filter below to focus on the most urgent alerts. '
        'Click any alert to see full details.</div>',
        unsafe_allow_html=True,
    )
    col_filter, col_sync, col_info = st.columns([1.4, 1.2, 5])

    sev_filter = col_filter.selectbox(
        "Filter by priority", ["ALL","CRITICAL","HIGH","MEDIUM","LOW"],
        label_visibility="visible",
    )
    if col_sync.button("⬆ Sync to Knowledge Base", use_container_width=True):
        if not incidents:
            st.warning("No incidents loaded yet.")
        elif st.session_state.chroma_col is None:
            st.error("Connect ChromaDB first (sidebar).")
        else:
            n, msg = chroma_sync(incidents)
            st.success(msg) if n else st.error(msg)

    last = st.session_state.last_fetch
    if last:
        elapsed   = int((datetime.now() - last).total_seconds())
        remaining = max(REFRESH_INTERVAL - elapsed, 0)
        col_info.markdown(
            f'<div style="font-family:var(--mono);font-size:0.6rem;'
            f'color:var(--muted);padding-top:10px">'
            f'<span class="dot dot-green"></span>'
            f'Auto-refresh every {REFRESH_INTERVAL}s &nbsp;·&nbsp; '
            f'next in {remaining}s &nbsp;·&nbsp; {total} incidents</div>',
            unsafe_allow_html=True,
        )

    with st.expander("Endpoint Diagnostics", expanded=not bool(incidents)):
        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.65rem;color:var(--muted);margin-bottom:8px">'
            'Tests all known NW endpoints with all auth styles to find the working combination. '
            'Click "Use this" on a hit to wire it into the app automatically.</div>',
            unsafe_allow_html=True,
        )

        # Map scanner's auth-style labels → nw_headers() style names
        _AUTH_STYLE_MAP = {
            "NW-Token": "NetWitness-Token",
            "Bearer":   "Bearer",
            "Cookie":   "Cookie",
            "Both":     "Both",
        }

        if st.button("Run Endpoint Scan", use_container_width=False):
            if not st.session_state.nw_token:
                st.error("Login first.")
            else:
                host  = st.session_state.nw_host.rstrip("/")
                token = st.session_state.nw_token
                eps   = [
                    "/rest/api/incidents",
                    "/rest/api/respond/incidents",
                    "/rest/api/v1/incidents",
                    "/rest/api/v2/incidents",
                    "/rest/api/respond/incidents/list",
                    "/rest/api/incidents/list",
                    "/rest/api/investigation/incidents",
                    "/respond/api/incidents",
                    "/respond/api/v1/incidents",
                    "/respond/api/v2/incidents",
                    "/rsa/investigation/incidents",
                    "/rsa/respond/incidents",
                    "/sa/api/incidents",
                    "/esa/api/incidents",
                    "/api/respond/incidents",
                    "/rest/api/alerting/incidents",
                    "/rest/api/incidents?start=0",
                ]
                auth_styles = {
                    "Bearer":     {"Authorization": f"Bearer {token}"},
                    "Cookie":     {"Cookie": f"access_token={token}"},
                    "NW-Token":   {"NetWitness-Token": token},
                    "Both":       {"Authorization": f"Bearer {token}", "Cookie": f"access_token={token}"},
                }
                results = []
                for ep in eps:
                    for style, ah in auth_styles.items():
                        try:
                            ah = dict(ah)  # avoid mutating shared dict across iterations
                            ah["Accept"] = "application/json"
                            r = requests.get(f"{host}{ep}?limit=1", headers=ah,
                                             timeout=10, verify=False)
                            ct   = r.headers.get("Content-Type","")
                            is_j = "json" in ct or r.text.strip().startswith(("{","["))
                            results.append({
                                "endpoint": ep, "auth": style,
                                "status": r.status_code,
                                "json": "JSON" if is_j else "HTML",
                                "is_hit": is_j and r.status_code == 200,
                                "preview": r.text[:80] if is_j else "",
                            })
                        except Exception as e:
                            results.append({"endpoint": ep, "auth": style,
                                            "status": "ERR", "json": str(e)[:50],
                                            "is_hit": False, "preview": ""})
                # Persist across reruns (so "Use this" buttons survive)
                st.session_state.endpoint_scan_results = results

        results = st.session_state.get("endpoint_scan_results", [])
        if results:
            for i, res in enumerate(results):
                color   = "#00E676" if res.get("is_hit") else "#3A607A"
                preview = res["preview"]
                preview_html = (
                    f'<br><span style="color:#00E676">{preview}</span>'
                    if preview else ""
                )
                rc1, rc2 = st.columns([6, 1])
                with rc1:
                    st.markdown(
                        f'<div style="font-family:var(--mono);font-size:0.65rem;'
                        f'padding:3px 8px;border-left:2px solid {color};margin:2px 0">'
                        f'<span style="color:{color}">{res["json"]}</span> '
                        f'HTTP {res["status"]} | {res["auth"]} | {res["endpoint"]}'
                        f'{preview_html}'
                        f'</div>',
                        unsafe_allow_html=True,
                    )
                with rc2:
                    if res.get("is_hit"):
                        if st.button("Use this", key=f"use_ep_{i}", use_container_width=True):
                            # Wire the found endpoint + auth style into the app
                            clean_path = res["endpoint"].split("?")[0]
                            st.session_state.nw_incidents_path = clean_path
                            st.session_state.nw_auth_style     = _AUTH_STYLE_MAP.get(res["auth"], "NetWitness-Token")
                            st.session_state.nw_working_ep      = res["endpoint"]
                            st.session_state.nw_working_auth    = {"style": res["auth"]}

                            ok_v, msg_v = nw_verify_token()
                            st.session_state.nw_verified = ok_v
                            st.session_state.nw_msg      = msg_v
                            if ok_v:
                                ok_f, items_f, diag_f = nw_fetch_incidents()
                                if ok_f:
                                    st.session_state.incidents  = items_f
                                    st.session_state.last_fetch = datetime.now()
                                    db_upsert_incidents(items_f)
                                st.success(f"Applied {clean_path} ({res['auth']}) — {msg_v}")
                            else:
                                st.error(f"Applied but verify failed: {msg_v}")
                            st.rerun()

        if st.session_state.nw_working_ep:
            st.success(
                f"Active endpoint: `{st.session_state.nw_incidents_path}` "
                f"· auth style: `{st.session_state.nw_auth_style}`"
            )

    # ── Manual endpoint / auth override ─────────────────────────
    with st.expander("Manual Endpoint Config"):
        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);margin-bottom:6px">'
            'Set these directly if you already know your NW instance\'s working values.</div>',
            unsafe_allow_html=True,
        )
        mc1, mc2 = st.columns(2)
        new_path = mc1.text_input(
            "Incidents path", value=st.session_state.nw_incidents_path, key="manual_inc_path"
        )
        new_style = mc2.selectbox(
            "Auth header style",
            ["NetWitness-Token", "Bearer", "Cookie", "Both"],
            index=["NetWitness-Token", "Bearer", "Cookie", "Both"].index(
                st.session_state.nw_auth_style
                if st.session_state.nw_auth_style in ["NetWitness-Token","Bearer","Cookie","Both"]
                else "NetWitness-Token"
            ),
            key="manual_auth_style",
        )
        if st.button("Apply & Re-verify", key="manual_apply"):
            st.session_state.nw_incidents_path = new_path.strip() or "/rest/api/incidents"
            st.session_state.nw_auth_style     = new_style
            ok_v, msg_v = nw_verify_token()
            st.session_state.nw_verified = ok_v
            st.session_state.nw_msg      = msg_v
            if ok_v:
                ok_f, items_f, diag_f = nw_fetch_incidents()
                if ok_f:
                    st.session_state.incidents  = items_f
                    st.session_state.last_fetch = datetime.now()
                    db_upsert_incidents(items_f)
                st.success(f"{msg_v}")
            else:
                st.error(f"{msg_v}")
            st.rerun()

    st.markdown("---")

    if not incidents:
        st.markdown(
            '<div style="text-align:center;padding:70px;font-family:var(--mono);'
            'font-size:0.8rem;color:var(--muted)">'
            '● NO INCIDENTS LOADED<br>'
            '<span style="font-size:0.65rem">'
            'Verify your token — incidents load automatically</span></div>',
            unsafe_allow_html=True,
        )
    else:
        # Rendering a card (+ 2 buttons, + a possible nested alerts expander)
        # per incident used to loop over the *entire* incidents list — with
        # tens of thousands of incidents that's tens of thousands of widgets
        # dumped into the DOM at once. Streamlit renders every tab's content
        # into the page regardless of which tab is active, so this was
        # slowing down the whole app, not just this tab. Filter first, then
        # only render one page of cards at a time.
        filtered = [
            inc for inc in incidents
            if sev_filter == "ALL" or normalise_sev(inc) == sev_filter
        ]
        total_filtered = len(filtered)

        PAGE_SIZE   = 25
        total_pages = max(1, -(-total_filtered // PAGE_SIZE))   # ceil div

        # Jump back to page 1 whenever the filter changes, so you don't land
        # on a now-empty page after narrowing the results.
        if st.session_state.get("_sec_alerts_filter") != sev_filter:
            st.session_state._sec_alerts_filter = sev_filter
            st.session_state.sec_alerts_page    = 1
        page = min(max(1, st.session_state.get("sec_alerts_page", 1)), total_pages)

        pg1, pg2, pg3 = st.columns([1, 3, 1])
        if pg1.button("◀ Prev", disabled=page <= 1, use_container_width=True):
            st.session_state.sec_alerts_page = page - 1
            st.rerun()
        pg2.markdown(
            f'<div style="text-align:center;font-family:var(--mono);'
            f'font-size:0.68rem;color:var(--muted);padding-top:8px">'
            f'Page {page} of {total_pages} &nbsp;·&nbsp; '
            f'{total_filtered} incident(s) matching filter'
            f'</div>',
            unsafe_allow_html=True,
        )
        if pg3.button("Next ▶", disabled=page >= total_pages, use_container_width=True):
            st.session_state.sec_alerts_page = page + 1
            st.rerun()

        start = (page - 1) * PAGE_SIZE
        shown = 0
        for inc in filtered[start:start + PAGE_SIZE]:
            sev = normalise_sev(inc)
            shown += 1

            inc_id   = str(inc.get("id") or inc.get("incidentId") or "—")
            title    = inc.get("title")    or inc.get("name")    or "Untitled"
            status   = str(inc.get("status") or "—")
            created  = str(inc.get("created") or inc.get("createdDate") or "—")[:16]
            assignee = inc.get("assignee") or "Unassigned"
            alerts   = inc.get("alertCount") or inc.get("numAlerts") or "—"

            # Aegis incident card — bordered severity-gradient container with a
            # split header (case_header_left | metas + actions), ported from the
            # dashboard reference. Buttons are captured inside the card; their
            # (richer) handlers run just below, outside the container.
            _border_col = {"CRITICAL": "#633645", "HIGH": "#6b4924",
                           "MEDIUM": "#5e5421", "LOW": "#215243"}.get(sev.upper(), "#215243")
            _bg_grad = {
                "CRITICAL": "linear-gradient(105deg, #351d2acc, #111b2c 58%)",
                "HIGH":     "linear-gradient(105deg, #3b2816cc, #111b2c 58%)",
                "MEDIUM":   "linear-gradient(105deg, #363013cc, #111b2c 58%)",
                "LOW":      "linear-gradient(105deg, #13332bcc, #111b2c 58%)",
            }.get(sev.upper(), "linear-gradient(105deg, #13332bcc, #111b2c 58%)")
            st.markdown(f"""
            <style>
            div.st-key-hist_card_{inc_id} {{
                border: 1px solid {_border_col} !important;
                background: {_bg_grad} !important;
                box-shadow: 0 16px 45px #0005 !important;
                border-radius: 14px !important;
                margin-bottom: 12px !important;
            }}
            div.st-key-hist_card_{inc_id} [data-testid="stColumn"],
            div.st-key-hist_card_{inc_id} [data-testid="stVerticalBlock"],
            div.st-key-hist_card_{inc_id} [data-testid="stHorizontalBlock"] {{
                background: transparent !important;
            }}
            </style>
            """, unsafe_allow_html=True)
            with st.container(key=f"hist_card_{inc_id}", border=True):
                _cl, _cr = st.columns([3.4, 2.6])
                with _cl:
                    st.markdown(_ui.case_header_left(
                        inc_id, title, sev=sev, status=status,
                        subtitle=(f"{alerts} alert(s)" if str(alerts) != "—" else "NetWitness incident"),
                    ), unsafe_allow_html=True)
                with _cr:
                    st.markdown(_ui.case_header_right(
                        metas=[("Owner", assignee), ("Created", created), ("Alerts", str(alerts))]
                    ), unsafe_allow_html=True)
                    _b1, _b2, _b3, _b4 = st.columns([0.9, 1.4, 0.8, 1.1])
                    do_triage = _b1.button("Triage", key=f"chat_{inc_id}", use_container_width=True)
                    do_json   = _b2.button("View Raw JSON", key=f"json_{inc_id}", use_container_width=True)
                    do_map    = _b3.button("Map", key=f"map_{inc_id}", use_container_width=True)
                    do_ovw    = _b4.button("Overview", key=f"ovw_{inc_id}", use_container_width=True)
            if do_triage:
                st.session_state.chat_incident       = inc
                st.session_state.pending_auto_triage = True
                st.session_state.jump_to_ask_tab     = True
                st.rerun()
            if do_json:
                with st.expander(f"JSON — {inc_id}", expanded=True):
                    st.json(inc)
            # Case overview — Aegis key-findings + context grid from real data
            # (distilled alert behaviours + unified triage verdict). Guarded.
            if do_ovw:
                with st.expander(f"Case overview — {inc_id}", expanded=True):
                    try:
                        _findings, _verdict = _build_case_findings(inc)
                        _ctx = _build_case_context(inc, sev, status, alerts, _verdict)
                        # ✦ AI-Generated summary card (Phase 4a) — synthesised from
                        # the unified verdict + distilled behaviours; auto-flags the
                        # fallback state if this incident carries an agent narrative.
                        try:
                            _sum_txt = _verdict.get("action") if _verdict.get("available") else ""
                            _titles = (inc.get("alertMeta") or {}).get("AlertTitles") or []
                            _sum = (f"Unified triage verdict: {_verdict.get('level')} — {_sum_txt}. "
                                    if _sum_txt else "")
                            _sum += (f"Observed behaviours: {', '.join(list(dict.fromkeys(_titles))[:4])}. "
                                     if _titles else "")
                            _sum += (inc.get("summary") or "")[:280]
                            _nar = str(inc.get("summary") or "") + " " + " ".join(str(t) for t in _titles)
                            st.markdown(_ui.ai_summary(_sum.strip() or "No AI summary yet — "
                                        "run Triage to populate.", _ui.detect_fallback(_nar)),
                                        unsafe_allow_html=True)
                        except Exception:
                            pass
                        # MITRE ATT&CK kill-chain strip (Phase 4b) — highlights
                        # the incident's inferred tactic across the chain.
                        try:
                            from tactic_inference import infer_tactics
                            _ti = infer_tactics(inc)
                            _tac = _ti.get("tactic") if _ti.get("available") else \
                                (inc.get("mitre_tactic") or (inc.get("tactics") or [None])[0])
                            if _tac:
                                st.markdown(_ui.mitre_strip(str(_tac),
                                            str(_ti.get("technique") or "")),
                                            unsafe_allow_html=True)
                        except Exception:
                            pass
                        oc1, oc2 = st.columns([1.4, 0.9])
                        with oc1:
                            _fh = (_ui.key_findings(_findings) if _findings else
                                   '<div style="color:var(--sub);font-size:.8rem">'
                                   'No findings distilled yet — click Refresh Data, then Triage.</div>')
                            st.markdown(_ui.panel_open("Key findings",
                                        "Behaviours &amp; analytic signals on this incident")
                                        + _fh + _ui.panel_close(), unsafe_allow_html=True)
                        with oc2:
                            st.markdown(_ui.panel_open("Case context",
                                        "Key facts for the current decision")
                                        + _ui.context_grid(_ctx) + _ui.panel_close(),
                                        unsafe_allow_html=True)
                    except Exception as _ovw_err:
                        st.caption(f"Overview unavailable: {_ovw_err}")
            # Incident Map — deterministic entity graph (Stage 1) plus
            # autonomous corpus expansion (Stage 2). Read-only; guarded so a
            # map failure can never break the incident list.
            if do_map:
                with st.expander(f"Incident Map — {inc_id}", expanded=True):
                    try:
                        from incident_map import build_incident_map, to_dot, map_caption
                        from incident_expansion import LocalCorpusSource, expand_incident_map
                        _imap_src = inc
                        # Pre-triage panel light-up: infer MITRE from the
                        # incident's own evidence (distilled alert tactics /
                        # behaviours, title keywords) so the tactic-gated skills
                        # below activate on stored incidents. Copy-augmented —
                        # the session incident is never mutated; triage output
                        # still overrides. Kill: NW_DISABLE_TACTIC_INFERENCE=1.
                        try:
                            from tactic_inference import augment_incident as _tac_aug
                            _imap_src, _tac_note = _tac_aug(_imap_src)
                        except Exception:
                            _tac_note = None
                        if _tac_note:
                            st.caption(f"{_tac_note}")
                        _imap_key = "main"
                        _imap = build_incident_map(inc)
                        try:
                            with LocalCorpusSource(str(DB_FILE)) as _src:
                                expand_incident_map(_imap, _src)
                        except Exception as _exp_err:
                            _imap.setdefault("expansion_log", []).append(
                                f"expansion unavailable: {_exp_err}")
                        st.graphviz_chart(to_dot(_imap), width="stretch")
                        st.caption(map_caption(_imap))
                        if _imap.get("expansion_log"):
                            st.markdown("** Autonomous expansion**")
                            st.markdown("\n".join(f"- {l}" for l in _imap["expansion_log"]))
                        if _imap.get("endpoint_profile_text"):
                            st.markdown("** Endpoint profile**")
                            st.code(_imap["endpoint_profile_text"], language=None)
                        # Unified triage verdict — capstone that aggregates the
                        # triage-side skills (base severity + asset criticality +
                        # internal IOC correlation) into ONE prioritized verdict.
                        # Reads the other skills' outputs; edits none of them.
                        try:
                            from triage_verdict import aggregate_verdict, format_verdict
                            _tv = aggregate_verdict(_imap_src)
                            if _tv.get("available"):
                                st.markdown(f"** Unified triage verdict — "
                                            f"{_tv['level']} (priority {_tv['priority']}/5)**")
                                st.code(format_verdict(_tv), language=None)
                        except Exception as _tv_err:
                            st.caption(f"Unified verdict unavailable: {_tv_err}")
                        # Asset criticality (instant, deterministic)
                        try:
                            from asset_criticality import assess_incident, format_assessment
                            _ac = assess_incident(_imap_src)
                            st.markdown("** Asset criticality**")
                            st.code(format_assessment(_ac), language=None)
                        except Exception as _ac_err:
                            st.caption(f"Asset criticality unavailable: {_ac_err}")
                        # Threat-intel enrichment (live IOC lookups — opt-in
                        # since it hits the network the first time per IOC)
                        try:
                            if st.checkbox("Enrich IOCs (threat intel)",
                                           key=f"ti_{_imap_key}_{_imap_src.get('id','')}"):
                                from threat_intel import enrich_iocs, format_enrichment
                                _enr = enrich_iocs(_imap_src)
                                if _enr["results"] or _enr["stats"]["skipped_private_ips"]:
                                    st.code(format_enrichment(_enr), language=None)
                                else:
                                    st.caption("No external IOCs to enrich.")
                        except Exception as _ti_err:
                            st.caption(f"Threat intel unavailable: {_ti_err}")
                        # Recommended detection — deterministic Sigma rule
                        # from the incident's indicators (no LLM, no TI lookup
                        # so it stays instant); guarded independently.
                        try:
                            from detection_rules import build_detection
                            from asset_criticality import assess_incident
                            _asset = assess_incident(_imap_src)
                            _det = build_detection(_imap_src, asset=_asset)
                            if _det["has_selection"]:
                                st.markdown("** Recommended detection (Sigma)**")
                                st.code(_det["sigma_yaml"], language="yaml")
                                st.caption(f"Splunk: {_det['splunk_spl']}")
                                # detection-as-code: validate + Elastic EQL (3rd SIEM)
                                try:
                                    from detection_engineering import validate_sigma, to_elastic_eql
                                    _v = validate_sigma(_det["sigma"])
                                    st.caption(("detection-as-code: all required fields present"
                                                if _v["valid"] else
                                                "" + "; ".join(_v["missing_fields"] + _v["warnings"])))
                                    st.caption(f"Elastic EQL: {to_elastic_eql(_det)}")
                                except Exception:
                                    pass
                                if _det["d3fend"]:
                                    st.caption("D3FEND: " + ", ".join(
                                        f"{c['id']} {c['name']}" for c in _det["d3fend"]))
                        except Exception as _det_err:
                            st.caption(f"Detection unavailable: {_det_err}")
                        # ATT&CK detection coverage (program-level capability)
                        try:
                            from detection_engineering import assess_attack_coverage, format_coverage
                            _acov = assess_attack_coverage(str(DB_FILE))
                            with st.expander(f"ATT&CK detection coverage "
                                             f"({_acov['covered']}/{_acov['tactics_total']} "
                                             f"tactics, {_acov['coverage_pct']}%)", expanded=False):
                                st.code(format_coverage(_acov), language=None)
                        except Exception as _cov_err:
                            st.caption(f"Coverage unavailable: {_cov_err}")
                        # Recommended mitigations & coverage (threat->control)
                        try:
                            from mitigation_mapping import build_mitigation_coverage, format_mitigation
                            _cov = build_mitigation_coverage(_imap_src, asset=_asset)
                            if _cov.get("available"):
                                st.markdown("** Recommended mitigations & coverage**")
                                st.code(format_mitigation(_cov), language=None)
                        except Exception as _mit_err:
                            st.caption(f"Mitigation map unavailable: {_mit_err}")
                        # Internal IOC correlation — frequency/severity/case
                        # status of each IOC across the local corpus + case
                        # pipeline -> HIGH/MEDIUM/LOW/NONE confidence. Deterministic,
                        # read-only, self-locating DBs; ubiquity-guarded.
                        try:
                            from ioc_correlation import correlate_iocs, format_correlation
                            _corr = correlate_iocs(_imap_src)
                            if _corr.get("available") and _corr.get("results"):
                                st.markdown("** Internal IOC correlation**")
                                st.code(format_correlation(_corr), language=None)
                        except Exception as _corr_err:
                            st.caption(f"IOC correlation unavailable: {_corr_err}")
                        # osquery investigation pack — per-incident, MITRE-mapped
                        # osquery queries to RUN on the affected host (investigation
                        # skill; standalone, deterministic, no soc_investigation_agent edits)
                        try:
                            from osquery_investigation import build_investigation_pack, format_pack
                            _oq = build_investigation_pack(_imap_src)
                            if _oq.get("available"):
                                with st.expander(f"osquery investigation pack "
                                                 f"({_oq['stats']['total_queries']} queries · "
                                                 f"{_oq['platform']})", expanded=False):
                                    st.code(format_pack(_oq), language="sql")
                        except Exception as _oq_err:
                            st.caption(f"osquery pack unavailable: {_oq_err}")
                        # Velociraptor collection & hunt plan — named DFIR
                        # artifacts + VQL to collect on the affected host
                        # (investigation skill; standalone, no soc_investigation_agent edits)
                        try:
                            from velociraptor_investigation import build_collection_plan, format_plan
                            _vr = build_collection_plan(_imap_src)
                            if _vr.get("available"):
                                with st.expander(f"Velociraptor collection plan "
                                                 f"({_vr['stats']['artifacts']} artifacts · "
                                                 f"{_vr['stats']['total_vql']} VQL)", expanded=False):
                                    st.code(format_plan(_vr), language="sql")
                        except Exception as _vr_err:
                            st.caption(f"Velociraptor plan unavailable: {_vr_err}")
                        # Diamond Model of Intrusion Analysis — structures the
                        # incident into Adversary/Capability/Infrastructure/Victim
                        # (investigation skill; standalone, no soc_investigation_agent edits)
                        try:
                            from diamond_model import build_diamond, format_diamond, to_dot
                            _dm = build_diamond(_imap_src)
                            if _dm.get("available"):
                                with st.expander(f"Diamond Model "
                                                 f"({_dm['stats']['completeness_pct']}% complete)",
                                                 expanded=False):
                                    st.graphviz_chart(to_dot(_dm), width="stretch")
                                    st.code(format_diamond(_dm), language=None)
                        except Exception as _dm_err:
                            st.caption(f"Diamond Model unavailable: {_dm_err}")
                        # Proactive threat hunting + statistical anomalies
                        # (standalone module; needs a MITRE tactic on the incident)
                        try:
                            from threat_hunting import build_hunt_package, format_hunt
                            _pkg = build_hunt_package(_imap_src, None, str(DB_FILE))
                            if _pkg.get("available"):
                                st.markdown("** Proactive threat hunting**")
                                st.code(format_hunt(_pkg), language=None)
                        except Exception as _hunt_err:
                            st.caption(f"Threat hunting unavailable: {_hunt_err}")
                        # Incident Response SOP / runbook — the actionable
                        # containment→recovery procedure (standalone reporting-agent
                        # skill; also folded into the written report via
                        # skills_sidecar). Deterministic, guarded.
                        try:
                            from reporting_sop import build_incident_sop, format_sop
                            _sop = build_incident_sop(_imap_src)
                            if _sop.get("available"):
                                _sv = _sop.get("validation", {})
                                with st.expander(
                                    f"Response SOP ({_sop['stats']['steps']} steps · "
                                    f"{_sop['meta']['scenario']} · validation "
                                    f"{'PASS' if _sv.get('valid') else 'REVIEW'})",
                                    expanded=False):
                                    st.markdown(format_sop(_sop))
                        except Exception as _sop_err:
                            st.caption(f"Response SOP unavailable: {_sop_err}")
                        if _imap["timeline"]:
                            st.markdown("** Timeline**")
                            st.markdown("\n".join(
                                f"- `{t['time'][:19]}` — {t['event']}"
                                for t in _imap["timeline"][:12]))
                    except Exception as _map_err:
                        st.caption(f"Map unavailable: {_map_err}")
            if inc.get("alerts_fetch_error"):
                st.warning(_alerts_fetch_warning(inc))

            # Associated Alerts / Logs
            alerts_list = inc.get("alerts")
            if alerts_list:
                with st.expander(f"Associated Alerts ({len(alerts_list)})", expanded=False):
                    for alert in alerts_list:
                        a_title = alert.get("title") or alert.get("name") or "Untitled Alert"
                        a_id = alert.get("id") or ""
                        a_source = alert.get("source") or "Unknown"
                        a_type = alert.get("type") or "Unknown"
                        a_created = alert.get("created") or alert.get("receivedTime") or ""
                        st.markdown(
                            f'<div style="background:#091624;padding:8px 12px;border-radius:4px;margin-bottom:6px;border-left:3px solid var(--accent)">'
                            f'<div style="display:flex;justify-content:between;align-items:center">'
                            f'<strong>{a_title}</strong>'
                            f'<code style="color:var(--muted);font-size:0.7rem;margin-left:auto">{a_id}</code>'
                            f'</div>'
                            f'<div style="font-size:0.72rem;color:var(--muted);margin-top:4px">'
                            f'Incident ID: <strong style="color:var(--accent)">{inc_id}</strong> &nbsp;·&nbsp; '
                            f'Source: {a_source} &nbsp;·&nbsp; Type: {a_type} &nbsp;·&nbsp; Time: {a_created}'
                            f'</div>'
                            f'</div>',
                            unsafe_allow_html=True
                        )
                        # Nested event details
                        events = alert.get("events")
                        if events:
                            for idx, ev in enumerate(events):
                                ev_src = ev.get("source", {})
                                ev_dst = ev.get("destination", {})
                                src_ip = ev_src.get("device", {}).get("ipAddress") or ev_src.get("ipAddress") or "—"
                                dst_ip = ev_dst.get("device", {}).get("ipAddress") or ev_dst.get("ipAddress") or "—"
                                user = ev_src.get("user", {}).get("username") or ev_src.get("username") or "—"
                                ev_proto = ev.get("ip_proto") or ev.get("protocol") or "—"
                                ev_port = ev_dst.get("device", {}).get("port") or ev_dst.get("port") or "—"
                                st.markdown(
                                    f'<div style="margin-left:15px;padding:4px 8px;font-family:var(--mono);font-size:0.7rem;color:var(--muted);border-left:1px dashed #1B4A62">'
                                    f'Event {idx+1}: {user} | ➡ {src_ip} → {dst_ip} (Port {ev_port}, Proto {ev_proto})'
                                    f'</div>',
                                    unsafe_allow_html=True
                                )

        if shown == 0:
            st.info(f"No incidents match filter: {sev_filter}")


# ─────────────────────────────────────────────────────────────
# TAB 3 — CHAT
# ─────────────────────────────────────────────────────────────
with tab_chat:
    st.markdown(
        '<div class="info-box"><div class="title"> Ask a Question</div>'
        'You can ask plain-language questions about your security alerts here. '
        'For example: <em>"What are the most critical incidents today?"</em> or '
        '<em>"Summarise the latest high-priority alerts."</em> '
        'You do not need any technical knowledge to use this.</div>',
        unsafe_allow_html=True,
    )

    # ── Global CSS for the spinning phase icon (injected once, never reset) ──
    st.markdown("""
    <style>
    @keyframes soc-phase-spin { to { transform: rotate(360deg); } }
    .soc-phase-spinner {
        display: inline-block;
        animation: soc-phase-spin 1.2s linear infinite;
        color: #E8623A;
        font-size: 0.95rem;
        line-height: 1;
        margin-right: 8px;
        vertical-align: middle;
    }
    </style>
    """, unsafe_allow_html=True)

    # ── Helper: parse uploaded file into an incident dict ──────
    def _parse_uploaded_file(uploaded_file) -> tuple[dict, str]:
        """
        Parse a Streamlit UploadedFile into an incident dict.
        Supports: .json, .csv, .txt, .log
        Returns (incident_dict, error_message).  error_message is "" on success.
        """
        import io
        name = uploaded_file.name
        ext  = name.rsplit(".", 1)[-1].lower()
        raw  = uploaded_file.read()

        try:
            if ext == "json":
                data = _json.loads(raw.decode("utf-8", errors="replace"))
                # Accept a list → take first item; accept a dict → use directly
                if isinstance(data, list):
                    if not data:
                        return {}, "JSON array is empty."
                    incident = data[0] if isinstance(data[0], dict) else {"raw": data[0]}
                elif isinstance(data, dict):
                    # NetWitness envelope: {"items": [...]}
                    if "items" in data and isinstance(data["items"], list) and data["items"]:
                        incident = data["items"][0]
                    else:
                        incident = data
                else:
                    return {}, "JSON root must be an object or array."

            elif ext == "csv":
                import csv, io as _io
                text    = raw.decode("utf-8", errors="replace")
                reader  = list(csv.DictReader(_io.StringIO(text)))
                if not reader:
                    return {}, "CSV has no data rows."
                # Use the first row as the incident dict
                incident = dict(reader[0])
                # Attach all rows as an alert list for richer context
                incident["_csv_alerts"] = reader

            elif ext in ("txt", "log"):
                text = raw.decode("utf-8", errors="replace")
                # Build a minimal incident dict wrapping the raw text
                incident = {
                    "id":          name,
                    "title":       f"Uploaded log — {name}",
                    "description": text[:4000],   # cap to avoid token overflow
                    "raw_log":     text,
                    "source":      "file_upload",
                }

            else:
                return {}, f"Unsupported file type '.{ext}'. Upload a .json, .csv, .txt, or .log file."

        except Exception as exc:
            return {}, f"Failed to parse file: {exc}"

        # Multi-source alert normalization (defensive-security skill's
        # analyze_alert): a SIEM/EDR/NDR/syslog alert that isn't already
        # NetWitness-shaped (no alertMeta) gets normalized into the incident
        # schema — extracted IPs into alertMeta, a triage verdict, MITRE —
        # so the whole pipeline (map/TI/detection/triage) runs on it. Purely
        # additive: NetWitness-shaped uploads keep their alertMeta untouched.
        try:
            if not isinstance(incident.get("alertMeta"), dict) or not incident.get("alertMeta"):
                from alert_triage import normalize_to_incident, validate_alert
                # for a raw txt/log upload, feed the text as the message
                if ext in ("txt", "log") and "message" not in incident:
                    incident["message"] = incident.get("raw_log", "")[:8000]
                    incident.setdefault("timestamp", datetime.now().isoformat())
                    incident.setdefault("source", "file_upload")
                if validate_alert(incident)["ok"]:
                    ctx = ("edr" if ext in ("txt", "log") else "siem")
                    incident = normalize_to_incident(incident, ctx)
        except Exception:
            pass  # normalization is best-effort; never block an upload

        # Ensure there's always an id and title
        if "id" not in incident:
            incident["id"] = name
        if "title" not in incident:
            incident["title"] = name

        return incident, ""

    # ── Resolve which incident the agent will use ──────────────
    # Priority: NW incident (from Incidents tab) > uploaded file
    nw_inc       = st.session_state.chat_incident
    up_inc       = st.session_state.uploaded_incident
    active_inc   = nw_inc if nw_inc else up_inc   # what the agent receives

    # ── Context banner ─────────────────────────────────────────
    st.markdown(
        '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
        'letter-spacing:2px;margin-bottom:10px">■ INCIDENT CONTEXT</div>',
        unsafe_allow_html=True,
    )

    ctx_col, clr_col = st.columns([5, 1])

    with ctx_col:
        if nw_inc:
            sev = normalise_sev(nw_inc)
            st.markdown(
                f'<div style="background:#050F1A;border:1px solid var(--accent);'
                f'border-radius:6px;padding:10px 16px;font-family:var(--mono);'
                f'font-size:0.72rem">'
                f'<span class="badge badge-{sev.lower()}">{sev}</span>'
                f'&nbsp;'
                f'<strong>{nw_inc.get("id","?")}</strong> — {nw_inc.get("title","?")}'
                f'<span style="color:var(--muted);font-size:0.6rem;margin-left:10px">'
                f'● from NetWitness</span>'
                f'</div>',
                unsafe_allow_html=True,
            )
            if nw_inc.get("alerts_fetch_error"):
                st.warning(_alerts_fetch_warning(nw_inc))
        elif up_inc:
            st.markdown(
                f'<div style="background:#050F1A;border:1px solid var(--warn);'
                f'border-radius:6px;padding:10px 16px;font-family:var(--mono);'
                f'font-size:0.72rem">'
                f'<strong>{st.session_state.uploaded_filename}</strong>'
                f'<span style="color:var(--muted);font-size:0.6rem;margin-left:10px">'
                f'● from uploaded file</span>'
                f'</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                '<div style="background:#07080F;border:1px solid var(--border);'
                'border-radius:6px;padding:10px 16px;font-family:var(--mono);'
                'font-size:0.68rem;color:var(--muted)">'
                'No incident context — select one from the Incidents tab '
                'or upload a file below.'
                '</div>',
                unsafe_allow_html=True,
            )

    with clr_col:
        if nw_inc or up_inc:
            if st.button("✕ Clear", use_container_width=True, key="clear_ctx"):
                st.session_state.chat_incident    = None
                st.session_state.uploaded_incident = None
                st.session_state.uploaded_filename = ""
                st.rerun()

    st.markdown("---")

    # ── AGENT BOARD — live thinking + outputs for all 3 agents ─────────────
    st.markdown(
        '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
        'letter-spacing:2px;margin-bottom:8px">■ AGENT BOARD — CLICK AN AGENT '
        'TO SEE ITS THINKING &amp; OUTPUT</div>',
        unsafe_allow_html=True,
    )

    _AGENTS = [
        ("triage",        "", "Triage Agent",        "#00D4FF"),
        ("investigation", "", "Investigation Agent", "#FF7700"),
        ("reporting",     "", "Reporting Agent",     "#A78BFA"),
    ]
    _BOARD_BADGES = {
        "idle":    ("○ IDLE",    "#3A607A"),
        "queued":  ("QUEUED",  "#8B9DC3"),
        "running": ("RUNNING", "#FFB700"),
        "done":    ("DONE",    "#0AF0A0"),
        "cached":  ("CACHED",  "#0AF0A0"),
        "skipped": ("SKIPPED", "#3A607A"),
        "failed":  ("FAILED",  "#FF3B3B"),
    }

    # ── Background workflow sync: worker state → board, finalize once ──────
    _wf_active = _workflow_store().get("run")
    if _wf_active:
        # Adopt the worker's live panel dicts (same objects — subsequent
        # worker writes are visible on every poll rerun).
        for _ag in ("investigation", "reporting"):
            if _ag in _wf_active.get("panels", {}):
                st.session_state.agent_board[_ag] = _wf_active["panels"][_ag]
        if _wf_active.get("done") and not _wf_active.get("finalized"):
            _wf_active["finalized"] = True
            for _stage, _rec in _wf_active.get("chroma_queue", []):
                try:
                    pipeline_chroma_insert(_stage, _rec)
                except Exception:
                    pass
            st.session_state.chat_history.append(
                {"role": "assistant",
                 "content": "\n".join(_wf_active.get("wf_md") or
                                      ["Workflow finished."]),
                 "ts": datetime.now().strftime("%H:%M:%S")})
            st.session_state.setdefault("_surfaced_runs", []).append(
                _wf_active.get("run_id"))
            _workflow_store()["run"] = None
    else:
        # No active worker — sweep zombie states from interrupted sessions.
        for _ag in ("investigation", "reporting"):
            if st.session_state.agent_board[_ag]["status"] in ("running", "queued"):
                st.session_state.agent_board[_ag]["status"] = "failed"
                st.session_state.agent_board[_ag]["thinking"].append(
                    f"[{datetime.now().strftime('%H:%M:%S')}] previous run "
                    "was interrupted before completion")

    # ── Disk fallback: surface a completed run this session never saw ──────
    # The in-memory finalize above only works while the session's poll loop
    # stays alive for the whole run. The worker also persists its finished
    # results to disk; if a recent run (< 30 min) hasn't been surfaced in
    # THIS session, deliver its board panels + chat summary now.
    try:
        _lwr_path = SOC_DB_DIR / "last_workflow_result.json"
        if _lwr_path.exists():
            _lwr = _json.loads(_lwr_path.read_text(encoding="utf-8"))
            _surf = st.session_state.setdefault("_surfaced_runs", [])
            _fresh = (time.time() - float(_lwr.get("finished_at") or 0)) < 1800
            if (_lwr.get("done") and _fresh
                    and _lwr.get("run_id") not in _surf
                    and not (_workflow_store().get("run") or {})):
                for _ag in ("investigation", "reporting"):
                    _p = (_lwr.get("panels") or {}).get(_ag)
                    if _p:
                        st.session_state.agent_board[_ag] = _p
                st.session_state.chat_history.append(
                    {"role": "assistant",
                     "content": "\n".join(_lwr.get("wf_md") or
                                          ["Workflow finished."]),
                     "ts": datetime.now().strftime("%H:%M:%S")})
                _surf.append(_lwr.get("run_id"))
    except Exception:
        pass

    _ANSI_RE = re.compile(r"\x1b\[[0-9;]*m")

    def _board_set(agent: str, status: str | None = None,
                   think: str | None = None, output: str | None = None,
                   progress: int | None = None) -> None:
        """Update an agent's board state (+ live card/detail refresh mid-run)."""
        panel = st.session_state.agent_board[agent]
        if status is not None:
            panel["status"] = status
            if status in ("done", "cached", "skipped"):
                panel["progress"] = 100
        if progress is not None:
            panel["progress"] = max(0, min(100, int(progress)))
        if think is not None:
            clean = _ANSI_RE.sub("", str(think))   # subprocess logs carry colours
            panel["thinking"].append(f"[{datetime.now().strftime('%H:%M:%S')}] {clean}")
            panel["thinking"] = panel["thinking"][-60:]   # keep the tail
        if output is not None:
            panel["output"] = output
        panel["updated"] = datetime.now().strftime("%H:%M:%S")
        # Live refresh of this agent's card + open detail panel during a run.
        slot = _board_live.get(agent)
        if slot is not None:
            try:
                _render_board_card(slot, agent)
            except Exception:
                pass
        try:
            _render_board_detail(agent)
        except Exception:
            pass

    class _BoardTee:
        """Duck-typed st.empty() that mirrors writes to several containers —
        used to send the triage LLM token stream to both the in-status panel
        and the agent board's detail view."""
        def __init__(self, *targets):
            self.targets = [t for t in targets if t is not None]

        def markdown(self, *a, **k):
            for t in self.targets:
                try:
                    t.markdown(*a, **k)
                except Exception:
                    pass

        def empty(self):
            for t in self.targets:
                try:
                    t.empty()
                except Exception:
                    pass

    def _render_board_card(container, agent: str) -> None:
        icon, name, color = {a: (i, n, c) for a, i, n, c in _AGENTS}[agent]
        panel  = st.session_state.agent_board[agent]
        status = panel["status"]
        badge, bcolor = _BOARD_BADGES.get(status, _BOARD_BADGES["idle"])
        upd = panel.get("updated") or "—"
        tail = ""
        if status == "running" and panel["thinking"]:
            last = panel["thinking"][-1]
            tail = (f'<div style="font-family:var(--mono);font-size:0.52rem;'
                    f'color:var(--muted);margin-top:5px;white-space:nowrap;'
                    f'overflow:hidden;text-overflow:ellipsis">{last[-80:]}</div>')
        pct = int(panel.get("progress", 0) or 0)
        _bar = ""
        if status in ("running", "done", "cached") or pct:
            _bar = (f'<div style="height:3px;background:#12202f;border-radius:2px;'
                    f'margin-top:7px;overflow:hidden"><div style="width:{pct}%;'
                    f'height:100%;background:{color};border-radius:2px;'
                    f'transition:width .4s ease"></div></div>')
        container.markdown(
            f'<div style="background:#060C16;border:1px solid {color}44;'
            f'border-left:3px solid {color};border-radius:7px;'
            f'padding:10px 12px;min-height:74px">'
            f'<div style="display:flex;align-items:center;gap:8px">'
            f'<span style="font-size:1.15rem">{icon}</span>'
            f'<strong style="flex:1;font-size:0.8rem;color:{color}">{name}</strong>'
            f'<span style="background:{bcolor}22;color:{bcolor};'
            f'border:1px solid {bcolor}44;padding:1px 7px;border-radius:3px;'
            f'font-family:var(--mono);font-size:0.55rem">{badge}</span></div>'
            f'<div style="font-family:var(--mono);font-size:0.52rem;'
            f'color:var(--muted);margin-top:5px">last activity: {upd}'
            f'{(" · " + str(pct) + "%") if (status == "running" or pct) else ""}</div>'
            f'{_bar}{tail}</div>',
            unsafe_allow_html=True,
        )

    # HITL manual-review toggle — OFF (default) keeps the pipeline auto-chaining
    # exactly as before; ON pauses after triage and after investigation for the
    # analyst to review and click Approve before the next agent runs.
    # Mirror the toggle into a PLAIN session_state key (mr_pref). Widget-keyed
    # state (key="…") is garbage-collected when the widget isn't rendered as part
    # of the triggering interaction — which happens when the pipeline is kicked
    # off from ANOTHER tab (the Overview "Triage now" button, alerts tab, etc.):
    # the toggle key silently reset to False mid-flow, so a genuinely-ON toggle
    # read OFF and the pipeline auto-ran. A plain key we own is never GC'd, so it
    # survives cross-tab reruns. Everything downstream reads mr_pref, not a widget.
    st.session_state.setdefault("mr_pref", False)
    # Re-seed the widget key from our stable pref (setdefault is a no-op while the
    # widget key survives; it only re-seeds after Streamlit GC'd it on a cross-tab
    # rerun). mr_pref is mutated ONLY by the on_change callback below — i.e. only
    # on a REAL user toggle — so the dispatch always reads the true intent even if
    # the widget's own state was garbage-collected OR desynced by a cross-tab run.
    st.session_state.setdefault("mr_toggle", st.session_state.mr_pref)
    def _mr_sync():
        st.session_state.mr_pref = bool(st.session_state.get("mr_toggle"))
    st.toggle("Manual review — approve each hand-off",
              key="mr_toggle", on_change=_mr_sync,
              help="When on, the pipeline pauses after triage and after "
                   "investigation. Review each agent's output on this board, "
                   "then click Approve to hand off to the next agent.")
    if st.session_state.mr_pref:
        st.caption("Manual review is **armed** — the pipeline will pause after "
                   "triage and wait for your approval before each agent runs.")

    _board_live: dict = {}
    _b_cols = st.columns(3)
    for _bi, (_ag, _icon, _name, _color) in enumerate(_AGENTS):
        with _b_cols[_bi]:
            _slot = st.empty()
            _board_live[_ag] = _slot
            _render_board_card(_slot, _ag)
            if st.button("View", key=f"board_view_{_ag}", use_container_width=True):
                st.session_state.agent_board_sel = (
                    None if st.session_state.agent_board_sel == _ag else _ag)
                st.rerun()

    # ── HITL dispatch: decide auto-vs-manual for a freshly-triaged run. Runs on
    # a rerun AFTER the chat submission, so the toggle is reliably committed —
    # and it reads the SAME st.session_state.mr_pref the "armed" caption above
    # reads, so the decision can never disagree with what the user sees. ──
    _disp = _workflow_store().get("run")
    if (_disp and _disp.get("awaiting") == "dispatch"
            and not _disp.get("_spawned") and not _disp.get("done")):
        if st.session_state.get("mr_pref"):
            _disp["manual_review"] = True
            _disp["awaiting"] = ("investigate" if _disp.get("investigate")
                                 else "report_only")
        else:
            _disp["manual_review"] = False
            _dg = _disp.get("gate_report")
            if _dg is not None:
                _dg.set()          # auto mode — the worker never pauses
            _disp["_spawned"] = True
            _disp["awaiting"] = None
            import threading as _thd
            _thd.Thread(target=_workflow_worker,
                        args=(_disp, _disp.get("_tri"), _disp.get("_incident")),
                        daemon=True).start()
            st.rerun()

    # ── HITL approval controls (shown only when a run awaits the analyst) ──
    _hitl_run = _workflow_store().get("run")
    if _hitl_run and _hitl_run.get("manual_review") and not _hitl_run.get("done"):
        _await = _hitl_run.get("awaiting")
        if _await in ("investigate", "report_only") and not _hitl_run.get("_spawned"):
            if _await == "investigate":
                st.info("**Triage complete.** Review it above, then approve to "
                        "hand off to the Investigation agent.")
                _lbl = "Approve → Investigate"
            else:
                st.info("**Triage complete** (no investigation needed). Review it "
                        "above, then approve to generate the report.")
                _lbl = "Approve → Generate report"
            if st.button(_lbl, type="primary", use_container_width=True,
                         key="hitl_go_start"):
                import threading as _th2
                _hitl_run["_spawned"] = True
                _hitl_run["awaiting"] = None
                if _await == "report_only":
                    # No investigation to gate — this click IS the pre-report
                    # approval, so let the worker go straight to reporting.
                    _g = _hitl_run.get("gate_report")
                    if _g is not None:
                        _g.set()
                _th2.Thread(target=_workflow_worker,
                            args=(_hitl_run, _hitl_run["_tri"],
                                  _hitl_run["_incident"]),
                            daemon=True).start()
                st.rerun()
        elif _await == "report":
            st.info("**Investigation complete.** Review it above, then approve "
                    "to generate the final report.")
            if st.button("Approve → Generate report", type="primary",
                         use_container_width=True, key="hitl_go_rep"):
                _g = _hitl_run.get("gate_report")
                if _g is not None:
                    _g.set()
                _hitl_run["awaiting"] = None
                st.rerun()

    # ── Detail panel for the selected agent (live slots) ───────────────────
    # The thinking/output areas are st.empty() slots registered in
    # _board_live_detail; _board_set rewrites them mid-run, so the panel
    # updates in real time while the workflow executes further down the page.
    _board_live_detail: dict = {}

    def _render_board_detail(agent: str) -> None:
        slots = _board_live_detail.get(agent)
        if not slots:
            return
        panel = st.session_state.agent_board[agent]
        _acolor = {a: c for a, _i, _n, c in _AGENTS}.get(agent, "#00D4FF")
        if panel["thinking"]:
            # Chat-style live transcript: one row per thinking line, newest at the
            # bottom. Show the recent tail so the latest activity is always in view
            # (Streamlit strips <script>, so we can't auto-scroll a tall box).
            _rows = []
            for _ln in panel["thinking"][-16:]:
                _ts, _msg = "", _ln
                if _ln.startswith("[") and "]" in _ln:
                    _ts = _ln[1:_ln.index("]")]
                    _msg = _ln[_ln.index("]") + 1:].strip()
                _esc = (_msg.replace("&", "&amp;").replace("<", "&lt;")
                            .replace(">", "&gt;"))
                _rows.append(
                    f'<div style="display:flex;gap:9px;padding:5px 0;'
                    f'border-bottom:1px solid #0e1826">'
                    f'<span style="font-family:var(--mono);font-size:0.55rem;'
                    f'color:var(--faint);flex-shrink:0;min-width:52px">{_ts}</span>'
                    f'<span style="font-size:0.73rem;color:var(--text);'
                    f'line-height:1.45">{_esc}</span></div>')
            slots["think"].markdown(
                f'<div style="max-height:300px;overflow-y:auto;'
                f'padding:4px 12px;background:#070d16;border:1px solid {_acolor}33;'
                f'border-radius:9px;border-left:3px solid {_acolor}">'
                + "".join(_rows)
                + '</div>',
                unsafe_allow_html=True)
        else:
            slots["think"].caption("No activity yet — run a triage to see "
                                   "this agent think.")
        if panel["output"]:
            slots["out"].markdown(panel["output"], unsafe_allow_html=True)
        else:
            slots["out"].caption("No output yet.")

    # Auto-follow: surface the currently-running agent's live thinking without a
    # manual click (keeps the "live LLM chat" feel as work moves between agents).
    if st.session_state.agent_board_sel is None:
        for _ag2, _i2, _n2, _c2 in _AGENTS:
            if st.session_state.agent_board[_ag2]["status"] == "running":
                st.session_state.agent_board_sel = _ag2
                break

    _sel_ag = st.session_state.agent_board_sel
    if _sel_ag:
        _panel = st.session_state.agent_board[_sel_ag]
        _icon, _name, _color = {a: (i, n, c) for a, i, n, c in _AGENTS}[_sel_ag]
        st.markdown(
            f'<div style="font-family:var(--mono);font-size:0.6rem;color:{_color};'
            f'letter-spacing:2px;margin:10px 0 4px">{_icon} {_name.upper()} — DETAIL</div>',
            unsafe_allow_html=True,
        )
        with st.expander("Thinking process", expanded=True):
            _think_slot = st.empty()
            _token_slot = st.empty()   # live LLM token stream (triage phases)
        with st.expander("Output", expanded=bool(_panel["output"])):
            _out_slot = st.empty()
        _board_live_detail[_sel_ag] = {"think": _think_slot,
                                       "token": _token_slot, "out": _out_slot}
        _render_board_detail(_sel_ag)

    # ── File uploader (kept — tucked below the board) ──────────────────────
    with st.expander("Upload incident file (JSON · CSV · TXT · LOG)"):
        uploaded_file = st.file_uploader(
            "Upload incident file",
            type=["json", "csv", "txt", "log"],
            label_visibility="collapsed",
            help="Upload a JSON export, CSV log, or plain-text log file to use as incident context.",
        )

        if uploaded_file is not None:
            # Only re-parse if it's a new file
            if uploaded_file.name != st.session_state.uploaded_filename:
                parsed_inc, err = _parse_uploaded_file(uploaded_file)
                if err:
                    st.error(f"{err}")
                else:
                    st.session_state.uploaded_incident = parsed_inc
                    st.session_state.uploaded_filename  = uploaded_file.name
                    st.session_state.chat_incident      = None   # clear NW incident
                    st.rerun()

        # Preview of parsed file fields
        if st.session_state.uploaded_incident and not nw_inc:
            preview_inc = st.session_state.uploaded_incident
            preview_keys = [k for k in preview_inc if not k.startswith("_")][:12]
            preview_html = " &nbsp;·&nbsp; ".join(
                f'<span style="color:var(--accent)">{k}</span>'
                f':<span style="color:var(--text)"> '
                f'{str(preview_inc[k])[:40]}</span>'
                for k in preview_keys
            )
            st.markdown(
                f'<div style="background:#060C16;border:1px solid var(--border);'
                f'border-radius:5px;padding:8px 12px;font-family:var(--mono);'
                f'font-size:0.62rem;margin:6px 0;line-height:1.8">'
                f'Parsed fields: {preview_html}'
                f'</div>',
                unsafe_allow_html=True,
            )
            # Multi-source triage verdict when the upload was normalized from a
            # SIEM/EDR/NDR/log alert (defensive-security skill's analyze_alert)
            _av = preview_inc.get("_analyze_alert")
            if isinstance(_av, dict) and _av.get("classification") not in (None, "invalid"):
                _tp = "true positive" if _av["is_true_positive"] else "needs review"
                st.markdown(
                    f'<div style="background:#0A1A10;border:1px solid #1E5A38;'
                    f'border-radius:5px;padding:8px 12px;font-family:var(--mono);'
                    f'font-size:0.62rem;margin:6px 0;line-height:1.7">'
                    f'<b>Alert triage</b> ({preview_inc.get("_source_format","?").upper()} '
                    f'source): {_av["classification"]} · '
                    f'severity <b>{_av["severity"]}</b> · {_tp}'
                    + (" · MITRE " + ", ".join(m["technique"] for m in _av.get("mitre", [])[:4])
                       if _av.get("mitre") else "")
                    + '</div>',
                    unsafe_allow_html=True,
                )
            # (checkbox, not an expander — Streamlit forbids nesting expanders)
            if st.checkbox("View full parsed incident JSON", key="upl_json_view"):
                st.json({k: v for k, v in preview_inc.items() if k != "raw_log"})

    st.markdown("---")

    # ── Trigger hint ───────────────────────────────────────────
    if active_inc:
        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.6rem;color:var(--muted);'
            'margin-bottom:6px"> Type <strong style="color:var(--accent)">triage</strong>, '
            '<strong style="color:var(--accent)">ioc</strong>, '
            '<strong style="color:var(--accent)">classify</strong> or '
            '<strong style="color:var(--accent)">ticket</strong> to run the full triage pipeline.'
            '</div>',
            unsafe_allow_html=True,
        )

    # ── Chat input  (called first — always fixed at page bottom by Streamlit)
    user_input = st.chat_input("Ask the SOC agent…")

    # The "Triage" button sets this flag instead of the user typing a
    # message — synthesize the trigger word so it flows through the exact
    # same pipeline below, with no manual typing required.
    _triage_restart = False
    if not user_input and st.session_state.pending_auto_triage and active_inc:
        st.session_state.pending_auto_triage = False
        user_input = "Triage this incident"

    # Interaction-interrupt recovery: triage runs INLINE (kept there for live
    # token streaming), so ANY click mid-run — e.g. opening the board's View
    # panel to watch the thinking — kills the script run silently and the
    # board freezes at "Triage started" forever (observed live 2026-07-17
    # 20:30). The in-flight marker set when triage starts survives into this
    # rerun; if it's still here, the previous run died mid-triage → restart.
    elif not user_input and isinstance(st.session_state.get("triage_in_flight"), dict) and active_inc:
        _tif = st.session_state.triage_in_flight
        if str(active_inc.get("id") or "") != _tif.get("incident_id"):
            st.session_state.triage_in_flight = None   # different incident — stale marker
        elif _tif.get("attempts", 1) < 3:
            _tif["attempts"] = _tif.get("attempts", 1) + 1
            user_input = "Triage this incident"
            _triage_restart = True
        else:
            st.session_state.triage_in_flight = None
            _board_set("triage", status="failed",
                       think="triage was interrupted repeatedly — click Triage to run it again",
                       output="Triage did not complete: the run was interrupted "
                              "repeatedly by page interactions. Click Triage to rerun.")

    # Append user message immediately so it shows in the history render below
    # (not on auto-restart — the original message is already in the history)
    if user_input and not _triage_restart:
        now = datetime.now().strftime("%H:%M:%S")
        st.session_state.chat_history.append(
            {"role": "user", "content": user_input, "ts": now}
        )

    # ── Chat history ───────────────────────────────────────────
    st.markdown(
        '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
        'letter-spacing:2px;margin-bottom:8px">■ CONVERSATION</div>',
        unsafe_allow_html=True,
    )

    if not st.session_state.chat_history:
        st.markdown(
            '<div style="text-align:center;padding:40px;font-family:var(--mono);'
            'font-size:0.78rem;color:var(--muted)">'
            'SOC TRIAGE AGENT READY<br>'
            '<span style="font-size:0.62rem">'
            'Upload a file or select an incident, then type a message below</span></div>',
            unsafe_allow_html=True,
        )

    for msg in st.session_state.chat_history:
        ts = msg.get("ts", "")
        if msg["role"] == "user":
            st.markdown(
                f'<div class="bubble-user">'
                f'<div class="bubble-label" style="color:var(--muted)">YOU · {ts}</div>'
                f'{msg["content"]}</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                f'<div class="bubble-agent">'
                f'<div class="bubble-label">SOC AGENT · {ts}</div>'
                f'{msg["content"]}</div>',
                unsafe_allow_html=True,
            )

    # ── Thinking containers — in conversation flow, above fixed input ──────────
    # ── Agent execution ────────────────────────────────────────────────────────
    if user_input:

        reply = "No response was generated."   # safe default

        # ── Triage trigger ─────────────────────────────────────────────────────
        if active_inc and _TRIAGE_TRIGGER.search(user_input):

            ALL_PHASES = [
                "IOC — Availability",
                "IOC — Confidentiality",
                "IOC — Integrity",
                "Risk Rating",
                "SOC Classification",
            ]
            PHASE_DESC = {
                "IOC — Availability":
                    "Scanning availability indicators — CPU spikes, reboots, buffer saturation",
                "IOC — Confidentiality":
                    "Scanning confidentiality indicators — exfil, geo anomalies, permission changes",
                "IOC — Integrity":
                    "Scanning integrity indicators — unknown binaries, hash anomalies, processes",
                "Risk Rating":
                    "Calculating risk across initiation, occurrence & adverse impact dimensions",
                "SOC Classification":
                    "Classifying incident severity and generating recommended actions",
            }

            # Receives the structured triage result (metakeys_payload + ticket)
            # so the downstream investigation/reporting agents can consume it.
            _triage_sink: dict = {}

            # Fresh board run: clear the triage panel and go live.
            st.session_state.agent_board["triage"].update(
                {"thinking": [], "output": ""})
            _board_set("triage", status="running", think="Triage started")
            # In-flight marker — survives an interaction-interrupted run so
            # the recovery branch above can restart triage on the next rerun.
            _tif_cur = st.session_state.get("triage_in_flight")
            if not (isinstance(_tif_cur, dict) and
                    _tif_cur.get("incident_id") == str(active_inc.get("id") or "")):
                st.session_state.triage_in_flight = {
                    "incident_id": str(active_inc.get("id") or ""), "attempts": 1}
            if _triage_restart:
                _board_set("triage", think="previous run was interrupted by a "
                                           "page interaction — restarted automatically")

            # st.status() + st.write() is Streamlit's guaranteed progressive-update
            # pattern — st.write() calls inside the with-block appear live as each
            # phase fires; status.update(label=…) changes the header to the current phase.
            with st.status("Initialising triage pipeline…", expanded=True) as triage_status:

                thinking_panel = st.empty()   # token-by-token stream inside the status
                # Mirror the token stream into the agent board's detail view
                # (live "thinking" like the original single-agent experience).
                thinking_tee = _BoardTee(
                    thinking_panel,
                    (_board_live_detail.get("triage") or {}).get("token"))

                _tri_prog = {"done": 0}   # completed triage phases → progress %

                def on_progress(event: str, label: str, text: str = "") -> None:
                    key = next(
                        (p for p in ALL_PHASES
                         if label == p or label in p or p in label),
                        label,
                    )
                    if event == "phase_start":
                        desc = PHASE_DESC.get(key, label)
                        triage_status.update(label=f"{desc}…", expanded=True)
                        thinking_tee.empty()
                        _board_set("triage", think=f"▶ {desc}")
                    elif event == "phase_complete":
                        result = f" — {text}" if text else ""
                        st.write(f"**{key}**{result}")
                        thinking_tee.empty()
                        _tri_prog["done"] += 1
                        _board_set("triage", think=f"{key}{result}",
                                   progress=round(_tri_prog["done"]
                                                  / max(1, len(ALL_PHASES)) * 100))

                try:
                    reply = soc_triage_chat_respond(
                        user_input,
                        incident           = active_inc,
                        llm_config         = get_cisco_cfg(),
                        progress_fn        = on_progress,
                        thinking_container = thinking_tee,
                        result_sink        = _triage_sink,
                    )
                    if not reply:
                        reply = "Triage returned an empty response."
                    triage_status.update(
                        label="Triage complete", state="complete", expanded=False
                    )
                    _board_set(
                        "triage",
                        status=("cached" if (_triage_sink.get("result") or {})
                                .get("cached") else "done"),
                        think="Triage complete", output=reply)
                    st.session_state.triage_in_flight = None
                except Exception as exc:
                    _es = str(exc)
                    if "503" in _es or "unavailable" in _es.lower():
                        reply = ("**The LLM endpoint is asleep (HTTP 503).** "
                                 "Your HuggingFace endpoint scales to zero when "
                                 "idle and needs a few minutes to boot — this "
                                 "request has started waking it. **Wait 2–5 "
                                 "minutes and send `triage` again.** (Your "
                                 "token is fine: 503 means the server isn't "
                                 "running, not an auth problem.)")
                    elif "401" in _es or "unauthorized" in _es.lower():
                        reply = ("**The LLM endpoint rejected the "
                                 "credentials (HTTP 401).** Check the token "
                                 "and endpoint URL in the sidebar LLM "
                                 "settings, and that the endpoint still "
                                 "exists in your HuggingFace console.")
                    else:
                        reply = f"Triage error: {exc}"
                    triage_status.update(
                        label="Triage failed", state="error", expanded=True
                    )
                    _board_set("triage", status="failed",
                               think=f"{str(exc)[:150]}", output=reply)
                    # real failure (error already shown) — no auto-restart
                    st.session_state.triage_in_flight = None

        # ── Plain Q&A fallback ─────────────────────────────────────────────────
        else:
            try:
                with st.spinner("Agent thinking…"):
                    reply = chat_respond(user_input, incident=active_inc)
                if not reply:
                    reply = "Agent returned an empty response."
            except Exception as exc:
                reply = f"Error: {exc}"

        # ── Sequential agent workflow: triage → investigation → reporting ──
        # Driven by the STRUCTURED triage result (not keyword-matching the
        # rendered reply). Stage records land in the pipeline DB as each
        # agent completes, matching the Pipeline DB tab's 6 stages.
        if active_inc and _TRIAGE_TRIGGER.search(user_input):
            _inc_id  = str(active_inc.get("id") or active_inc.get("incidentId") or "")
            _title   = str(active_inc.get("title") or active_inc.get("name") or "Untitled")
            _sev     = normalise_sev(active_inc)
            _summary = reply[:500]

            pipeline_insert_full("alerts_to_triage", {
                "id": _inc_id or f"alert_{now.replace(':','')}",
                "incident_id": _inc_id, "title": _title,
                "severity": _sev, "summary": _summary})

            _tri = _triage_sink.get("result") if WORKFLOW_OK else None
            if _tri and not _tri.get("error"):
                # Sub-agents must inherit the SAME live credentials triage
                # just used (sidebar/session state) — the .env copy can be
                # stale or base64-encoded, which surfaces as LLM 401s.
                _cfg_live = get_cisco_cfg()
                if _cfg_live.api_key and _cfg_live.api_key != "changeme":
                    os.environ["CISCO_LLM_URL"]   = _cfg_live.base_url
                    os.environ["CISCO_LLM_KEY"]   = _cfg_live.api_key
                    os.environ["CISCO_LLM_MODEL"] = _cfg_live.model
                _ticket = _tri["ticket"]
                _cls    = _ticket.get("classification") or _sev
                _unc    = _ticket.get("unc") or f"#TKT_{_inc_id[:6]}"
                _wf_md  = ["", "---", "", "## Agent Workflow"]
                _run_started = datetime.now()
                _run_stamp   = _run_started.strftime("%Y%m%d-%H%M%S")
                if _tri.get("cached"):
                    _wf_md.append("- Triage served from cache — MITRE mapping "
                                  "and metakey extraction reflect the *original* "
                                  "run. Type **retriage** for a fresh analysis "
                                  "with the latest agent upgrades.")
                _n_alerts_reported = int(active_inc.get("alertCount") or 0)
                _n_alerts_attached = len(active_inc.get("alerts") or [])
                if _n_alerts_reported and not _n_alerts_attached:
                    _fetch_err = active_inc.get("alerts_fetch_error")
                    _wf_md.append(f"- NetWitness reports "
                                  f"{_n_alerts_reported} alert(s) for this "
                                  f"incident but none were attached"
                                  + (f" (alerts fetch: {_fetch_err})"
                                     if _fetch_err else "")
                                  + " — usernames/hosts may be missing from "
                                    "the investigation.")

                pipeline_insert_full("initial_ticket", {
                    "id": _unc, "incident_id": _inc_id,
                    "title": f"Ticket {_unc} — {_title}",
                    "severity": _cls, "summary": _ticket.get("summary") or ""})

                # ── Stages 2+3: hand off to the BACKGROUND WORKER ───────────
                # Inline execution died whenever the user interacted mid-run
                # (Streamlit kills the script at its next UI call). The
                # worker thread survives clicks/refreshes; the board polls it.
                import threading as _threading
                _investigate = wf_needs_investigation(_tri)
                if _investigate:
                    pipeline_insert_full("post_triage_investigate", {
                        "id": f"inv_{_inc_id}", "incident_id": _inc_id,
                        "title": _title, "severity": _cls, "summary": _summary})
                else:
                    pipeline_insert_full("post_triage_no_investigate", {
                        "id": f"noinv_{_inc_id}", "incident_id": _inc_id,
                        "title": _title, "severity": _cls, "summary": _summary})
                pipeline_insert_full("pending_ticket_report", {
                    "id": f"pending_{_unc}", "incident_id": _inc_id,
                    "title": f"[PENDING] {_title}", "severity": _cls,
                    "summary": "Handed off to reporting agent."})

                _run_rec = {
                    "run_id": _run_stamp, "incident_id": _inc_id,
                    "title": _title, "cls": _cls, "unc": _unc,
                    "investigate": _investigate,
                    "started_ts": _run_started.timestamp(),
                    "started_hms": _run_started.strftime("%H:%M:%S"),
                    "cached_triage": bool(_tri.get("cached")),
                    "done": False, "finalized": False,
                    "panels": {
                        "investigation": {"status": "queued", "thinking": [],
                                          "output": "", "updated": ""},
                        "reporting":     {"status": "queued", "thinking": [],
                                          "output": "", "updated": ""},
                    },
                    "wf_md": _wf_md,
                    "chroma_queue": [],
                }
                # HITL: DO NOT decide auto-vs-manual here. Reading the toggle on
                # the same rerun that submitted the chat is racy (the toggle's
                # committed state can lag the submission, so a genuinely-ON toggle
                # read as OFF and the pipeline auto-spawned). Instead we store the
                # run as "dispatch" (pending) with NO worker spawned; the Agent
                # Board's dispatch step decides on the NEXT rerun — reading the
                # SAME session_state the "armed" caption reads, so they always
                # agree. See the dispatch block in the board section.
                _run_rec["gate_report"] = _threading.Event()   # clear; set by dispatch/approval
                _run_rec["_tri"]        = _tri
                _run_rec["_incident"]   = active_inc
                _run_rec["awaiting"]    = "dispatch"
                _workflow_store()["run"] = _run_rec
                reply += ("\n\n---\n\n**Triage complete.** See the **Agent Board** "
                          "above — if Manual review is on it will wait for your "
                          "approval, otherwise investigation & reporting run "
                          "automatically.")
            else:
                # Triage failed or workflow module unavailable — record the
                # alert only; downstream agents need a valid triage result.
                pipeline_insert_full("pending_ticket_report", {
                    "id": f"pending_{_inc_id}", "incident_id": _inc_id,
                    "title": f"[PENDING] {_title}", "severity": _sev,
                    "summary": "Triage did not produce a structured result; "
                               "workflow not started."})

        st.session_state.chat_history.append(
            {"role": "assistant", "content": reply,
             "ts": datetime.now().strftime("%H:%M:%S")}
        )
        st.rerun()

    if st.session_state.chat_history:
        if st.button("Clear Chat", key="clear_chat"):
            st.session_state.chat_history = []
            st.rerun()


# ─────────────────────────────────────────────────────────────
# TAB 4 — CHROMADB
# ─────────────────────────────────────────────────────────────
with tab_chroma:
    st.markdown(
        '<div class="info-box"><div class="title"> Knowledge Base</div>'
        'This is where the AI stores its understanding of your security alerts. '
        'Once synced, the AI can answer questions much more accurately. '
        'Use the search box below to find specific incidents by description.</div>',
        unsafe_allow_html=True,
    )
    if st.session_state.chroma_col is None:
        st.warning("The Knowledge Base isn't connected yet. Connect it from the left panel under Knowledge Base.")
    else:
        col = st.session_state.chroma_col
        st.markdown(f"### Knowledge Base — {col.count()} incidents indexed")
        st.markdown("---")

        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
            'letter-spacing:2px;margin-bottom:8px">■ SEMANTIC SEARCH</div>',
            unsafe_allow_html=True,
        )
        sq, sn = st.columns([5,1])
        query = sq.text_input("Query", placeholder="e.g. ransomware lateral movement C2",
                               label_visibility="collapsed")
        top_n = sn.number_input("N", 1, 20, 5, label_visibility="collapsed")

        if st.button("Search"):
            st.session_state.search_results = chroma_search(query, n=top_n)

        for r in st.session_state.search_results:
            m = r["meta"]
            st.markdown(
                f'<div style="background:#060C16;border:1px solid var(--border);'
                f'border-radius:5px;padding:10px 14px;margin:4px 0;'
                f'font-family:var(--mono);font-size:0.7rem">'
                f'<span class="badge badge-info">{r["score"]}% match</span>'
                f'&nbsp;<strong>{r["id"]}</strong>'
                f'&nbsp;<span style="color:var(--muted)">'
                f'sev:{m.get("severity","?")} · {m.get("status","?")}</span><br>'
                f'<span style="font-size:0.67rem;color:var(--text)">'
                f'{r["text"][:220]}…</span></div>',
                unsafe_allow_html=True,
            )

        st.markdown("---")
        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
            'letter-spacing:2px;margin-bottom:8px">■ ACTIONS</div>',
            unsafe_allow_html=True,
        )
        a1, a2 = st.columns(2)

        if a1.button("⬆ Sync Incidents", use_container_width=True):
            if not incidents:
                st.warning("No incidents loaded yet.")
            else:
                n, msg = chroma_sync(incidents)
                st.success(msg) if n else st.error(msg)
                st.rerun()

        if a2.button("Wipe & Reset", use_container_width=True):
            try:
                st.session_state.chroma_client.delete_collection("soc_incidents")
                chroma_connect("./chroma_db")
                st.success("Collection wiped and recreated.")
                st.rerun()
            except Exception as e:
                st.error(str(e))

        st.markdown("---")
        st.markdown(
            '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
            'letter-spacing:2px;margin-bottom:8px">■ LANGCHAIN WIRING</div>',
            unsafe_allow_html=True,
        )
        st.code("""
# Replace chat_respond() body with your chain:

from langchain_community.vectorstores import Chroma
from langchain_anthropic import ChatAnthropic
from langchain.chains import RetrievalQA

vectorstore = Chroma(
    client=st.session_state.chroma_client,
    collection_name="soc_incidents",
    embedding_function=your_embedder,   # must match sync-time embedder
)
llm   = ChatAnthropic(model="claude-sonnet-4-20250514")
chain = RetrievalQA.from_chain_type(
    llm=llm,
    retriever=vectorstore.as_retriever(search_kwargs={"k": 5}),
)
return chain.invoke(user_msg)["result"]
        """, language="python")


# ─────────────────────────────────────────────────────────────
# TAB 5 — LOG HISTORY  (permanent SQLite store)
# ─────────────────────────────────────────────────────────────
with tab_log:
    st.markdown(
        '<div class="info-box"><div class="title"> History</div>'
        'All security alerts that have ever been loaded are saved here permanently, '
        'even after a restart. You can search, filter, export to Excel, or generate '
        'a report for any incident.</div>',
        unsafe_allow_html=True,
    )
    stats = db_stats()

    # ── Summary stats ──────────────────────────────────────────
    st.markdown(
        '<div style="font-size:0.75rem;font-weight:600;color:var(--muted);'
        'text-transform:uppercase;letter-spacing:1px;margin-bottom:12px"> Summary</div>',
        unsafe_allow_html=True,
    )
    ls1, ls2, ls3, ls4 = st.columns(4)
    ls1.metric("Total Logged",    stats["total"])
    ls2.metric("Total Fetches",   stats["fetches"])
    ls3.metric("Critical (ever)", stats["by_sev"].get("CRITICAL", 0))
    ls4.metric("High (ever)",     stats["by_sev"].get("HIGH", 0))

    st.markdown("---")

    # ── Filters ────────────────────────────────────────────────
    st.markdown(
        '<div style="font-family:var(--mono);font-size:0.65rem;color:var(--muted);'
        'letter-spacing:2px;margin-bottom:10px">■ FILTER & SEARCH</div>',
        unsafe_allow_html=True,
    )
    hf1, hf2, hf3, hf4 = st.columns([1.2, 1.2, 2, 1])

    hist_sev    = hf1.selectbox("Severity", ["ALL","CRITICAL","HIGH","MEDIUM","LOW"],
                                 key="hist_sev", label_visibility="collapsed")
    # Collect unique statuses from DB for filter
    with db_connect() as _con:
        _statuses = ["ALL"] + [
            r[0] for r in _con.execute(
                "SELECT DISTINCT status FROM incidents ORDER BY status"
            ).fetchall() if r[0]
        ]
    hist_status = hf2.selectbox("Status", _statuses,
                                 key="hist_status", label_visibility="collapsed")
    hist_search = hf3.text_input("Search title / ID / assignee", "",
                                  placeholder="Type to filter…",
                                  label_visibility="collapsed")
    hist_limit  = hf4.number_input("Limit", 10, 5000, 200,
                                    label_visibility="collapsed")

    # ── Load rows ──────────────────────────────────────────────
    log_rows = db_load_incidents(
        severity=hist_sev,
        status=hist_status,
        search=hist_search,
        limit=hist_limit,
    )

    st.markdown(
        f'<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
        f'margin:8px 0">{len(log_rows)} records shown</div>',
        unsafe_allow_html=True,
    )

    # ── Export button ──────────────────────────────────────────
    csv_data = db_export_csv()
    if csv_data:
        st.download_button(
            label="Export all as CSV",
            data=csv_data,
            file_name=f"soc_incidents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv",
        )

    st.markdown("---")

    # ── Incident rows ──────────────────────────────────────────
    if not log_rows:
        st.markdown(
            '<div style="text-align:center;padding:50px;font-family:var(--mono);'
            'font-size:0.78rem;color:var(--muted)">'
            '● NO RECORDS IN DATABASE<br>'
            '<span style="font-size:0.62rem">'
            'Incidents are saved automatically on every fetch</span></div>',
            unsafe_allow_html=True,
        )
    else:
        for row in log_rows:
            sev   = row.get("severity", "LOW")
            color = SEV_COLORS.get(sev, "#3A607A")
            inc_id    = row.get("id", "—")
            title     = row.get("title", "Untitled")
            status    = row.get("status", "—")
            assignee  = row.get("assignee", "—")
            alerts    = row.get("alert_count", "—")
            created   = str(row.get("created", ""))[:16]
            first_seen= str(row.get("first_seen",""))[:16]
            last_seen = str(row.get("last_seen", ""))[:16]

            st.markdown(
                f'<div class="card card-{sev.lower()}">'
                f'<div style="display:flex;align-items:center;gap:10px;flex-wrap:wrap">'
                f'<span class="badge badge-{sev.lower()}">{sev}</span>'
                f'<strong style="flex:1;font-size:0.88rem">{title}</strong>'
                f'<span style="color:var(--muted);font-size:0.62rem">Incident ID:</span>'
                f'<code style="color:var(--muted);font-size:0.68rem">{inc_id}</code>'
                f'</div>'
                f'<div style="margin-top:8px;font-size:0.73rem;color:var(--muted);'
                f'display:flex;gap:16px;flex-wrap:wrap">'
                f'<span> {status}</span>'
                f'<span> {assignee}</span>'
                f'<span> {alerts} alerts</span>'
                f'<span> created {created}</span>'
                f'<span style="color:#1A4A62">first seen {first_seen}</span>'
                f'<span style="color:#1A4A62">last seen {last_seen}</span>'
                f'</div></div>',
                unsafe_allow_html=True,
            )

            bj, bm, _ = st.columns([0.7, 0.9, 7.4])
            if bj.button("{ }", key=f"log_json_{inc_id}"):
                try:
                    raw = _json.loads(row.get("raw_json") or "{}")
                except Exception:
                    raw = row
                with st.expander(f"Full JSON — {inc_id}", expanded=True):
                    st.json(raw)
            # Incident Map — Stage 1 entity graph + Stage 2 corpus expansion
            if bm.button("Map", key=f"log_map_{inc_id}"):
                with st.expander(f"Incident Map — {inc_id}", expanded=True):
                    try:
                        from incident_map import build_incident_map, to_dot, map_caption
                        from incident_expansion import LocalCorpusSource, expand_incident_map
                        _imap_src = _json.loads(row.get("raw_json") or "{}")
                        # Pre-triage panel light-up (same as the live tab) —
                        # infer MITRE from stored evidence so skills activate.
                        try:
                            from tactic_inference import augment_incident as _tac_aug
                            _imap_src, _tac_note = _tac_aug(_imap_src)
                        except Exception:
                            _tac_note = None
                        if _tac_note:
                            st.caption(f"{_tac_note}")
                        _imap_key = "log"
                        _imap = build_incident_map(_imap_src)
                        try:
                            with LocalCorpusSource(str(DB_FILE)) as _src:
                                expand_incident_map(_imap, _src)
                        except Exception as _exp_err:
                            _imap.setdefault("expansion_log", []).append(
                                f"expansion unavailable: {_exp_err}")
                        st.graphviz_chart(to_dot(_imap), width="stretch")
                        st.caption(map_caption(_imap))
                        if _imap.get("expansion_log"):
                            st.markdown("** Autonomous expansion**")
                            st.markdown("\n".join(f"- {l}" for l in _imap["expansion_log"]))
                        if _imap.get("endpoint_profile_text"):
                            st.markdown("** Endpoint profile**")
                            st.code(_imap["endpoint_profile_text"], language=None)
                        # Unified triage verdict — capstone that aggregates the
                        # triage-side skills (base severity + asset criticality +
                        # internal IOC correlation) into ONE prioritized verdict.
                        # Reads the other skills' outputs; edits none of them.
                        try:
                            from triage_verdict import aggregate_verdict, format_verdict
                            _tv = aggregate_verdict(_imap_src)
                            if _tv.get("available"):
                                st.markdown(f"** Unified triage verdict — "
                                            f"{_tv['level']} (priority {_tv['priority']}/5)**")
                                st.code(format_verdict(_tv), language=None)
                        except Exception as _tv_err:
                            st.caption(f"Unified verdict unavailable: {_tv_err}")
                        # Asset criticality (instant, deterministic)
                        try:
                            from asset_criticality import assess_incident, format_assessment
                            _ac = assess_incident(_imap_src)
                            st.markdown("** Asset criticality**")
                            st.code(format_assessment(_ac), language=None)
                        except Exception as _ac_err:
                            st.caption(f"Asset criticality unavailable: {_ac_err}")
                        # Threat-intel enrichment (live IOC lookups — opt-in
                        # since it hits the network the first time per IOC)
                        try:
                            if st.checkbox("Enrich IOCs (threat intel)",
                                           key=f"ti_{_imap_key}_{_imap_src.get('id','')}"):
                                from threat_intel import enrich_iocs, format_enrichment
                                _enr = enrich_iocs(_imap_src)
                                if _enr["results"] or _enr["stats"]["skipped_private_ips"]:
                                    st.code(format_enrichment(_enr), language=None)
                                else:
                                    st.caption("No external IOCs to enrich.")
                        except Exception as _ti_err:
                            st.caption(f"Threat intel unavailable: {_ti_err}")
                        # Recommended detection — deterministic Sigma rule
                        # from the incident's indicators (no LLM, no TI lookup
                        # so it stays instant); guarded independently.
                        try:
                            from detection_rules import build_detection
                            from asset_criticality import assess_incident
                            _asset = assess_incident(_imap_src)
                            _det = build_detection(_imap_src, asset=_asset)
                            if _det["has_selection"]:
                                st.markdown("** Recommended detection (Sigma)**")
                                st.code(_det["sigma_yaml"], language="yaml")
                                st.caption(f"Splunk: {_det['splunk_spl']}")
                                # detection-as-code: validate + Elastic EQL (3rd SIEM)
                                try:
                                    from detection_engineering import validate_sigma, to_elastic_eql
                                    _v = validate_sigma(_det["sigma"])
                                    st.caption(("detection-as-code: all required fields present"
                                                if _v["valid"] else
                                                "" + "; ".join(_v["missing_fields"] + _v["warnings"])))
                                    st.caption(f"Elastic EQL: {to_elastic_eql(_det)}")
                                except Exception:
                                    pass
                                if _det["d3fend"]:
                                    st.caption("D3FEND: " + ", ".join(
                                        f"{c['id']} {c['name']}" for c in _det["d3fend"]))
                        except Exception as _det_err:
                            st.caption(f"Detection unavailable: {_det_err}")
                        # ATT&CK detection coverage (program-level capability)
                        try:
                            from detection_engineering import assess_attack_coverage, format_coverage
                            _acov = assess_attack_coverage(str(DB_FILE))
                            with st.expander(f"ATT&CK detection coverage "
                                             f"({_acov['covered']}/{_acov['tactics_total']} "
                                             f"tactics, {_acov['coverage_pct']}%)", expanded=False):
                                st.code(format_coverage(_acov), language=None)
                        except Exception as _cov_err:
                            st.caption(f"Coverage unavailable: {_cov_err}")
                        # Recommended mitigations & coverage (threat->control)
                        try:
                            from mitigation_mapping import build_mitigation_coverage, format_mitigation
                            _cov = build_mitigation_coverage(_imap_src, asset=_asset)
                            if _cov.get("available"):
                                st.markdown("** Recommended mitigations & coverage**")
                                st.code(format_mitigation(_cov), language=None)
                        except Exception as _mit_err:
                            st.caption(f"Mitigation map unavailable: {_mit_err}")
                        # Internal IOC correlation — frequency/severity/case
                        # status of each IOC across the local corpus + case
                        # pipeline -> HIGH/MEDIUM/LOW/NONE confidence. Deterministic,
                        # read-only, self-locating DBs; ubiquity-guarded.
                        try:
                            from ioc_correlation import correlate_iocs, format_correlation
                            _corr = correlate_iocs(_imap_src)
                            if _corr.get("available") and _corr.get("results"):
                                st.markdown("** Internal IOC correlation**")
                                st.code(format_correlation(_corr), language=None)
                        except Exception as _corr_err:
                            st.caption(f"IOC correlation unavailable: {_corr_err}")
                        # osquery investigation pack — per-incident, MITRE-mapped
                        # osquery queries to RUN on the affected host (investigation
                        # skill; standalone, deterministic, no soc_investigation_agent edits)
                        try:
                            from osquery_investigation import build_investigation_pack, format_pack
                            _oq = build_investigation_pack(_imap_src)
                            if _oq.get("available"):
                                with st.expander(f"osquery investigation pack "
                                                 f"({_oq['stats']['total_queries']} queries · "
                                                 f"{_oq['platform']})", expanded=False):
                                    st.code(format_pack(_oq), language="sql")
                        except Exception as _oq_err:
                            st.caption(f"osquery pack unavailable: {_oq_err}")
                        # Velociraptor collection & hunt plan — named DFIR
                        # artifacts + VQL to collect on the affected host
                        # (investigation skill; standalone, no soc_investigation_agent edits)
                        try:
                            from velociraptor_investigation import build_collection_plan, format_plan
                            _vr = build_collection_plan(_imap_src)
                            if _vr.get("available"):
                                with st.expander(f"Velociraptor collection plan "
                                                 f"({_vr['stats']['artifacts']} artifacts · "
                                                 f"{_vr['stats']['total_vql']} VQL)", expanded=False):
                                    st.code(format_plan(_vr), language="sql")
                        except Exception as _vr_err:
                            st.caption(f"Velociraptor plan unavailable: {_vr_err}")
                        # Diamond Model of Intrusion Analysis — structures the
                        # incident into Adversary/Capability/Infrastructure/Victim
                        # (investigation skill; standalone, no soc_investigation_agent edits)
                        try:
                            from diamond_model import build_diamond, format_diamond, to_dot
                            _dm = build_diamond(_imap_src)
                            if _dm.get("available"):
                                with st.expander(f"Diamond Model "
                                                 f"({_dm['stats']['completeness_pct']}% complete)",
                                                 expanded=False):
                                    st.graphviz_chart(to_dot(_dm), width="stretch")
                                    st.code(format_diamond(_dm), language=None)
                        except Exception as _dm_err:
                            st.caption(f"Diamond Model unavailable: {_dm_err}")
                        # Proactive threat hunting + statistical anomalies
                        # (standalone module; needs a MITRE tactic on the incident)
                        try:
                            from threat_hunting import build_hunt_package, format_hunt
                            _pkg = build_hunt_package(_imap_src, None, str(DB_FILE))
                            if _pkg.get("available"):
                                st.markdown("** Proactive threat hunting**")
                                st.code(format_hunt(_pkg), language=None)
                        except Exception as _hunt_err:
                            st.caption(f"Threat hunting unavailable: {_hunt_err}")
                        # Incident Response SOP / runbook — the actionable
                        # containment→recovery procedure (standalone reporting-agent
                        # skill; also folded into the written report via
                        # skills_sidecar). Deterministic, guarded.
                        try:
                            from reporting_sop import build_incident_sop, format_sop
                            _sop = build_incident_sop(_imap_src)
                            if _sop.get("available"):
                                _sv = _sop.get("validation", {})
                                with st.expander(
                                    f"Response SOP ({_sop['stats']['steps']} steps · "
                                    f"{_sop['meta']['scenario']} · validation "
                                    f"{'PASS' if _sv.get('valid') else 'REVIEW'})",
                                    expanded=False):
                                    st.markdown(format_sop(_sop))
                        except Exception as _sop_err:
                            st.caption(f"Response SOP unavailable: {_sop_err}")
                        if _imap["timeline"]:
                            st.markdown("** Timeline**")
                            st.markdown("\n".join(
                                f"- `{t['time'][:19]}` — {t['event']}"
                                for t in _imap["timeline"][:12]))
                    except Exception as _map_err:
                        st.caption(f"Map unavailable: {_map_err}")
            
            # Associated Alerts / Logs
            try:
                raw_dict = _json.loads(row.get("raw_json") or "{}")
            except Exception:
                raw_dict = {}
            alerts_list = raw_dict.get("alerts")
            if alerts_list:
                with st.expander(f"Associated Alerts ({len(alerts_list)})", expanded=False):
                    for alert in alerts_list:
                        a_title = alert.get("title") or alert.get("name") or "Untitled Alert"
                        a_id = alert.get("id") or ""
                        a_source = alert.get("source") or "Unknown"
                        a_type = alert.get("type") or "Unknown"
                        a_created = alert.get("created") or alert.get("receivedTime") or ""
                        st.markdown(
                            f'<div style="background:#091624;padding:8px 12px;border-radius:4px;margin-bottom:6px;border-left:3px solid var(--accent)">'
                            f'<div style="display:flex;justify-content:between;align-items:center">'
                            f'<strong>{a_title}</strong>'
                            f'<code style="color:var(--muted);font-size:0.7rem;margin-left:auto">{a_id}</code>'
                            f'</div>'
                            f'<div style="font-size:0.72rem;color:var(--muted);margin-top:4px">'
                            f'Incident ID: <strong style="color:var(--accent)">{inc_id}</strong> &nbsp;·&nbsp; '
                            f'Source: {a_source} &nbsp;·&nbsp; Type: {a_type} &nbsp;·&nbsp; Time: {a_created}'
                            f'</div>'
                            f'</div>',
                            unsafe_allow_html=True
                        )
                        # Nested event details
                        events = alert.get("events")
                        if events:
                            for idx, ev in enumerate(events):
                                ev_src = ev.get("source", {})
                                ev_dst = ev.get("destination", {})
                                src_ip = ev_src.get("device", {}).get("ipAddress") or ev_src.get("ipAddress") or "—"
                                dst_ip = ev_dst.get("device", {}).get("ipAddress") or ev_dst.get("ipAddress") or "—"
                                user = ev_src.get("user", {}).get("username") or ev_src.get("username") or "—"
                                ev_proto = ev.get("ip_proto") or ev.get("protocol") or "—"
                                ev_port = ev_dst.get("device", {}).get("port") or ev_dst.get("port") or "—"
                                st.markdown(
                                    f'<div style="margin-left:15px;padding:4px 8px;font-family:var(--mono);font-size:0.7rem;color:var(--muted);border-left:1px dashed #1B4A62">'
                                    f'Event {idx+1}: {user} | ➡ {src_ip} → {dst_ip} (Port {ev_port}, Proto {ev_proto})'
                                    f'</div>',
                                    unsafe_allow_html=True
                                )

# ═══════════════════════════════════════════════════════════════════════════════
# TAB 6 — PIPELINE DB
# Fully inline — single streamlit run app.py, no second process.
# st.stop() removed from tab_chroma so this tab always renders.
# ═══════════════════════════════════════════════════════════════════════════════
with tab_pipeline:
    st.markdown(
        '<div class="info-box"><div class="title"> Data Pipeline</div>'
        'This tool helps IT staff manage how security data flows through the system — '
        'from collection, to review, to archiving. If you are not sure what this is for, '
        'you likely do not need to use it.</div>',
        unsafe_allow_html=True,
    )

    # ── session defaults ──────────────────────────────────────────────────────
    if "pl_stage"      not in st.session_state: st.session_state.pl_stage      = None
    if "pl_cv_results" not in st.session_state: st.session_state.pl_cv_results = []

    st.markdown(
        '<div style="font-family:var(--mono);font-size:0.65rem;color:var(--muted);'
        'letter-spacing:2px;margin-bottom:14px">'
        '■ SOC PIPELINE — CLICK SELECT UNDER ANY STAGE TO OPEN ITS DB VIEWER</div>',
        unsafe_allow_html=True,
    )

    # ══════════════════════════════════════════════════════════════════════════
    # STAGE CARDS  (always visible — one column per stage)
    # ══════════════════════════════════════════════════════════════════════════
    _pl_cols = st.columns(len(PIPELINE_STAGES))
    for _idx, _stage in enumerate(PIPELINE_STAGES):
        _cnt    = pipeline_count(_stage)
        _icon   = PIPELINE_ICONS[_stage]
        _label  = PIPELINE_LABELS[_stage]
        _color  = PIPELINE_COLORS[_stage]
        _active = st.session_state.pl_stage == _stage
        _border = f"3px solid {_color}" if _active else f"1px solid {_color}44"
        _bg     = f"{_color}18"          if _active else "#060C16"
        _last = pipeline_last_write(_stage)
        with _pl_cols[_idx]:
            st.markdown(
                f'<div style="background:{_bg};border:{_border};border-radius:7px;'
                f'padding:10px 8px;text-align:center;transition:all 0.15s">'
                f'<div style="font-size:1.3rem">{_icon}</div>'
                f'<div style="font-family:var(--mono);font-size:1.5rem;'
                f'color:{_color};margin:4px 0">{_cnt}</div>'
                f'<div style="font-family:var(--mono);font-size:0.48rem;'
                f'color:var(--muted);letter-spacing:1px;line-height:1.4">'
                f'{_label.upper()}</div>'
                f'<div style="font-family:var(--mono);font-size:0.5rem;'
                f'color:{_color}99;margin-top:3px">last: {_last}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
            if st.button("Select", key=f"pl_sel_{_stage}", use_container_width=True):
                st.session_state.pl_stage      = None if _active else _stage
                st.session_state.pl_cv_results = []
                st.rerun()

    # ══════════════════════════════════════════════════════════════════════════
    # INLINE VIEWER  (expands below when a stage is selected)
    # ══════════════════════════════════════════════════════════════════════════
    _sel = st.session_state.pl_stage

    if not _sel:
        st.markdown(
            '<div style="text-align:center;padding:60px 20px;font-family:var(--mono);'
            'font-size:0.78rem;color:var(--muted)">'
            '↑ Click <strong style="color:var(--accent)">Select</strong> '
            'under any stage above to open its database viewer<br>'
            '<span style="font-size:0.62rem;margin-top:8px;display:block">'
            'Records auto-insert when the triage agent runs in the Chat tab.'
            '</span></div>',
            unsafe_allow_html=True,
        )
    else:
        _color = PIPELINE_COLORS[_sel]
        _icon  = PIPELINE_ICONS[_sel]
        _label = PIPELINE_LABELS[_sel]

        st.markdown("<br>", unsafe_allow_html=True)

        # ── header banner ─────────────────────────────────────────────────────
        st.markdown(
            f'<div style="background:#060C16;border:2px solid {_color};'
            f'border-left:8px solid {_color};border-radius:8px;'
            f'padding:18px 24px;margin-bottom:16px">'
            f'<div style="font-family:var(--mono);font-size:1.05rem;'
            f'color:{_color};letter-spacing:3px">{_icon} {_label.upper()}</div>'
            f'<div style="font-family:var(--mono);font-size:0.58rem;'
            f'color:var(--muted);margin-top:6px">'
            f'ChromaDB collection: <code style="color:var(--accent)">pipeline_{_sel}</code>'
            f'&nbsp;·&nbsp;SQLite table: <code style="color:var(--accent)">{_sel}</code>'
            f'</div></div>',
            unsafe_allow_html=True,
        )

        # ── metrics ───────────────────────────────────────────────────────────
        _sql_cnt = pipeline_count(_sel)
        _vec_cnt = pipeline_chroma_count(_sel)
        _vm1, _vm2, _vm3, _vm4 = st.columns(4)
        _vm1.metric("SQLite Records",   _sql_cnt)
        _vm2.metric("ChromaDB Vectors", _vec_cnt)
        _vm3.metric("Stage",            _label[:22])
        _vm4.metric("Collection",       f"pipeline_{_sel}"[:24])

        # ── action strip ──────────────────────────────────────────────────────
        _ac1, _ac2, _ac3, _ac4, _ = st.columns([1, 1.4, 1, 1.1, 2.5])

        if _ac1.button("Test Insert", key=f"pl_add_{_sel}"):
            import uuid as _uuid3
            pipeline_insert_full(_sel, {
                "id":          f"test_{str(_uuid3.uuid4())[:8]}",
                "incident_id": "DEMO-001",
                "title":       f"[Demo] Test record · {_label}",
                "severity":    "MEDIUM",
                "summary":     f"Auto-inserted test record into: {_label}",
            })
            st.success("Test record inserted!")
            st.rerun()

        if _sel == "pending_ticket_report":
            if _ac2.button("Finalize All", key="pl_finalize_btn"):
                _pending = pipeline_load("pending_ticket_report")
                for _prow in _pending:
                    _fin = dict(_prow)
                    _fin["id"]      = f"final_{_prow['id']}"
                    _fin["title"]   = (_fin.get("title") or "").replace("[PENDING]","[FINAL]")
                    _fin["summary"] = "Finalized. " + (_prow.get("summary") or "")
                    pipeline_insert_full("finalized_report", _fin)
                st.success(f"Finalized {len(_pending)} records → Finalized Report.")
                st.rerun()

        if _ac3.button("Clear Stage", key=f"pl_clear_{_sel}"):
            with _pl_con() as _c:
                _c.execute(f"DELETE FROM {_sel}")
                _c.commit()
            _wipe_col = _pl_chroma_col(_sel)
            if _wipe_col:
                try:
                    st.session_state.chroma_client.delete_collection(f"pipeline_{_sel}")
                except Exception:
                    pass
            st.success(f"Cleared {_label}.")
            st.rerun()

        if _ac4.button("✕ Close Viewer", key="pl_close_viewer"):
            st.session_state.pl_stage      = None
            st.session_state.pl_cv_results = []
            st.rerun()

        st.markdown("---")

        # ── sub-tabs ──────────────────────────────────────────────────────────
        _vtab_sql, _vtab_chroma = st.tabs(["SQLITE RECORDS", "CHROMADB SEARCH"])

        # ════════════════════ SQLite Records ════════════════════
        with _vtab_sql:
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;'
                'color:var(--muted);letter-spacing:2px;margin-bottom:10px">'
                '■ ALL RECORDS IN THIS STAGE</div>',
                unsafe_allow_html=True,
            )
            _fc1, _fc2 = st.columns([5, 1])
            _sql_srch = _fc1.text_input(
                "filter", placeholder="Filter by title / ID / summary…",
                key=f"sql_srch_{_sel}", label_visibility="collapsed")
            _sql_lim = int(_fc2.number_input(
                "limit", 10, 2000, 200,
                key=f"sql_lim_{_sel}", label_visibility="collapsed"))

            _rows = pipeline_load(_sel, limit=_sql_lim)
            if _sql_srch.strip():
                _s = _sql_srch.lower()
                _rows = [r for r in _rows if (
                    _s in (r.get("title")   or "").lower() or
                    _s in (r.get("id")      or "").lower() or
                    _s in (r.get("summary") or "").lower())]

            if not _rows:
                st.markdown(
                    '<div style="text-align:center;padding:50px;font-family:var(--mono);'
                    'font-size:0.78rem;color:var(--muted)">● NO RECORDS FOUND<br>'
                    '<span style="font-size:0.62rem">'
                    'Run a triage in the Chat tab, or click Test Insert above.'
                    '</span></div>',
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(
                    f'<div style="font-family:var(--mono);font-size:0.6rem;'
                    f'color:var(--muted);margin-bottom:8px">{len(_rows)} records</div>',
                    unsafe_allow_html=True,
                )
                # ── Aegis queue table (My Queue mockup) — compact default view.
                # Detail cards (JSON / CSV / DOCX / delete per record) stay fully
                # available behind the checkbox below; nothing was removed.
                try:
                    _qrows = [[
                        {"mono": str(_r.get("id", "—"))[:28]},
                        str(_r.get("title") or "Untitled")[:70],
                        {"mono": str(_r.get("incident_id") or "—")},
                        {"pill": str(_r.get("severity") or "—").upper(),
                         "kind": _ui.sev_class(_r.get("severity") or "")},
                        {"pill": _label, "kind": "stage"},
                        str(_r.get("created_at") or "")[:16],
                    ] for _r in _rows[:100]]
                    st.markdown(_ui.queue_table(
                        ["Record", "Case", "Incident", "Severity", "Stage", "Created"],
                        _qrows), unsafe_allow_html=True)
                    if len(_rows) > 100:
                        st.caption(f"Table shows the first 100 of {len(_rows)} records — "
                                   "use the filter above or the detail cards below.")
                except Exception:
                    pass
                _show_detail = st.checkbox(
                    "Show detail cards (JSON · exports · delete per record)",
                    value=False, key=f"pl_detail_{_sel}")
                _rows_detail = _rows if _show_detail else []
                for _row in _rows_detail:
                    _r_id  = _row.get("id", "—")
                    _r_ttl = _row.get("title", "Untitled")
                    _r_sev = str(_row.get("severity", "—")).upper()
                    _r_inc = _row.get("incident_id", "—")
                    _r_sum = (_row.get("summary") or "")[:200]
                    _r_ts  = str(_row.get("created_at", ""))[:16]
                    _sev_c = SEV_COLORS.get(_r_sev, "#3A607A")

                    # ── Stage-specific export badge ─────────────────────────
                    _is_csv_stage  = _sel in ("post_triage_investigate", "post_triage_no_investigate")
                    _is_docx_stage = _sel == "initial_ticket"
                    _is_report_stage = _sel == "finalized_report"
                    _is_postinv_stage = _sel == "post_investigation"
                    # Post-investigation records carry the analysis narrative
                    # (markdown) produced by the investigation agent.
                    _postinv_md = ""
                    if _is_postinv_stage:
                        try:
                            _rj_inv = _json.loads(_row.get("raw_json") or "{}")
                            _inv_j  = _rj_inv.get("investigation") or {}
                            _postinv_md = _inv_j.get("narrative_report") or ""
                            if not _postinv_md:
                                _md_path = (_inv_j.get("artifacts") or {}).get("report_markdown")
                                if _md_path and Path(str(_md_path)).exists():
                                    _postinv_md = Path(str(_md_path)).read_text(encoding="utf-8")
                        except Exception:
                            _postinv_md = ""
                    # Finalized reports carry real Word/PDF files generated by
                    # the reporting agent — resolve their paths from raw_json.
                    _report_exports = {}
                    if _is_report_stage:
                        try:
                            _rj_exp = _json.loads(_row.get("raw_json") or "{}")
                            _report_exports = ((_rj_exp.get("report") or {})
                                               .get("document_exports") or {})
                        except Exception:
                            _report_exports = {}
                        _report_exports = {
                            k: v for k, v in _report_exports.items()
                            if k in ("docx", "pdf") and v and Path(str(v)).exists()
                        }
                    _export_badge  = ""
                    if _is_csv_stage:
                        _export_badge = (
                            f'<span style="background:#00403A;color:#0AF0A0;'
                            f'border:1px solid #0AF0A044;padding:2px 8px;'
                            f'border-radius:3px;font-family:var(--mono);font-size:0.56rem;'
                            f'margin-left:6px"> CSV</span>'
                        )
                    elif _is_docx_stage:
                        _export_badge = (
                            f'<span style="background:#1A0040;color:#A78BFA;'
                            f'border:1px solid #A78BFA44;padding:2px 8px;'
                            f'border-radius:3px;font-family:var(--mono);font-size:0.56rem;'
                            f'margin-left:6px"> DOCX</span>'
                        )
                    elif _is_report_stage and _report_exports:
                        _fmt_txt = " · ".join(k.upper() for k in ("docx", "pdf")
                                              if k in _report_exports)
                        _export_badge = (
                            f'<span style="background:#1A0040;color:#A78BFA;'
                            f'border:1px solid #A78BFA44;padding:2px 8px;'
                            f'border-radius:3px;font-family:var(--mono);font-size:0.56rem;'
                            f'margin-left:6px"> {_fmt_txt}</span>'
                        )
                    elif _is_postinv_stage and _postinv_md:
                        _export_badge = (
                            f'<span style="background:#04342C;color:#2DD4BF;'
                            f'border:1px solid #2DD4BF44;padding:2px 8px;'
                            f'border-radius:3px;font-family:var(--mono);font-size:0.56rem;'
                            f'margin-left:6px"> REPORT</span>'
                        )

                    st.markdown(
                        f'<div style="background:#060C16;border:1px solid {_color}44;'
                        f'border-left:3px solid {_color};border-radius:6px;'
                        f'padding:10px 14px;margin:4px 0">'
                        f'<div style="display:flex;align-items:center;gap:10px;flex-wrap:wrap">'
                        f'<code style="font-family:var(--mono);font-size:0.65rem;color:{_color}">'
                        f'{_r_id}</code>'
                        f'<strong style="flex:1;font-size:0.85rem">{_r_ttl}</strong>'
                        f'{_export_badge}'
                        f'<span style="background:{_sev_c}22;color:{_sev_c};'
                        f'border:1px solid {_sev_c}44;padding:2px 8px;border-radius:3px;'
                        f'font-family:var(--mono);font-size:0.6rem">{_r_sev}</span>'
                        f'<span style="font-family:var(--mono);font-size:0.56rem;'
                        f'color:var(--muted)">{_r_ts}</span>'
                        f'</div>'
                        f'<div style="font-size:0.7rem;color:var(--muted);margin-top:6px">'
                        f'Incident: <code style="color:var(--accent)">{_r_inc}</code>'
                        f'{"&nbsp;·&nbsp;" + _r_sum if _r_sum else ""}'
                        f'</div></div>',
                        unsafe_allow_html=True,
                    )

                    # ── Button row: JSON | export | delete ──────────────────
                    _bj2b = None
                    if _is_csv_stage:
                        _bj1, _bj2, _bj3, _ = st.columns([0.6, 1.2, 0.28, 7])
                    elif _is_docx_stage:
                        _bj1, _bj2, _bj3, _ = st.columns([0.6, 1.3, 0.28, 7])
                    elif _is_report_stage and _report_exports:
                        _bj1, _bj2, _bj2b, _bj3, _ = st.columns([0.6, 1.1, 1.1, 0.28, 6])
                    elif _is_postinv_stage and _postinv_md:
                        _bj1, _bj2, _bj2b, _bj3, _ = st.columns([0.6, 1.1, 1.3, 0.28, 6])
                    else:
                        _bj1, _bj3, _ = st.columns([0.6, 0.28, 9])
                        _bj2 = None

                    if _bj1.button("{ } JSON", key=f"pl_json_{_sel}_{_r_id}"):
                        try:
                            _raw_j = _json.loads(_row.get("raw_json") or "{}")
                        except Exception:
                            _raw_j = _row
                        with st.expander(f"JSON — {_r_id}", expanded=True):
                            st.json(_raw_j)

                    # ── CSV download for post-triage stages ─────────────────
                    if _is_csv_stage and _bj2 is not None:
                        _safe_id = re.sub(r"[^A-Za-z0-9_\-]", "_", _r_id)[:40]
                        _csv_bytes = _make_csv_bytes(_row)
                        _bj2.download_button(
                            label="View as Sheet",
                            data=_csv_bytes,
                            file_name=f"triage_{_safe_id}.csv",
                            mime="text/csv",
                            key=f"pl_csv_{_sel}_{_r_id}",
                        )

                    # ── DOCX download for initial_ticket stage ──────────────
                    elif _is_docx_stage and _bj2 is not None:
                        _safe_id = re.sub(r"[^A-Za-z0-9_\-]", "_", _r_id)[:40]
                        _docx_bytes = _make_docx_bytes(_row)
                        # Detect if python-docx is available (bytes start with PK zip magic)
                        _is_real_docx = _docx_bytes[:2] == b'PK'
                        _bj2.download_button(
                            label="View as Ticket",
                            data=_docx_bytes,
                            file_name=f"ticket_{_safe_id}.{'docx' if _is_real_docx else 'txt'}",
                            mime=(
                                "application/vnd.openxmlformats-officedocument"
                                ".wordprocessingml.document"
                                if _is_real_docx else "text/plain"
                            ),
                            key=f"pl_docx_{_sel}_{_r_id}",
                        )

                    # ── Word/PDF downloads for finalized_report stage ───────
                    # Serves the actual documents generated by the reporting
                    # agent (outputs/<incident>/reports/exports/), not a
                    # reconstruction from the pipeline record.
                    elif _is_report_stage and _report_exports and _bj2 is not None:
                        _safe_id = re.sub(r"[^A-Za-z0-9_\-]", "_", _r_id)[:40]
                        if _report_exports.get("docx"):
                            try:
                                _bj2.download_button(
                                    label="Word Report",
                                    data=Path(str(_report_exports["docx"])).read_bytes(),
                                    file_name=f"incident_report_{_safe_id}.docx",
                                    mime=("application/vnd.openxmlformats-officedocument"
                                          ".wordprocessingml.document"),
                                    key=f"pl_repdocx_{_sel}_{_r_id}",
                                )
                            except Exception:
                                pass
                        if _report_exports.get("pdf") and _bj2b is not None:
                            try:
                                _bj2b.download_button(
                                    label="PDF Report",
                                    data=Path(str(_report_exports["pdf"])).read_bytes(),
                                    file_name=f"incident_report_{_safe_id}.pdf",
                                    mime="application/pdf",
                                    key=f"pl_reppdf_{_sel}_{_r_id}",
                                )
                            except Exception:
                                pass

                    # ── Investigation findings for post_investigation stage ──
                    elif _is_postinv_stage and _postinv_md and _bj2 is not None:
                        _safe_id = re.sub(r"[^A-Za-z0-9_\-]", "_", _r_id)[:40]
                        _bj2.download_button(
                            label="Report (MD)",
                            data=_postinv_md.encode("utf-8"),
                            file_name=f"investigation_{_safe_id}.md",
                            mime="text/markdown",
                            key=f"pl_invmd_{_sel}_{_r_id}",
                        )
                        if _bj2b is not None and _bj2b.button(
                                "View Findings", key=f"pl_invview_{_sel}_{_r_id}"):
                            with st.expander(f"Investigation findings — {_r_id}",
                                             expanded=True):
                                # ✦ AI-summary card (Phase 4a) — flags an
                                # error/fallback report with the amber tag so a
                                # failed LLM run is visible at a glance.
                                try:
                                    _fb = _ui.detect_fallback(_postinv_md)
                                    _first = next((ln.strip() for ln in _postinv_md.splitlines()
                                                   if ln.strip() and not ln.lstrip().startswith("#")),
                                                  "Investigation report")
                                    st.markdown(_ui.ai_summary(_first[:400], _fb,
                                                title="Investigation report"),
                                                unsafe_allow_html=True)
                                except Exception:
                                    pass
                                st.markdown(_postinv_md)

                    if _bj3.button("✕", key=f"pl_del_{_sel}_{_r_id}"):
                        pipeline_delete(_sel, _r_id)
                        st.rerun()

        # ════════════════════ ChromaDB Search ════════════════════
        with _vtab_chroma:
            _chroma_ok = CHROMA_OK and bool(st.session_state.get("chroma_client"))

            if not _chroma_ok:
                st.markdown(
                    '<div style="background:#0A0608;border:1px solid #2A1010;'
                    'border-radius:6px;padding:12px 16px;font-family:var(--mono);'
                    'font-size:0.68rem;margin-bottom:12px">'
                    '<span style="color:var(--warn)"> ChromaDB not connected</span><br>'
                    '<span style="color:var(--muted);font-size:0.6rem">'
                    'Connect ChromaDB from the sidebar. SQLite records persist regardless.'
                    '</span></div>',
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(
                    f'<div style="font-family:var(--mono);font-size:0.65rem;'
                    f'color:var(--muted);margin-bottom:10px">'
                    f'{pipeline_chroma_count(_sel)} vectors in '
                    f'<code style="color:var(--accent)">pipeline_{_sel}</code></div>',
                    unsafe_allow_html=True,
                )

            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
                'letter-spacing:2px;margin-bottom:8px">■ SEMANTIC SEARCH</div>',
                unsafe_allow_html=True,
            )
            _sq2, _sn2 = st.columns([5, 1])
            _cv_q = _sq2.text_input(
                "query", key=f"cv_q_{_sel}",
                placeholder="e.g. ransomware lateral movement C2",
                label_visibility="collapsed")
            _cv_n = int(_sn2.number_input(
                "topn", 1, 20, 5,
                key=f"cv_n_{_sel}", label_visibility="collapsed"))

            if st.button("Search", key=f"cv_srch_{_sel}"):
                if not _chroma_ok:
                    st.warning("Connect ChromaDB from the sidebar first.")
                elif not _cv_q.strip():
                    st.warning("Enter a query.")
                else:
                    st.session_state.pl_cv_results = pipeline_chroma_search(
                        _sel, _cv_q, n=_cv_n)

            for _r in st.session_state.pl_cv_results:
                _rm   = _r["meta"]
                _scol = PIPELINE_COLORS.get(_rm.get("stage", _sel), _color)
                st.markdown(
                    f'<div style="background:#060C16;border:1px solid {_scol}55;'
                    f'border-left:3px solid {_scol};border-radius:6px;'
                    f'padding:10px 14px;margin:4px 0">'
                    f'<div style="display:flex;align-items:center;gap:10px">'
                    f'<span style="background:{_scol}22;color:{_scol};'
                    f'border:1px solid {_scol}44;padding:2px 8px;border-radius:3px;'
                    f'font-family:var(--mono);font-size:0.6rem">{_r["score"]}%</span>'
                    f'<strong style="flex:1">{_r["id"]}</strong>'
                    f'<span style="font-family:var(--mono);font-size:0.58rem;'
                    f'color:var(--muted)">'
                    f'sev:{_rm.get("severity","?")} · {_rm.get("created","")[:10]}'
                    f'</span></div>'
                    f'<div style="font-size:0.72rem;color:var(--text);margin-top:6px">'
                    f'{str(_r["doc"])[:300]}</div></div>',
                    unsafe_allow_html=True,
                )

            st.markdown("---")
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
                'letter-spacing:2px;margin-bottom:8px">■ BROWSE ALL VECTORS</div>',
                unsafe_allow_html=True,
            )
            if st.button("Load All Vectors", key=f"cv_all_{_sel}"):
                if not _chroma_ok:
                    st.warning("Connect ChromaDB from the sidebar first.")
                else:
                    _all_v = pipeline_chroma_all(_sel)
                    if not _all_v:
                        st.info("No vectors in this collection yet.")
                    for _v in _all_v:
                        _vm  = _v["meta"]
                        _vc  = PIPELINE_COLORS.get(_vm.get("stage", _sel), _color)
                        st.markdown(
                            f'<div style="background:#060C16;border:1px solid {_vc}44;'
                            f'border-left:3px solid {_vc};border-radius:6px;'
                            f'padding:8px 14px;margin:3px 0">'
                            f'<div style="display:flex;gap:10px;align-items:center">'
                            f'<code style="font-family:var(--mono);font-size:0.62rem;'
                            f'color:{_vc}">{_v["id"]}</code>'
                            f'<span style="font-family:var(--mono);font-size:0.56rem;'
                            f'color:var(--muted)">'
                            f'sev:{_vm.get("severity","?")} · {_vm.get("created","")[:10]}'
                            f'</span></div>'
                            f'<div style="font-size:0.7rem;color:var(--text);margin-top:5px">'
                            f'{str(_v["doc"])[:250]}</div></div>',
                            unsafe_allow_html=True,
                        )

            st.markdown("---")
            st.markdown(
                '<div style="font-family:var(--mono);font-size:0.62rem;color:var(--muted);'
                'letter-spacing:2px;margin-bottom:8px">■ COLLECTION ACTIONS</div>',
                unsafe_allow_html=True,
            )
            if st.button(f"Wipe Chroma: pipeline_{_sel}", key=f"cv_wipe_{_sel}"):
                if not _chroma_ok:
                    st.warning("Connect ChromaDB from the sidebar first.")
                else:
                    try:
                        st.session_state.chroma_client.delete_collection(f"pipeline_{_sel}")
                        _pl_chroma_col(_sel)
                        st.success(f"Wiped pipeline_{_sel}.")
                        st.rerun()
                    except Exception as _we:
                        st.error(str(_we))

# ══════════════════════════════════════════════════════════════════════════════
# BACKGROUND WORKFLOW POLL — last statement so every tab renders first.
# While the worker thread runs, rerun every ~1.5s so the agent board (and any
# open detail panel) refreshes with the worker's latest thinking/output.
# ══════════════════════════════════════════════════════════════════════════════
try:
    _wf_poll = _workflow_store().get("run")
    if _wf_poll is not None:
        if not _wf_poll.get("done"):
            time.sleep(1.5)
        st.rerun()
except Exception:
    pass
